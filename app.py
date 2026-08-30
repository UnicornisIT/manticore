import os
import re
import csv
import copy
import json
import secrets
import tempfile
import shutil
import shlex
import subprocess
import difflib
import pandas as pd
from flask import Flask, render_template, request, send_file, redirect, url_for, session, flash, has_request_context, jsonify
from urllib.parse import urlparse
from werkzeug.exceptions import RequestEntityTooLarge
from werkzeug.security import check_password_hash, generate_password_hash
import sqlite3
import io
import hmac
from string import Formatter
from functools import wraps
from datetime import date, datetime, timezone
from pathlib import Path
from time import time
import update_app
try:
    from dotenv import load_dotenv
except ImportError:
    def load_dotenv(*args, **kwargs):
        return False

# Load environment variables from .env file
load_dotenv()

TRUTHY_ENV_VALUES = {'1', 'true', 'yes', 'on'}
FALSY_SETTING_VALUES = {'0', 'false', 'no', 'off'}

def env_bool(name, default=False):
    value = os.environ.get(name)
    if value is None:
        return default
    return value.lower() in TRUTHY_ENV_VALUES

def setting_bool(value, default=False):
    if isinstance(value, bool):
        return value
    if isinstance(value, int):
        return value != 0
    if value is None:
        return default
    normalized = str(value).strip().casefold()
    if normalized in TRUTHY_ENV_VALUES:
        return True
    if normalized in FALSY_SETTING_VALUES:
        return False
    return default

def env_int(name, default, minimum=None):
    try:
        value = int(os.environ.get(name, default))
    except (TypeError, ValueError):
        value = default
    if minimum is not None:
        value = max(minimum, value)
    return value

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY') or os.urandom(32).hex()
MAX_UPLOAD_SIZE_MB = env_int('MAX_UPLOAD_SIZE_MB', 16, minimum=1)
MAX_UPLOAD_BYTES = env_int('MAX_CONTENT_LENGTH', MAX_UPLOAD_SIZE_MB * 1024 * 1024, minimum=1024)
app.config.update(
    MAX_CONTENT_LENGTH=MAX_UPLOAD_BYTES,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE=os.environ.get('SESSION_COOKIE_SAMESITE', 'Lax'),
    SESSION_COOKIE_SECURE=env_bool('SESSION_COOKIE_SECURE', False),
)
app.config['UPLOAD_FOLDER'] = os.environ.get('UPLOAD_FOLDER', 'uploads')
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

DB_FILENAME = os.environ.get('DB_FILENAME', 'baze.db')
DB_PATH = os.path.join(app.config['UPLOAD_FOLDER'], DB_FILENAME)
APP_DIR = Path(__file__).resolve().parent
try:
    APP_VERSION = (APP_DIR / 'VERSION').read_text(encoding='utf-8-sig').strip()
except OSError:
    APP_VERSION = os.environ.get('APP_VERSION', '1.1.1')
APP_UPDATE_ENABLED = env_bool('APP_UPDATE_ENABLED', os.name != 'nt')
APP_UPDATE_STALE_SECONDS = env_int('APP_UPDATE_STALE_SECONDS', 3600, minimum=60)
ENABLE_HSTS = env_bool('ENABLE_HSTS', False)
HSTS_MAX_AGE = env_int('HSTS_MAX_AGE', 31536000, minimum=0)
HSTS_INCLUDE_SUBDOMAINS = env_bool('HSTS_INCLUDE_SUBDOMAINS', False)
HSTS_PRELOAD = env_bool('HSTS_PRELOAD', False)
ABITURIENT_UPLOAD_EXTENSIONS = {'xlsx', 'xls', 'csv'}
STUDENTS_UPLOAD_EXTENSIONS = {'xlsx', 'xls', 'csv'}
ENROLLMENT_ORDER_UPLOAD_EXTENSIONS = {'xlsx', 'xls', 'csv', 'docx', 'pdf'}
GROUPS_UPLOAD_EXTENSIONS = {'csv'}
STUDENT_TRANSFER_ORDER_EXTENSIONS = {'pdf'}
EMAIL_SOURCE_REQUIRED_COLUMNS = {
    'surname': {'фамилия'},
    'name_patronymic': {'имя отчество'},
    'email': {'почта'},
}
PENDING_ABITURIENTS_IMPORT_PREFIX = 'pending_abiturients_'
ABITURIENTS_IMPORT_RESULT_PREFIX = 'abiturients_result_'
ABITURIENTS_IMPORT_RESULT_SESSION_KEY = 'abiturients_import_result'
PENDING_STUDENTS_IMPORT_PREFIX = 'pending_students_'
PENDING_ENROLLMENT_ORDER_IMPORT_PREFIX = 'pending_enrollment_order_'
PENDING_EMAIL_SOURCE_IMPORT_PREFIX = 'pending_email_source_'
STUDENT_TRANSFER_ORDER_DIR = 'student_transfer_orders'
ENROLLMENT_FIO_SUGGESTION_THRESHOLD = 0.86
DB_BACKUP_PREFIX = 'baze_backup_'
ABITURIENT_REQUIRED_COLUMNS = {'ФИО', 'Договор'}
ABITURIENT_RESULT_COLUMNS = [
    'campaign_year', 'ФИО', 'Договор', 'login', 'Фамилия',
    'Имя_Отчество', 'import_action', 'import_status'
]
STUDENT_PREVIEW_COLUMNS = [
    '_row_number', 'username', 'email', 'firstname', 'lastname', 'cohort1', 'cohort2',
    'source_dogovor', 'source_campaign_year', 'import_action', 'import_status'
]
ENROLLMENT_ORDER_PREVIEW_COLUMNS = [
    '_row_number', 'ФИО', 'Специальность', 'Группа', 'Номер приказа',
    'Дата приказа', 'import_action', 'import_status'
]
UPLOAD_REPORT_LIMIT = 40
STUDENT_UPLOAD_REQUIRED_COLUMNS = ['username', 'password', 'email', 'firstname', 'lastname', 'cohort1']
STUDENT_UPLOAD_FIELD_LABELS = {
    'username': 'Логин',
    'password': 'Пароль',
    'email': 'Email',
    'firstname': 'Имя',
    'lastname': 'Фамилия',
    'cohort1': 'Академическая группа',
    'cohort2': 'Глобальная группа курса',
}
ENROLLMENT_ORDER_COLUMN_ALIASES = {
    'fio': {
        'фио', 'ф.и.о.', 'фамилия имя отчество', 'абитуриент', 'студент',
        'обучающийся', 'зачисляемый', 'зачислен',
    },
    'specialty': {
        'специальность', 'наименование специальности', 'направление',
        'образовательная программа', 'программа', 'профессия',
        'код специальности', 'specialty', 'spec',
    },
    'group_name': {
        'группа', 'академическая группа', 'учебная группа', 'group', 'cohort1',
    },
    'order_number': {
        'приказ', 'номер приказа', '№ приказа', 'n приказа', 'order',
        'order_number',
    },
    'order_date': {
        'дата приказа', 'дата', 'order_date',
    },
}
ENROLLMENT_ORDER_FIELD_LABELS = {
    'fio': 'ФИО',
    'specialty': 'Специальность',
    'group_name': 'Группа',
    'order_number': 'Номер приказа',
    'order_date': 'Дата приказа',
}
ROLE_LABELS = {
    'admin': 'Администратор',
    'manager': 'Куратор',
    'operator': 'Оператор',
    'assistant': 'Ассистент',
    'viewer': 'Только просмотр',
}
ARCHIVED_CAMPAIGN_MESSAGE = 'Кампания архивирована. Изменения в ней недоступны.'
WITHDRAWN_LOGIN_PREFIX = 'del'

class UploadValidationError(ValueError):
    pass

def is_withdrawn_login(login):
    return str(login or '').strip().casefold().startswith(WITHDRAWN_LOGIN_PREFIX)

def next_withdrawn_login(login, campaign_year=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    original_login = str(login or '').strip()
    used_logins = get_used_logins(campaign_year)
    number = 1
    while True:
        candidate = f'{WITHDRAWN_LOGIN_PREFIX}{number}_{original_login}'
        if candidate not in used_logins:
            return candidate
        number += 1

def format_upload_size(size_bytes):
    size_mb = size_bytes / (1024 * 1024)
    if size_mb >= 1:
        return f'{size_mb:.0f} МБ'
    return f'{max(1, size_bytes // 1024)} КБ'

def format_display_date(value):
    text = str(value or '').strip()
    if not text:
        return '-'
    for date_format in ('%Y-%m-%d %H:%M:%S', '%Y-%m-%d', '%d.%m.%Y %H:%M:%S', '%d.%m.%Y'):
        try:
            return datetime.strptime(text, date_format).strftime('%d.%m.%Y')
        except ValueError:
            pass
    return text[:10] if len(text) > 10 else text

def allowed_extensions_text(allowed_extensions):
    return ', '.join(f'.{extension}' for extension in sorted(allowed_extensions))

def get_upload_extension(file_storage):
    filename = file_storage.filename if file_storage else ''
    return os.path.splitext(filename or '')[1].lstrip('.').lower()

def validate_uploaded_file(file_storage, allowed_extensions):
    if not file_storage or not file_storage.filename:
        raise UploadValidationError('Выберите файл для загрузки.')

    extension = get_upload_extension(file_storage)
    if extension not in allowed_extensions:
        raise UploadValidationError(
            f'Неверный тип файла. Разрешены: {allowed_extensions_text(allowed_extensions)}.'
        )
    return extension

def make_temp_upload_path(extension, prefix='upload_'):
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    fd, path = tempfile.mkstemp(
        prefix=prefix,
        suffix=f'.{extension}',
        dir=app.config['UPLOAD_FOLDER']
    )
    os.close(fd)
    return path

def save_upload_to_temp(file_storage, allowed_extensions, prefix='upload_'):
    extension = validate_uploaded_file(file_storage, allowed_extensions)
    temp_path = make_temp_upload_path(extension, prefix=prefix)
    file_storage.save(temp_path)
    return temp_path

def cleanup_temp_files(*paths):
    for path in paths:
        if not path:
            continue
        try:
            os.remove(path)
        except FileNotFoundError:
            continue
        except OSError:
            app.logger.warning('Could not remove temporary file: %s', path)

def send_temp_download(file_path, download_name, mimetype):
    with open(file_path, 'rb') as file_obj:
        output = io.BytesIO(file_obj.read())
    output.seek(0)
    cleanup_temp_files(file_path)
    return send_file(output, as_attachment=True, download_name=download_name, mimetype=mimetype)

def read_csv_dataframe(file_path):
    last_error = None
    for encoding in ('utf-8-sig', 'cp1251', 'utf-8'):
        try:
            return pd.read_csv(file_path, sep=None, engine='python', encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc
            continue
    if last_error:
        raise ValueError('Не удалось прочитать CSV в кодировке UTF-8 или Windows-1251')
    raise ValueError('Не удалось прочитать CSV-файл')

def read_tabular_upload(file_path):
    extension = os.path.splitext(file_path)[1].lower()
    if extension == '.csv':
        return read_csv_dataframe(file_path)
    if extension == '.xls':
        return pd.read_excel(file_path)
    return pd.read_excel(file_path, engine="openpyxl")

_campaign_year_re = re.compile(r'^20\d{2}$')
_dogovor_year_re = re.compile(r'20\d{2}')
_email_re = re.compile(r'^[^@\s]+@[^@\s]+\.[^@\s]+$')

def is_valid_email(value):
    value = str(value or '').strip()
    return not value or bool(_email_re.fullmatch(value))

def clean_campaign_year(value, fallback):
    value = str(value or '').strip()
    if _campaign_year_re.fullmatch(value):
        return value
    return fallback

DEFAULT_CAMPAIGN_YEAR = clean_campaign_year(os.environ.get('DEFAULT_CAMPAIGN_YEAR'), str(date.today().year))
LEGACY_CAMPAIGN_YEAR = clean_campaign_year(os.environ.get('LEGACY_CAMPAIGN_YEAR'), '2026')
BASE_CAMPAIGN_YEARS = [str(y) for y in range(2020, 2031)]
MAX_GROUP_STUDENTS = 25
GROUPS_TEMPLATE_EXAMPLES = (
    ('ФМ', '11'),
    ('СД', '9'),
    ('ЛД', '11'),
)

_group_name_re = re.compile(r'^\d{2}[A-Za-zА-Яа-яЁё]+-(?:\d{1,2}(?:[A-Za-zА-Яа-яЁё])?|[A-Za-zА-Яа-яЁё]+)-\d+$')
_group_head_re = re.compile(r'^(\d{2})([A-Za-zА-Яа-яЁё]+)$')
_group_year_code_re = re.compile(r'^\s*(\d{2})')
_specialty_aliases = {
    'ФМ': 'ФМ',
    'ФАРМАЦИЯ': 'ФМ',
    '33.02.01': 'ФМ',
    'СД': 'СД',
    'СЕСТРИНСКОЕДЕЛО': 'СД',
    '34.02.01': 'СД',
    'СТО': 'СтО',
    'СТОМАТОЛОГИЯОРТОПЕДИЧЕСКАЯ': 'СтО',
    '31.02.05': 'СтО',
    'СТП': 'СтП',
    'СТОМАТОЛОГИЯПРОФИЛАКТИЧЕСКАЯ': 'СтП',
    '31.02.06': 'СтП',
    'СТД': 'СтД',
    'СТОМАТОЛОГИЧЕСКОЕДЕЛО': 'СтД',
    'СТОМАТОЛОГИЯДЕТСКАЯ': 'СтД',
    '31.02.07': 'СтД',
    'СТПР': 'СтПр',
    'АД': 'АД',
    'АКУШЕРСКОЕДЕЛО': 'АД',
    '31.02.02': 'АД',
    'ЛД': 'ЛД',
    'ЛЕЧЕБНОЕДЕЛО': 'ЛД',
    '31.02.01': 'ЛД',
    'ЛАБД': 'ЛабД',
    'ЛАБОРАТОРНАЯДИАГНОСТИКА': 'ЛабД',
    '31.02.03': 'ЛабД',
}
SPECIALTY_DISPLAY_NAMES = {
    'АД': '31.02.02 «Акушерское дело»',
    'ЛД': '31.02.01 «Лечебное дело»',
    'СД': '34.02.01 «Сестринское дело»',
    'ФМ': '33.02.01 «Фармация»',
    'СтО': '31.02.05 «Стоматология ортопедическая»',
    'СтП': '31.02.06 «Стоматология профилактическая»',
    'СтПр': '31.02.06 «Стоматология профилактическая»',
    'СтД': '31.02.07 «Стоматологическое дело»',
    'ЛабД': '31.02.03 «Лабораторная диагностика»',
}

def normalize_campaign_year(value, fallback=None):
    return clean_campaign_year(value, fallback or DEFAULT_CAMPAIGN_YEAR)

def require_campaign_year(value):
    value = str(value or '').strip()
    if not _campaign_year_re.fullmatch(value):
        raise ValueError('Введите год кампании в формате 2026.')
    return value

def normalize_group_year(value, fallback=None):
    fallback = fallback or DEFAULT_CAMPAIGN_YEAR
    value = str(value or '').strip()
    if re.fullmatch(r'\d{2}', value):
        value = f'20{value}'
    return normalize_campaign_year(value, fallback)

def infer_campaign_year(dogovor, fallback=LEGACY_CAMPAIGN_YEAR):
    match = _dogovor_year_re.search(str(dogovor or ''))
    if match:
        return normalize_campaign_year(match.group(0), fallback)
    return fallback

def infer_group_year(group_name, fallback=None):
    fallback = normalize_group_year(fallback, DEFAULT_CAMPAIGN_YEAR)
    match = _group_year_code_re.match(str(group_name or ''))
    if match:
        return normalize_group_year(match.group(1), fallback)
    return fallback

def normalize_specialty(value):
    value = re.sub(r'\s+', '', str(value or ''))
    key = value.upper().replace('Ё', 'Е')
    return _specialty_aliases.get(key, value)

def normalize_group_base(value):
    value = re.sub(r'\s+', '', str(value or '')).upper()
    return value.replace('I', 'И').replace('M', 'М')

def normalize_group_name(value):
    value = str(value or '').strip()
    value = value.replace('–', '-').replace('—', '-').replace('−', '-')
    value = re.sub(r'\s+', '', value)
    parts = value.split('-')
    if len(parts) < 2:
        return value

    head_match = _group_head_re.fullmatch(parts[0])
    if head_match:
        year_code, specialty = head_match.groups()
        parts[0] = f'{year_code}{normalize_specialty(specialty)}'

    parts[1] = normalize_group_base(parts[1])
    return '-'.join(parts)

SUPPORTED_COHORT2_SPECIALTIES = {'АД', 'ЛД', 'СД', 'СтД', 'СтО', 'СтП', 'ФМ'}
SUPPORTED_COHORT2_BASES = {'9', '11'}
COHORT2_BASE_ALIASES = {
    '9': '9',
    '9И': '9',
    '11': '11',
    '11И': '11',
    'М': '11',
}
SUPPORTED_COHORT2_VALUES = {
    f'{specialty}-{base}'
    for specialty in SUPPORTED_COHORT2_SPECIALTIES
    for base in SUPPORTED_COHORT2_BASES
}

def normalize_cohort2_base(value):
    return COHORT2_BASE_ALIASES.get(normalize_group_base(value), '')

def parse_academic_group_for_cohort2(cohort1):
    normalized = normalize_group_name(cohort1)
    parts = normalized.split('-')
    if len(parts) != 3 or not parts[2].isdigit():
        return None

    head_match = _group_head_re.fullmatch(parts[0])
    if not head_match:
        return None

    year_code, specialty = head_match.groups()
    return {
        'group_name': normalized,
        'year_code': year_code,
        'specialty': normalize_specialty(specialty),
        'base': normalize_group_base(parts[1]),
        'subgroup': parts[2],
    }

def derive_cohort2(cohort1):
    parts = parse_academic_group_for_cohort2(cohort1)
    if not parts:
        return None
    if parts['specialty'] not in SUPPORTED_COHORT2_SPECIALTIES:
        return None
    cohort2_base = normalize_cohort2_base(parts['base'])
    if not cohort2_base:
        return None
    return f"{parts['specialty']}-{cohort2_base}"

def normalize_cohort2(value):
    value = str(value or '').strip()
    if not value:
        return ''
    value = value.replace('–', '-').replace('—', '-').replace('−', '-')
    value = re.sub(r'\s+', '', value)
    parts = value.split('-')
    if len(parts) != 2:
        return value
    specialty = normalize_specialty(parts[0])
    base = normalize_cohort2_base(parts[1]) or normalize_group_base(parts[1])
    return f'{specialty}-{base}' if specialty and base else value

def is_supported_cohort2(value):
    return normalize_cohort2(value) in SUPPORTED_COHORT2_VALUES

def cohort2_quality_issue(cohort1, cohort2, rules=None):
    if not are_course_groups_enabled(rules):
        return None
    expected = derive_cohort2(cohort1)
    actual = normalize_cohort2(cohort2)
    if not expected:
        return None
    if actual != expected:
        return {
            'expected': expected,
            'actual': actual,
        }
    return None

spec_codes = {
    "ЛД": "1", "АД": "2", "СД": "3", "СтО": "4",
    "СтПр": "5", "СтП": "5", "ФМ": "6", "ЛабД": "7", "СтД": "8"
}

base_codes = {
    "2НМ": "inm", "2М": "im",
    "НМ": "nm", "М": "im",
    "11И": "11i", "9И": "9i",
    "11": "11", "9": "9", 
}

LOGIN_GROUP_SPECIALTY_CODES = {
    '1': 'ЛД',
    '2': 'АД',
    '3': 'СД',
    '4': 'СтО',
    '5': 'СтП',
    '6': 'ФМ',
    '7': 'ЛабД',
    '8': 'СтД',
}

LOGIN_GROUP_BASE_CODES = {
    '11i': '11И',
    '9i': '9И',
    '11': '11',
    '9': '9',
    'im': 'М',
    'm': 'М',
}

LOGIN_RULES_SETTINGS_ID = 1
LOGIN_TEMPLATE_FIELDS = {'yyyy', 'yy', 'spec', 'base', 'seq', 'num'}
LOGIN_MATCH_MODES = {'last_part', 'anywhere'}

DEFAULT_LOGIN_GENERATION_RULES = {
    'mode': 'standard',
    'template': '{yy}{spec}{base}{seq}',
    'number_width': 3,
    'error_prefix': 'error',
    'duplicate_prefix': 'dubl',
    'unique_scope': 'campaign',
    'year_regex': r'20\d{2}',
    'base_match_mode': 'last_part',
    'require_enrollment_order': True,
    'use_course_groups': True,
    'spec_codes': spec_codes,
    'base_codes': base_codes,
}

def get_default_login_generation_rules():
    return copy.deepcopy(DEFAULT_LOGIN_GENERATION_RULES)

def canonicalize_base_label(label):
    label = str(label or '').strip()
    folded = label.casefold()
    if folded == '11и':
        return '11И'
    if folded == '9и':
        return '9И'
    return label

def normalize_base_codes_mapping(mapping):
    normalized = {}
    for label, code in (mapping or {}).items():
        label = canonicalize_base_label(label)
        if label:
            normalized[label] = str(code).strip()
    return normalized

def merge_login_generation_rules(raw_rules=None):
    rules = get_default_login_generation_rules()
    if isinstance(raw_rules, dict):
        for key, value in raw_rules.items():
            if key == 'base_codes' and isinstance(value, dict):
                rules[key] = normalize_base_codes_mapping(value)
            elif key == 'spec_codes' and isinstance(value, dict):
                rules[key] = {str(k).strip(): str(v).strip() for k, v in value.items() if str(k).strip()}
            elif key in rules:
                rules[key] = value
    rules['base_codes'] = normalize_base_codes_mapping(rules.get('base_codes'))
    rules['mode'] = 'custom' if rules.get('mode') == 'custom' else 'standard'
    rules['template'] = str(rules.get('template') or DEFAULT_LOGIN_GENERATION_RULES['template']).strip()
    try:
        rules['number_width'] = max(1, min(8, int(rules.get('number_width', 3))))
    except (TypeError, ValueError):
        rules['number_width'] = 3
    rules['error_prefix'] = str(rules.get('error_prefix') or 'error').strip() or 'error'
    rules['duplicate_prefix'] = str(rules.get('duplicate_prefix') or 'dubl').strip() or 'dubl'
    rules['unique_scope'] = 'global' if rules.get('unique_scope') == 'global' else 'campaign'
    rules['year_regex'] = str(rules.get('year_regex') or r'20\d{2}').strip() or r'20\d{2}'
    rules['base_match_mode'] = rules.get('base_match_mode') if rules.get('base_match_mode') in LOGIN_MATCH_MODES else 'last_part'
    rules['require_enrollment_order'] = setting_bool(rules.get('require_enrollment_order'), True)
    rules['use_course_groups'] = setting_bool(rules.get('use_course_groups'), True)
    if not rules.get('spec_codes'):
        rules['spec_codes'] = copy.deepcopy(spec_codes)
    if not rules.get('base_codes'):
        rules['base_codes'] = copy.deepcopy(base_codes)
    return rules

def create_login_generation_settings_table(conn):
    conn.execute('''
        CREATE TABLE IF NOT EXISTS login_generation_settings (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            settings_json TEXT NOT NULL,
            setup_completed INTEGER NOT NULL DEFAULT 0,
            updated_at TEXT DEFAULT (datetime('now', 'localtime')),
            updated_by TEXT
        )
    ''')

def get_login_generation_settings_row():
    try:
        with sqlite3.connect(DB_PATH) as conn:
            create_login_generation_settings_table(conn)
            return conn.execute(
                '''
                SELECT settings_json, setup_completed, updated_at, updated_by
                FROM login_generation_settings
                WHERE id=?
                ''',
                (LOGIN_RULES_SETTINGS_ID,)
            ).fetchone()
    except sqlite3.Error:
        return None

def get_login_generation_rules():
    row = get_login_generation_settings_row()
    if not row:
        return get_default_login_generation_rules()
    try:
        stored_rules = json.loads(row[0] or '{}')
    except (TypeError, ValueError):
        stored_rules = {}
    return merge_login_generation_rules(stored_rules)

def is_enrollment_order_required(rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    return bool(rules.get('require_enrollment_order'))

def are_course_groups_enabled(rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    return bool(rules.get('use_course_groups'))

def get_student_course_group(cohort1, rules=None):
    if not are_course_groups_enabled(rules):
        return ''
    return derive_cohort2(cohort1) or ''

def is_login_generation_setup_completed():
    row = get_login_generation_settings_row()
    return bool(row and row[1])

def save_login_generation_settings(rules, setup_completed=True, updated_by=''):
    normalized_rules = merge_login_generation_rules(rules)
    validate_login_generation_rules(normalized_rules)
    with sqlite3.connect(DB_PATH) as conn:
        create_login_generation_settings_table(conn)
        conn.execute(
            '''
            INSERT INTO login_generation_settings (id, settings_json, setup_completed, updated_at, updated_by)
            VALUES (?, ?, ?, datetime('now', 'localtime'), ?)
            ON CONFLICT(id) DO UPDATE SET
                settings_json=excluded.settings_json,
                setup_completed=excluded.setup_completed,
                updated_at=excluded.updated_at,
                updated_by=excluded.updated_by
            ''',
            (
                LOGIN_RULES_SETTINGS_ID,
                json.dumps(normalized_rules, ensure_ascii=False, sort_keys=True),
                1 if setup_completed else 0,
                updated_by,
            )
        )
        conn.commit()
    return normalized_rules

def mapping_to_text(mapping):
    return '\n'.join(f'{key} = {value}' for key, value in (mapping or {}).items())

def parse_mapping_text(text, field_label):
    mapping = {}
    for line_number, raw_line in enumerate(str(text or '').splitlines(), start=1):
        line = raw_line.strip()
        if not line:
            continue
        separator = '=' if '=' in line else ':' if ':' in line else ''
        if not separator:
            raise ValueError(f'{field_label}: строка {line_number} должна быть в формате "значение = код".')
        key, value = [part.strip() for part in line.split(separator, 1)]
        if not key or not value:
            raise ValueError(f'{field_label}: строка {line_number} содержит пустое значение или код.')
        if re.search(r'\s', value):
            raise ValueError(f'{field_label}: код в строке {line_number} не должен содержать пробелы.')
        mapping[key] = value
    if not mapping:
        raise ValueError(f'{field_label}: добавьте хотя бы одну строку.')
    return mapping

def validate_login_template(template):
    template = str(template or '').strip()
    if not template:
        raise ValueError('Шаблон логина не может быть пустым.')
    formatter = Formatter()
    try:
        parsed_parts = list(formatter.parse(template))
    except ValueError as exc:
        raise ValueError(f'Шаблон логина заполнен некорректно: {exc}') from exc
    used_fields = set()
    for _literal, field_name, format_spec, conversion in parsed_parts:
        if field_name is None:
            continue
        if conversion:
            raise ValueError('Шаблон логина не должен использовать преобразования вида !r или !s.')
        if field_name not in LOGIN_TEMPLATE_FIELDS:
            allowed = ', '.join(sorted(f'{{{field}}}' for field in LOGIN_TEMPLATE_FIELDS))
            raise ValueError(f'Неизвестное поле {{{field_name}}}. Доступны: {allowed}.')
        used_fields.add(field_name)
        if format_spec:
            for _spec_literal, nested_field, _nested_spec, _nested_conversion in formatter.parse(format_spec):
                if nested_field:
                    raise ValueError('Шаблон логина не должен содержать вложенные поля форматирования.')
    if 'seq' not in used_fields and 'num' not in used_fields:
        raise ValueError('Шаблон логина должен содержать {seq} или {num}.')
    return template

def validate_login_generation_rules(rules):
    rules = merge_login_generation_rules(rules)
    validate_login_template(rules['template'])
    re.compile(rules['year_regex'])
    if not rules['spec_codes']:
        raise ValueError('Добавьте хотя бы один код специальности.')
    if not rules['base_codes']:
        raise ValueError('Добавьте хотя бы один код базы.')
    for prefix_label, prefix in (('Префикс ошибок', rules['error_prefix']), ('Префикс дублей', rules['duplicate_prefix'])):
        if re.search(r'\s', prefix):
            raise ValueError(f'{prefix_label} не должен содержать пробелы.')
    render_login_template(
        rules,
        {'year': '2026', 'year_code': '26', 'spec_code': '1', 'base_code': '11'},
        1
    )
    return True

def form_checkbox_bool(form, name, default=False):
    values = form.getlist(name) if hasattr(form, 'getlist') else []
    if values:
        return any(setting_bool(value, False) for value in values)
    return default

def build_login_rules_from_form(form, mode='custom'):
    if mode == 'standard':
        rules = get_default_login_generation_rules()
        rules['require_enrollment_order'] = form_checkbox_bool(form, 'require_enrollment_order', False)
        rules['use_course_groups'] = form_checkbox_bool(form, 'use_course_groups', True)
        return rules
    return merge_login_generation_rules({
        'mode': 'custom',
        'template': form.get('template'),
        'number_width': form.get('number_width'),
        'error_prefix': form.get('error_prefix'),
        'duplicate_prefix': form.get('duplicate_prefix'),
        'unique_scope': form.get('unique_scope'),
        'year_regex': form.get('year_regex'),
        'base_match_mode': form.get('base_match_mode'),
        'require_enrollment_order': form_checkbox_bool(form, 'require_enrollment_order', False),
        'use_course_groups': form_checkbox_bool(form, 'use_course_groups', True),
        'spec_codes': parse_mapping_text(form.get('spec_codes_text'), 'Коды специальностей'),
        'base_codes': parse_mapping_text(form.get('base_codes_text'), 'Коды базы'),
    })

def render_login_template(rules, parts, number):
    rules = merge_login_generation_rules(rules)
    width = rules['number_width']
    number = int(number)
    context = {
        'yyyy': parts['year'],
        'yy': parts['year_code'],
        'spec': parts['spec_code'],
        'base': parts['base_code'],
        'seq': f'{number:0{width}d}',
        'num': number,
    }
    return rules['template'].format(**context)

def get_login_rules_form_context(rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    sample_spec = next(iter(rules.get('spec_codes') or {'ФМ': '6'}))
    sample_base = next(iter(rules.get('base_codes') or {'11': '11'}))
    sample_dogovor = f'2026-{sample_spec}-0001-{sample_base}'
    sample_parts = parse_dogovor_parts(sample_dogovor, rules)
    sample_login = render_login_template(rules, sample_parts, 1) if sample_parts else ''
    return {
        'rules': rules,
        'setup_completed': is_login_generation_setup_completed(),
        'spec_codes_text': mapping_to_text(rules.get('spec_codes')),
        'base_codes_text': mapping_to_text(rules.get('base_codes')),
        'sample_dogovor': sample_dogovor,
        'sample_login': sample_login,
        'enrollment_order_required': is_enrollment_order_required(rules),
        'course_groups_enabled': are_course_groups_enabled(rules),
    }

_dogovor_latin_lookalikes = str.maketrans({
    'A': 'А',
    'B': 'В',
    'C': 'С',
    'E': 'Е',
    'H': 'Н',
    'I': 'И',
    'K': 'К',
    'M': 'М',
    'O': 'О',
    'P': 'Р',
    'T': 'Т',
    'X': 'Х',
})
_dogovor_dash_re = re.compile(r'[\u2010-\u2015\u2212]')
_dogovor_individual_base_re = re.compile(r'\b(9|11)[иi]\b', re.IGNORECASE)

def normalize_dogovor_text(dogovor):
    normalized = str(dogovor or '').strip()
    normalized = _dogovor_dash_re.sub('-', normalized).replace(' ', '-')
    return normalized.upper().translate(_dogovor_latin_lookalikes)

def normalize_dogovor_storage_text(dogovor):
    value = str(dogovor or '').strip()
    return _dogovor_individual_base_re.sub(lambda match: f'{match.group(1)}И', value)

def find_mapped_code(normalized_text, mapping):
    for label, code in sorted((mapping or {}).items(), key=lambda item: len(item[0]), reverse=True):
        if str(label).upper() in normalized_text:
            return str(label), str(code)
    return None, None

def find_base_code(normalized_text, mapping, match_mode='last_part'):
    if match_mode == 'anywhere':
        return find_mapped_code(normalized_text, mapping)

    parts = normalized_text.split('-')
    if len(parts) < 2:
        return None, None
    last_part = parts[-1].strip()
    for label, code in sorted((mapping or {}).items(), key=lambda item: len(item[0]), reverse=True):
        if last_part == str(label).upper():
            return str(label), str(code)
    return None, None

def parse_dogovor_parts(dogovor, rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    normalized = normalize_dogovor_text(dogovor)
    year_match = re.search(rules['year_regex'], normalized)
    spec_label, spec_code = find_mapped_code(normalized, rules['spec_codes'])
    base_label, base_code = find_base_code(normalized, rules['base_codes'], rules['base_match_mode'])

    if not (year_match and spec_code and base_code):
        return None

    year = year_match.group()
    return {
        'year': year,
        'year_code': year[-2:],
        'spec_label': spec_label,
        'spec_code': spec_code,
        'base_label': base_label,
        'base_code': base_code,
    }

def build_login_from_parts(parts, number, rules=None):
    return render_login_template(rules or get_login_generation_rules(), parts, number)

def next_login_from_parts(parts, existing_logins, rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    used_login_keys = {
        str(existing_login or '').strip().casefold()
        for existing_login in existing_logins
        if str(existing_login or '').strip()
    }
    number = 1
    while True:
        login = build_login_from_parts(parts, number, rules)
        if login.casefold() not in used_login_keys:
            return login
        number += 1

def parse_dogovor(dogovor, rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    parts = parse_dogovor_parts(dogovor, rules)
    if not parts:
        return rules['error_prefix']
    return f"{parts['year_code']}{parts['spec_code']}{parts['base_code']}"

def split_fio(fio):
    fio = ' '.join(str(fio or '').split())
    if not fio:
        return '', '', ''
    fam, imotch = fio.split(' ', 1) if ' ' in fio else (fio, '')
    return fio, fam, imotch

def get_table_columns(conn, table):
    cur = conn.execute(f'PRAGMA table_info({table})')
    return [row[1] for row in cur.fetchall()]

def table_exists(conn, table):
    cur = conn.execute(
        "SELECT 1 FROM sqlite_master WHERE type='table' AND name=?",
        (table,)
    )
    return cur.fetchone() is not None

def get_unique_table_name(conn, base_name):
    existing_tables = {
        row[0] for row in conn.execute("SELECT name FROM sqlite_master WHERE type='table'")
    }
    if base_name not in existing_tables:
        return base_name

    suffix = 2
    while f'{base_name}_{suffix}' in existing_tables:
        suffix += 1
    return f'{base_name}_{suffix}'

def create_abiturients_table(conn):
    conn.execute(f'''
        CREATE TABLE IF NOT EXISTS abiturients (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            fio TEXT,
            dogovor TEXT,
            login TEXT,
            campaign_year TEXT NOT NULL DEFAULT '{LEGACY_CAMPAIGN_YEAR}',
            fam TEXT,
            imotch TEXT,
            email TEXT,
            paid INTEGER DEFAULT 0,
            comment TEXT,
            created_at TEXT DEFAULT (datetime('now', 'localtime'))
        )
    ''')

def migrate_existing_enrollment_orders_to_uploads(conn):
    if not table_exists(conn, 'enrollment_orders') or not table_exists(conn, 'enrollment_order_upload_rows'):
        return
    campaigns = [
        row[0]
        for row in conn.execute(
            '''
            SELECT DISTINCT campaign_year
            FROM enrollment_orders
            WHERE campaign_year IS NOT NULL AND campaign_year <> ''
            '''
        ).fetchall()
        if row[0]
    ]
    for campaign_year in campaigns:
        if conn.execute(
            'SELECT 1 FROM enrollment_order_upload_rows WHERE campaign_year=? LIMIT 1',
            (campaign_year,)
        ).fetchone():
            continue
        rows = conn.execute(
            '''
            SELECT id, fio, specialty, group_name, order_number, order_date,
                   fio_key, specialty_key, group_key
            FROM enrollment_orders
            WHERE campaign_year=?
            ORDER BY id
            ''',
            (campaign_year,)
        ).fetchall()
        if not rows:
            continue
        order_numbers = []
        order_dates = []
        for row in rows:
            order_number = str(row[4] or '').strip()
            order_date = str(row[5] or '').strip()
            if order_number and order_number not in order_numbers:
                order_numbers.append(order_number)
            if order_date and order_date not in order_dates:
                order_dates.append(order_date)
        cur = conn.execute(
            '''
            INSERT INTO enrollment_order_uploads
                (campaign_year, original_filename, total_count, import_count,
                 order_numbers, order_dates)
            VALUES (?, ?, ?, ?, ?, ?)
            ''',
            (
                campaign_year, 'Ранее загруженные приказы', len(rows), len(rows),
                ', '.join(order_numbers), ', '.join(order_dates)
            )
        )
        upload_id = cur.lastrowid
        for row_number, row in enumerate(rows, start=2):
            _row_id, fio, specialty, group_name, order_number, order_date, fio_key, specialty_key, group_key = row
            conn.execute(
                '''
                INSERT INTO enrollment_order_upload_rows
                    (upload_id, campaign_year, row_number, fio, specialty, group_name,
                     order_number, order_date, fio_key, specialty_key, group_key,
                     has_candidate, import_action, import_status)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0, 'import', ?)
                ''',
                (
                    upload_id, campaign_year, row_number, fio, specialty, group_name,
                    order_number, order_date, fio_key, specialty_key, group_key,
                    'Перенесено из прежней версии журнала приказов'
                )
            )
        conn.execute(
            'UPDATE enrollment_orders SET upload_id=? WHERE campaign_year=? AND upload_id IS NULL',
            (upload_id, campaign_year)
        )

def create_enrollment_orders_table(conn):
    conn.execute(f'''
        CREATE TABLE IF NOT EXISTS enrollment_orders (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            upload_id INTEGER,
            campaign_year TEXT NOT NULL DEFAULT '{DEFAULT_CAMPAIGN_YEAR}',
            fio TEXT,
            specialty TEXT,
            group_name TEXT,
            order_number TEXT,
            order_date TEXT,
            fio_key TEXT,
            specialty_key TEXT,
            group_key TEXT,
            created_at TEXT DEFAULT (datetime('now', 'localtime'))
        )
    ''')
    conn.execute('''
        CREATE INDEX IF NOT EXISTS idx_enrollment_orders_campaign_keys
        ON enrollment_orders (campaign_year, fio_key, specialty_key)
    ''')
    columns = get_table_columns(conn, 'enrollment_orders')
    if 'upload_id' not in columns:
        conn.execute('ALTER TABLE enrollment_orders ADD COLUMN upload_id INTEGER')
    create_enrollment_order_uploads_table(conn)
    migrate_existing_enrollment_orders_to_uploads(conn)

def create_enrollment_order_uploads_table(conn):
    conn.execute(f'''
        CREATE TABLE IF NOT EXISTS enrollment_order_uploads (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            campaign_year TEXT NOT NULL DEFAULT '{DEFAULT_CAMPAIGN_YEAR}',
            original_filename TEXT,
            stored_filename TEXT,
            uploaded_by TEXT,
            total_count INTEGER DEFAULT 0,
            import_count INTEGER DEFAULT 0,
            matched_count INTEGER DEFAULT 0,
            unmatched_count INTEGER DEFAULT 0,
            duplicate_count INTEGER DEFAULT 0,
            skipped_count INTEGER DEFAULT 0,
            order_numbers TEXT,
            order_dates TEXT,
            created_at TEXT DEFAULT (datetime('now', 'localtime'))
        )
    ''')
    conn.execute(f'''
        CREATE TABLE IF NOT EXISTS enrollment_order_upload_rows (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            upload_id INTEGER NOT NULL,
            campaign_year TEXT NOT NULL DEFAULT '{DEFAULT_CAMPAIGN_YEAR}',
            row_number INTEGER,
            fio TEXT,
            specialty TEXT,
            group_name TEXT,
            order_number TEXT,
            order_date TEXT,
            fio_key TEXT,
            specialty_key TEXT,
            group_key TEXT,
            has_candidate INTEGER DEFAULT 0,
            import_action TEXT,
            import_status TEXT,
            created_at TEXT DEFAULT (datetime('now', 'localtime'))
        )
    ''')
    conn.execute('''
        CREATE INDEX IF NOT EXISTS idx_enrollment_order_uploads_campaign
        ON enrollment_order_uploads (campaign_year, created_at)
    ''')
    conn.execute('''
        CREATE INDEX IF NOT EXISTS idx_enrollment_order_upload_rows_upload
        ON enrollment_order_upload_rows (upload_id)
    ''')
    conn.execute('''
        CREATE INDEX IF NOT EXISTS idx_enrollment_order_upload_rows_campaign_keys
        ON enrollment_order_upload_rows (campaign_year, fio_key, specialty_key)
    ''')
    row_columns = get_table_columns(conn, 'enrollment_order_upload_rows')
    row_extra_columns = {
        'fio_review_status': "TEXT DEFAULT ''",
        'fio_review_candidate_id': 'INTEGER',
        'fio_reviewed_at': 'TEXT',
        'fio_reviewed_by': 'TEXT',
    }
    for column, definition in row_extra_columns.items():
        if column not in row_columns:
            conn.execute(
                f'ALTER TABLE enrollment_order_upload_rows ADD COLUMN {column} {definition}'
            )

def create_enrollment_candidates_table(conn):
    conn.execute(f'''
        CREATE TABLE IF NOT EXISTS enrollment_candidates (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            abiturient_id INTEGER,
            campaign_year TEXT NOT NULL DEFAULT '{DEFAULT_CAMPAIGN_YEAR}',
            fio TEXT,
            dogovor TEXT,
            login TEXT,
            fam TEXT,
            imotch TEXT,
            email TEXT,
            specialty TEXT,
            specialty_key TEXT,
            base_label TEXT,
            verification_status TEXT DEFAULT 'waiting_order',
            matched_order_id INTEGER,
            order_group_name TEXT,
            verified_at TEXT,
            created_at TEXT DEFAULT (datetime('now', 'localtime'))
        )
    ''')
    columns = get_table_columns(conn, 'enrollment_candidates')
    extra_columns = {
        'abiturient_id': 'INTEGER',
        'fam': 'TEXT',
        'imotch': 'TEXT',
        'email': 'TEXT',
        'specialty': 'TEXT',
        'specialty_key': 'TEXT',
        'base_label': 'TEXT',
        'verification_status': "TEXT DEFAULT 'waiting_order'",
        'matched_order_id': 'INTEGER',
        'order_group_name': 'TEXT',
        'verified_at': 'TEXT',
    }
    for column, column_type in extra_columns.items():
        if column not in columns:
            conn.execute(f'ALTER TABLE enrollment_candidates ADD COLUMN {column} {column_type}')
    conn.execute('''
        CREATE INDEX IF NOT EXISTS idx_enrollment_candidates_campaign_login
        ON enrollment_candidates (campaign_year, login)
    ''')
    conn.execute('''
        CREATE INDEX IF NOT EXISTS idx_enrollment_candidates_campaign_status
        ON enrollment_candidates (campaign_year, verification_status)
    ''')

def migrate_abiturients_table(conn):
    columns = get_table_columns(conn, 'abiturients')
    if not columns:
        create_abiturients_table(conn)
        return

    table_sql_row = conn.execute(
        "SELECT sql FROM sqlite_master WHERE type='table' AND name='abiturients'"
    ).fetchone()
    table_sql = re.sub(r'\s+', ' ', (table_sql_row[0] if table_sql_row else '').lower())
    has_global_login_unique = 'login text unique' in table_sql

    if 'campaign_year' in columns and not has_global_login_unique:
        if 'paid' not in columns:
            conn.execute('ALTER TABLE abiturients ADD COLUMN paid INTEGER DEFAULT 0')
        conn.execute(
            "UPDATE abiturients SET campaign_year=? WHERE campaign_year IS NULL OR campaign_year=''",
            (LEGACY_CAMPAIGN_YEAR,)
        )
        conn.execute('''
            CREATE UNIQUE INDEX IF NOT EXISTS idx_abiturients_campaign_login
            ON abiturients (campaign_year, login)
        ''')
        return

    existing_tables = {
        row[0] for row in conn.execute("SELECT name FROM sqlite_master WHERE type='table'")
    }
    backup_table = 'abiturients_legacy_campaign_migration'
    suffix = 1
    while backup_table in existing_tables:
        suffix += 1
        backup_table = f'abiturients_legacy_campaign_migration_{suffix}'

    conn.execute(f'ALTER TABLE abiturients RENAME TO {backup_table}')
    create_abiturients_table(conn)

    old_columns = get_table_columns(conn, backup_table)
    copy_columns = [
        'id', 'fio', 'dogovor', 'login', 'campaign_year',
        'fam', 'imotch', 'email', 'paid', 'comment', 'created_at'
    ]
    selectable_columns = [column for column in copy_columns if column in old_columns]
    if selectable_columns:
        cur = conn.execute(f'SELECT {", ".join(selectable_columns)} FROM {backup_table}')
        for values in cur.fetchall():
            row = dict(zip(selectable_columns, values))
            campaign_year = normalize_campaign_year(
                row.get('campaign_year'),
                infer_campaign_year(row.get('dogovor'))
            )
            conn.execute(
                '''
                INSERT OR IGNORE INTO abiturients
                    (id, fio, dogovor, login, campaign_year, fam, imotch, email, comment, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    row.get('id'),
                    row.get('fio'),
                    row.get('dogovor'),
                    row.get('login'),
                    campaign_year,
                    row.get('fam'),
                    row.get('imotch'),
                    row.get('email'),
                    row.get('comment'),
                    row.get('created_at'),
                )
            )
    conn.execute(f'DROP TABLE {backup_table}')
    conn.execute('''
        CREATE UNIQUE INDEX IF NOT EXISTS idx_abiturients_campaign_login
        ON abiturients (campaign_year, login)
    ''')

def migrate_legacy_students_abiturients_table(conn):
    if not table_exists(conn, 'students'):
        return

    columns = get_table_columns(conn, 'students')
    legacy_columns = {'fio', 'dogovor', 'login', 'fam', 'imotch'}
    moodle_columns = {'username', 'password', 'firstname', 'lastname', 'cohort1'}
    if not legacy_columns.issubset(columns) or moodle_columns.issubset(columns):
        return

    backup_table = get_unique_table_name(conn, 'students_legacy_abiturients_backup')
    conn.execute(f'ALTER TABLE students RENAME TO {backup_table}')

    old_columns = get_table_columns(conn, backup_table)
    copy_columns = [
        'id', 'fio', 'dogovor', 'login', 'fam',
        'imotch', 'email', 'paid', 'comment', 'created_at'
    ]
    selectable_columns = [column for column in copy_columns if column in old_columns]
    if not selectable_columns:
        return

    preserve_ids = conn.execute('SELECT COUNT(*) FROM abiturients').fetchone()[0] == 0
    cur = conn.execute(f'SELECT {", ".join(selectable_columns)} FROM {backup_table}')
    for values in cur.fetchall():
        row = dict(zip(selectable_columns, values))
        campaign_year = infer_campaign_year(row.get('dogovor'))
        if preserve_ids and 'id' in selectable_columns:
            conn.execute(
                '''
                INSERT OR IGNORE INTO abiturients
                    (id, fio, dogovor, login, campaign_year, fam, imotch, email, comment, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    row.get('id'),
                    row.get('fio'),
                    row.get('dogovor'),
                    row.get('login'),
                    campaign_year,
                    row.get('fam'),
                    row.get('imotch'),
                    row.get('email'),
                    row.get('comment'),
                    row.get('created_at'),
                )
            )
        else:
            conn.execute(
                '''
                INSERT OR IGNORE INTO abiturients
                    (fio, dogovor, login, campaign_year, fam, imotch, email, comment, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    row.get('fio'),
                    row.get('dogovor'),
                    row.get('login'),
                    campaign_year,
                    row.get('fam'),
                    row.get('imotch'),
                    row.get('email'),
                    row.get('comment'),
                    row.get('created_at'),
                )
            )

def ensure_campaign_column(conn, table):
    columns = get_table_columns(conn, table)
    if not columns:
        return

    added_column = False
    if 'campaign_year' not in columns:
        conn.execute(
            f"ALTER TABLE {table} ADD COLUMN campaign_year TEXT DEFAULT '{LEGACY_CAMPAIGN_YEAR}'"
        )
        added_column = True

    cur = conn.execute(f'SELECT id, dogovor, campaign_year FROM {table}')
    for row_id, dogovor, campaign_year in cur.fetchall():
        if added_column:
            fixed_year = infer_campaign_year(dogovor)
        else:
            fixed_year = normalize_campaign_year(campaign_year, infer_campaign_year(dogovor))
        if fixed_year != campaign_year:
            conn.execute(f'UPDATE {table} SET campaign_year=? WHERE id=?', (fixed_year, row_id))

def ensure_students_origin_columns(conn):
    columns = get_table_columns(conn, 'students')
    if not columns:
        return

    student_columns = {
        'cohort2': 'TEXT',
        'source_campaign_year': 'TEXT',
        'source_dogovor': 'TEXT',
        'source_fio': 'TEXT',
    }
    for column, column_type in student_columns.items():
        if column not in columns:
            conn.execute(f'ALTER TABLE students ADD COLUMN {column} {column_type}')
            columns.append(column)
    backfill_students_cohort2(conn)

def backfill_students_cohort2(conn):
    columns = get_table_columns(conn, 'students')
    if not {'cohort1', 'cohort2'}.issubset(columns):
        return

    rows = conn.execute(
        '''
        SELECT rowid, cohort1
        FROM students
        WHERE (cohort2 IS NULL OR TRIM(cohort2)='')
          AND cohort1 IS NOT NULL
          AND TRIM(cohort1)<>''
        '''
    ).fetchall()
    for rowid, cohort1 in rows:
        cohort2 = derive_cohort2(cohort1)
        if cohort2:
            conn.execute('UPDATE students SET cohort2=? WHERE rowid=?', (cohort2, rowid))

def ensure_students_duplicates_columns(conn):
    columns = get_table_columns(conn, 'students_duplicates')
    if not columns:
        return
    if 'cohort2' not in columns:
        conn.execute('ALTER TABLE students_duplicates ADD COLUMN cohort2 TEXT')

def ensure_group_year_column(conn):
    if not table_exists(conn, 'groups'):
        return

    columns = get_table_columns(conn, 'groups')
    if not columns:
        return

    if 'group_year' not in columns:
        conn.execute('ALTER TABLE groups ADD COLUMN group_year TEXT')
    if 'is_hidden' not in columns:
        conn.execute('ALTER TABLE groups ADD COLUMN is_hidden INTEGER DEFAULT 0')

    cur = conn.execute('SELECT id, name, group_year, is_hidden FROM groups')
    for row_id, name, group_year, is_hidden in cur.fetchall():
        fixed_year = normalize_group_year(group_year, infer_group_year(name, DEFAULT_CAMPAIGN_YEAR))
        if fixed_year != group_year:
            conn.execute('UPDATE groups SET group_year=? WHERE id=?', (fixed_year, row_id))
        if is_hidden not in (0, 1):
            conn.execute('UPDATE groups SET is_hidden=0 WHERE id=?', (row_id,))

    conn.execute('CREATE INDEX IF NOT EXISTS idx_groups_group_year_name ON groups (group_year, name)')

PASSWORD_HASH_PREFIXES = ('scrypt:', 'pbkdf2:', 'argon2:')
MIN_PASSWORD_LENGTH = 8
LOGIN_MAX_ATTEMPTS = int(os.environ.get('LOGIN_MAX_ATTEMPTS', '5'))
LOGIN_WINDOW_SECONDS = int(os.environ.get('LOGIN_WINDOW_SECONDS', '600'))
CSRF_SESSION_KEY = 'csrf_token'
CSRF_FORM_FIELD = 'csrf_token'

@app.context_processor
def inject_template_globals():
    def url_for_with_query(endpoint, **updates):
        args = request.args.to_dict(flat=True) if has_request_context() else {}
        for key, value in updates.items():
            if value is None:
                args.pop(key, None)
            else:
                args[key] = value
        args = {key: value for key, value in args.items() if value not in (None, '')}
        return url_for(endpoint, **args)

    return {
        'app_version': APP_VERSION,
        'csrf_token': get_csrf_token,
        'role_labels': ROLE_LABELS,
        'is_campaign_archived': is_campaign_archived,
        'url_for_with_query': url_for_with_query,
        'format_display_date': format_display_date,
        'course_groups_enabled': are_course_groups_enabled(),
        'is_withdrawn_login': is_withdrawn_login,
    }

def request_uses_https():
    forwarded_proto = request.headers.get('X-Forwarded-Proto', '')
    forwarded_proto = forwarded_proto.split(',')[0].strip().lower()
    return request.is_secure or forwarded_proto == 'https'

def build_hsts_header():
    parts = [f'max-age={HSTS_MAX_AGE}']
    if HSTS_INCLUDE_SUBDOMAINS:
        parts.append('includeSubDomains')
    if HSTS_PRELOAD:
        parts.append('preload')
    return '; '.join(parts)

@app.after_request
def add_security_headers(response):
    response.headers.setdefault('X-Content-Type-Options', 'nosniff')
    response.headers.setdefault('X-Frame-Options', 'SAMEORIGIN')
    response.headers.setdefault('Referrer-Policy', 'strict-origin-when-cross-origin')
    if ENABLE_HSTS and request_uses_https():
        response.headers.setdefault('Strict-Transport-Security', build_hsts_header())
    if request.endpoint != 'static' and response.mimetype == 'text/html':
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    return response

@app.errorhandler(RequestEntityTooLarge)
def handle_upload_too_large(error):
    flash(f'Файл слишком большой. Максимальный размер: {format_upload_size(MAX_UPLOAD_BYTES)}.', 'error')
    return redirect(get_safe_referrer(default_endpoint='file_work'), code=303)

def get_safe_referrer(default_endpoint='index'):
    referrer = request.referrer
    if referrer:
        parsed = urlparse(referrer)
        if not parsed.netloc or parsed.netloc == request.host:
            return referrer
    return url_for(default_endpoint if 'user' in session else 'login')

def get_csrf_token():
    token = session.get(CSRF_SESSION_KEY)
    if not token:
        token = secrets.token_urlsafe(32)
        session[CSRF_SESSION_KEY] = token
    return token

def refresh_csrf_token():
    session[CSRF_SESSION_KEY] = secrets.token_urlsafe(32)
    return session[CSRF_SESSION_KEY]

def validate_csrf_token():
    expected = session.get(CSRF_SESSION_KEY)
    actual = request.form.get(CSRF_FORM_FIELD) or request.headers.get('X-CSRF-Token') or ''
    return bool(expected) and hmac.compare_digest(expected, actual)

def get_login_csrf_token():
    return get_csrf_token()

def refresh_login_csrf_token():
    return refresh_csrf_token()

def validate_login_csrf_token():
    return validate_csrf_token()

@app.before_request
def protect_post_requests_with_csrf():
    if request.method != 'POST':
        return None
    if validate_csrf_token():
        return None
    refresh_csrf_token()
    flash('Сессия формы устарела. Попробуйте ещё раз.', 'error')
    return redirect(get_safe_referrer(default_endpoint='index'), code=303)

SETUP_ALLOWED_ENDPOINTS = {
    'static', 'login', 'logout', 'register', 'login_generation_setup', 'documentation'
}

@app.before_request
def require_login_generation_setup():
    if request.endpoint in SETUP_ALLOWED_ENDPOINTS or not session.get('user'):
        return None
    if is_login_generation_setup_completed():
        return None
    if session.get('role') == 'admin':
        return redirect(url_for('login_generation_setup'), code=303)
    session.clear()
    flash('Администратор должен завершить первичную настройку правил логинов.', 'error')
    return redirect(url_for('login'), code=303)

def get_client_ip():
    forwarded_for = request.headers.get('X-Forwarded-For', '')
    if forwarded_for:
        return forwarded_for.split(',')[0].strip()
    return request.remote_addr or 'unknown'

def sanitize_backup_reason(reason):
    safe_reason = re.sub(r'[^a-zA-Z0-9_-]+', '_', str(reason or 'manual')).strip('_')
    return safe_reason[:60] or 'manual'

def create_database_backup(reason='manual'):
    if not os.path.exists(DB_PATH):
        return None
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
    safe_reason = sanitize_backup_reason(reason)
    backup_name = f'{DB_BACKUP_PREFIX}{safe_reason}_{timestamp}.db'
    backup_path = os.path.join(app.config['UPLOAD_FOLDER'], backup_name)
    shutil.copy2(DB_PATH, backup_path)
    return backup_path

def list_database_backups():
    upload_folder = app.config['UPLOAD_FOLDER']
    if not os.path.isdir(upload_folder):
        return []
    backups = []
    for name in os.listdir(upload_folder):
        if not name.startswith(DB_BACKUP_PREFIX) or not name.endswith('.db'):
            continue
        path = os.path.join(upload_folder, name)
        if not os.path.isfile(path):
            continue
        stat = os.stat(path)
        backups.append({
            'name': name,
            'size': stat.st_size,
            'size_text': format_upload_size(stat.st_size),
            'created_at': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
        })
    return sorted(backups, key=lambda item: item['created_at'], reverse=True)

def get_backup_path(backup_name):
    backup_name = os.path.basename(str(backup_name or ''))
    if not backup_name.startswith(DB_BACKUP_PREFIX) or not backup_name.endswith('.db'):
        raise ValueError('Некорректное имя резервной копии')
    upload_root = os.path.abspath(app.config['UPLOAD_FOLDER'])
    backup_path = os.path.abspath(os.path.join(upload_root, backup_name))
    if os.path.commonpath([upload_root, backup_path]) != upload_root:
        raise ValueError('Некорректный путь резервной копии')
    if not os.path.exists(backup_path):
        raise FileNotFoundError('Резервная копия не найдена')
    return backup_path

def create_audit_log_table(conn):
    conn.execute('''
        CREATE TABLE IF NOT EXISTS audit_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT,
            action TEXT NOT NULL,
            entity_type TEXT,
            entity_id TEXT,
            details TEXT,
            ip_address TEXT,
            created_at TEXT DEFAULT (datetime('now', 'localtime'))
        )
    ''')

def create_student_group_transfers_table(conn):
    conn.execute('''
        CREATE TABLE IF NOT EXISTS student_group_transfers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            movement_type TEXT DEFAULT 'transfer',
            old_cohort1 TEXT,
            old_cohort2 TEXT,
            new_cohort1 TEXT NOT NULL,
            new_cohort2 TEXT,
            enrollment_order_id INTEGER,
            enrollment_order_upload_id INTEGER,
            order_number TEXT,
            order_date TEXT,
            order_source TEXT,
            order_filename TEXT,
            order_original_filename TEXT,
            order_mime_type TEXT,
            order_size INTEGER,
            created_by TEXT,
            created_at TEXT DEFAULT (datetime('now', 'localtime'))
        )
    ''')
    conn.execute('''
        CREATE INDEX IF NOT EXISTS idx_student_group_transfers_username_created
        ON student_group_transfers (username, created_at)
    ''')
    columns = get_table_columns(conn, 'student_group_transfers')
    extra_columns = {
        'movement_type': "TEXT DEFAULT 'transfer'",
        'enrollment_order_id': 'INTEGER',
        'enrollment_order_upload_id': 'INTEGER',
        'order_number': 'TEXT',
        'order_date': 'TEXT',
        'order_source': 'TEXT',
    }
    for column, column_type in extra_columns.items():
        if column not in columns:
            conn.execute(f'ALTER TABLE student_group_transfers ADD COLUMN {column} {column_type}')
    conn.execute(
        '''
        UPDATE student_group_transfers
        SET movement_type='transfer'
        WHERE movement_type IS NULL OR movement_type=''
        '''
    )

def create_campaign_settings_table(conn):
    conn.execute('''
        CREATE TABLE IF NOT EXISTS campaign_settings (
            campaign_year TEXT PRIMARY KEY,
            is_archived INTEGER DEFAULT 0,
            archived_at TEXT,
            archived_by TEXT,
            created_at TEXT,
            created_by TEXT,
            is_active INTEGER DEFAULT 0,
            active_at TEXT,
            active_by TEXT
        )
    ''')
    columns = get_table_columns(conn, 'campaign_settings')
    if 'created_at' not in columns:
        conn.execute('ALTER TABLE campaign_settings ADD COLUMN created_at TEXT')
    if 'created_by' not in columns:
        conn.execute('ALTER TABLE campaign_settings ADD COLUMN created_by TEXT')
    if 'is_active' not in columns:
        conn.execute('ALTER TABLE campaign_settings ADD COLUMN is_active INTEGER DEFAULT 0')
        conn.execute('UPDATE campaign_settings SET is_active=0 WHERE is_active IS NULL')
    if 'active_at' not in columns:
        conn.execute('ALTER TABLE campaign_settings ADD COLUMN active_at TEXT')
    if 'active_by' not in columns:
        conn.execute('ALTER TABLE campaign_settings ADD COLUMN active_by TEXT')
    conn.execute('''
        CREATE UNIQUE INDEX IF NOT EXISTS idx_campaign_settings_single_active
        ON campaign_settings (is_active)
        WHERE is_active=1
    ''')

def log_action(action, entity_type='', entity_id='', details='', conn=None):
    username = session.get('user', '') if has_request_context() else ''
    ip_address = get_client_ip() if has_request_context() else ''
    should_close = conn is None
    if should_close:
        conn = sqlite3.connect(DB_PATH)
    try:
        create_audit_log_table(conn)
        conn.execute(
            '''
            INSERT INTO audit_logs (username, action, entity_type, entity_id, details, ip_address)
            VALUES (?, ?, ?, ?, ?, ?)
            ''',
            (username, action, entity_type, str(entity_id or ''), str(details or ''), ip_address)
        )
        if should_close:
            conn.commit()
    finally:
        if should_close:
            conn.close()

def get_audit_logs(limit=200):
    limit = max(1, min(int(limit or 200), 1000))
    with sqlite3.connect(DB_PATH) as conn:
        create_audit_log_table(conn)
        cur = conn.execute(
            '''
            SELECT username, action, entity_type, entity_id, details, ip_address, created_at
            FROM audit_logs
            ORDER BY id DESC
            LIMIT ?
            ''',
            (limit,)
        )
        columns = [desc[0] for desc in cur.description]
        return [dict(zip(columns, row)) for row in cur.fetchall()]

def is_campaign_archived(campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, DEFAULT_CAMPAIGN_YEAR)
    with sqlite3.connect(DB_PATH) as conn:
        create_campaign_settings_table(conn)
        row = conn.execute(
            'SELECT is_archived FROM campaign_settings WHERE campaign_year=?',
            (campaign_year,)
        ).fetchone()
    return bool(row and row[0])

def ensure_campaign_open(campaign_year):
    if is_campaign_archived(campaign_year):
        flash(ARCHIVED_CAMPAIGN_MESSAGE, 'error')
        return False
    return True

def get_campaign_settings():
    years = get_campaign_years()
    with sqlite3.connect(DB_PATH) as conn:
        create_campaign_settings_table(conn)
        rows = {
            row[0]: {
                'is_archived': bool(row[1]),
                'archived_at': row[2],
                'archived_by': row[3],
                'created_at': row[4],
                'created_by': row[5],
                'is_active': bool(row[6]),
                'active_at': row[7],
                'active_by': row[8],
            }
            for row in conn.execute(
                '''
                SELECT campaign_year, is_archived, archived_at, archived_by,
                       created_at, created_by, is_active, active_at, active_by
                FROM campaign_settings
                '''
            )
        }
    settings = []
    for year in sorted(years, reverse=True):
        values = rows.get(year, {})
        changed_at = values.get('archived_at') or values.get('created_at') or ''
        changed_by = values.get('archived_by') or values.get('created_by') or ''
        settings.append({
            'campaign_year': year,
            'is_archived': bool(values.get('is_archived')),
            'archived_at': values.get('archived_at') or '',
            'archived_by': values.get('archived_by') or '',
            'created_at': values.get('created_at') or '',
            'created_by': values.get('created_by') or '',
            'is_active': bool(values.get('is_active')),
            'active_at': values.get('active_at') or '',
            'active_by': values.get('active_by') or '',
            'changed_at': changed_at,
            'changed_by': changed_by,
            'is_explicit': year in rows,
        })
    return settings

def get_dashboard_data(campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, DEFAULT_CAMPAIGN_YEAR)
    enrollment_order_required = is_enrollment_order_required()
    with sqlite3.connect(DB_PATH) as conn:
        create_campaign_settings_table(conn)
        ab_total = conn.execute('SELECT COUNT(*) FROM abiturients WHERE campaign_year=?', (campaign_year,)).fetchone()[0]
        no_email = conn.execute(
            "SELECT COUNT(*) FROM abiturients WHERE campaign_year=? AND LOWER(COALESCE(login, '')) NOT LIKE 'del%' AND (email IS NULL OR email='')",
            (campaign_year,)
        ).fetchone()[0]
        unpaid = conn.execute(
            "SELECT COUNT(*) FROM abiturients WHERE campaign_year=? AND LOWER(COALESCE(login, '')) NOT LIKE 'del%' AND COALESCE(paid, 0)=0",
            (campaign_year,)
        ).fetchone()[0]
        ready = conn.execute(
            "SELECT COUNT(*) FROM abiturients WHERE campaign_year=? AND LOWER(COALESCE(login, '')) NOT LIKE 'del%' AND email IS NOT NULL AND email<>'' AND COALESCE(paid, 0)=1",
            (campaign_year,)
        ).fetchone()[0]
        duplicates = conn.execute('SELECT COUNT(*) FROM pending_duplicates WHERE campaign_year=?', (campaign_year,)).fetchone()[0]
        conflicts = conn.execute('SELECT COUNT(*) FROM login_conflicts WHERE campaign_year=?', (campaign_year,)).fetchone()[0]
        students_total = conn.execute(
            'SELECT COUNT(*) FROM students WHERE source_campaign_year=?',
            (campaign_year,)
        ).fetchone()[0]
        students_without_campaign = conn.execute(
            "SELECT COUNT(*) FROM students WHERE source_campaign_year IS NULL OR source_campaign_year=''"
        ).fetchone()[0]
        students_without_dogovor = conn.execute(
            "SELECT COUNT(*) FROM students WHERE source_campaign_year=? AND (source_dogovor IS NULL OR source_dogovor='')",
            (campaign_year,)
        ).fetchone()[0]
        groups = get_groups_with_counts(conn, campaign_year)
        refresh_enrollment_candidate_statuses(conn, campaign_year)
        candidate_total = conn.execute(
            'SELECT COUNT(*) FROM enrollment_candidates WHERE campaign_year=?',
            (campaign_year,)
        ).fetchone()[0]
        candidate_verified = conn.execute(
            "SELECT COUNT(*) FROM enrollment_candidates WHERE campaign_year=? AND verification_status='verified'",
            (campaign_year,)
        ).fetchone()[0]
        candidate_missing_order = conn.execute(
            "SELECT COUNT(*) FROM enrollment_candidates WHERE campaign_year=? AND verification_status<>'verified'",
            (campaign_year,)
        ).fetchone()[0]
        if not enrollment_order_required:
            candidate_missing_order = 0
        order_total = conn.execute(
            'SELECT COUNT(*) FROM enrollment_orders WHERE campaign_year=?',
            (campaign_year,)
        ).fetchone()[0]
        ready_sources = conn.execute(
            '''
            SELECT id, login
            FROM abiturients
            WHERE campaign_year=?
              AND LOWER(COALESCE(login, '')) NOT LIKE 'del%'
              AND email IS NOT NULL AND email<>'' AND COALESCE(paid, 0)=1
            ''',
            (campaign_year,)
        ).fetchall()
        candidate_sources = conn.execute(
            '''
            SELECT abiturient_id, login
            FROM enrollment_candidates
            WHERE campaign_year=?
            ''',
            (campaign_year,)
        ).fetchall()

    full_groups = [group for group in groups if group['is_full']]
    almost_full_groups = [
        group for group in groups
        if not group['is_full'] and group['capacity'] - group['count'] <= 3
    ]
    alerts = []
    if is_campaign_archived(campaign_year):
        alerts.append(('Архив', ARCHIVED_CAMPAIGN_MESSAGE))
    if conflicts:
        alerts.append(('Конфликты', f'Конфликтов логинов: {conflicts}'))
    if duplicates:
        alerts.append(('Дубли', f'Записей в дублях: {duplicates}'))
    if no_email:
        alerts.append(('Почта', f'Без почты: {no_email}'))
    if unpaid:
        alerts.append(('Оплата', f'Не оплачены: {unpaid}'))
    if ready and not candidate_total:
        alerts.append(('Кандидаты', 'Готовые абитуриенты ещё не перенесены в кандидаты к зачислению'))
    if enrollment_order_required and candidate_total and not order_total:
        alerts.append(('Приказ', 'Приказ о зачислении ещё не загружен'))
    if enrollment_order_required and candidate_missing_order:
        alerts.append(('Сверка', f'Кандидатов без совпадения с приказом: {candidate_missing_order}'))
    if full_groups:
        alerts.append(('Группы', f'Заполненных групп: {len(full_groups)}'))
    if students_without_dogovor:
        alerts.append(('Договоры', f'Студентов без договора: {students_without_dogovor}'))
    if students_without_campaign:
        alerts.append(('Студенты', f'Студентов без привязки к кампании: {students_without_campaign}'))

    data_quality = get_data_quality_report(campaign_year)
    candidate_source_ids = {row[0] for row in candidate_sources if row[0] is not None}
    candidate_source_logins = {row[1] for row in candidate_sources if row[1]}
    ready_not_candidates = sum(
        1
        for abiturient_id, login in ready_sources
        if abiturient_id not in candidate_source_ids and login not in candidate_source_logins
    )
    task_counts = {
        'no_email': no_email,
        'unpaid': unpaid,
        'ready_not_candidates': ready_not_candidates,
        'candidate_missing_order': candidate_missing_order,
        'duplicates': duplicates,
        'conflicts': conflicts,
        'students_without_dogovor': students_without_dogovor,
        'other_data_quality': max(
            0,
            data_quality['total_issues'] - no_email - unpaid - ready_not_candidates
            - candidate_missing_order - duplicates - conflicts - students_without_dogovor
        ),
    }

    return {
        'campaign_year': campaign_year,
        'is_archived': is_campaign_archived(campaign_year),
        'abiturients_total': ab_total,
        'no_email': no_email,
        'unpaid': unpaid,
        'ready': ready,
        'candidate_total': candidate_total,
        'candidate_verified': candidate_verified,
        'candidate_missing_order': candidate_missing_order,
        'order_total': order_total,
        'ready_to_students': candidate_verified if enrollment_order_required else candidate_total,
        'enrollment_order_required': enrollment_order_required,
        'duplicates': duplicates,
        'conflicts': conflicts,
        'students_total': students_total,
        'students_without_dogovor': students_without_dogovor,
        'students_without_campaign': students_without_campaign,
        'groups': groups,
        'full_groups': full_groups,
        'almost_full_groups': almost_full_groups,
        'alerts': alerts,
        'tasks': build_dashboard_tasks(task_counts),
    }

def make_data_check(check_id, title, count, description, action_url='', action_label='Открыть', samples=None, tone='warning'):
    return {
        'id': check_id,
        'title': title,
        'count': int(count or 0),
        'description': description,
        'action_url': action_url,
        'action_label': action_label,
        'samples': samples or [],
        'tone': tone if count else 'success',
    }

def make_sample(title, detail='', url=''):
    return {
        'title': title or 'Без названия',
        'detail': detail or '',
        'url': url,
    }

def abiturient_sample(row):
    abiturient_id, fio, dogovor, login, email = row[:5]
    detail_parts = [part for part in (dogovor, login, email) if part]
    return make_sample(fio or login or f'Запись {abiturient_id}', ' · '.join(detail_parts), url_for('person_card', kind='abiturient', record_id=abiturient_id))

def student_sample(row):
    username, email, firstname, lastname, cohort1, cohort2, source_dogovor = row[:7]
    fio = ' '.join(part for part in (lastname, firstname) if part).strip() or username
    detail_parts = [username, cohort1]
    if are_course_groups_enabled():
        detail_parts.append(cohort2)
    detail_parts.extend([source_dogovor, email])
    detail_parts = [part for part in detail_parts if part]
    return make_sample(fio, ' · '.join(detail_parts), url_for('person_card', kind='student', record_id=username))

def student_cohort2_sample(row):
    sample = student_sample(row)
    issue = cohort2_quality_issue(row[4], row[5])
    if issue:
        actual = issue['actual'] or 'не указана'
        sample['detail'] = f"{sample['detail']} · ожидалось {issue['expected']}, сейчас {actual}"
    return sample

def enrollment_candidate_sample(row):
    candidate_id, _abiturient_id, fio, dogovor, login, specialty, _specialty_key, status, order_group_name = row[:9]
    detail_parts = [part for part in (specialty, order_group_name, login, dogovor, status) if part]
    return make_sample(fio or login or f'Кандидат {candidate_id}', ' · '.join(detail_parts), url_for('abiturients_to_students'))

def enrollment_order_sample(row):
    order_id, fio, specialty, group_name, _fio_key, _specialty_key = row[:6]
    detail_parts = [part for part in (specialty, group_name) if part]
    return make_sample(fio or f'Строка приказа {order_id}', ' · '.join(detail_parts), file_work_url('orders'))

def enrollment_order_roster_sample(row):
    detail_parts = [
        row.get('group_name'),
        row.get('login'),
        row.get('status'),
        row.get('group_assignment_status'),
        (
            'ФИО исправлено'
            if row.get('fio_review_status') == 'fixed'
            else 'Опечатка в приказе — логин присвоен'
            if row.get('fio_review_status') == 'linked'
            else 'Сверка ФИО пропущена'
            if row.get('fio_review_status') == 'skipped'
            else 'Возможная опечатка в ФИО'
            if row.get('fio_review_kind') == 'fio_typo'
            else 'Конфликт специальности'
            if row.get('fio_review_kind') == 'specialty_conflict'
            else ''
        ),
    ]
    return make_sample(
        row.get('fio') or f"Строка приказа {row.get('row_number') or '-'}",
        ' · '.join(str(part) for part in detail_parts if part),
        row.get('person_url') or '',
    )

def duplicate_group_samples(groups, title_index=1, detail_index=2, limit=5):
    samples = []
    for grouped_rows in groups[:limit]:
        first = grouped_rows[0]
        title = first[detail_index] or first[title_index] or 'Повтор'
        names = ', '.join((row[title_index] or row[3] or 'без имени') for row in grouped_rows[:3])
        if len(grouped_rows) > 3:
            names += '...'
        samples.append(make_sample(title, f'{len(grouped_rows)} записи: {names}'))
    return samples

def collect_duplicate_groups(rows, key_index):
    grouped = {}
    for row in rows:
        key = normalize_dogovor_key(row[key_index]) if key_index in {2, 5} else normalize_fio_key(row[key_index])
        if key:
            grouped.setdefault(key, []).append(row)
    return [items for items in grouped.values() if len(items) > 1]

def get_invalid_abiturient_email_rows(rows):
    return [row for row in rows if row[4] and not is_valid_email(row[4])]

def get_invalid_student_email_rows(rows):
    return [row for row in rows if row[1] and not is_valid_email(row[1])]

def get_data_quality_report(campaign_year=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    login_rules = get_login_generation_rules()
    enrollment_order_required = is_enrollment_order_required()
    course_groups_enabled = are_course_groups_enabled(login_rules)
    sample_limit = 5
    with sqlite3.connect(DB_PATH) as conn:
        ab_rows = conn.execute(
            '''
            SELECT id, fio, dogovor, login, email, paid
            FROM abiturients
            WHERE campaign_year=?
            ORDER BY fio
            ''',
            (campaign_year,)
        ).fetchall()
        pending_duplicates = conn.execute(
            'SELECT id, fio, dogovor, login, fam, imotch, campaign_year FROM pending_duplicates WHERE campaign_year=? ORDER BY fio',
            (campaign_year,)
        ).fetchall()
        login_conflict_rows = conn.execute(
            'SELECT id, fio, dogovor, login, fam, imotch, campaign_year, conflict_time FROM login_conflicts WHERE campaign_year=? ORDER BY conflict_time DESC',
            (campaign_year,)
        ).fetchall()
        student_rows = conn.execute(
            '''
            SELECT username, email, firstname, lastname, cohort1, cohort2, source_dogovor, source_campaign_year
            FROM students
            WHERE source_campaign_year=? OR source_campaign_year IS NULL OR source_campaign_year=''
            ORDER BY lastname, firstname, username
            ''',
            (campaign_year,)
        ).fetchall()
        refresh_enrollment_candidate_statuses(conn, campaign_year)
        candidate_rows = conn.execute(
            '''
            SELECT id, abiturient_id, fio, dogovor, login, specialty, specialty_key,
                   verification_status, order_group_name
            FROM enrollment_candidates
            WHERE campaign_year=?
            ORDER BY specialty, fio
            ''',
            (campaign_year,)
        ).fetchall()
        order_rows = conn.execute(
            '''
            SELECT id, fio, specialty, group_name, fio_key, specialty_key
            FROM enrollment_orders
            WHERE campaign_year=?
            ORDER BY specialty, fio
            ''',
            (campaign_year,)
        ).fetchall()

    withdrawn_ab_rows = [row for row in ab_rows if is_withdrawn_login(row[3])]
    ab_rows = [row for row in ab_rows if not is_withdrawn_login(row[3])]
    ab_without_email = [row for row in ab_rows if not str(row[4] or '').strip()]
    ab_unpaid = [row for row in ab_rows if not is_paid_person_value(row[5])]
    ab_without_dogovor = [row for row in ab_rows if not normalize_dogovor_key(row[2])]
    ab_invalid_email = get_invalid_abiturient_email_rows(ab_rows)
    ab_duplicate_dogovors = collect_duplicate_groups(ab_rows, 2)
    ab_same_fio = collect_duplicate_groups(ab_rows, 1)

    current_students = [row for row in student_rows if row[7] == campaign_year]
    students_without_group = [row for row in current_students if not str(row[4] or '').strip()]
    students_without_dogovor = [row for row in current_students if not normalize_dogovor_key(row[6])]
    students_without_campaign = [row for row in student_rows if not str(row[7] or '').strip()]
    students_invalid_email = get_invalid_student_email_rows(current_students)
    student_duplicate_dogovors = collect_duplicate_groups(current_students, 6)
    students_cohort2_issues = [
        row for row in current_students
        if cohort2_quality_issue(row[4], row[5], login_rules)
    ]
    ready_abiturients = [
        row for row in ab_rows
        if str(row[4] or '').strip() and is_paid_person_value(row[5])
    ]
    candidate_source_ids = {row[1] for row in candidate_rows if row[1] is not None}
    candidate_source_logins = {row[4] for row in candidate_rows if row[4]}
    ready_not_candidates = [
        row for row in ready_abiturients
        if row[0] not in candidate_source_ids and row[3] not in candidate_source_logins
    ]
    candidates_missing_order = [
        row for row in candidate_rows
        if enrollment_order_required and row[7] != 'verified'
    ]
    candidate_keys = {
        make_enrollment_match_key(row[2], row[6])
        for row in candidate_rows
    }
    candidate_keys.discard(None)
    withdrawn_order_keys = {
        make_abiturient_enrollment_match_key(row[1], row[2], login_rules)
        for row in withdrawn_ab_rows
    }
    withdrawn_order_keys.discard(None)
    order_without_candidate = [
        row for row in order_rows
        if (row[4], row[5]) not in candidate_keys
        and (row[4], row[5]) not in withdrawn_order_keys
    ]

    student_checks = [
        make_data_check('students-without-group', 'Без академической группы', len(students_without_group), 'Студент есть в базе, но не привязан к группе.', url_for('students_list'), 'Открыть студентов', [student_sample(row) for row in students_without_group[:sample_limit]]),
        make_data_check('students-without-dogovor', 'Без договора при поступлении', len(students_without_dogovor), 'Без договора сложнее проверить, от какого абитуриента появился студент.', url_for('students_list'), 'Открыть студентов', [student_sample(row) for row in students_without_dogovor[:sample_limit]]),
        make_data_check('students-without-campaign', 'Без кампании поступления', len(students_without_campaign), 'У студента не указан год кампании, поэтому он выпадает из отчетов по кампании.', url_for('students_list'), 'Открыть студентов', [student_sample(row) for row in students_without_campaign[:sample_limit]]),
        make_data_check('students-invalid-email', 'Некорректная почта', len(students_invalid_email), 'Почта студента заполнена, но похожа на ошибочную.', url_for('students_list'), 'Открыть студентов', [student_sample(row) for row in students_invalid_email[:sample_limit]]),
    ]
    if course_groups_enabled:
        student_checks.append(
            make_data_check('students-cohort2-mismatch', 'Проблемы глобальной группы курса', len(students_cohort2_issues), 'Для этих студентов глобальная группа курса отсутствует или не совпадает с академической группой.', url_for('students_list'), 'Открыть студентов', [student_cohort2_sample(row) for row in students_cohort2_issues[:sample_limit]])
        )
    student_checks.append(
        make_data_check('students-duplicate-dogovor', 'Повторяющиеся договоры', sum(len(group) for group in student_duplicate_dogovors), 'Один договор при поступлении найден у нескольких студентов.', url_for('students_list'), 'Открыть студентов', duplicate_group_samples(student_duplicate_dogovors, title_index=3, detail_index=6))
    )

    order_roster_checks = []
    for upload in get_enrollment_order_uploads(campaign_year):
        roster = build_enrollment_order_student_roster(upload['id'])
        if not roster:
            continue
        summary = roster['summary']
        title_parts = []
        if upload.get('order_numbers'):
            title_parts.append(f"№ {upload['order_numbers']}")
        if upload.get('order_dates'):
            title_parts.append(f"от {upload['order_dates']}")
        order_title = ' '.join(title_parts) or upload.get('original_filename') or f"Загрузка #{upload['id']}"
        check = make_data_check(
            f"enrollment-order-roster-{upload['id']}",
            f"Предварительный список: {order_title}",
            summary['attention_count'],
            (
                f"Всего в приказе: {summary['total_count']}; групп: {summary['group_count']}; "
                f"справочных групп: {summary['virtual_group_count']}; "
                f"не распределено: {summary['unassigned_count']}; "
                f"уже студентов: {summary['student_count']}; готовы к переносу: {summary['ready_count']}; "
                f"без почты: {summary['missing_email_count']}; без оплаты: {summary['unpaid_count']}; "
                f"не найдены в базе: {summary['not_found_count']}; "
                f"сверить ФИО: {summary['fio_review_count']}; "
                f"логин присвоен после сверки: {summary['fio_review_linked_count']}; "
                f"конфликтов специальности: {summary['specialty_conflict_count']}."
            ),
            url_for('enrollment_order_student_roster', upload_id=upload['id']),
            'Открыть полный список',
            [
                enrollment_order_roster_sample(row)
                for row in roster['rows']
                if row['needs_attention']
            ][:sample_limit],
        )
        check.update({
            'expandable': True,
            'download_url': url_for('download_enrollment_order_student_roster', upload_id=upload['id']),
            'download_label': 'Скачать Excel',
            'upload_filename': upload.get('original_filename') or '',
        })
        order_roster_checks.append(check)
    if not order_roster_checks:
        order_roster_checks.append(
            make_data_check(
                'enrollment-order-rosters-empty',
                'Предварительные списки по приказам',
                0,
                'Загрузите приказ о зачислении, чтобы сформировать общий список с группами, логинами, почтой, оплатой и состоянием переноса.',
                file_work_url('orders'),
                'Загрузить приказ',
                tone='info',
            )
        )

    sections = [
        {
            'title': 'Абитуриенты',
            'checks': [
                make_data_check('abiturients-without-email', 'Без почты', len(ab_without_email), 'Не получится восстановить доступ и выполнить миграцию без почты.', url_for('abiturients', has_email='0', withdrawn='0'), 'Открыть список', [abiturient_sample(row) for row in ab_without_email[:sample_limit]]),
                make_data_check('abiturients-unpaid', 'Не оплачены', len(ab_unpaid), 'Эти записи не готовы к миграции, пока оплата не отмечена.', url_for('abiturients', has_paid='0', withdrawn='0'), 'Открыть список', [abiturient_sample(row) for row in ab_unpaid[:sample_limit]]),
                make_data_check('abiturients-invalid-email', 'Некорректная почта', len(ab_invalid_email), 'Почта заполнена, но похожа на ошибочную.', url_for('abiturients', withdrawn='0'), 'Открыть абитуриентов', [abiturient_sample(row) for row in ab_invalid_email[:sample_limit]]),
                make_data_check('abiturients-without-dogovor', 'Без договора', len(ab_without_dogovor), 'Без договора сложно отличать тёзок и проверять повторы.', url_for('abiturients', withdrawn='0'), 'Открыть абитуриентов', [abiturient_sample(row) for row in ab_without_dogovor[:sample_limit]]),
                make_data_check('abiturients-duplicate-dogovor', 'Повторяющиеся договоры', sum(len(group) for group in ab_duplicate_dogovors), 'Один договор найден в нескольких записях абитуриентов.', url_for('abiturients', withdrawn='0'), 'Открыть абитуриентов', duplicate_group_samples(ab_duplicate_dogovors)),
                make_data_check('abiturients-same-fio', 'Одинаковое ФИО', sum(len(group) for group in ab_same_fio), 'Это могут быть дубли или тёзки. Лучше сверить договоры.', url_for('abiturients', withdrawn='0'), 'Открыть абитуриентов', duplicate_group_samples(ab_same_fio, title_index=1, detail_index=1)),
            ],
        },
        {
            'title': 'Отозванные документы',
            'informational': True,
            'checks': [
                make_data_check(
                    'withdrawn-abiturients',
                    'Абитуриенты, отозвавшие документы',
                    len(withdrawn_ab_rows),
                    'Справочный список. Эти записи исключены из замечаний и подготовки к зачислению, а исходные логины освобождены.',
                    url_for('abiturients', withdrawn='1'),
                    'Открыть список',
                    [abiturient_sample(row) for row in withdrawn_ab_rows[:sample_limit]],
                    tone='info'
                ),
            ],
        },
        {
            'title': 'Зачисление и приказ',
            'checks': [
                make_data_check('ready-not-candidates', 'Готовы, но не в кандидатах', len(ready_not_candidates), 'Есть абитуриенты с почтой и оплатой, которых нужно перенести на первый этап зачисления.', url_for('abiturients_to_students'), 'Подготовить кандидатов', [abiturient_sample(row) for row in ready_not_candidates[:sample_limit]]),
                make_data_check('candidates-missing-order', 'Кандидаты без приказа', len(candidates_missing_order), 'Эти кандидаты не будут перенесены в студенты, пока не совпадут с итоговым приказом.', url_for('abiturients_to_students'), 'Открыть сверку', [enrollment_candidate_sample(row) for row in candidates_missing_order[:sample_limit]]),
                make_data_check('order-without-candidate', 'В приказе без кандидата', len(order_without_candidate), 'В приказе есть человек, но среди кандидатов к зачислению он не найден. Проверьте ФИО, специальность, почту и оплату.', file_work_url('orders'), 'Проверить приказ', [enrollment_order_sample(row) for row in order_without_candidate[:sample_limit]]),
            ],
        },
        {
            'title': 'Предварительные списки по приказам',
            'informational': True,
            'checks': order_roster_checks,
        },
        {
            'title': 'Студенты',
            'checks': student_checks,
        },
        {
            'title': 'Миграция и конфликты',
            'checks': [
                make_data_check('pending-duplicates', 'Дублирующие записи абитуриентов', len(pending_duplicates), 'Эти записи ждут решения: подтвердить или отклонить.', url_for('duplicates_abiturients'), 'Разобрать дубли', [make_sample(row[1], f'{row[2]} · {row[3]}', url_for('person_card', kind='duplicate', record_id=row[0])) for row in pending_duplicates[:sample_limit]]),
                make_data_check('login-conflicts', 'Конфликты логинов', len(login_conflict_rows), 'Система не смогла безопасно назначить логин.', url_for('login_conflicts'), 'Разобрать конфликты', [make_sample(row[1], f'{row[2]} · {row[3]}', url_for('person_card', kind='conflict', record_id=row[0])) for row in login_conflict_rows[:sample_limit]]),
            ],
        },
    ]
    total_issues = sum(
        check['count']
        for section in sections
        if not section.get('informational')
        for check in section['checks']
    )
    return {
        'campaign_year': campaign_year,
        'sections': sections,
        'total_issues': total_issues,
    }

def build_dashboard_tasks(counts):
    tasks = [
        {
            'title': 'Абитуриенты без почты',
            'count': counts.get('no_email', 0),
            'description': 'Нужно добавить почту перед миграцией.',
            'url': url_for('abiturients', has_email='0', withdrawn='0'),
            'label': 'Открыть',
        },
        {
            'title': 'Не оплачены',
            'count': counts.get('unpaid', 0),
            'description': 'Эти абитуриенты пока не готовы к миграции.',
            'url': url_for('abiturients', has_paid='0', withdrawn='0'),
            'label': 'Открыть',
        },
        {
            'title': 'Готовы к кандидатам',
            'count': counts.get('ready_not_candidates', 0),
            'description': 'Есть абитуриенты с почтой и оплатой, которых нужно перенести в кандидаты к зачислению.',
            'url': url_for('abiturients_to_students'),
            'label': 'Подготовить',
        },
        {
            'title': 'Не сверены с приказом',
            'count': counts.get('candidate_missing_order', 0),
            'description': 'Кандидаты есть, но они ещё не совпали с итоговым приказом.',
            'url': url_for('abiturients_to_students'),
            'label': 'Сверить',
        },
        {
            'title': 'Дубли',
            'count': counts.get('duplicates', 0),
            'description': 'Нужно подтвердить или отклонить записи.',
            'url': url_for('duplicates_abiturients'),
            'label': 'Разобрать',
        },
        {
            'title': 'Конфликты логинов',
            'count': counts.get('conflicts', 0),
            'description': 'Нужно назначить корректный логин.',
            'url': url_for('login_conflicts'),
            'label': 'Разобрать',
        },
        {
            'title': 'Студенты без договора',
            'count': counts.get('students_without_dogovor', 0),
            'description': 'Лучше проверить договор, чтобы отличать тёзок.',
            'url': url_for('data_checks') + '#students-without-dogovor',
            'label': 'Проверить',
        },
        {
            'title': 'Другие замечания',
            'count': counts.get('other_data_quality', 0),
            'description': 'Есть дополнительные ошибки или подозрительные записи.',
            'url': url_for('data_checks'),
            'label': 'Проверить',
        },
    ]
    active_tasks = [task for task in tasks if task['count']]
    if active_tasks:
        return active_tasks
    return [{
        'title': 'Критичных задач нет',
        'count': 0,
        'description': 'По текущей кампании явных проблем не найдено.',
        'url': url_for('data_checks'),
        'label': 'Посмотреть проверку',
    }]

def like_pattern(value):
    return f"%{str(value or '').strip()}%"

def global_search_records(query, campaign_year=None, limit=80):
    query = str(query or '').strip()
    if not query:
        return []
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    course_groups_enabled = are_course_groups_enabled()
    pattern = like_pattern(query)
    results = []
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            '''
            SELECT id, fio, dogovor, login, email, campaign_year
            FROM abiturients
            WHERE campaign_year=? AND (fio LIKE ? OR dogovor LIKE ? OR login LIKE ? OR email LIKE ?)
            ORDER BY fio
            LIMIT ?
            ''',
            (campaign_year, pattern, pattern, pattern, pattern, limit)
        )
        for row in cur.fetchall():
            results.append({
                'kind': 'abiturient',
                'id': row[0],
                'title': row[1],
                'subtitle': f'{row[3]} · {row[2]}',
                'status': 'Абитуриент',
            })

        cur = conn.execute(
            '''
            SELECT username, email, firstname, lastname, cohort1, cohort2, source_dogovor
            FROM students
            WHERE username LIKE ? OR email LIKE ? OR firstname LIKE ? OR lastname LIKE ? OR cohort1 LIKE ? OR cohort2 LIKE ? OR source_dogovor LIKE ?
            ORDER BY lastname, firstname
            LIMIT ?
            ''',
            (pattern, pattern, pattern, pattern, pattern, pattern, pattern, limit)
        )
        for row in cur.fetchall():
            fio = ' '.join(part for part in (row[3], row[2]) if part).strip() or row[0]
            subtitle_parts = [row[0], row[4] or '-']
            if course_groups_enabled:
                subtitle_parts.append(row[5] or '-')
            subtitle_parts.append(row[6] or 'без договора')
            results.append({
                'kind': 'student',
                'id': row[0],
                'title': fio,
                'subtitle': ' · '.join(subtitle_parts),
                'status': 'Студент',
            })

        cur = conn.execute(
            '''
            SELECT id, fio, dogovor, login
            FROM pending_duplicates
            WHERE campaign_year=? AND (fio LIKE ? OR dogovor LIKE ? OR login LIKE ?)
            ORDER BY fio
            LIMIT ?
            ''',
            (campaign_year, pattern, pattern, pattern, limit)
        )
        for row in cur.fetchall():
            results.append({
                'kind': 'duplicate',
                'id': row[0],
                'title': row[1],
                'subtitle': f'{row[3]} · {row[2]}',
                'status': 'Дубль',
            })

        cur = conn.execute(
            '''
            SELECT id, fio, dogovor, login
            FROM login_conflicts
            WHERE campaign_year=? AND (fio LIKE ? OR dogovor LIKE ? OR login LIKE ?)
            ORDER BY conflict_time DESC
            LIMIT ?
            ''',
            (campaign_year, pattern, pattern, pattern, limit)
        )
        for row in cur.fetchall():
            results.append({
                'kind': 'conflict',
                'id': row[0],
                'title': row[1],
                'subtitle': f'{row[3]} · {row[2]}',
                'status': 'Конфликт',
            })
    return results[:limit]

def get_person_record(kind, record_id):
    kind = str(kind or '').strip()
    with sqlite3.connect(DB_PATH) as conn:
        if kind == 'abiturient':
            cur = conn.execute('SELECT * FROM abiturients WHERE id=?', (record_id,))
        elif kind == 'student':
            cur = conn.execute('SELECT * FROM students WHERE username=?', (record_id,))
        elif kind == 'duplicate':
            cur = conn.execute('SELECT * FROM pending_duplicates WHERE id=?', (record_id,))
        elif kind == 'conflict':
            cur = conn.execute('SELECT * FROM login_conflicts WHERE id=?', (record_id,))
        else:
            return None
        row = cur.fetchone()
        if not row:
            return None
        columns = [desc[0] for desc in cur.description]
    return {
        'kind': kind,
        'id': record_id,
        'fields': dict(zip(columns, row)),
    }

PERSON_KIND_LABELS = {
    'abiturient': 'Абитуриент',
    'student': 'Студент',
    'duplicate': 'Дублирующая запись',
    'conflict': 'Конфликт логина',
}

PERSON_KIND_TITLES = {
    'abiturient': 'Карточка абитуриента',
    'student': 'Карточка студента',
    'duplicate': 'Карточка возможного дубля',
    'conflict': 'Карточка конфликта логина',
}

PERSON_FIELD_LABELS = {
    'id': 'Номер записи',
    'fio': 'ФИО',
    'dogovor': 'Номер договора',
    'login': 'Логин Moodle',
    'username': 'Логин Moodle',
    'password': 'Пароль Moodle',
    'campaign_year': 'Приемная кампания',
    'source_campaign_year': 'Кампания поступления',
    'fam': 'Фамилия',
    'imotch': 'Имя и отчество',
    'firstname': 'Имя',
    'lastname': 'Фамилия',
    'email': 'Электронная почта',
    'paid': 'Оплата договора',
    'comment': 'Комментарий',
    'created_at': 'Дата добавления',
    'conflict_time': 'Дата конфликта',
    'cohort1': 'Академическая группа',
    'cohort2': 'Глобальная группа курса',
    'source_dogovor': 'Договор при поступлении',
    'source_fio': 'ФИО при поступлении',
}

PERSON_FIELD_HELP = {
    'id': 'Внутренний номер записи для точного поиска и поддержки.',
    'fio': 'Полное имя человека в текущей записи.',
    'dogovor': 'Номер договора, по нему удобнее всего отличать тезок.',
    'login': 'Учетная запись, с которой человек входит в Moodle.',
    'username': 'Учетная запись, с которой человек входит в Moodle.',
    'password': 'Пароль к учетной записи Moodle.',
    'campaign_year': 'Год приемной кампании, к которой относится запись.',
    'source_campaign_year': 'Год приемной кампании, из которой пришел студент.',
    'fam': 'Фамилия отдельно, используется при формировании логина и списков.',
    'imotch': 'Имя и отчество отдельно, используется при формировании логина и списков.',
    'firstname': 'Имя студента в Moodle.',
    'lastname': 'Фамилия студента в Moodle.',
    'email': 'Почта для связи и восстановления доступа.',
    'paid': 'Показывает, отмечена ли оплата договора.',
    'comment': 'Заметка сотрудника по этой записи.',
    'created_at': 'Когда запись была добавлена в систему.',
    'conflict_time': 'Когда система обнаружила конфликт логина.',
    'cohort1': 'Группа, к которой сейчас привязан студент.',
    'cohort2': 'Автоматически определяемая глобальная группа курса для назначения набора курсов.',
    'source_dogovor': 'Договор, по которому студент был найден при миграции.',
    'source_fio': 'ФИО из исходной записи абитуриента.',
}

PERSON_SECTION_FIELDS = {
    'abiturient': [
        ('Основные данные', ['fio', 'dogovor', 'campaign_year', 'paid']),
        ('Контакты и доступ', ['login', 'email']),
        ('ФИО по частям', ['fam', 'imotch']),
        ('Дополнительно', ['comment', 'created_at', 'id']),
    ],
    'student': [
        ('Основные данные', ['lastname', 'firstname', 'source_fio', 'cohort1', 'cohort2']),
        ('Контакты и доступ', ['username', 'password', 'email']),
        ('Данные при поступлении', ['source_dogovor', 'source_campaign_year']),
        ('Служебная информация', ['id']),
    ],
    'duplicate': [
        ('Основные данные', ['fio', 'dogovor', 'campaign_year']),
        ('Контакты и доступ', ['login']),
        ('ФИО по частям', ['fam', 'imotch']),
        ('Служебная информация', ['id']),
    ],
    'conflict': [
        ('Основные данные', ['fio', 'dogovor', 'campaign_year']),
        ('Контакты и доступ', ['login']),
        ('ФИО по частям', ['fam', 'imotch']),
        ('Служебная информация', ['conflict_time', 'id']),
    ],
}

PERSON_SUMMARY_FIELDS = {
    'abiturient': ['fio', 'dogovor', 'login', 'paid'],
    'student': ['source_fio', 'cohort1', 'cohort2', 'username', 'source_dogovor'],
    'duplicate': ['fio', 'dogovor', 'login'],
    'conflict': ['fio', 'dogovor', 'login'],
}

def is_blank_person_value(value):
    return value is None or str(value).strip() in {'', '-'}

def is_paid_person_value(value):
    return str(value).strip().casefold() in {'1', 'true', 'yes', 'да', 'оплачен', 'оплачено', 'paid'}

def humanize_person_field_name(key):
    return str(key or '').replace('_', ' ').strip().capitalize() or 'Поле'

def format_person_field_value(key, value):
    if key == 'paid':
        return 'Договор оплачен' if is_paid_person_value(value) else 'Договор не оплачен'
    if key == 'password' and value == '******':
        return 'Скрыт для безопасности'
    if key == 'email' and is_blank_person_value(value):
        return 'Почта не указана'
    if key == 'comment' and is_blank_person_value(value):
        return 'Комментария нет'
    if is_blank_person_value(value):
        return 'Не указано'
    return str(value)

def get_person_field_state(key, value):
    if key == 'paid':
        return 'success' if is_paid_person_value(value) else 'warning'
    if is_blank_person_value(value):
        return 'muted'
    if key == 'email':
        return 'success'
    return ''

def build_person_card_item(key, value):
    return {
        'key': key,
        'label': PERSON_FIELD_LABELS.get(key, humanize_person_field_name(key)),
        'help': PERSON_FIELD_HELP.get(key, 'Дополнительная информация по записи.'),
        'value': format_person_field_value(key, value),
        'state': get_person_field_state(key, value),
    }

def get_student_display_name(fields):
    source_fio = fields.get('source_fio')
    if not is_blank_person_value(source_fio):
        return str(source_fio)
    fio = ' '.join(part for part in (fields.get('lastname'), fields.get('firstname')) if not is_blank_person_value(part)).strip()
    return fio or fields.get('username') or 'Студент'

def get_person_display_name(kind, fields):
    if kind == 'student':
        return get_student_display_name(fields)
    return fields.get('fio') or fields.get('login') or fields.get('username') or PERSON_KIND_LABELS.get(kind, 'Запись')

def build_person_card_view(record):
    kind = record.get('kind')
    fields = record.get('fields') or {}
    hidden_fields = {'cohort2'} if kind == 'student' and not are_course_groups_enabled() else set()
    seen = set()
    sections = []

    for section_title, keys in PERSON_SECTION_FIELDS.get(kind, [('Данные записи', list(fields.keys()))]):
        items = []
        for key in keys:
            if key in fields and key not in hidden_fields:
                seen.add(key)
                items.append(build_person_card_item(key, fields.get(key)))
        if items:
            sections.append({'title': section_title, 'items': items})

    remaining_items = [
        build_person_card_item(key, value)
        for key, value in fields.items()
        if key not in seen and key not in hidden_fields
    ]
    if remaining_items:
        sections.append({'title': 'Дополнительные поля', 'items': remaining_items})

    summary = [
        {
            'label': 'Тип записи',
            'value': PERSON_KIND_LABELS.get(kind, 'Запись'),
            'state': 'info',
        }
    ]
    for key in PERSON_SUMMARY_FIELDS.get(kind, []):
        if key in fields and key not in hidden_fields:
            item = build_person_card_item(key, fields.get(key))
            summary.append({
                'label': item['label'],
                'value': item['value'],
                'state': item['state'],
            })

    return {
        'title': PERSON_KIND_TITLES.get(kind, 'Карточка записи'),
        'subtitle': get_person_display_name(kind, fields),
        'summary': summary,
        'sections': sections,
    }

def parse_paid_value(value):
    value = str(value or '').strip().casefold()
    if value in {'1', 'true', 'yes', 'да', 'оплачен', 'оплачено', 'paid'}:
        return 1
    if value in {'0', 'false', 'no', 'нет', 'не оплачен', 'не оплачено', 'unpaid'}:
        return 0
    return None

def sync_enrollment_candidate_from_abiturient(conn, abiturient_id, campaign_year=None, login_rules=None, refresh_status=True):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    login_rules = login_rules or get_login_generation_rules()
    row = conn.execute(
        '''
        SELECT id, fio, dogovor, login, fam, imotch, email, paid
        FROM abiturients
        WHERE id=? AND campaign_year=?
        ''',
        (abiturient_id, campaign_year)
    ).fetchone()
    if not row:
        return {'action': 'missing_abiturient', 'display_name': str(abiturient_id)}

    abiturient_id, fio, dogovor, login, fam, imotch, email, paid = row
    display_name = fio or login or str(abiturient_id)
    candidate_params = [campaign_year, abiturient_id]
    candidate_where = ['abiturient_id=?']
    if login:
        candidate_where.append('login=?')
        candidate_params.append(login)
    candidate_ids = [
        item[0]
        for item in conn.execute(
            f'''
            SELECT id
            FROM enrollment_candidates
            WHERE campaign_year=? AND ({' OR '.join(candidate_where)})
            ORDER BY CASE WHEN abiturient_id=? THEN 0 ELSE 1 END, id
            ''',
            candidate_params + [abiturient_id]
        ).fetchall()
    ]
    candidate_id = candidate_ids[0] if candidate_ids else None

    def remove_candidate(action):
        if candidate_ids:
            placeholders = ','.join('?' for _ in candidate_ids)
            conn.execute(f'DELETE FROM enrollment_candidates WHERE id IN ({placeholders})', candidate_ids)
        if refresh_status:
            refresh_enrollment_candidate_statuses(conn, campaign_year)
        return {'action': action, 'display_name': display_name, 'removed': len(candidate_ids)}

    if is_withdrawn_login(login):
        return remove_candidate('skipped_withdrawn')
    email = (email or '').strip()
    if not email:
        return remove_candidate('skipped_without_email')
    if not is_paid_person_value(paid):
        return remove_candidate('skipped_unpaid')
    if login and conn.execute('SELECT 1 FROM students WHERE username=?', (login,)).fetchone():
        return remove_candidate('skipped_existing_student')

    parts = parse_dogovor_parts(dogovor, login_rules)
    if not parts:
        return remove_candidate('skipped_without_specialty')

    specialty_code = normalize_specialty(parts['spec_label'])
    specialty = get_specialty_display_name(specialty_code)
    specialty_key = normalize_specialty_key(specialty_code)
    base_label = parts.get('base_label') or ''

    if candidate_id:
        conn.execute(
            '''
            UPDATE enrollment_candidates
            SET abiturient_id=?, fio=?, dogovor=?, login=?, fam=?, imotch=?,
                email=?, specialty=?, specialty_key=?, base_label=?
            WHERE id=?
            ''',
            (
                abiturient_id, fio, dogovor, login, fam, imotch, email,
                specialty, specialty_key, base_label, candidate_id
            )
        )
        duplicate_ids = candidate_ids[1:]
        if duplicate_ids:
            placeholders = ','.join('?' for _ in duplicate_ids)
            conn.execute(f'DELETE FROM enrollment_candidates WHERE id IN ({placeholders})', duplicate_ids)
        action = 'updated'
    else:
        conn.execute(
            '''
            INSERT INTO enrollment_candidates
                (abiturient_id, campaign_year, fio, dogovor, login, fam, imotch,
                 email, specialty, specialty_key, base_label)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''',
            (
                abiturient_id, campaign_year, fio, dogovor, login, fam, imotch,
                email, specialty, specialty_key, base_label
            )
        )
        action = 'created'

    if refresh_status:
        refresh_enrollment_candidate_statuses(conn, campaign_year)
    return {'action': action, 'display_name': display_name, 'removed': max(0, len(candidate_ids) - 1)}

def sync_enrollment_candidates_for_abiturients(conn, abiturient_ids, campaign_year=None, login_rules=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    summary = {
        'created': 0,
        'updated': 0,
        'removed': 0,
        'skipped_without_email': [],
        'skipped_unpaid': [],
        'skipped_without_specialty': [],
        'skipped_existing_students': [],
        'skipped_withdrawn': [],
    }
    for abiturient_id in dict.fromkeys(str(item) for item in abiturient_ids if str(item).isdigit()):
        result = sync_enrollment_candidate_from_abiturient(
            conn,
            int(abiturient_id),
            campaign_year,
            login_rules=login_rules,
            refresh_status=False
        )
        action = result.get('action')
        if action == 'created':
            summary['created'] += 1
        elif action == 'updated':
            summary['updated'] += 1
        elif action == 'skipped_without_email':
            summary['skipped_without_email'].append(result.get('display_name', abiturient_id))
            summary['removed'] += result.get('removed', 0)
        elif action == 'skipped_unpaid':
            summary['skipped_unpaid'].append(result.get('display_name', abiturient_id))
            summary['removed'] += result.get('removed', 0)
        elif action == 'skipped_without_specialty':
            summary['skipped_without_specialty'].append(result.get('display_name', abiturient_id))
            summary['removed'] += result.get('removed', 0)
        elif action == 'skipped_existing_student':
            summary['skipped_existing_students'].append(result.get('display_name', abiturient_id))
            summary['removed'] += result.get('removed', 0)
        elif action == 'skipped_withdrawn':
            summary['skipped_withdrawn'].append(result.get('display_name', abiturient_id))
            summary['removed'] += result.get('removed', 0)
    refresh_enrollment_candidate_statuses(conn, campaign_year)
    return summary

def find_row_value_casefold(row, aliases):
    normalized = {str(key).strip().casefold(): key for key in row.index}
    for alias in aliases:
        key = normalized.get(alias.casefold())
        if key is not None:
            return row.get(key)
    return None

def find_dataframe_column_casefold(columns, aliases):
    normalized_aliases = {str(alias).strip().casefold() for alias in aliases}
    for column in columns:
        if str(column).strip().casefold() in normalized_aliases:
            return column
    return None

def normalize_email_source_fio_key(value):
    return normalize_fio_key(value).replace('ё', 'е')

def build_email_source_person_context(campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    context = {
        'abiturients': {},
        'students': {},
        'records': [],
    }

    def add_record(kind, record_id, fio, email, detail):
        display_fio = ' '.join(str(fio or '').split())
        fio_key = normalize_email_source_fio_key(display_fio)
        similarity_key = normalize_fio_similarity_key(display_fio)
        if not fio_key or not similarity_key:
            return
        record = {
            'kind': kind,
            'id': record_id,
            'fio': display_fio,
            'fio_key': fio_key,
            'similarity_key': similarity_key,
            'email': str(email or '').strip(),
            'detail': str(detail or '').strip(),
        }
        context[kind].setdefault(fio_key, []).append(record)
        context['records'].append(record)

    with sqlite3.connect(DB_PATH) as conn:
        for abiturient_id, fio, fam, imotch, email, dogovor in conn.execute(
            '''
            SELECT id, fio, fam, imotch, email, dogovor
            FROM abiturients
            WHERE campaign_year=?
            ''',
            (campaign_year,)
        ).fetchall():
            stored_fio = str(fio or '').strip() or ' '.join(
                part for part in (fam, imotch) if str(part or '').strip()
            ).strip()
            add_record('abiturients', abiturient_id, stored_fio, email, dogovor)

        for student_id, username, email, firstname, lastname, source_fio in conn.execute(
            '''
            SELECT id, username, email, firstname, lastname, source_fio
            FROM students
            '''
        ).fetchall():
            stored_fio = str(source_fio or '').strip() or ' '.join(
                part for part in (lastname, firstname) if str(part or '').strip()
            ).strip()
            add_record('students', student_id, stored_fio, email, username)

    return context

def find_email_source_fio_suggestion(fio, context):
    source_key = normalize_fio_similarity_key(fio)
    if not source_key:
        return None

    suggestions_by_fio = {}
    for record in context.get('records', []):
        candidate_key = record.get('similarity_key') or ''
        if not candidate_key:
            continue
        score = difflib.SequenceMatcher(None, source_key, candidate_key).ratio()
        if score < ENROLLMENT_FIO_SUGGESTION_THRESHOLD:
            continue
        display_key = normalize_email_source_fio_key(record.get('fio'))
        suggestion = suggestions_by_fio.setdefault(display_key, {
            'fio': record.get('fio') or '',
            'score': score,
            'kinds': set(),
        })
        suggestion['score'] = max(suggestion['score'], score)
        suggestion['kinds'].add(record.get('kind'))

    suggestions = sorted(
        suggestions_by_fio.values(),
        key=lambda item: (-item['score'], normalize_fio_key(item['fio']))
    )
    if not suggestions:
        return None

    best = suggestions[0]
    best['similarity'] = int(round(best['score'] * 100))
    best['ambiguous'] = len(suggestions) > 1 and best['score'] - suggestions[1]['score'] < 0.03
    if best['ambiguous']:
        best['alternatives'] = [item['fio'] for item in suggestions[:2]]
    return best

def email_source_kind_label(kind):
    return 'абитуриент' if kind == 'abiturients' else 'студент'

def email_source_preview_row(row_number, fio, email):
    return {
        'row': row_number,
        'fio': fio,
        'new_email': email,
        'match_checked': False,
        'abiturient_match_count': 0,
        'abiturient_current_email': '',
        'student_match_count': 0,
        'student_current_email': '',
        'import_action': 'skip',
        'action_label': 'Пропустить',
        'badge_class': 'status-danger',
        'status': '',
    }

def set_email_source_preview_status(preview, action, label, badge_class, status):
    preview.update({
        'import_action': action,
        'action_label': label,
        'badge_class': badge_class,
        'status': status,
    })

def build_email_source_update_plan(file_path, campaign_year=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    df = read_tabular_upload(file_path)
    df.columns = [str(column).strip() for column in df.columns]

    source_columns = {}
    missing_columns = []
    for field, aliases in EMAIL_SOURCE_REQUIRED_COLUMNS.items():
        column = find_dataframe_column_casefold(df.columns, aliases)
        if column is None:
            missing_columns.append(next(iter(aliases)).title())
        else:
            source_columns[field] = column
    if missing_columns:
        raise ValueError(
            'В дополнительном файле с почтой нужны столбцы: Фамилия, Имя Отчество, Почта. '
            f"Не найдены: {', '.join(missing_columns)}."
        )

    issues = []
    preview_rows = []
    valid_rows = []
    for row_number, (_, row) in enumerate(df.iterrows(), start=2):
        surname = clean_upload_text(row.get(source_columns['surname']))
        name_patronymic = clean_upload_text(row.get(source_columns['name_patronymic']))
        email = clean_upload_text(row.get(source_columns['email']))
        fio = ' '.join(part for part in (surname, name_patronymic) if part)
        preview = email_source_preview_row(row_number, fio, email)
        preview_rows.append(preview)

        missing_values = []
        if not surname:
            missing_values.append('Фамилия')
        if not name_patronymic:
            missing_values.append('Имя Отчество')
        if not email:
            missing_values.append('Почта')
        if missing_values:
            message = f"Не заполнено: {', '.join(missing_values)}. Почта не изменена."
            issues.append(upload_report_item(row_number, 'Исходный файл', message))
            set_email_source_preview_status(preview, 'skip', 'Пропустить', 'status-danger', message)
            continue
        if len(name_patronymic.split()) < 2:
            message = (
                f'Для «{fio}» не указано отчество. Сопоставление выполняется по фамилии, '
                'имени и отчеству; почта не изменена.'
            )
            issues.append(upload_report_item(row_number, 'ФИО', message))
            set_email_source_preview_status(preview, 'skip', 'Пропустить', 'status-danger', message)
            continue
        if not is_valid_email(email):
            message = f'Почта выглядит некорректно: {email}. Изменения не внесены.'
            issues.append(upload_report_item(row_number, 'Почта', message))
            set_email_source_preview_status(preview, 'skip', 'Пропустить', 'status-danger', message)
            continue

        valid_rows.append({
            'row': row_number,
            'fio': fio,
            'fio_key': normalize_email_source_fio_key(fio),
            'email': email,
            'preview': preview,
        })

    rows_by_fio = {}
    for source_row in valid_rows:
        rows_by_fio.setdefault(source_row['fio_key'], []).append(source_row)

    unique_rows = []
    duplicate_count = 0
    conflicting_source_count = 0
    for fio_rows in rows_by_fio.values():
        distinct_emails = {row['email'].casefold() for row in fio_rows}
        if len(distinct_emails) > 1:
            conflicting_source_count += 1
            for row in fio_rows:
                message = f"Для ФИО «{row['fio']}» в файле указаны разные адреса. Изменения не внесены."
                issues.append(upload_report_item(row['row'], 'Почта', message))
                set_email_source_preview_status(
                    row['preview'], 'skip', 'Конфликт', 'status-danger', message
                )
            continue
        unique_rows.append(fio_rows[0])
        for row in fio_rows[1:]:
            duplicate_count += 1
            message = f"ФИО «{row['fio']}» повторяется с той же почтой. Повторная строка пропущена."
            issues.append(upload_report_item(row['row'], 'ФИО', message))
            set_email_source_preview_status(
                row['preview'], 'skip', 'Повтор', 'status-warning', message
            )

    context = build_email_source_person_context(campaign_year)
    operations = []
    matched_rows = 0
    matched_abiturient_rows = 0
    matched_student_rows = 0
    unchanged_email_count = 0
    corrected_invalid_student_emails = 0
    fio_warning_count = 0
    not_found_count = 0
    ambiguous_count = 0

    for source_row in unique_rows:
        preview = source_row['preview']
        fio_key = source_row['fio_key']
        matches = {
            kind: context[kind].get(fio_key, [])
            for kind in ('abiturients', 'students')
        }
        preview['match_checked'] = True
        preview['abiturient_match_count'] = len(matches['abiturients'])
        preview['student_match_count'] = len(matches['students'])
        if len(matches['abiturients']) == 1:
            preview['abiturient_current_email'] = matches['abiturients'][0]['email']
        if len(matches['students']) == 1:
            preview['student_current_email'] = matches['students'][0]['email']

        ambiguous_kinds = [kind for kind, records in matches.items() if len(records) > 1]
        if ambiguous_kinds:
            ambiguous_count += 1
            labels = ', '.join(
                f"{email_source_kind_label(kind)}ов: {len(matches[kind])}"
                for kind in ambiguous_kinds
            )
            message = (
                f"ФИО «{source_row['fio']}» найдено в нескольких записях ({labels}). "
                'Нельзя выбрать запись однозначно; почта не изменена.'
            )
            issues.append(upload_report_item(source_row['row'], 'ФИО', message))
            set_email_source_preview_status(
                preview, 'skip', 'Неоднозначно', 'status-warning', message
            )
            continue

        exact_records = [records[0] for records in matches.values() if records]
        if not exact_records:
            suggestion = find_email_source_fio_suggestion(source_row['fio'], context)
            fio_warning_count += 1
            if suggestion and suggestion.get('ambiguous'):
                alternatives = '» или «'.join(suggestion.get('alternatives') or [])
                message = (
                    f"ФИО «{source_row['fio']}» не совпало точно. Похожи записи «{alternatives}». "
                    'Возможна ошибка ручного ввода; почта не изменена.'
                )
            elif suggestion:
                sources = ' и '.join(sorted(
                    email_source_kind_label(kind) for kind in suggestion.get('kinds', set())
                ))
                message = (
                    f"ФИО «{source_row['fio']}» не совпало точно. Похоже на «{suggestion['fio']}» "
                    f"({suggestion['similarity']}%, {sources}). Возможна ошибка ручного ввода; "
                    'почта не изменена.'
                )
            else:
                not_found_count += 1
                message = (
                    f"ФИО «{source_row['fio']}» не найдено среди абитуриентов активной кампании "
                    'и студентов. Почта не изменена.'
                )
            issues.append(upload_report_item(source_row['row'], 'ФИО', message))
            set_email_source_preview_status(
                preview, 'fio_review', 'Проверить ФИО', 'status-warning', message
            )
            continue

        matched_rows += 1
        matched_abiturient_rows += int(bool(matches['abiturients']))
        matched_student_rows += int(bool(matches['students']))
        updated_kinds = []
        unchanged_kinds = []
        for record in exact_records:
            kind_label = email_source_kind_label(record['kind'])
            if record['email'].casefold() == source_row['email'].casefold():
                unchanged_email_count += 1
                unchanged_kinds.append(kind_label)
                continue
            if (
                record['kind'] == 'students'
                and record['email']
                and not is_valid_email(record['email'])
            ):
                corrected_invalid_student_emails += 1
            operations.append({
                'kind': record['kind'],
                'id': record['id'],
                'email': source_row['email'],
            })
            updated_kinds.append(kind_label)

        if updated_kinds:
            message = f"Будет обновлена почта: {', '.join(updated_kinds)}."
            if unchanged_kinds:
                message += f" Уже совпадает: {', '.join(unchanged_kinds)}."
            set_email_source_preview_status(
                preview, 'update', 'Обновить почту', 'status-success', message
            )
        else:
            message = f"Почта уже совпадает: {', '.join(unchanged_kinds)}."
            set_email_source_preview_status(
                preview, 'unchanged', 'Без изменений', 'status-success', message
            )

    updated_abiturients = sum(1 for operation in operations if operation['kind'] == 'abiturients')
    updated_students = sum(1 for operation in operations if operation['kind'] == 'students')
    return {
        'total': int(len(df)),
        'matched_rows': matched_rows,
        'matched_abiturient_rows': matched_abiturient_rows,
        'matched_student_rows': matched_student_rows,
        'updated_abiturients': updated_abiturients,
        'updated_students': updated_students,
        'planned_update_count': len(operations),
        'unchanged_email_count': unchanged_email_count,
        'corrected_invalid_student_emails': corrected_invalid_student_emails,
        'fio_warning_count': fio_warning_count,
        'not_found_count': not_found_count,
        'ambiguous_count': ambiguous_count,
        'duplicate_count': duplicate_count,
        'conflicting_source_count': conflicting_source_count,
        'updated_candidates': 0,
        'issues': issues,
        'preview_rows': preview_rows,
        '_operations': operations,
    }

def apply_email_source_updates(file_path, campaign_year=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    summary = build_email_source_update_plan(file_path, campaign_year)
    operations = summary.pop('_operations', [])
    backup_path = create_database_backup('before_email_source_updates') if operations else None
    updated_candidates = 0
    with sqlite3.connect(DB_PATH) as conn:
        for operation in operations:
            table = 'abiturients' if operation['kind'] == 'abiturients' else 'students'
            conn.execute(
                f'UPDATE {table} SET email=? WHERE id=?',
                (operation['email'], operation['id'])
            )
            if operation['kind'] == 'abiturients':
                updated_candidates += conn.execute(
                    '''
                    UPDATE enrollment_candidates
                    SET email=?
                    WHERE abiturient_id=? AND campaign_year=?
                    ''',
                    (operation['email'], operation['id'], campaign_year)
                ).rowcount
        log_action(
            'email_source_import',
            'campaign',
            campaign_year,
            (
                f"rows={summary['total']}; matched={summary['matched_rows']}; "
                f"abiturients={summary['updated_abiturients']}; students={summary['updated_students']}; "
                f"fio_warnings={summary['fio_warning_count']}; ambiguous={summary['ambiguous_count']}; "
                f"issues={len(summary['issues'])}; "
                f"backup={os.path.basename(backup_path) if backup_path else ''}"
            ),
            conn
        )
    summary['updated_candidates'] = updated_candidates
    return summary

def process_email_source_updates(file_path, campaign_year=None):
    return apply_email_source_updates(file_path, campaign_year)

def build_email_source_upload_report(summary, applied=False):
    action_word = 'Обновлено' if applied else 'Будет обновлено'
    return build_upload_report(
        'Отчет по дополнительному источнику почты',
        summary['total'],
        summary['issues'],
        [
            f"Обработано строк: {summary['total']}",
            f"Точно совпало по ФИО: {summary['matched_rows']}",
            f"Найдено среди абитуриентов: {summary['matched_abiturient_rows']}",
            f"Найдено среди студентов: {summary['matched_student_rows']}",
            f"{action_word} почт абитуриентов: {summary['updated_abiturients']}",
            f"{action_word} почт студентов: {summary['updated_students']}",
            f"Уже совпадали: {summary['unchanged_email_count']}",
        ]
    )

def process_abiturients_updates(file_path, campaign_year=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    df = read_tabular_upload(file_path)
    df.columns = [str(column).strip() for column in df.columns]
    if not any(column.casefold() in {'договор', 'dogovor', 'source_dogovor'} for column in df.columns):
        raise ValueError('В файле обновлений нужен столбец Договор')

    with sqlite3.connect(DB_PATH) as conn:
        rows = conn.execute(
            'SELECT id, dogovor FROM abiturients WHERE campaign_year=?',
            (campaign_year,)
        ).fetchall()
        dogovor_to_id = {
            normalize_dogovor_key(dogovor): abiturient_id
            for abiturient_id, dogovor in rows
            if normalize_dogovor_key(dogovor)
        }

    updated_email = 0
    updated_paid = 0
    not_found = []
    not_found_rows = []
    errors = []
    touched_abiturient_ids = set()
    backup_path = create_database_backup('before_abiturients_updates')
    with sqlite3.connect(DB_PATH) as conn:
        for row_number, (_, row) in enumerate(df.iterrows(), start=2):
            dogovor = find_row_value_casefold(row, {'Договор', 'dogovor', 'source_dogovor'})
            dogovor_key = normalize_dogovor_key(clean_upload_text(dogovor))
            if not dogovor_key:
                errors.append({
                    'row': row_number,
                    'field': 'Договор',
                    'message': 'Не указан номер договора. Строка пропущена.',
                })
                continue
            abiturient_id = dogovor_to_id.get(dogovor_key)
            if not abiturient_id:
                dogovor_text = clean_upload_text(dogovor)
                not_found.append(dogovor_text)
                not_found_rows.append({
                    'row': row_number,
                    'field': 'Договор',
                    'message': f'Договор не найден в кампании {campaign_year}: {dogovor_text}',
                })
                continue

            email = clean_upload_text(find_row_value_casefold(row, {'Email', 'email', 'Почта', 'почта'}))
            if email and not is_valid_email(email):
                errors.append({
                    'row': row_number,
                    'field': 'Email',
                    'message': f'Почта выглядит некорректно: {email}',
                })
                email = ''
            paid_raw = find_row_value_casefold(row, {'paid', 'Оплата', 'оплата', 'Оплачен', 'оплачен'})
            paid_text = clean_upload_text(paid_raw)
            paid_value = parse_paid_value(paid_raw)
            if paid_text and paid_value is None:
                errors.append({
                    'row': row_number,
                    'field': 'Оплата',
                    'message': f'Не удалось распознать значение оплаты: {paid_text}. Используйте да/нет, 1/0, оплачен/не оплачен.',
                })
            if email:
                conn.execute('UPDATE abiturients SET email=? WHERE id=?', (email, abiturient_id))
                updated_email += 1
                touched_abiturient_ids.add(abiturient_id)
            if paid_value is not None:
                conn.execute('UPDATE abiturients SET paid=? WHERE id=?', (paid_value, abiturient_id))
                updated_paid += 1
                touched_abiturient_ids.add(abiturient_id)

        sync_summary = sync_enrollment_candidates_for_abiturients(conn, touched_abiturient_ids, campaign_year)

        log_action(
            'abiturients_updates_import',
            'campaign',
            campaign_year,
            (
                f"rows={len(df)}; email={updated_email}; paid={updated_paid}; "
                f"candidates_created={sync_summary['created']}; candidates_updated={sync_summary['updated']}; "
                f"candidates_removed={sync_summary['removed']}; "
                f"not_found={len(not_found)}; errors={len(errors)}; "
                f"backup={os.path.basename(backup_path) if backup_path else ''}"
            ),
            conn
        )
    return {
        'total': int(len(df)),
        'updated_email': updated_email,
        'updated_paid': updated_paid,
        'candidate_sync': sync_summary,
        'not_found': not_found,
        'not_found_rows': not_found_rows,
        'errors': errors,
    }

def build_abiturients_updates_template():
    output = io.BytesIO()
    template_df = pd.DataFrame(columns=['Договор', 'Email', 'Оплата'])
    help_df = pd.DataFrame([
        {
            'Поле': 'Договор',
            'Что указать': 'Номер договора абитуриента. Это обязательное поле, по нему система ищет запись.',
            'Пример': '2026-СД-0001-11И',
        },
        {
            'Поле': 'Email',
            'Что указать': 'Новая электронная почта. Можно оставить пустым, если почту обновлять не нужно.',
            'Пример': 'student@example.ru',
        },
        {
            'Поле': 'Оплата',
            'Что указать': 'Статус оплаты договора. Подойдут значения: да/нет, 1/0, оплачен/не оплачен.',
            'Пример': 'да',
        },
    ])
    example_df = pd.DataFrame([
        {
            'Договор': '2026-СД-0001-11И',
            'Email': 'student@example.ru',
            'Оплата': 'да',
        },
        {
            'Договор': '2026-ЛД-0002-9И',
            'Email': '',
            'Оплата': 'нет',
        },
    ])
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        template_df.to_excel(writer, sheet_name='Шаблон', index=False)
        help_df.to_excel(writer, sheet_name='Подсказка', index=False)
        example_df.to_excel(writer, sheet_name='Пример', index=False)
    output.seek(0)
    return output

def get_login_attempt_key(username):
    return (get_client_ip(), str(username or '').strip().lower())

def prune_login_attempts(conn, now=None):
    now = now or time()
    conn.execute(
        'DELETE FROM login_attempts WHERE attempted_at < ?',
        (now - LOGIN_WINDOW_SECONDS,)
    )

def get_recent_login_attempts(key):
    now = time()
    ip_address, username = key
    with sqlite3.connect(DB_PATH) as conn:
        prune_login_attempts(conn, now)
        cur = conn.execute(
            '''
            SELECT attempted_at
            FROM login_attempts
            WHERE ip_address=? AND username=? AND attempted_at >= ?
            ORDER BY attempted_at ASC
            ''',
            (ip_address, username, now - LOGIN_WINDOW_SECONDS)
        )
        return [row[0] for row in cur.fetchall()]

def get_login_lockout(username):
    key = get_login_attempt_key(username)
    attempts = get_recent_login_attempts(key)
    if len(attempts) < LOGIN_MAX_ATTEMPTS:
        return 0
    return max(1, int(LOGIN_WINDOW_SECONDS - (time() - attempts[0])))

def record_login_failure(username):
    now = time()
    key = get_login_attempt_key(username)
    ip_address, normalized_username = key
    with sqlite3.connect(DB_PATH) as conn:
        prune_login_attempts(conn, now)
        conn.execute(
            'INSERT INTO login_attempts (ip_address, username, attempted_at) VALUES (?, ?, ?)',
            (ip_address, normalized_username, now)
        )

def clear_login_failures(username):
    ip_address, normalized_username = get_login_attempt_key(username)
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            'DELETE FROM login_attempts WHERE ip_address=? AND username=?',
            (ip_address, normalized_username)
        )

def is_password_hash(value):
    return str(value or '').startswith(PASSWORD_HASH_PREFIXES)

def hash_user_password(password):
    return generate_password_hash(str(password or ''))

def verify_user_password(stored_password, candidate_password):
    if stored_password in (None, '') or candidate_password in (None, ''):
        return False
    stored_password = str(stored_password or '')
    candidate_password = str(candidate_password or '')
    if is_password_hash(stored_password):
        return check_password_hash(stored_password, candidate_password)
    return hmac.compare_digest(stored_password, candidate_password)

def migrate_user_passwords(conn):
    if not table_exists(conn, 'users'):
        return

    cur = conn.execute('SELECT id, password FROM users')
    for user_id, password in cur.fetchall():
        if password is not None and not is_password_hash(password):
            conn.execute(
                'UPDATE users SET password=? WHERE id=?',
                (hash_user_password(password), user_id)
            )

def init_db():
    with sqlite3.connect(DB_PATH) as conn:
        create_abiturients_table(conn)
        create_enrollment_orders_table(conn)
        create_enrollment_candidates_table(conn)
        migrate_abiturients_table(conn)
        migrate_legacy_students_abiturients_table(conn)
        conn.execute(f'''
            CREATE TABLE IF NOT EXISTS pending_duplicates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                fio TEXT,
                dogovor TEXT,
                login TEXT,
                campaign_year TEXT NOT NULL DEFAULT '{LEGACY_CAMPAIGN_YEAR}',
                fam TEXT,
                imotch TEXT
            )
        ''')
        conn.execute(f'''
            CREATE TABLE IF NOT EXISTS login_conflicts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                fio TEXT,
                dogovor TEXT,
                login TEXT,
                campaign_year TEXT NOT NULL DEFAULT '{LEGACY_CAMPAIGN_YEAR}',
                fam TEXT,
                imotch TEXT,
                conflict_time TEXT DEFAULT (datetime('now', 'localtime'))
            )
        ''')
        conn.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE,
                password TEXT,
                fio TEXT,
                position TEXT,
                role TEXT,
                approved INTEGER DEFAULT 0
            )
        ''')
        conn.execute('''
            CREATE TABLE IF NOT EXISTS login_attempts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ip_address TEXT NOT NULL,
                username TEXT NOT NULL,
                attempted_at REAL NOT NULL
            )
        ''')
        migrate_user_passwords(conn)
        conn.execute('''
            CREATE TABLE IF NOT EXISTS students (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE,
                password TEXT,
                email TEXT,
                firstname TEXT,
                lastname TEXT,
                cohort1 TEXT,
                cohort2 TEXT,
                source_campaign_year TEXT,
                source_dogovor TEXT,
                source_fio TEXT
            )
        ''')
        conn.execute('''
            CREATE TABLE IF NOT EXISTS students_duplicates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT,
                password TEXT,
                email TEXT,
                firstname TEXT,
                lastname TEXT,
                cohort1 TEXT,
                cohort2 TEXT
            )
        ''')
        conn.execute(f'''
            CREATE TABLE IF NOT EXISTS groups (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT UNIQUE,
                group_year TEXT NOT NULL DEFAULT '{DEFAULT_CAMPAIGN_YEAR}',
                is_hidden INTEGER DEFAULT 0
            )
        ''')
        ensure_group_year_column(conn)
        ensure_campaign_column(conn, 'pending_duplicates')
        ensure_campaign_column(conn, 'login_conflicts')
        ensure_students_origin_columns(conn)
        ensure_students_duplicates_columns(conn)
        create_enrollment_orders_table(conn)
        create_enrollment_candidates_table(conn)
        create_audit_log_table(conn)
        create_student_group_transfers_table(conn)
        create_campaign_settings_table(conn)
        create_login_generation_settings_table(conn)
        conn.execute('CREATE INDEX IF NOT EXISTS idx_abiturients_campaign_year ON abiturients (campaign_year)')
        conn.execute('CREATE INDEX IF NOT EXISTS idx_pending_duplicates_campaign_year ON pending_duplicates (campaign_year)')
        conn.execute('CREATE INDEX IF NOT EXISTS idx_login_conflicts_campaign_year ON login_conflicts (campaign_year)')
        conn.execute('CREATE INDEX IF NOT EXISTS idx_login_attempts_key_time ON login_attempts (ip_address, username, attempted_at)')

init_db()

# Initialize default admin user from local environment only.
_default_admin_password = os.environ.get('ADMIN_DEFAULT_PASSWORD')
with sqlite3.connect(DB_PATH) as conn:
    admin_exists = conn.execute(
        "SELECT 1 FROM users WHERE username=?",
        ("admin",)
    ).fetchone()
    if not admin_exists:
        if not _default_admin_password:
            raise RuntimeError(
                "ADMIN_DEFAULT_PASSWORD is not set. Copy .env.example to .env "
                "and set a strong local admin password before first launch."
            )
        conn.execute(
            "INSERT INTO users (username, password, role, approved) VALUES (?, ?, ?, ?)",
            ("admin", hash_user_password(_default_admin_password), "admin", 1)
        )

def get_campaign_years():
    years = set(BASE_CAMPAIGN_YEARS)
    with sqlite3.connect(DB_PATH) as conn:
        create_campaign_settings_table(conn)
        cur = conn.execute(
            "SELECT campaign_year FROM campaign_settings WHERE campaign_year IS NOT NULL AND campaign_year != ''"
        )
        years.update(str(row[0]) for row in cur.fetchall() if row[0])
        for table in ('abiturients', 'pending_duplicates', 'login_conflicts', 'enrollment_candidates', 'enrollment_orders', 'enrollment_order_uploads'):
            if 'campaign_year' not in get_table_columns(conn, table):
                continue
            cur = conn.execute(
                f"SELECT DISTINCT campaign_year FROM {table} WHERE campaign_year IS NOT NULL AND campaign_year != ''"
            )
            years.update(str(row[0]) for row in cur.fetchall() if row[0])
        if table_exists(conn, 'students') and 'source_campaign_year' in get_table_columns(conn, 'students'):
            cur = conn.execute(
                "SELECT DISTINCT source_campaign_year FROM students WHERE source_campaign_year IS NOT NULL AND source_campaign_year != ''"
            )
            years.update(str(row[0]) for row in cur.fetchall() if row[0])
    return sorted(years)

def get_latest_campaign_year():
    years = []
    with sqlite3.connect(DB_PATH) as conn:
        create_campaign_settings_table(conn)
        cur = conn.execute(
            "SELECT campaign_year FROM campaign_settings WHERE campaign_year IS NOT NULL AND campaign_year != ''"
        )
        years.extend(str(row[0]) for row in cur.fetchall() if row[0])
        for table in ('abiturients', 'pending_duplicates', 'login_conflicts', 'enrollment_candidates', 'enrollment_orders', 'enrollment_order_uploads'):
            if 'campaign_year' not in get_table_columns(conn, table):
                continue
            cur = conn.execute(
                f"SELECT DISTINCT campaign_year FROM {table} WHERE campaign_year IS NOT NULL AND campaign_year != ''"
            )
            years.extend(str(row[0]) for row in cur.fetchall() if row[0])
        if table_exists(conn, 'students') and 'source_campaign_year' in get_table_columns(conn, 'students'):
            cur = conn.execute(
                "SELECT DISTINCT source_campaign_year FROM students WHERE source_campaign_year IS NOT NULL AND source_campaign_year != ''"
            )
            years.extend(str(row[0]) for row in cur.fetchall() if row[0])
    return max(years) if years else DEFAULT_CAMPAIGN_YEAR

def get_pinned_campaign_year():
    with sqlite3.connect(DB_PATH) as conn:
        create_campaign_settings_table(conn)
        row = conn.execute(
            '''
            SELECT campaign_year
            FROM campaign_settings
            WHERE is_active=1
            ORDER BY active_at DESC, campaign_year DESC
            LIMIT 1
            '''
        ).fetchone()
    if not row or not _campaign_year_re.fullmatch(str(row[0] or '')):
        return None
    return str(row[0])

def get_default_campaign_year():
    return get_pinned_campaign_year() or get_latest_campaign_year()

def get_active_campaign_year():
    if not has_request_context():
        return DEFAULT_CAMPAIGN_YEAR
    requested_year = request.values.get('campaign_year')
    fallback_year = session.get('campaign_year') or get_default_campaign_year()
    campaign_year = normalize_campaign_year(requested_year or fallback_year, fallback_year)
    if session.get('user'):
        session['campaign_year'] = campaign_year
    return campaign_year

def get_group_years(selected_year=None, include_base=False):
    years = set(BASE_CAMPAIGN_YEARS if include_base else [])
    if selected_year:
        years.add(normalize_group_year(selected_year, DEFAULT_CAMPAIGN_YEAR))

    with sqlite3.connect(DB_PATH) as conn:
        create_campaign_settings_table(conn)
        cur = conn.execute(
            "SELECT campaign_year FROM campaign_settings WHERE campaign_year IS NOT NULL AND campaign_year != ''"
        )
        years.update(normalize_group_year(row[0], DEFAULT_CAMPAIGN_YEAR) for row in cur.fetchall() if row[0])
        if table_exists(conn, 'groups'):
            columns = get_table_columns(conn, 'groups')
            if 'group_year' in columns:
                cur = conn.execute(
                    "SELECT DISTINCT group_year FROM groups WHERE group_year IS NOT NULL AND group_year != ''"
                )
                years.update(normalize_group_year(row[0], DEFAULT_CAMPAIGN_YEAR) for row in cur.fetchall() if row[0])
            else:
                cur = conn.execute("SELECT name FROM groups WHERE name IS NOT NULL AND name != ''")
                years.update(infer_group_year(row[0], DEFAULT_CAMPAIGN_YEAR) for row in cur.fetchall() if row[0])

    return sorted(years)

def get_used_logins(campaign_year, rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    used_logins = set()
    with sqlite3.connect(DB_PATH) as conn:
        for table in ('abiturients', 'pending_duplicates', 'login_conflicts'):
            if rules['unique_scope'] == 'global':
                cur = conn.execute(f"SELECT login FROM {table} WHERE login IS NOT NULL")
            else:
                cur = conn.execute(
                    f"SELECT login FROM {table} WHERE campaign_year=? AND login IS NOT NULL",
                    (campaign_year,)
                )
            used_logins.update(str(row[0]).strip() for row in cur.fetchall() if str(row[0]).strip())
        if table_exists(conn, 'students') and 'username' in get_table_columns(conn, 'students'):
            cur = conn.execute(
                "SELECT username FROM students WHERE username IS NOT NULL AND username != ''"
            )
            used_logins.update(str(row[0]).strip() for row in cur.fetchall() if str(row[0]).strip())
    return used_logins

def get_prefixed_logins(table, prefix, campaign_year, rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    with sqlite3.connect(DB_PATH) as conn:
        if rules['unique_scope'] == 'global':
            cur = conn.execute(
                f"SELECT login FROM {table} WHERE login LIKE ?",
                (f'{prefix}%',)
            )
        else:
            cur = conn.execute(
                f"SELECT login FROM {table} WHERE campaign_year=? AND login LIKE ?",
                (campaign_year, f'{prefix}%')
            )
        return set(row[0] for row in cur.fetchall())

def next_numbered_login(prefix, existing_logins, width=3):
    try:
        width = max(1, min(8, int(width)))
    except (TypeError, ValueError):
        width = 3
    number = 1
    while True:
        login = f"{prefix}{number:0{width}d}"
        if login not in existing_logins:
            return login
        number += 1

@app.context_processor
def inject_campaign_context():
    if not has_request_context():
        return {}
    return {
        'campaign_years': get_campaign_years(),
        'active_campaign_year': get_active_campaign_year(),
    }

def is_login_exists(login, campaign_year=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    login = str(login or '').strip()
    if not login:
        return False
    rules = get_login_generation_rules()
    with sqlite3.connect(DB_PATH) as conn:
        if rules['unique_scope'] == 'global':
            cur = conn.execute('SELECT 1 FROM abiturients WHERE login=?', (login,))
        else:
            cur = conn.execute(
                'SELECT 1 FROM abiturients WHERE login=? AND campaign_year=?',
                (login, campaign_year)
            )
        if cur.fetchone() is not None:
            return True
        if table_exists(conn, 'students') and 'username' in get_table_columns(conn, 'students'):
            cur = conn.execute('SELECT 1 FROM students WHERE username=?', (login,))
            if cur.fetchone() is not None:
                return True
        return False

def is_fio_duplicate(fam, campaign_year=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            'SELECT fio FROM abiturients WHERE fam=? AND campaign_year=?',
            (fam, campaign_year)
        )
        return cur.fetchall()

def save_abiturient(fio, dogovor, login, fam, imotch, campaign_year=None):
    dogovor = normalize_dogovor_storage_text(dogovor)
    campaign_year = normalize_campaign_year(campaign_year, infer_campaign_year(dogovor))
    if is_login_exists(login, campaign_year):
        raise sqlite3.IntegrityError(f'Login already exists: {login}')
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            'INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch) VALUES (?, ?, ?, ?, ?, ?)',
            (fio, dogovor, login, campaign_year, fam, imotch)
        )
        conn.commit()

def save_pending_duplicate(fio, dogovor, login, fam, imotch, campaign_year=None):
    dogovor = normalize_dogovor_storage_text(dogovor)
    campaign_year = normalize_campaign_year(campaign_year, infer_campaign_year(dogovor))
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            'INSERT INTO pending_duplicates (fio, dogovor, login, campaign_year, fam, imotch) VALUES (?, ?, ?, ?, ?, ?)',
            (fio, dogovor, login, campaign_year, fam, imotch)
        )
        conn.commit()

def save_login_conflict(fio, dogovor, login, fam, imotch, campaign_year=None):
    dogovor = normalize_dogovor_storage_text(dogovor)
    campaign_year = normalize_campaign_year(campaign_year, infer_campaign_year(dogovor))
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            'INSERT INTO login_conflicts (fio, dogovor, login, campaign_year, fam, imotch) VALUES (?, ?, ?, ?, ?, ?)',
            (fio, dogovor, login, campaign_year, fam, imotch)
        )
        conn.commit()

def clean_upload_text(value):
    if pd.isna(value):
        return ''
    return str(value).strip()

def normalize_fio_key(value):
    return ' '.join(str(value or '').split()).casefold()

def normalize_fio_similarity_key(value):
    value = normalize_fio_key(value).replace('ё', 'е')
    return re.sub(r'[^0-9a-zа-я]+', '', value)

def split_fio_for_storage(fio):
    parts = [part for part in str(fio or '').split() if part]
    fam = parts[0] if parts else ''
    imotch = ' '.join(parts[1:]) if len(parts) > 1 else ''
    return fam, imotch

def normalize_dogovor_key(value):
    normalized = normalize_dogovor_text(value)
    return normalized if normalized else ''

def normalize_specialty_key(value):
    raw_value = str(value or '').strip()
    compact = re.sub(r'[^0-9A-Za-zА-Яа-яЁё.]+', '', raw_value).upper().replace('Ё', 'Е')
    if not compact:
        return ''
    alias = _specialty_aliases.get(compact)
    if alias:
        return re.sub(r'[^0-9A-Za-zА-Яа-яЁё.]+', '', str(alias)).upper().replace('Ё', 'Е').casefold()
    for alias_key, alias_value in sorted(_specialty_aliases.items(), key=lambda item: len(item[0]), reverse=True):
        if (len(alias_key) >= 5 or re.fullmatch(r'\d{2}\.\d{2}\.\d{2}', alias_key)) and alias_key in compact:
            return re.sub(r'[^0-9A-Za-zА-Яа-яЁё.]+', '', str(alias_value)).upper().replace('Ё', 'Е').casefold()
    return compact.casefold()

def get_dogovor_specialty_key(dogovor, rules=None):
    parts = parse_dogovor_parts(dogovor, rules or get_login_generation_rules())
    if not parts:
        return ''
    return normalize_specialty_key(parts.get('spec_label'))

def make_enrollment_match_key(fio, specialty):
    fio_key = normalize_fio_key(fio)
    specialty_key = normalize_specialty_key(specialty)
    if not fio_key or not specialty_key:
        return None
    return fio_key, specialty_key

def make_abiturient_enrollment_match_key(fio, dogovor, rules=None):
    return make_enrollment_match_key(fio, get_dogovor_specialty_key(dogovor, rules))

def get_enrollment_order_map(campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    with sqlite3.connect(DB_PATH) as conn:
        rows = conn.execute(
            '''
            SELECT
                o.id, o.upload_id, o.fio_key, o.specialty_key, o.group_name,
                o.specialty, o.order_number, o.order_date,
                u.original_filename, u.stored_filename
            FROM enrollment_orders o
            LEFT JOIN enrollment_order_uploads u ON u.id=o.upload_id
            WHERE o.campaign_year=?
            ''',
            (campaign_year,)
        ).fetchall()
    order_map = {}
    for (
        order_id, upload_id, fio_key, specialty_key, group_name,
        specialty, order_number, order_date, original_filename, stored_filename
    ) in rows:
        if fio_key and specialty_key:
            order_map[(fio_key, specialty_key)] = {
                'id': order_id,
                'upload_id': upload_id,
                'group_name': group_name or '',
                'specialty': specialty or '',
                'order_number': order_number or '',
                'order_date': order_date or '',
                'original_filename': original_filename or '',
                'stored_filename': stored_filename or '',
            }
    return order_map

def get_enrollment_order_match_for_abiturient(fio, dogovor, campaign_year, order_map=None, rules=None):
    order_map = order_map if order_map is not None else get_enrollment_order_map(campaign_year)
    match_key = make_abiturient_enrollment_match_key(fio, dogovor, rules)
    if not match_key:
        return None
    return order_map.get(match_key)

def is_abiturient_in_enrollment_order(fio, dogovor, campaign_year, order_map=None, rules=None):
    return bool(get_enrollment_order_match_for_abiturient(fio, dogovor, campaign_year, order_map, rules))

def get_abiturient_order_keys(campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    rules = get_login_generation_rules()
    keys = set()
    with sqlite3.connect(DB_PATH) as conn:
        rows = conn.execute(
            '''
            SELECT fio, dogovor
            FROM abiturients
            WHERE campaign_year=?
            ''',
            (campaign_year,)
        ).fetchall()
    for fio, dogovor in rows:
        match_key = make_abiturient_enrollment_match_key(fio, dogovor, rules)
        if match_key:
            keys.add(match_key)
    return keys

def get_existing_person_keys(campaign_year):
    keys = set()
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            '''
            SELECT fio FROM abiturients
            WHERE campaign_year=? AND fio IS NOT NULL AND fio <> ''
            ''',
            (campaign_year,)
        )
        keys.update(normalize_fio_key(row[0]) for row in cur.fetchall() if normalize_fio_key(row[0]))

        cur = conn.execute(
            '''
            SELECT source_fio, lastname, firstname FROM students
            WHERE source_campaign_year=?
            ''',
            (campaign_year,)
        )
        for source_fio, lastname, firstname in cur.fetchall():
            fio_key = normalize_fio_key(source_fio)
            if not fio_key:
                fio_key = normalize_fio_key(' '.join(part for part in (lastname, firstname) if part))
            if fio_key:
                keys.add(fio_key)
    return keys

def get_abiturient_import_reference_records(campaign_year):
    references = {
        'people': {},
        'dogovors': {
            'abiturients': {},
            'pending_duplicates': {},
            'students': {},
            'login_conflicts': {},
        },
    }

    def add_record(source_key, source_label, fio, dogovor, login, include_person=False):
        record = {
            'source': source_label,
            'fio': ' '.join(str(fio or '').split()),
            'dogovor': str(dogovor or '').strip(),
            'login': str(login or '').strip(),
        }
        dogovor_key = normalize_dogovor_key(record['dogovor'])
        if dogovor_key:
            references['dogovors'][source_key].setdefault(dogovor_key, []).append(record)
        fio_key = normalize_fio_key(record['fio'])
        if include_person and fio_key:
            references['people'].setdefault(fio_key, []).append(record)

    with sqlite3.connect(DB_PATH) as conn:
        for source_key, source_label in (
            ('abiturients', 'список абитуриентов'),
            ('pending_duplicates', 'раздел дублей'),
            ('login_conflicts', 'конфликты логинов'),
        ):
            rows = conn.execute(
                f'''
                SELECT fio, dogovor, login
                FROM {source_key}
                WHERE campaign_year=?
                ''',
                (campaign_year,),
            ).fetchall()
            for fio, dogovor, login in rows:
                add_record(
                    source_key,
                    source_label,
                    fio,
                    dogovor,
                    login,
                    include_person=source_key == 'abiturients',
                )

        rows = conn.execute(
            '''
            SELECT source_fio, lastname, firstname, source_dogovor, username
            FROM students
            WHERE source_campaign_year=?
            ''',
            (campaign_year,),
        ).fetchall()
        for source_fio, lastname, firstname, source_dogovor, username in rows:
            fio = str(source_fio or '').strip() or ' '.join(
                part for part in (lastname, firstname) if str(part or '').strip()
            ).strip()
            add_record(
                'students',
                'список студентов',
                fio,
                source_dogovor,
                username,
                include_person=True,
            )
    return references

def format_abiturient_import_reference_hint(records, prefix):
    details = []
    for record in list(records or [])[:3]:
        details.append(
            (
                f"{record.get('source') or 'база'}: "
                f"ФИО — {record.get('fio') or 'не указано'}, "
                f"договор — {record.get('dogovor') or 'не указан'}, "
                f"логин — {record.get('login') or 'не указан'}"
            )
        )
    hidden_count = max(0, len(records or []) - len(details))
    if hidden_count:
        details.append(f'и ещё записей: {hidden_count}')
    return f"{prefix} {'; '.join(details)}." if details else prefix

def get_existing_dogovor_keys(campaign_year):
    dogovor_keys = {
        'abiturients': set(),
        'pending_duplicates': set(),
        'students': set(),
        'login_conflicts': set(),
    }
    with sqlite3.connect(DB_PATH) as conn:
        sources = (
            ('abiturients', 'SELECT dogovor FROM abiturients WHERE campaign_year=?'),
            ('pending_duplicates', 'SELECT dogovor FROM pending_duplicates WHERE campaign_year=?'),
            ('login_conflicts', 'SELECT dogovor FROM login_conflicts WHERE campaign_year=?'),
        )
        for source_name, query in sources:
            cur = conn.execute(query, (campaign_year,))
            dogovor_keys[source_name].update(
                key for key in (normalize_dogovor_key(row[0]) for row in cur.fetchall()) if key
            )

        cur = conn.execute(
            '''
            SELECT source_dogovor FROM students
            WHERE source_campaign_year=? AND source_dogovor IS NOT NULL AND source_dogovor <> ''
            ''',
            (campaign_year,)
        )
        dogovor_keys['students'].update(
            key for key in (normalize_dogovor_key(row[0]) for row in cur.fetchall()) if key
        )
    return dogovor_keys

def summarize_abiturients_import(df, campaign_year):
    action_counts = df['import_action'].value_counts().to_dict() if not df.empty else {}
    status_counts = df['import_status'].value_counts().to_dict() if not df.empty else {}
    return {
        'campaign_year': campaign_year,
        'total': int(len(df)),
        'ready_count': int(action_counts.get('create', 0)),
        'duplicate_count': int(action_counts.get('duplicate', 0)),
        'conflict_count': int(action_counts.get('conflict', 0)),
        'warning_count': int(df['has_warning'].sum()) if 'has_warning' in df else 0,
        'status_counts': status_counts,
    }

def dataframe_preview_rows(df):
    preview_df = df.copy()
    preview_df = preview_df.where(pd.notnull(preview_df), '')
    rows = preview_df[ABITURIENT_RESULT_COLUMNS].to_dict(orient='records')
    warning_values = preview_df['has_warning'].tolist() if 'has_warning' in preview_df else [False] * len(rows)
    status_hints = preview_df['status_hint'].tolist() if 'status_hint' in preview_df else [''] * len(rows)
    for row, has_warning, status_hint in zip(rows, warning_values, status_hints):
        row['has_warning'] = bool(has_warning)
        row['status_hint'] = str(status_hint or '').strip()
    return rows

def upload_report_item(row, field, message):
    return {
        'row': row,
        'field': field,
        'message': message,
    }

def build_upload_report(title, total, items, summary=None, limit=UPLOAD_REPORT_LIMIT):
    report_items = list(items or [])
    return {
        'title': title,
        'total': int(total or 0),
        'summary': list(summary or []),
        'issue_count': len(report_items),
        'items': report_items[:limit],
        'hidden_count': max(0, len(report_items) - limit),
    }

ABITURIENT_PREVIEW_REPORT_MESSAGES = {
    'Пустое ФИО': 'Не заполнено ФИО. Строка попадет в конфликты, пока ФИО не исправят.',
    'Ошибка договора': 'Не удалось разобрать номер договора. Проверьте год, специальность и базу 9/11.',
    'Договор уже есть у абитуриента': 'Такой договор уже есть в списке абитуриентов. Строка попадет в конфликты.',
    'Договор уже есть у студента': 'Такой договор уже есть у студента. Строка попадет в конфликты, чтобы не создать повтор.',
    'Договор уже ожидает проверки в дублях': 'Такой договор уже находится в дублях. Строка попадет на ручную проверку.',
    'Договор уже есть в конфликтах': 'Такой договор уже есть в конфликтах. Сначала разберите существующий конфликт.',
    'Договор повторяется в файле импорта': 'Такой договор повторяется внутри загруженного файла. Повторная строка попадет в конфликты.',
    'Возможный тёзка, договор другой; будет добавлен': 'ФИО уже встречается в системе, но договор другой. Это может быть тёзка, проверьте перед подтверждением.',
}

def build_abiturients_preview_report(df, summary):
    items = []
    for _, row in df.iterrows():
        status = clean_upload_text(row.get('import_status', ''))
        action = clean_upload_text(row.get('import_action', ''))
        has_warning = bool(row.get('has_warning', False))
        if action == 'create' and not has_warning:
            continue
        field = 'ФИО' if status == 'Пустое ФИО' or 'тёз' in status else 'Договор'
        message = ABITURIENT_PREVIEW_REPORT_MESSAGES.get(status, status)
        items.append(upload_report_item(int(row.get('_row_number', 0)), field, message))

    if not items:
        return None
    return build_upload_report(
        'Отчет по проверке абитуриентов',
        summary['total'],
        items,
        [
            f"К добавлению: {summary['ready_count']}",
            f"В дубли: {summary['duplicate_count']}",
            f"В конфликты: {summary['conflict_count']}",
            f"Возможных тёзок: {summary['warning_count']}",
        ]
    )

def build_abiturients_import_plan(file_path, campaign_year=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    df = read_tabular_upload(file_path)
    df.columns = [str(column).strip() for column in df.columns]
    missing_columns = sorted(ABITURIENT_REQUIRED_COLUMNS - set(df.columns))
    if missing_columns:
        raise ValueError(f"В файле отсутствуют обязательные столбцы: {', '.join(missing_columns)}")
    if df.empty:
        raise ValueError('Файл не содержит строк для импорта')

    df = df.copy()
    df['_row_number'] = range(2, len(df) + 2)
    df['ФИО'] = df['ФИО'].apply(clean_upload_text)
    df['Договор'] = df['Договор'].apply(lambda value: normalize_dogovor_storage_text(clean_upload_text(value)))

    fio_split = df['ФИО'].str.split(' ', n=2, expand=True)
    for column_index in range(3):
        if column_index not in fio_split.columns:
            fio_split[column_index] = ''
    df['Фамилия'] = fio_split[0].fillna('').astype(str).str.strip()
    second_name = fio_split[1].fillna('').astype(str).str.strip()
    third_name = fio_split[2].fillna('').astype(str).str.strip()
    df['Имя_Отчество'] = (second_name + ' ' + third_name).str.strip()
    login_rules = get_login_generation_rules()
    df['login_parts'] = df['Договор'].apply(lambda value: parse_dogovor_parts(value, login_rules))
    df['campaign_year'] = campaign_year

    used_logins = get_used_logins(campaign_year, login_rules)
    error_prefix = login_rules['error_prefix']
    duplicate_prefix = login_rules['duplicate_prefix']
    number_width = login_rules['number_width']
    used_duplicate_logins = set(get_prefixed_logins('pending_duplicates', duplicate_prefix, campaign_year, login_rules))
    reference_records = get_abiturient_import_reference_records(campaign_year)
    known_person_records = reference_records['people']
    known_person_keys = set(known_person_records)
    existing_dogovor_records = reference_records['dogovors']
    existing_dogovor_keys = {
        source: set(records)
        for source, records in existing_dogovor_records.items()
    }
    planned_dogovor_records = {}

    logins = []
    actions = []
    statuses = []
    status_hints = []
    is_duplicate_values = []
    has_warning_values = []

    for _, row in df.iterrows():
        fio = row['ФИО']
        fam = row['Фамилия']
        fio_key = normalize_fio_key(fio)
        dogovor_key = normalize_dogovor_key(row['Договор'])
        login_parts = row['login_parts']

        if not fio or not fam:
            login = next_numbered_login(error_prefix, used_logins, number_width)
            used_logins.add(login)
            logins.append(login)
            actions.append('conflict')
            statuses.append('Пустое ФИО')
            status_hints.append(
                f"Расхождение найдено в загруженном файле, строка {row['_row_number']}: поле ФИО не заполнено."
            )
            is_duplicate_values.append(False)
            has_warning_values.append(False)
            continue

        if not login_parts:
            login = next_numbered_login(error_prefix, used_logins, number_width)
            used_logins.add(login)
            logins.append(login)
            actions.append('conflict')
            statuses.append('Ошибка договора')
            status_hints.append(
                (
                    f"Расхождение найдено в загруженном файле, строка {row['_row_number']}: "
                    f"договор «{row['Договор'] or 'пусто'}» не соответствует правилам формирования логина."
                )
            )
            is_duplicate_values.append(False)
            has_warning_values.append(False)
            continue

        if dogovor_key in existing_dogovor_keys['abiturients']:
            login = next_numbered_login(error_prefix, used_logins, number_width)
            used_logins.add(login)
            logins.append(login)
            actions.append('conflict')
            statuses.append('Договор уже есть у абитуриента')
            status_hints.append(format_abiturient_import_reference_hint(
                existing_dogovor_records['abiturients'].get(dogovor_key, []),
                'Совпадение договора найдено в базе.',
            ))
            is_duplicate_values.append(False)
            has_warning_values.append(False)
            continue

        if dogovor_key in existing_dogovor_keys['students']:
            login = next_numbered_login(error_prefix, used_logins, number_width)
            used_logins.add(login)
            logins.append(login)
            actions.append('conflict')
            statuses.append('Договор уже есть у студента')
            status_hints.append(format_abiturient_import_reference_hint(
                existing_dogovor_records['students'].get(dogovor_key, []),
                'Совпадение договора найдено в базе.',
            ))
            is_duplicate_values.append(False)
            has_warning_values.append(False)
            continue

        if dogovor_key in existing_dogovor_keys['pending_duplicates']:
            login = next_numbered_login(duplicate_prefix, used_duplicate_logins, number_width)
            used_duplicate_logins.add(login)
            used_logins.add(login)
            logins.append(login)
            actions.append('duplicate')
            statuses.append('Договор уже ожидает проверки в дублях')
            status_hints.append(format_abiturient_import_reference_hint(
                existing_dogovor_records['pending_duplicates'].get(dogovor_key, []),
                'Совпадение договора найдено в базе.',
            ))
            is_duplicate_values.append(True)
            has_warning_values.append(False)
            continue

        if dogovor_key in existing_dogovor_keys['login_conflicts']:
            login = next_numbered_login(error_prefix, used_logins, number_width)
            used_logins.add(login)
            logins.append(login)
            actions.append('conflict')
            statuses.append('Договор уже есть в конфликтах')
            status_hints.append(format_abiturient_import_reference_hint(
                existing_dogovor_records['login_conflicts'].get(dogovor_key, []),
                'Совпадение договора найдено в базе.',
            ))
            is_duplicate_values.append(False)
            has_warning_values.append(False)
            continue

        if dogovor_key in planned_dogovor_records:
            login = next_numbered_login(error_prefix, used_logins, number_width)
            used_logins.add(login)
            logins.append(login)
            actions.append('conflict')
            statuses.append('Договор повторяется в файле импорта')
            status_hints.append(format_abiturient_import_reference_hint(
                [planned_dogovor_records[dogovor_key]],
                'Совпадение договора найдено ранее в загруженном файле.',
            ))
            is_duplicate_values.append(False)
            has_warning_values.append(False)
            continue

        login = next_login_from_parts(login_parts, used_logins, login_rules)

        used_logins.add(login)
        namesake_records = list(known_person_records.get(fio_key, []))
        is_possible_namesake = fio_key in known_person_keys
        known_person_keys.add(fio_key)
        logins.append(login)
        actions.append('create')
        statuses.append('Возможный тёзка, договор другой; будет добавлен' if is_possible_namesake else 'Будет добавлен')
        if is_possible_namesake:
            status_hints.append(format_abiturient_import_reference_hint(
                namesake_records,
                f"Совпадение ФИО найдено, но договор «{row['Договор']}» отличается.",
            ))
        else:
            status_hints.append('Расхождений с существующими записями и другими строками файла не найдено.')
        is_duplicate_values.append(False)
        has_warning_values.append(is_possible_namesake)

        planned_record = {
            'source': f"загруженный файл, строка {row['_row_number']}",
            'fio': fio,
            'dogovor': row['Договор'],
            'login': login,
        }
        planned_dogovor_records[dogovor_key] = planned_record
        known_person_records.setdefault(fio_key, []).append(planned_record)

    df['login'] = logins
    df['import_action'] = actions
    df['import_status'] = statuses
    df['status_hint'] = status_hints
    df['is_duplicate'] = is_duplicate_values
    df['has_warning'] = has_warning_values

    return df, summarize_abiturients_import(df, campaign_year)

def create_abiturients_result_file(df):
    output_path = make_temp_upload_path('xlsx', prefix=ABITURIENTS_IMPORT_RESULT_PREFIX)
    result_df = df[ABITURIENT_RESULT_COLUMNS].copy()
    result_df.to_excel(output_path, index=False)
    return output_path

def apply_abiturients_import(file_path, campaign_year=None):
    df, summary = build_abiturients_import_plan(file_path, campaign_year)
    backup_path = create_database_backup('before_abiturients_import')
    with sqlite3.connect(DB_PATH) as conn:
        for _, row in df.iterrows():
            action = row['import_action']
            values = (
                row['ФИО'], row['Договор'], row['login'], row['campaign_year'],
                row['Фамилия'], row['Имя_Отчество']
            )
            if action == 'create':
                conn.execute(
                    'INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch) VALUES (?, ?, ?, ?, ?, ?)',
                    values
                )
            elif action == 'duplicate':
                conn.execute(
                    'INSERT INTO pending_duplicates (fio, dogovor, login, campaign_year, fam, imotch) VALUES (?, ?, ?, ?, ?, ?)',
                    values
                )
            else:
                conn.execute(
                    'INSERT INTO login_conflicts (fio, dogovor, login, campaign_year, fam, imotch) VALUES (?, ?, ?, ?, ?, ?)',
                    values
                )
        log_action(
            'abiturients_import',
            'campaign',
            summary['campaign_year'],
            (
                f"rows={summary['total']}; create={summary['ready_count']}; "
                f"duplicates={summary['duplicate_count']}; conflicts={summary['conflict_count']}; "
                f"backup={os.path.basename(backup_path) if backup_path else ''}"
            ),
            conn
        )

    return create_abiturients_result_file(df), summary

def process_excel(file_path, campaign_year=None):
    output_path, _summary = apply_abiturients_import(file_path, campaign_year)
    return output_path

def summarize_students_import(df):
    action_counts = df['import_action'].value_counts().to_dict() if not df.empty else {}
    return {
        'total': int(len(df)),
        'ready_count': int(action_counts.get('create', 0)),
        'duplicate_count': int(action_counts.get('duplicate', 0)),
        'skipped_count': int(action_counts.get('skip', 0)),
    }

def build_students_import_plan(file_path):
    course_groups_enabled = are_course_groups_enabled()
    df = read_tabular_upload(file_path)
    df.columns = [str(column).strip() for column in df.columns]
    missing_columns = [column for column in STUDENT_UPLOAD_REQUIRED_COLUMNS if column not in df.columns]
    if missing_columns:
        readable_columns = ', '.join(STUDENT_UPLOAD_FIELD_LABELS.get(column, column) for column in missing_columns)
        raise ValueError(f"В файле студентов не хватает столбцов: {readable_columns}")
    if df.empty:
        raise ValueError('Файл студентов не содержит строк для загрузки')

    df = df.copy()
    df['_row_number'] = range(2, len(df) + 2)
    for column in STUDENT_UPLOAD_REQUIRED_COLUMNS:
        df[column] = df[column].apply(clean_upload_text)
    df['cohort1'] = df['cohort1'].apply(normalize_group_name)
    if 'cohort2' not in df.columns:
        df['cohort2'] = ''
    df['cohort2'] = df['cohort2'].apply(lambda value: normalize_cohort2(clean_upload_text(value)))
    for column in ('source_dogovor', 'source_fio'):
        if column not in df.columns:
            df[column] = ''
        df[column] = df[column].apply(clean_upload_text)
    df['source_campaign_year'] = df['source_dogovor'].apply(
        lambda value: infer_campaign_year(value, DEFAULT_CAMPAIGN_YEAR) if value else ''
    )

    with sqlite3.connect(DB_PATH) as conn:
        existing_usernames = {
            clean_upload_text(row[0])
            for row in conn.execute("SELECT username FROM students WHERE username IS NOT NULL AND username<>''")
        }

    planned_usernames = set()
    actions = []
    statuses = []
    errors = []

    for row_index, row in df.iterrows():
        row_number = int(row['_row_number'])
        cohort2_warning = ''
        missing_values = [
            STUDENT_UPLOAD_FIELD_LABELS[column]
            for column in STUDENT_UPLOAD_REQUIRED_COLUMNS
            if not row[column]
        ]
        if missing_values:
            actions.append('skip')
            statuses.append(f"Не заполнено: {', '.join(missing_values)}")
            errors.append(upload_report_item(
                row_number,
                'Обязательные поля',
                f"Не заполнено: {', '.join(missing_values)}. Строка будет пропущена."
            ))
            continue

        if not is_valid_email(row['email']):
            actions.append('skip')
            statuses.append('Некорректная почта')
            errors.append(upload_report_item(
                row_number,
                'Email',
                f"Почта выглядит некорректно: {row['email']}. Строка будет пропущена."
            ))
            continue

        expected_cohort2 = (derive_cohort2(row['cohort1']) or '') if course_groups_enabled else ''
        uploaded_cohort2 = normalize_cohort2(row.get('cohort2'))
        if not course_groups_enabled:
            df.at[row_index, 'cohort2'] = ''
        else:
            if uploaded_cohort2 and not is_supported_cohort2(uploaded_cohort2):
                actions.append('skip')
                statuses.append('Некорректная глобальная группа курса')
                errors.append(upload_report_item(
                    row_number,
                    'Глобальная группа курса',
                    f"Глобальная группа курса {uploaded_cohort2} не входит в поддерживаемый список. Строка будет пропущена."
                ))
                continue
            if expected_cohort2:
                if uploaded_cohort2 and uploaded_cohort2 != expected_cohort2:
                    actions.append('skip')
                    statuses.append('cohort2 не соответствует академической группе')
                    errors.append(upload_report_item(
                        row_number,
                        'Глобальная группа курса',
                        (
                            f"Для академической группы {row['cohort1']} ожидается {expected_cohort2}, "
                            f"но в файле указано {uploaded_cohort2}. Строка будет пропущена."
                        )
                    ))
                    continue
                df.at[row_index, 'cohort2'] = expected_cohort2
            elif uploaded_cohort2:
                df.at[row_index, 'cohort2'] = ''
                cohort2_warning = 'cohort2 не применена'
                errors.append(upload_report_item(
                    row_number,
                    'Глобальная группа курса',
                    (
                        f"Для академической группы {row['cohort1']} нет правила автоматического определения "
                        "глобальной группы курса. Значение cohort2 из файла не будет записано."
                    )
                ))

        username = row['username']
        if username in existing_usernames:
            actions.append('duplicate')
            statuses.append('Логин уже есть у студента')
            errors.append(upload_report_item(
                row_number,
                'Логин',
                f"Логин {username} уже есть у студента. Строка будет перенесена в дубли студентов."
            ))
            continue

        if username in planned_usernames:
            actions.append('duplicate')
            statuses.append('Логин повторяется в файле')
            errors.append(upload_report_item(
                row_number,
                'Логин',
                f"Логин {username} повторяется в загруженном файле. Повторная строка будет перенесена в дубли студентов."
            ))
            continue

        planned_usernames.add(username)
        actions.append('create')
        statuses.append('Будет добавлен' if not cohort2_warning else f'Будет добавлен; {cohort2_warning}')

    df['import_action'] = actions
    df['import_status'] = statuses

    summary = summarize_students_import(df)
    summary['errors'] = errors
    return df, summary

def student_preview_rows(df):
    preview_df = df.copy()
    preview_df = preview_df.where(pd.notnull(preview_df), '')
    rows = preview_df[STUDENT_PREVIEW_COLUMNS].to_dict(orient='records')
    for row in rows:
        action = row.get('import_action')
        if action == 'create':
            row['action_label'] = 'Будет добавлен'
            row['badge_class'] = 'status-success'
        elif action == 'duplicate':
            row['action_label'] = 'В дубли'
            row['badge_class'] = 'status-warning'
        else:
            row['action_label'] = 'Пропущен'
            row['badge_class'] = 'status-danger'
    return rows

def build_students_preview_report(summary):
    report_items = summary.get('errors') or []
    if not report_items:
        return None
    return build_upload_report(
        'Отчет по проверке студентов',
        summary['total'],
        report_items,
        [
            f"К добавлению: {summary['ready_count']}",
            f"В дубли студентов: {summary['duplicate_count']}",
            f"Будет пропущено: {summary['skipped_count']}",
        ]
    )

def apply_students_import(file_path):
    df, summary = build_students_import_plan(file_path)
    backup_path = create_database_backup('before_students_import')
    with sqlite3.connect(DB_PATH) as conn:
        for _, row in df.iterrows():
            action = row['import_action']
            if action == 'create':
                conn.execute(
                    '''
                    INSERT INTO students
                        (username, password, email, firstname, lastname, cohort1, cohort2, source_campaign_year, source_dogovor, source_fio)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''',
                    (
                        row["username"], row["password"], row["email"],
                        row["firstname"], row["lastname"], row["cohort1"], row["cohort2"],
                        row["source_campaign_year"], row["source_dogovor"], row["source_fio"]
                    )
                )
            elif action == 'duplicate':
                conn.execute(
                    '''INSERT INTO students_duplicates (username, password, email, firstname, lastname, cohort1, cohort2)
                       VALUES (?, ?, ?, ?, ?, ?, ?)''',
                    (
                        row["username"], row["password"], row["email"],
                        row["firstname"], row["lastname"], row["cohort1"], row["cohort2"]
                    )
                )

        log_action(
            'students_import',
            'students',
            '',
            (
                f"rows={summary['total']}; inserted={summary['ready_count']}; "
                f"duplicates={summary['duplicate_count']}; skipped={summary['skipped_count']}; "
                f"backup={os.path.basename(backup_path) if backup_path else ''}"
            ),
            conn
        )

    return {
        'total': summary['total'],
        'inserted_count': summary['ready_count'],
        'duplicate_count': summary['duplicate_count'],
        'skipped_count': summary['skipped_count'],
        'errors': summary.get('errors') or [],
    }

def process_students_excel(file_path):
    return apply_students_import(file_path)

def build_group_name(year_code, specialty, base, subgroup='1'):
    year_code = re.sub(r'\D+', '', str(year_code or '').strip())
    if len(year_code) == 4 and year_code.startswith('20'):
        year_code = year_code[-2:]

    specialty = normalize_specialty(specialty)
    base = normalize_group_base(base)
    subgroup = re.sub(r'\D+', '', str(subgroup or '').strip()) or '1'
    if not year_code or not specialty or not base:
        return ''

    return normalize_group_name(f'{year_code}{specialty}-{base}-{subgroup}')

def get_specialty_display_name(value):
    specialty_key = normalize_specialty_key(value)
    for code, display_name in SPECIALTY_DISPLAY_NAMES.items():
        if normalize_specialty_key(code) == specialty_key:
            return display_name
    specialty = normalize_specialty(value)
    return SPECIALTY_DISPLAY_NAMES.get(specialty, specialty or str(value or '').strip())

def group_specialty_key(group_name):
    normalized = normalize_group_name(group_name)
    head = normalized.split('-', 1)[0]
    match = _group_head_re.fullmatch(head)
    if not match:
        return ''
    return normalize_specialty_key(match.group(2))

def group_matches_specialty(group_name, specialty_key):
    specialty_key = str(specialty_key or '').strip()
    return bool(specialty_key and group_specialty_key(group_name) == specialty_key)

def group_matches_order_group(group_name, order_group_name):
    order_group_name = str(order_group_name or '').strip()
    if not order_group_name:
        return True
    return normalize_group_name(group_name).casefold() == normalize_group_name(order_group_name).casefold()

def build_groups_template_csv(group_year=None):
    group_year = normalize_group_year(group_year, DEFAULT_CAMPAIGN_YEAR)
    year_code = group_year[-2:]
    rows = ['group_year;group_name']
    for specialty, base in GROUPS_TEMPLATE_EXAMPLES:
        rows.append(f'{group_year};{build_group_name(year_code, specialty, base)}')
    return '\n'.join(rows) + '\n'

def is_valid_group_name(group_name):
    return bool(_group_name_re.fullmatch(group_name or ''))

def base_group_name(group_name):
    parts = (group_name or '').split('-')
    if len(parts) > 2 and parts[-1].isdigit():
        return '-'.join(parts[:-1])
    return group_name

def group_subgroup_index(group_name):
    parts = (group_name or '').split('-')
    if len(parts) > 2 and parts[-1].isdigit():
        return int(parts[-1])
    return 1

def subgroup_name(root_group, index):
    return root_group if index == 1 else f'{root_group}-{index}'

def find_row_value(row, aliases):
    normalized_row = {
        str(key or '').strip().casefold(): value
        for key, value in row.items()
    }
    for alias in aliases:
        value = normalized_row.get(alias.casefold())
        if value is not None and str(value).strip():
            return str(value).strip()
    return ''

def read_groups_csv(file_path):
    last_error = None
    for encoding in ('utf-8-sig', 'cp1251', 'utf-8'):
        try:
            with open(file_path, newline='', encoding=encoding) as csv_file:
                sample = csv_file.read(4096)
                csv_file.seek(0)
                delimiter = ';' if sample.count(';') >= sample.count(',') else ','
                reader = csv.DictReader(csv_file, delimiter=delimiter)
                if not reader.fieldnames:
                    raise ValueError('в CSV не найдены заголовки')
                return list(reader)
        except UnicodeDecodeError as exc:
            last_error = exc
            continue

    if last_error:
        raise ValueError('не удалось прочитать CSV в кодировке UTF-8 или Windows-1251')
    return []

def group_exists_casefold(existing_groups, group_name):
    return group_name.casefold() in existing_groups

def process_groups_csv(file_path, fallback_group_year=None):
    fallback_group_year = normalize_group_year(fallback_group_year, get_active_campaign_year())
    rows = read_groups_csv(file_path)
    created_groups = []
    skipped_groups = []
    errors = []

    with sqlite3.connect(DB_PATH) as conn:
        existing_groups = {
            row[0].casefold(): row[0]
            for row in conn.execute('SELECT name FROM groups')
        }

        for row_number, row in enumerate(rows, start=2):
            group_name = find_row_value(row, ['group_name', 'name', 'group', 'группа', 'название', 'название_группы'])
            if group_name:
                group_name = normalize_group_name(group_name)
            else:
                group_name = build_group_name(
                    find_row_value(row, ['year_code', 'year', 'год', 'год_поступления']),
                    find_row_value(row, ['specialty', 'spec', 'специальность', 'направление']),
                    find_row_value(row, ['base', 'база', 'база_классов']),
                    find_row_value(row, ['subgroup', 'subgroup_number', 'подгруппа', 'номер_подгруппы']),
                )

            group_year_value = find_row_value(row, ['group_year', 'folder_year', 'year_folder', 'папка', 'год_папки', 'год_групп', 'год_группы'])
            if group_year_value:
                group_year = normalize_group_year(group_year_value, infer_group_year(group_name, fallback_group_year))
                if str(group_year_value).strip() not in (group_year, group_year[-2:]):
                    errors.append(f'строка {row_number}: неверный год папки "{group_year_value}"')
                    continue
            else:
                group_year = infer_group_year(group_name, fallback_group_year)

            if not group_name:
                errors.append(f'строка {row_number}: не указана группа')
                continue
            if not is_valid_group_name(group_name):
                errors.append(f'строка {row_number}: неверный формат группы "{group_name}"')
                continue
            if infer_group_year(group_name, group_year) != group_year:
                errors.append(f'строка {row_number}: группа "{group_name}" не соответствует папке {group_year}')
                continue

            if group_exists_casefold(existing_groups, group_name):
                skipped_groups.append(group_name)
                continue
            conn.execute('INSERT INTO groups (name, group_year) VALUES (?, ?)', (group_name, group_year))
            existing_groups[group_name.casefold()] = group_name
            created_groups.append(f'{group_year}: {group_name}')

    return {
        'created': created_groups,
        'skipped': skipped_groups,
        'errors': errors,
    }

def get_group_student_count(conn, group_name):
    cur = conn.execute('SELECT COUNT(*) FROM students WHERE cohort1=?', (group_name,))
    return cur.fetchone()[0]

def get_next_subgroup_name(conn, group_name, group_year=None):
    root_group = base_group_name(group_name)
    current_index = group_subgroup_index(group_name)
    existing_indices = {current_index}

    if group_year:
        rows = conn.execute('SELECT name FROM groups WHERE group_year=?', (group_year,))
    else:
        rows = conn.execute('SELECT name FROM groups')

    for row in rows:
        existing_name = row[0]
        if base_group_name(existing_name).casefold() == root_group.casefold():
            existing_indices.add(group_subgroup_index(existing_name))

    next_index = current_index + 1
    while next_index in existing_indices:
        next_index += 1
    return f'{root_group}-{next_index}'

def is_last_subgroup(conn, group_name, group_year=None):
    root_group = base_group_name(group_name)
    current_index = group_subgroup_index(group_name)
    max_index = current_index

    if group_year:
        rows = conn.execute('SELECT name FROM groups WHERE group_year=?', (group_year,))
    else:
        rows = conn.execute('SELECT name FROM groups')

    for row in rows:
        existing_name = row[0]
        if base_group_name(existing_name).casefold() == root_group.casefold():
            max_index = max(max_index, group_subgroup_index(existing_name))

    return current_index == max_index

def get_groups_with_counts(conn, group_year=None, include_hidden=False):
    group_year = normalize_group_year(group_year, get_active_campaign_year()) if group_year else None
    course_groups_enabled = are_course_groups_enabled()
    groups = []
    if group_year:
        if include_hidden:
            rows = conn.execute(
                'SELECT name, group_year, is_hidden FROM groups WHERE group_year=? ORDER BY is_hidden, name',
                (group_year,)
            )
        else:
            rows = conn.execute(
                'SELECT name, group_year, is_hidden FROM groups WHERE group_year=? AND COALESCE(is_hidden, 0)=0 ORDER BY name',
                (group_year,)
            )
    else:
        if include_hidden:
            rows = conn.execute('SELECT name, group_year, is_hidden FROM groups ORDER BY group_year, is_hidden, name')
        else:
            rows = conn.execute(
                'SELECT name, group_year, is_hidden FROM groups WHERE COALESCE(is_hidden, 0)=0 ORDER BY group_year, name'
            )

    for row in rows:
        name = row[0]
        row_group_year = row[1] or infer_group_year(name, DEFAULT_CAMPAIGN_YEAR)
        is_hidden = bool(row[2])
        count = get_group_student_count(conn, name)
        is_full = count >= MAX_GROUP_STUDENTS
        can_create_next = not is_hidden and is_full and is_last_subgroup(conn, name, row_group_year)
        groups.append({
            'name': name,
            'group_year': row_group_year,
            'specialty_key': group_specialty_key(name),
            'cohort2': (derive_cohort2(name) or '') if course_groups_enabled else '',
            'is_hidden': is_hidden,
            'count': count,
            'capacity': MAX_GROUP_STUDENTS,
            'fill': f'{count}/{MAX_GROUP_STUDENTS}',
            'is_full': is_full,
            'can_create_next': can_create_next,
            'next_name': get_next_subgroup_name(conn, name, row_group_year) if can_create_next else '',
        })
    return groups

def get_candidate_group_options(groups, specialty_key=None):
    if not specialty_key:
        return groups
    return [group for group in groups if group_matches_specialty(group['name'], specialty_key)]

def login_group_specialty_options(rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    options = {
        str(code).strip().casefold(): label
        for code, label in LOGIN_GROUP_SPECIALTY_CODES.items()
    }
    for label, code in (rules.get('spec_codes') or {}).items():
        code_key = str(code or '').strip().casefold()
        specialty = normalize_specialty(label)
        if code_key and code_key not in options:
            options[code_key] = specialty
    return sorted(options.items(), key=lambda item: len(item[0]), reverse=True)

def login_group_base_preference(base_label):
    base_label = normalize_group_base(base_label)
    preferred = {'11': 0, '9': 0, '11И': 0, '9И': 0, 'М': 0}
    return preferred.get(base_label, 1), len(base_label), base_label

def login_group_base_options(rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    options = {
        code.casefold(): normalize_group_base(base)
        for code, base in LOGIN_GROUP_BASE_CODES.items()
    }
    for label, code in (rules.get('base_codes') or {}).items():
        code_key = str(code or '').strip().casefold()
        base_label = normalize_group_base(canonicalize_base_label(label))
        if not code_key or not base_label:
            continue
        current = options.get(code_key)
        if current is None or login_group_base_preference(base_label) < login_group_base_preference(current):
            options[code_key] = base_label
    return sorted(options.items(), key=lambda item: len(item[0]), reverse=True)

def parse_login_group_target(login, campaign_year=None, rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    value = re.sub(r'[\s\-]+', '', str(login or '')).strip()
    folded = value.casefold()
    match = re.fullmatch(r'(?P<year>\d{2})(?P<body>[0-9a-zа-яё]+)', folded)
    if not match:
        return {
            'ok': False,
            'error': 'логин не похож на учебный формат',
        }

    year_code = match.group('year')
    remainder = match.group('body')
    for spec_code, specialty_label in login_group_specialty_options(rules):
        if not remainder.startswith(spec_code):
            continue
        after_spec = remainder[len(spec_code):]
        for base_code, base_label in login_group_base_options(rules):
            if not after_spec.startswith(base_code):
                continue
            sequence = after_spec[len(base_code):]
            if not sequence or not re.fullmatch(r'\d+', sequence):
                continue
            group_root = normalize_group_name(f'{year_code}{specialty_label}-{base_label}')
            result = {
                'ok': True,
                'year_code': year_code,
                'specialty_code': spec_code,
                'specialty': specialty_label,
                'specialty_key': normalize_specialty_key(specialty_label),
                'base_label': base_label,
                'sequence': sequence,
                'group_root': group_root,
                'target_group': normalize_group_name(f'{group_root}-1'),
            }
            if campaign_year and year_code != str(campaign_year)[-2:]:
                result['warning'] = f'год в логине {year_code}, кампания {campaign_year}'
            return result

    return {
        'ok': False,
        'year_code': year_code,
        'error': 'не удалось разобрать специальность или базу образования в логине',
    }

def get_group_assignment_state(conn, group_year):
    rows = conn.execute(
        'SELECT name, group_year, COALESCE(is_hidden, 0) FROM groups'
    ).fetchall()
    visible = {}
    hidden = {}
    other_year = {}
    counts = {}
    for name, row_group_year, is_hidden in rows:
        normalized_name = normalize_group_name(name)
        group_key = normalized_name.casefold()
        normalized_year = normalize_group_year(row_group_year, infer_group_year(normalized_name, group_year))
        if normalized_year != group_year:
            other_year[group_key] = normalized_year
            continue
        if is_hidden:
            hidden[group_key] = normalized_name
            continue
        visible[group_key] = normalized_name
        counts[group_key] = get_group_student_count(conn, name)
    return {
        'visible': visible,
        'hidden': hidden,
        'other_year': other_year,
        'counts': counts,
    }

def reserve_exact_group_assignment(group_name, group_year, state, planned_counts, reserve=True):
    group_name = normalize_group_name(group_name)
    group_key = group_name.casefold()
    if not group_name or not is_valid_group_name(group_name):
        return None, False, 'группа имеет неверный формат'
    if infer_group_year(group_name, group_year) != group_year:
        return None, False, f'группа не относится к папке {group_year}'
    if group_key in state['hidden']:
        return None, False, f'группа {group_name} скрыта'
    if group_key in state['other_year']:
        return None, False, f'группа {group_name} уже есть в другой папке'

    target_group = state['visible'].get(group_key, group_name)
    current_count = state['counts'].get(group_key, 0)
    planned_count = planned_counts.get(group_key, 0)
    if current_count + planned_count >= MAX_GROUP_STUDENTS:
        return None, False, f'в группе {target_group} нет свободных мест'

    if reserve:
        planned_counts[group_key] = planned_count + 1
    return target_group, group_key not in state['visible'], ''

def reserve_group_by_root_assignment(root_group, group_year, state, planned_counts, reserve=True):
    root_group = normalize_group_name(root_group)
    if not root_group:
        return None, False, 'не удалось определить корень группы'
    index = 1
    while index < 100:
        group_name = normalize_group_name(f'{root_group}-{index}')
        group_key = group_name.casefold()
        if group_key in state['hidden']:
            index += 1
            continue
        if group_key in state['other_year']:
            return None, False, f'группа {group_name} уже есть в другой папке'
        target_group = state['visible'].get(group_key, group_name)
        current_count = state['counts'].get(group_key, 0)
        planned_count = planned_counts.get(group_key, 0)
        if current_count + planned_count < MAX_GROUP_STUDENTS:
            if reserve:
                planned_counts[group_key] = planned_count + 1
            return target_group, group_key not in state['visible'], ''
        index += 1
    return None, False, f'для {root_group} не удалось подобрать свободную подгруппу'

def get_selected_enrollment_candidate_rows(conn, campaign_year, ids):
    ids = [str(item) for item in ids if str(item).isdigit()]
    if not ids:
        return []
    placeholders = ','.join('?' for _ in ids)
    cur = conn.execute(
        f'''
        SELECT
            c.id AS candidate_id,
            c.abiturient_id,
            c.fio,
            c.dogovor,
            c.login,
            c.fam,
            c.imotch,
            c.email AS candidate_email,
            c.specialty,
            c.specialty_key,
            c.verification_status,
            c.order_group_name,
            a.id AS source_id,
            a.email AS source_email,
            a.paid AS source_paid
        FROM enrollment_candidates c
        LEFT JOIN abiturients a
            ON a.id=c.abiturient_id AND a.campaign_year=c.campaign_year
        WHERE c.campaign_year=? AND c.id IN ({placeholders})
        ORDER BY c.fio
        ''',
        [campaign_year] + ids
    )
    columns = [description[0] for description in cur.description]
    return [dict(zip(columns, row)) for row in cur.fetchall()]

def enrollment_candidate_readiness_issues(abiturient_id, email, paid):
    if not abiturient_id:
        return ['Исходная карточка абитуриента не найдена']
    issues = []
    if not str(email or '').strip():
        issues.append('В карточке абитуриента не указана почта')
    if not is_paid_person_value(paid):
        issues.append('В карточке абитуриента не отмечена оплата')
    return issues

def build_login_group_distribution_plan(conn, candidate_rows, campaign_year, group_year):
    group_year = normalize_group_year(group_year, campaign_year)
    rules = get_login_generation_rules()
    enrollment_order_required = is_enrollment_order_required(rules)
    course_groups_enabled = are_course_groups_enabled(rules)
    order_map = get_enrollment_order_map(campaign_year) if enrollment_order_required else {}
    state = get_group_assignment_state(conn, group_year)
    planned_counts = {}
    rows = []
    create_groups = set()

    sorted_rows = sorted(
        candidate_rows,
        key=lambda row: (normalize_fio_key(row.get('fio')), str(row.get('login') or ''))
    )
    for row in sorted_rows:
        issues = []
        notes = []
        target_group = ''
        source = ''
        will_create_group = False
        display_name = row.get('fio') or row.get('login') or str(row.get('candidate_id'))
        order_match = None
        if enrollment_order_required:
            order_match = get_enrollment_order_match_for_abiturient(
                row.get('fio'),
                row.get('dogovor'),
                campaign_year,
                order_map=order_map,
                rules=rules
            )

        if enrollment_order_required and (row.get('verification_status') != 'verified' or not order_match):
            issues.append('нет подтвержденного совпадения с приказом')

        email = ((row.get('source_email') or row.get('candidate_email') or '')).strip()
        issues.extend(enrollment_candidate_readiness_issues(row.get('source_id'), email, row.get('source_paid')))

        username = str(row.get('login') or '').strip()
        if not username:
            issues.append('не указан логин')
        elif conn.execute('SELECT 1 FROM students WHERE username=?', (username,)).fetchone():
            issues.append('такой логин уже есть в студентах')

        required_group = normalize_group_name(order_match.get('group_name') if order_match else '')
        parsed_login = parse_login_group_target(username, campaign_year, rules) if username else {'ok': False}
        if required_group:
            source = 'Приказ'
            should_reserve = not issues
            target_group, will_create_group, group_error = reserve_exact_group_assignment(
                required_group,
                group_year,
                state,
                planned_counts,
                reserve=should_reserve
            )
            notes.append('группа указана в приказе')
            if group_error:
                issues.append(group_error)
        else:
            source = 'Логин'
            if not parsed_login.get('ok'):
                issues.append(parsed_login.get('error') or 'не удалось разобрать логин')
            else:
                if parsed_login.get('warning'):
                    issues.append(parsed_login['warning'])
                if row.get('specialty_key') and parsed_login.get('specialty_key') != row.get('specialty_key'):
                    issues.append(
                        f"специальность логина {parsed_login.get('specialty')} не совпадает с кандидатом"
                    )
                should_reserve = not issues
                target_group, will_create_group, group_error = reserve_group_by_root_assignment(
                    parsed_login.get('group_root'),
                    group_year,
                    state,
                    planned_counts,
                    reserve=should_reserve
                )
                if group_error:
                    issues.append(group_error)

        if target_group and row.get('specialty_key') and not group_matches_specialty(target_group, row.get('specialty_key')):
            issues.append('предложенная группа не совпадает со специальностью кандидата')

        if will_create_group and target_group:
            create_groups.add(target_group)
            notes.append('группа будет создана при подтверждении')
        elif target_group:
            notes.append('группа уже есть в справочнике')

        can_migrate = not issues and bool(target_group)
        rows.append({
            'candidate_id': row.get('candidate_id'),
            'abiturient_id': row.get('abiturient_id'),
            'fio': display_name,
            'dogovor': row.get('dogovor') or '',
            'login': username,
            'email': email,
            'firstname': row.get('imotch') or '',
            'lastname': row.get('fam') or '',
            'specialty': row.get('specialty') or '',
            'target_group': target_group or '-',
            'cohort2': (derive_cohort2(target_group) or '-') if course_groups_enabled else '',
            'source': source or '-',
            'order_match': order_match or {},
            'will_create_group': will_create_group,
            'can_migrate': can_migrate,
            'status_label': 'Готов' if can_migrate else 'Проверить',
            'badge_class': 'status-success' if can_migrate else 'status-danger',
            'status_note': '; '.join(issues or notes),
            'sort_status': 'готов' if can_migrate else 'проверить',
        })

    ready_count = sum(1 for row in rows if row['can_migrate'])
    issue_count = len(rows) - ready_count
    return {
        'rows': rows,
        'summary': {
            'total': len(rows),
            'ready_count': ready_count,
            'issue_count': issue_count,
            'create_group_count': len(create_groups),
            'login_group_count': sum(1 for row in rows if row['source'] == 'Логин' and row['can_migrate']),
            'order_group_count': sum(1 for row in rows if row['source'] == 'Приказ' and row['can_migrate']),
            'can_confirm': bool(rows) and issue_count == 0,
        },
    }

def get_enrollment_candidate_keys(campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    keys = set()
    with sqlite3.connect(DB_PATH) as conn:
        rows = conn.execute(
            '''
            SELECT fio, specialty_key
            FROM enrollment_candidates
            WHERE campaign_year=?
            ''',
            (campaign_year,)
        ).fetchall()
    for fio, specialty_key in rows:
        match_key = make_enrollment_match_key(fio, specialty_key)
        if match_key:
            keys.add(match_key)
    return keys

def get_enrollment_candidate_match_context(campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    rows = []
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            '''
            SELECT c.id, c.fio, c.specialty, c.specialty_key, c.login, c.dogovor,
                   c.abiturient_id, a.fio
            FROM enrollment_candidates c
            LEFT JOIN abiturients a
                ON a.id=c.abiturient_id AND a.campaign_year=c.campaign_year
            WHERE c.campaign_year=?
            ''',
            (campaign_year,)
        )
        for candidate_id, fio, specialty, specialty_key, login, dogovor, abiturient_id, source_fio in cur.fetchall():
            fio_key = normalize_fio_key(fio)
            specialty_key = normalize_specialty_key(specialty_key or specialty)
            if not fio_key or not specialty_key:
                continue
            rows.append({
                'id': candidate_id,
                'fio': fio or '',
                'fio_key': fio_key,
                'fio_similarity_key': normalize_fio_similarity_key(fio),
                'source_fio': source_fio or '',
                'source_fio_key': normalize_fio_key(source_fio),
                'source_fio_similarity_key': normalize_fio_similarity_key(source_fio),
                'specialty': specialty or '',
                'specialty_key': specialty_key,
                'login': login or '',
                'dogovor': dogovor or '',
                'abiturient_id': abiturient_id,
            })
    keys = {(row['fio_key'], row['specialty_key']) for row in rows}
    by_specialty = {}
    by_key = {}
    for row in rows:
        by_specialty.setdefault(row['specialty_key'], []).append(row)
        by_key.setdefault((row['fio_key'], row['specialty_key']), row)
    return {
        'keys': keys,
        'by_key': by_key,
        'by_specialty': by_specialty,
    }

def find_fio_candidate_suggestion(order_fio, specialty_key, candidate_context):
    order_key = normalize_fio_similarity_key(order_fio)
    if not order_key or not specialty_key:
        return None
    best_candidate = None
    best_score = 0.0
    for candidate in candidate_context.get('by_specialty', {}).get(specialty_key, []):
        candidate_key = candidate.get('fio_similarity_key') or ''
        if not candidate_key or candidate_key == order_key:
            continue
        score = difflib.SequenceMatcher(None, order_key, candidate_key).ratio()
        if score > best_score:
            best_score = score
            best_candidate = candidate
    if not best_candidate or best_score < ENROLLMENT_FIO_SUGGESTION_THRESHOLD:
        return None
    suggestion = dict(best_candidate)
    suggestion['similarity'] = int(round(best_score * 100))
    return suggestion

def build_abiturient_fio_review_context(abiturient_rows, rules=None):
    rules = rules or get_login_generation_rules()
    rows = []
    for row in abiturient_rows:
        abiturient_id, fio, dogovor, login, email, paid, fam, imotch = row
        specialty_key = get_dogovor_specialty_key(dogovor, rules)
        item = {
            'id': abiturient_id,
            'fio': str(fio or '').strip(),
            'fio_key': normalize_fio_key(fio),
            'fio_similarity_key': normalize_fio_similarity_key(fio),
            'dogovor': str(dogovor or '').strip(),
            'login': str(login or '').strip(),
            'email': str(email or '').strip(),
            'paid': paid,
            'fam': str(fam or '').strip(),
            'imotch': str(imotch or '').strip(),
            'specialty_key': specialty_key,
            'specialty': get_specialty_display_name(specialty_key) if specialty_key else '',
        }
        rows.append(item)
    by_specialty = {}
    by_fio = {}
    by_id = {}
    for row in rows:
        by_id[row['id']] = row
        if row['specialty_key']:
            by_specialty.setdefault(row['specialty_key'], []).append(row)
        if row['fio_key']:
            by_fio.setdefault(row['fio_key'], []).append(row)
    return {
        'rows': rows,
        'by_specialty': by_specialty,
        'by_fio': by_fio,
        'by_id': by_id,
    }

def find_abiturient_fio_review(order_fio, specialty_key, context):
    fio_key = normalize_fio_key(order_fio)
    similarity_key = normalize_fio_similarity_key(order_fio)
    specialty_key = normalize_specialty_key(specialty_key)
    if not fio_key or not similarity_key or not specialty_key:
        return None

    exact_fio_rows = context.get('by_fio', {}).get(fio_key, [])
    specialty_conflicts = [
        row for row in exact_fio_rows
        if row.get('specialty_key') and row.get('specialty_key') != specialty_key
    ]
    if len(exact_fio_rows) == 1 and len(specialty_conflicts) == 1:
        review = dict(specialty_conflicts[0])
        review.update({
            'kind': 'specialty_conflict',
            'similarity': 100,
        })
        return review

    scored = []
    for candidate in context.get('by_specialty', {}).get(specialty_key, []):
        candidate_key = candidate.get('fio_similarity_key') or ''
        if not candidate_key or candidate_key == similarity_key:
            continue
        score = difflib.SequenceMatcher(None, similarity_key, candidate_key).ratio()
        if score >= ENROLLMENT_FIO_SUGGESTION_THRESHOLD:
            scored.append((score, candidate))
    scored.sort(key=lambda item: (-item[0], normalize_fio_key(item[1].get('fio')), item[1].get('id') or 0))
    if not scored:
        return None
    best_score, best_candidate = scored[0]
    if len(scored) > 1 and best_score - scored[1][0] < 0.03:
        return None
    review = dict(best_candidate)
    review.update({
        'kind': 'fio_typo',
        'similarity': int(round(best_score * 100)),
    })
    return review

def enrollment_order_storage_dir():
    path = os.path.join(app.config['UPLOAD_FOLDER'], 'enrollment_orders')
    os.makedirs(path, exist_ok=True)
    return path

def get_stored_enrollment_order_path(stored_filename):
    stored_filename = os.path.basename(str(stored_filename or ''))
    if not stored_filename:
        return ''
    storage_root = os.path.abspath(enrollment_order_storage_dir())
    stored_path = os.path.abspath(os.path.join(storage_root, stored_filename))
    if os.path.commonpath([storage_root, stored_path]) != storage_root:
        return ''
    return stored_path

def summarize_order_upload_values(df, column_name):
    if column_name not in df:
        return ''
    values = []
    seen = set()
    for value in df[column_name].fillna('').astype(str):
        value = value.strip()
        if not value or value in seen:
            continue
        seen.add(value)
        values.append(value)
    return ', '.join(values)

def rebuild_enrollment_orders_from_upload_rows(conn, campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    conn.execute('DELETE FROM enrollment_orders WHERE campaign_year=?', (campaign_year,))
    rows = conn.execute(
        '''
        SELECT upload_id, fio, specialty, group_name, order_number, order_date,
               fio_key, specialty_key, group_key
        FROM enrollment_order_upload_rows
        WHERE campaign_year=? AND import_action='import'
        ORDER BY upload_id, row_number, id
        ''',
        (campaign_year,)
    ).fetchall()
    active_rows = {}
    for row in rows:
        upload_id, fio, specialty, group_name, order_number, order_date, fio_key, specialty_key, group_key = row
        if fio_key and specialty_key:
            active_rows[(fio_key, specialty_key)] = row

    for row in active_rows.values():
        upload_id, fio, specialty, group_name, order_number, order_date, fio_key, specialty_key, group_key = row
        conn.execute(
            '''
            INSERT INTO enrollment_orders
                (upload_id, campaign_year, fio, specialty, group_name, order_number, order_date,
                 fio_key, specialty_key, group_key)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''',
            (
                upload_id, campaign_year, fio, specialty, group_name, order_number, order_date,
                fio_key, specialty_key, group_key
            )
        )

def get_enrollment_order_uploads(campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    with sqlite3.connect(DB_PATH) as conn:
        rows = conn.execute(
            '''
            SELECT
                u.id, u.original_filename, u.stored_filename, u.uploaded_by,
                u.total_count, u.import_count, u.matched_count, u.unmatched_count,
                u.duplicate_count, u.skipped_count, u.order_numbers, u.order_dates,
                u.created_at,
                COALESCE(active.active_count, 0) AS active_count,
                COALESCE(all_rows.row_count, 0) AS row_count
            FROM enrollment_order_uploads u
            LEFT JOIN (
                SELECT upload_id, COUNT(*) AS active_count
                FROM enrollment_orders
                WHERE campaign_year=?
                GROUP BY upload_id
            ) active ON active.upload_id=u.id
            LEFT JOIN (
                SELECT upload_id, COUNT(*) AS row_count
                FROM enrollment_order_upload_rows
                GROUP BY upload_id
            ) all_rows ON all_rows.upload_id=u.id
            WHERE u.campaign_year=?
            ORDER BY u.created_at DESC, u.id DESC
            ''',
            (campaign_year, campaign_year)
        ).fetchall()
    columns = [
        'id', 'original_filename', 'stored_filename', 'uploaded_by',
        'total_count', 'import_count', 'matched_count', 'unmatched_count',
        'duplicate_count', 'skipped_count', 'order_numbers', 'order_dates',
        'created_at', 'active_count', 'row_count'
    ]
    return [dict(zip(columns, row)) for row in rows]

def get_enrollment_order_upload(upload_id):
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            '''
            SELECT id, campaign_year, original_filename, stored_filename, uploaded_by,
                   total_count, import_count, matched_count, unmatched_count,
                   duplicate_count, skipped_count, order_numbers, order_dates, created_at
            FROM enrollment_order_uploads
            WHERE id=?
            ''',
            (upload_id,)
        )
        row = cur.fetchone()
    if not row:
        return None
    columns = [
        'id', 'campaign_year', 'original_filename', 'stored_filename', 'uploaded_by',
        'total_count', 'import_count', 'matched_count', 'unmatched_count',
        'duplicate_count', 'skipped_count', 'order_numbers', 'order_dates', 'created_at'
    ]
    return dict(zip(columns, row))

def get_enrollment_order_upload_rows(upload_id):
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            '''
            SELECT row_number, fio, specialty, group_name, order_number, order_date,
                   has_candidate, import_action, import_status
            FROM enrollment_order_upload_rows
            WHERE upload_id=?
            ORDER BY row_number, id
            ''',
            (upload_id,)
        )
        rows = cur.fetchall()
    columns = [
        'row_number', 'fio', 'specialty', 'group_name', 'order_number', 'order_date',
        'has_candidate', 'import_action', 'import_status'
    ]
    result = []
    for row in rows:
        item = dict(zip(columns, row))
        action = item.get('import_action')
        if action == 'import' and item.get('has_candidate'):
            item['action_label'] = 'Совпадение'
            item['badge_class'] = 'status-success'
        elif action == 'import':
            item['action_label'] = 'Нет кандидата'
            item['badge_class'] = 'status-warning'
        elif action == 'duplicate':
            item['action_label'] = 'Повтор'
            item['badge_class'] = 'status-warning'
        else:
            item['action_label'] = 'Пропуск'
            item['badge_class'] = 'status-danger'
        result.append(item)
    return result

def assign_virtual_groups_to_enrollment_order_roster(conn, roster_rows, campaign_year):
    group_year = normalize_group_year(campaign_year, campaign_year)
    rules = get_login_generation_rules()
    course_groups_enabled = are_course_groups_enabled(rules)
    state = get_group_assignment_state(conn, group_year)
    planned_counts = {}

    for row in roster_rows:
        row['group_name'] = 'Не распределено'
        row['cohort2'] = ''
        row['group_source'] = 'Не определена'
        row['group_is_virtual'] = False
        row['group_assignment_status'] = ''
        if row['record_type'] == 'Студент' and row.get('actual_group_name'):
            row['group_name'] = row['actual_group_name']
            row['cohort2'] = row.get('actual_cohort2') or (
                derive_cohort2(row['actual_group_name']) if course_groups_enabled else ''
            ) or ''
            row['group_source'] = 'Фактическая группа'
            row['group_assignment_status'] = 'Уже зачислен'

    pending_rows = sorted(
        (row for row in roster_rows if row['record_type'] != 'Студент'),
        key=lambda row: (normalize_fio_key(row.get('fio')), str(row.get('login') or '')),
    )
    for row in pending_rows:
        target_group = ''
        will_be_virtual = False
        group_error = ''
        order_group = normalize_group_name(row.get('order_group_name'))
        specialty_key = normalize_specialty_key(row.get('specialty'))

        if order_group:
            row['group_source'] = 'Приказ'
            target_group, will_be_virtual, group_error = reserve_exact_group_assignment(
                order_group,
                group_year,
                state,
                planned_counts,
                reserve=True,
            )
        else:
            row['group_source'] = 'Логин'
            parsed_login = parse_login_group_target(row.get('login'), campaign_year, rules)
            if not parsed_login.get('ok'):
                group_error = parsed_login.get('error') or 'не удалось разобрать логин'
            elif specialty_key and parsed_login.get('specialty_key') != specialty_key:
                group_error = 'специальность логина не совпадает с приказом'
            else:
                target_group, will_be_virtual, group_error = reserve_group_by_root_assignment(
                    parsed_login.get('group_root'),
                    group_year,
                    state,
                    planned_counts,
                    reserve=True,
                )

        if target_group and specialty_key and not group_matches_specialty(target_group, specialty_key):
            group_error = 'предложенная группа не совпадает со специальностью'
            target_group = ''
        if target_group:
            row['group_name'] = target_group
            row['cohort2'] = (derive_cohort2(target_group) or '') if course_groups_enabled else ''
            row['group_is_virtual'] = will_be_virtual
            row['group_assignment_status'] = (
                'Справочная группа — в базе не создаётся'
                if will_be_virtual
                else 'Группа уже есть в справочнике'
            )
        else:
            row['group_assignment_status'] = group_error or 'группа не определена'

    for row in roster_rows:
        if row['group_name'] == 'Не распределено':
            row['needs_attention'] = True

    return roster_rows

def build_enrollment_order_student_roster(upload_id):
    upload = get_enrollment_order_upload(upload_id)
    if not upload:
        return None
    campaign_year = upload['campaign_year']
    rules = get_login_generation_rules()

    with sqlite3.connect(DB_PATH) as conn:
        order_rows = conn.execute(
            '''
            SELECT id, row_number, fio, specialty, group_name, order_number, order_date,
                   fio_key, specialty_key, import_action, import_status,
                   fio_review_status, fio_review_candidate_id, fio_reviewed_at, fio_reviewed_by
            FROM enrollment_order_upload_rows
            WHERE upload_id=? AND TRIM(COALESCE(fio, ''))<>''
            ORDER BY row_number, id
            ''',
            (upload_id,)
        ).fetchall()
        abiturient_rows = conn.execute(
            '''
            SELECT id, fio, dogovor, login, email, paid, fam, imotch
            FROM abiturients
            WHERE campaign_year=?
            ORDER BY id
            ''',
            (campaign_year,)
        ).fetchall()
        student_rows = conn.execute(
            '''
            SELECT username, email, firstname, lastname, cohort1, cohort2,
                   source_dogovor, source_fio, source_campaign_year
            FROM students
            WHERE source_campaign_year=? OR source_campaign_year IS NULL OR source_campaign_year=''
            ORDER BY username
            ''',
            (campaign_year,)
        ).fetchall()
        candidate_rows = conn.execute(
            '''
            SELECT id, fio, dogovor, login, email, specialty_key,
                   verification_status, order_group_name, fam, imotch
            FROM enrollment_candidates
            WHERE campaign_year=?
            ORDER BY id
            ''',
            (campaign_year,)
        ).fetchall()
        movement_rows = conn.execute(
            '''
            SELECT movement.username, orders.fio_key, orders.specialty_key
            FROM student_group_transfers movement
            LEFT JOIN enrollment_orders orders ON orders.id=movement.enrollment_order_id
            WHERE movement.movement_type='enrollment'
              AND movement.enrollment_order_upload_id=?
            ORDER BY movement.id
            ''',
            (upload_id,)
        ).fetchall()

    abiturients_by_key = {}
    abiturients_by_id = {}
    for row in abiturient_rows:
        abiturients_by_id[row[0]] = row
        key = make_abiturient_enrollment_match_key(row[1], row[2], rules)
        if key:
            abiturients_by_key.setdefault(key, row)
    abiturient_review_context = build_abiturient_fio_review_context(abiturient_rows, rules)

    students_by_username = {row[0]: row for row in student_rows}
    students_by_key = {}
    for row in student_rows:
        username, _email, firstname, lastname, _cohort1, _cohort2, source_dogovor, source_fio, _source_year = row
        fio = str(source_fio or '').strip()
        if not fio:
            fio = ' '.join(part for part in (lastname, firstname) if str(part or '').strip()).strip()
        key = make_abiturient_enrollment_match_key(fio, source_dogovor, rules)
        if key:
            students_by_key.setdefault(key, row)
    for username, fio_key, specialty_key in movement_rows:
        student = students_by_username.get(username)
        if student and fio_key and specialty_key:
            students_by_key.setdefault((fio_key, specialty_key), student)

    candidates_by_key = {}
    for row in candidate_rows:
        key = make_enrollment_match_key(row[1], row[5])
        if key:
            candidates_by_key.setdefault(key, row)

    roster_rows = []
    for order_row in order_rows:
        (
            row_id, row_number, fio, specialty, group_name, order_number, order_date,
            fio_key, specialty_key, import_action, import_status,
            fio_review_status, fio_review_candidate_id, fio_reviewed_at, fio_reviewed_by
        ) = order_row
        key = (fio_key, specialty_key) if fio_key and specialty_key else None
        student = students_by_key.get(key) if key else None
        abiturient = abiturients_by_key.get(key) if key else None
        candidate = candidates_by_key.get(key) if key else None
        fio_review = None
        if (
            not student
            and not abiturient
            and not candidate
            and fio_review_status == 'linked'
            and fio_review_candidate_id
        ):
            linked_abiturient = abiturients_by_id.get(int(fio_review_candidate_id))
            if linked_abiturient:
                linked_specialty_key = get_dogovor_specialty_key(linked_abiturient[2], rules)
                if linked_specialty_key == normalize_specialty_key(specialty_key or specialty):
                    abiturient = linked_abiturient
                    linked_review = abiturient_review_context['by_id'].get(linked_abiturient[0])
                    if linked_review:
                        fio_review = dict(linked_review)
                        fio_review.update({
                            'kind': 'fio_typo',
                            'similarity': int(round(difflib.SequenceMatcher(
                                None,
                                normalize_fio_similarity_key(fio),
                                linked_review.get('fio_similarity_key') or '',
                            ).ratio() * 100)),
                        })
        if not student and not abiturient and not candidate:
            fio_review = find_abiturient_fio_review(
                fio,
                specialty_key or specialty,
                abiturient_review_context,
            )

        login = ''
        email = ''
        dogovor = ''
        payment_status = 'Нет данных'
        record_type = 'Не найден в базе'
        source_id = ''
        person_url = ''
        missing_email = False
        unpaid = False
        status = 'Абитуриент не найден'
        badge_class = 'status-danger'
        lastname, firstname = split_fio_for_storage(fio)
        actual_group_name = ''
        actual_cohort2 = ''

        if student:
            (
                username, email, student_firstname, student_lastname, cohort1, cohort2,
                source_dogovor, _source_fio, _source_year
            ) = student
            login = username or ''
            dogovor = source_dogovor or ''
            firstname = student_firstname or firstname
            lastname = student_lastname or lastname
            actual_group_name = str(cohort1 or '').strip()
            actual_cohort2 = str(cohort2 or '').strip()
            payment_status = 'Подтверждено при переносе'
            record_type = 'Студент'
            source_id = username or ''
            missing_email = not str(email or '').strip()
            status = 'Перенесён в студенты' if not missing_email else 'Перенесён, но нет почты'
            badge_class = 'status-success' if not missing_email else 'status-warning'
            if has_request_context() and username:
                person_url = url_for('person_card', kind='student', record_id=username)
        elif abiturient:
            abiturient_id, _ab_fio, dogovor, login, email, paid, ab_fam, ab_imotch = abiturient
            lastname = ab_fam or lastname
            firstname = ab_imotch or firstname
            source_id = abiturient_id
            missing_email = not str(email or '').strip()
            unpaid = not is_paid_person_value(paid)
            payment_status = 'Оплачено' if not unpaid else 'Не оплачено'
            record_type = 'Абитуриент'
            if is_withdrawn_login(login):
                status = 'Документы отозваны'
                badge_class = 'status-info'
            elif missing_email and unpaid:
                status = 'Нет почты и оплаты'
                badge_class = 'status-warning'
            elif missing_email:
                status = 'Нет почты'
                badge_class = 'status-warning'
            elif unpaid:
                status = 'Нет оплаты'
                badge_class = 'status-warning'
            else:
                status = 'Готов к переносу'
                badge_class = 'status-success'
            if has_request_context():
                person_url = url_for('person_card', kind='abiturient', record_id=abiturient_id)
        elif candidate:
            (
                candidate_id, _candidate_fio, dogovor, login, email, _candidate_specialty,
                verification_status, _order_group, candidate_fam, candidate_imotch
            ) = candidate
            lastname = candidate_fam or lastname
            firstname = candidate_imotch or firstname
            source_id = candidate_id
            missing_email = not str(email or '').strip()
            payment_status = 'Подтверждено при подготовке'
            record_type = 'Кандидат'
            status = 'Кандидат к зачислению' if not missing_email else 'Кандидат без почты'
            if verification_status == 'verified' and not missing_email:
                badge_class = 'status-success'
            else:
                badge_class = 'status-warning'
            if has_request_context():
                person_url = url_for('abiturients_to_students')

        is_blocked = bool(missing_email or unpaid)
        needs_attention = bool(is_blocked or (not student and not abiturient and not candidate))
        roster_rows.append({
            'id': row_id,
            'row_number': row_number,
            'group_name': 'Не распределено',
            'order_group_name': str(group_name or '').strip(),
            'actual_group_name': actual_group_name,
            'actual_cohort2': actual_cohort2,
            'fio': str(fio or '').strip(),
            'firstname': str(firstname or '').strip(),
            'lastname': str(lastname or '').strip(),
            'specialty': str(specialty or '').strip(),
            'dogovor': str(dogovor or '').strip(),
            'login': str(login or '').strip(),
            'email': str(email or '').strip(),
            'email_status': 'Есть' if str(email or '').strip() else 'Нет',
            'payment_status': payment_status,
            'record_type': record_type,
            'source_id': source_id,
            'person_url': person_url,
            'status': status,
            'badge_class': badge_class,
            'missing_email': missing_email,
            'unpaid': unpaid,
            'is_blocked': is_blocked,
            'needs_attention': needs_attention,
            'order_number': str(order_number or '').strip(),
            'order_date': str(order_date or '').strip(),
            'order_row_action': str(import_action or '').strip(),
            'order_row_status': str(import_status or '').strip(),
            'fio_review_status': str(fio_review_status or '').strip(),
            'fio_reviewed_at': str(fio_reviewed_at or '').strip(),
            'fio_reviewed_by': str(fio_reviewed_by or '').strip(),
            'has_fio_review': bool(fio_review),
            'fio_review_kind': fio_review.get('kind', '') if fio_review else '',
            'suggested_abiturient_id': fio_review.get('id', '') if fio_review else '',
            'suggested_abiturient_fio': fio_review.get('fio', '') if fio_review else '',
            'suggested_abiturient_login': fio_review.get('login', '') if fio_review else '',
            'suggested_abiturient_dogovor': fio_review.get('dogovor', '') if fio_review else '',
            'suggested_abiturient_specialty': fio_review.get('specialty', '') if fio_review else '',
            'fio_similarity': fio_review.get('similarity', '') if fio_review else '',
            'reviewed_candidate_matches': bool(
                fio_review
                and fio_review_candidate_id
                and int(fio_review_candidate_id) == int(fio_review.get('id') or 0)
            ),
        })

    with sqlite3.connect(DB_PATH) as conn:
        assign_virtual_groups_to_enrollment_order_roster(
            conn,
            roster_rows,
            campaign_year,
        )
    roster_rows.sort(key=lambda row: (
        natural_text_sort_key(row['group_name']),
        normalize_fio_key(row['fio']),
        row['row_number'] or 0,
    ))
    grouped_rows = {}
    for row in roster_rows:
        grouped_rows.setdefault(row['group_name'], []).append(row)
    groups = [
        {
            'name': group_name,
            'rows': rows,
            'count': len(rows),
            'attention_count': sum(1 for row in rows if row['needs_attention']),
            'is_virtual': any(row['group_is_virtual'] for row in rows),
        }
        for group_name, rows in grouped_rows.items()
    ]

    summary = {
        'total_count': len(roster_rows),
        'group_count': sum(1 for group in groups if group['name'] != 'Не распределено'),
        'virtual_group_count': len({
            row['group_name'] for row in roster_rows if row['group_is_virtual']
        }),
        'unassigned_count': sum(1 for row in roster_rows if row['group_name'] == 'Не распределено'),
        'student_count': sum(1 for row in roster_rows if row['record_type'] == 'Студент'),
        'ready_count': sum(
            1 for row in roster_rows
            if row['status'] in {'Готов к переносу', 'Кандидат к зачислению'}
        ),
        'missing_email_count': sum(1 for row in roster_rows if row['missing_email']),
        'unpaid_count': sum(1 for row in roster_rows if row['unpaid']),
        'blocked_count': sum(1 for row in roster_rows if row['is_blocked']),
        'not_found_count': sum(1 for row in roster_rows if row['record_type'] == 'Не найден в базе'),
        'attention_count': sum(1 for row in roster_rows if row['needs_attention']),
        'fio_review_count': sum(
            1 for row in roster_rows
            if row['fio_review_kind'] == 'fio_typo'
            and row['fio_review_status'] not in {'skipped', 'linked'}
        ),
        'fio_review_skipped_count': sum(
            1 for row in roster_rows if row['fio_review_status'] == 'skipped'
        ),
        'fio_review_linked_count': sum(
            1 for row in roster_rows if row['fio_review_status'] == 'linked'
        ),
        'specialty_conflict_count': sum(
            1 for row in roster_rows if row['fio_review_kind'] == 'specialty_conflict'
        ),
    }
    return {
        'upload': upload,
        'rows': roster_rows,
        'groups': groups,
        'summary': summary,
    }

def enrollment_order_roster_export_rows(roster):
    return [
        {
            'username': row['login'],
            'password': 'cron',
            'email': row['email'],
            'firstname': row['firstname'],
            'lastname': row['lastname'],
            'cohort1': row['group_name'] if row['group_name'] != 'Не распределено' else '',
            'cohort2': row['cohort2'],
            'ФИО': row['fio'],
            'Специальность': row['specialty'],
            'Договор': row['dogovor'],
            'Наличие почты': row['email_status'],
            'Оплата': row['payment_status'],
            'Состояние': row['status'],
            'Тип записи': row['record_type'],
            'Источник группы': row['group_source'],
            'Статус распределения': row['group_assignment_status'],
            'Сверка ФИО': (
                'Исправлено по приказу'
                if row['fio_review_status'] == 'fixed'
                else 'Опечатка в приказе — логин присвоен'
                if row['fio_review_status'] == 'linked'
                else 'Пропущено без исправления'
                if row['fio_review_status'] == 'skipped'
                else 'Возможная опечатка'
                if row['fio_review_kind'] == 'fio_typo'
                else 'ФИО совпадает, специальность отличается'
                if row['fio_review_kind'] == 'specialty_conflict'
                else ''
            ),
            'ФИО в базе для сверки': row['suggested_abiturient_fio'],
            'Сходство ФИО, %': row['fio_similarity'],
            'Номер приказа': row['order_number'],
            'Дата приказа': row['order_date'],
            'Строка в приказе': row['row_number'],
            'Статус строки приказа': row['order_row_status'],
        }
        for row in roster['rows']
    ]

def fix_abiturient_fio_from_enrollment_order_roster(
    upload_id,
    row_id,
    abiturient_id,
    updated_by='',
):
    backup_path = ''
    with sqlite3.connect(DB_PATH) as conn:
        order_row = conn.execute(
            '''
            SELECT rows.id, rows.campaign_year, rows.fio, rows.specialty_key
            FROM enrollment_order_upload_rows rows
            WHERE rows.id=? AND rows.upload_id=?
            ''',
            (row_id, upload_id)
        ).fetchone()
        if not order_row:
            raise ValueError('Строка приказа больше не найдена.')
        _row_id, campaign_year, order_fio, order_specialty_key = order_row
        order_fio = ' '.join(str(order_fio or '').split())
        order_specialty_key = normalize_specialty_key(order_specialty_key)
        abiturient = conn.execute(
            '''
            SELECT id, fio, dogovor
            FROM abiturients
            WHERE id=? AND campaign_year=?
            ''',
            (abiturient_id, campaign_year)
        ).fetchone()
        if not abiturient:
            raise ValueError('Предложенная запись абитуриента больше не найдена.')
        _abiturient_id, old_fio, dogovor = abiturient
        abiturient_specialty_key = get_dogovor_specialty_key(dogovor)
        if not order_specialty_key or abiturient_specialty_key != order_specialty_key:
            raise ValueError('Специальность в приказе не совпадает со специальностью договора.')
        similarity = difflib.SequenceMatcher(
            None,
            normalize_fio_similarity_key(order_fio),
            normalize_fio_similarity_key(old_fio),
        ).ratio()
        if similarity < ENROLLMENT_FIO_SUGGESTION_THRESHOLD:
            raise ValueError('ФИО слишком сильно различаются для безопасного исправления.')

        new_key = make_enrollment_match_key(order_fio, order_specialty_key)
        other_abiturients = conn.execute(
            '''
            SELECT id, fio, dogovor
            FROM abiturients
            WHERE campaign_year=? AND id<>?
            ''',
            (campaign_year, abiturient_id)
        ).fetchall()
        for other_id, other_fio, other_dogovor in other_abiturients:
            if make_abiturient_enrollment_match_key(other_fio, other_dogovor) == new_key:
                raise ValueError(
                    f'Нельзя исправить ФИО: такая запись уже есть в базе ({other_fio}).'
                )

        linked_candidates = conn.execute(
            '''
            SELECT id
            FROM enrollment_candidates
            WHERE abiturient_id=? AND campaign_year=?
            ''',
            (abiturient_id, campaign_year)
        ).fetchall()
        linked_candidate_ids = {row[0] for row in linked_candidates}
        other_candidates = conn.execute(
            '''
            SELECT id, fio, specialty_key
            FROM enrollment_candidates
            WHERE campaign_year=?
            ''',
            (campaign_year,)
        ).fetchall()
        for candidate_id, candidate_fio, candidate_specialty_key in other_candidates:
            if candidate_id in linked_candidate_ids:
                continue
            if make_enrollment_match_key(candidate_fio, candidate_specialty_key) == new_key:
                raise ValueError(
                    f'Нельзя исправить ФИО: такой кандидат уже существует ({candidate_fio}).'
                )

        backup_path = create_database_backup('before_enrollment_order_roster_fio_fix')
        fam, imotch = split_fio_for_storage(order_fio)
        conn.execute(
            '''
            UPDATE abiturients
            SET fio=?, fam=?, imotch=?
            WHERE id=? AND campaign_year=?
            ''',
            (order_fio, fam, imotch, abiturient_id, campaign_year)
        )
        conn.execute(
            '''
            UPDATE enrollment_candidates
            SET fio=?, fam=?, imotch=?
            WHERE abiturient_id=? AND campaign_year=?
            ''',
            (order_fio, fam, imotch, abiturient_id, campaign_year)
        )
        conn.execute(
            '''
            UPDATE enrollment_order_upload_rows
            SET fio_review_status='fixed', fio_review_candidate_id=?,
                fio_reviewed_at=datetime('now', 'localtime'), fio_reviewed_by=?
            WHERE id=? AND upload_id=?
            ''',
            (abiturient_id, updated_by or '', row_id, upload_id)
        )
        refresh_enrollment_candidate_statuses(conn, campaign_year)
        log_action(
            'abiturient_fio_fixed_from_enrollment_order_roster',
            'abiturient',
            abiturient_id,
            (
                f'upload_id={upload_id}; row_id={row_id}; old_fio={old_fio}; '
                f'new_fio={order_fio}; similarity={int(round(similarity * 100))}; '
                f'backup={os.path.basename(backup_path) if backup_path else ""}'
            ),
            conn,
        )
    return {
        'old_fio': old_fio,
        'new_fio': order_fio,
        'campaign_year': campaign_year,
    }

def set_enrollment_order_roster_fio_review_status(
    upload_id,
    row_id,
    status,
    abiturient_id=None,
    updated_by='',
):
    if status not in {'skipped', 'linked', ''}:
        raise ValueError('Неизвестное действие сверки ФИО.')
    with sqlite3.connect(DB_PATH) as conn:
        row = conn.execute(
            '''
            SELECT id, campaign_year, fio, specialty_key
            FROM enrollment_order_upload_rows
            WHERE id=? AND upload_id=?
            ''',
            (row_id, upload_id)
        ).fetchone()
        if not row:
            raise ValueError('Строка приказа больше не найдена.')
        if status in {'skipped', 'linked'}:
            candidate = conn.execute(
                '''
                SELECT id, fio, dogovor, login
                FROM abiturients
                WHERE id=? AND campaign_year=?
                ''',
                (abiturient_id, row[1])
            ).fetchone()
            if not candidate:
                raise ValueError('Предложенная запись абитуриента больше не найдена.')
            if status == 'linked':
                _candidate_id, candidate_fio, candidate_dogovor, candidate_login = candidate
                order_specialty_key = normalize_specialty_key(row[3])
                candidate_specialty_key = get_dogovor_specialty_key(candidate_dogovor)
                if not order_specialty_key or candidate_specialty_key != order_specialty_key:
                    raise ValueError('Специальность в приказе не совпадает со специальностью договора.')
                similarity = difflib.SequenceMatcher(
                    None,
                    normalize_fio_similarity_key(row[2]),
                    normalize_fio_similarity_key(candidate_fio),
                ).ratio()
                if similarity < ENROLLMENT_FIO_SUGGESTION_THRESHOLD:
                    raise ValueError('ФИО слишком сильно различаются для безопасного присвоения логина.')
                if not str(candidate_login or '').strip():
                    raise ValueError('У найденного абитуриента не указан логин.')
        conn.execute(
            '''
            UPDATE enrollment_order_upload_rows
            SET fio_review_status=?, fio_review_candidate_id=?,
                fio_reviewed_at=CASE WHEN ?='' THEN NULL ELSE datetime('now', 'localtime') END,
                fio_reviewed_by=CASE WHEN ?='' THEN NULL ELSE ? END
            WHERE id=? AND upload_id=?
            ''',
            (
                status,
                abiturient_id if status else None,
                status,
                status,
                updated_by or '',
                row_id,
                upload_id,
            )
        )
        log_action(
            (
                'enrollment_order_roster_fio_review_linked'
                if status == 'linked'
                else 'enrollment_order_roster_fio_review_skipped'
                if status == 'skipped'
                else 'enrollment_order_roster_fio_review_reset'
            ),
            'enrollment_order_upload_row',
            row_id,
            f'upload_id={upload_id}; abiturient_id={abiturient_id or ""}',
            conn,
        )

def delete_enrollment_order_upload(upload_id, campaign_year=None):
    backup_path = create_database_backup('before_enrollment_order_upload_delete')
    stored_path = ''
    deleted = None
    with sqlite3.connect(DB_PATH) as conn:
        row = conn.execute(
            '''
            SELECT id, campaign_year, original_filename, stored_filename
            FROM enrollment_order_uploads
            WHERE id=?
            ''',
            (upload_id,)
        ).fetchone()
        if not row:
            raise ValueError('Загрузка приказа не найдена.')
        upload_id, row_campaign_year, original_filename, stored_filename = row
        row_campaign_year = normalize_campaign_year(row_campaign_year, get_active_campaign_year())
        if campaign_year and row_campaign_year != normalize_campaign_year(campaign_year, get_active_campaign_year()):
            raise ValueError('Загрузка приказа относится к другой приемной кампании.')
        stored_path = get_stored_enrollment_order_path(stored_filename)
        conn.execute('DELETE FROM enrollment_order_upload_rows WHERE upload_id=?', (upload_id,))
        conn.execute('DELETE FROM enrollment_order_uploads WHERE id=?', (upload_id,))
        rebuild_enrollment_orders_from_upload_rows(conn, row_campaign_year)
        refresh_enrollment_candidate_statuses(conn, row_campaign_year)
        log_action(
            'enrollment_order_upload_deleted',
            'campaign',
            row_campaign_year,
            f"upload_id={upload_id}; filename={original_filename}; backup={os.path.basename(backup_path) if backup_path else ''}",
            conn
        )
        deleted = {
            'id': upload_id,
            'campaign_year': row_campaign_year,
            'original_filename': original_filename,
        }
    cleanup_temp_files(stored_path)
    return deleted

def cleanup_not_ready_enrollment_candidates(conn, campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    rows = conn.execute(
        '''
        SELECT c.id, c.login, a.id, a.email, a.paid, a.login
        FROM enrollment_candidates c
        LEFT JOIN abiturients a
            ON a.id=c.abiturient_id AND a.campaign_year=c.campaign_year
        WHERE c.campaign_year=?
        ''',
        (campaign_year,)
    ).fetchall()
    remove_ids = []
    for candidate_id, login, source_id, email, paid, source_login in rows:
        if is_withdrawn_login(login) or is_withdrawn_login(source_login):
            remove_ids.append(candidate_id)
            continue
        if not source_id or not str(email or '').strip() or not is_paid_person_value(paid):
            remove_ids.append(candidate_id)
            continue
        if login and conn.execute('SELECT 1 FROM students WHERE username=?', (login,)).fetchone():
            remove_ids.append(candidate_id)
    if remove_ids:
        placeholders = ','.join('?' for _ in remove_ids)
        conn.execute(f'DELETE FROM enrollment_candidates WHERE id IN ({placeholders})', remove_ids)
    return len(remove_ids)

def refresh_enrollment_candidate_statuses(conn, campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    cleanup_not_ready_enrollment_candidates(conn, campaign_year)
    order_rows = conn.execute(
        '''
        SELECT id, fio_key, specialty_key, group_name
        FROM enrollment_orders
        WHERE campaign_year=?
        ORDER BY id
        ''',
        (campaign_year,)
    ).fetchall()
    order_map = {}
    for order_id, fio_key, specialty_key, group_name in order_rows:
        if fio_key and specialty_key and (fio_key, specialty_key) not in order_map:
            order_map[(fio_key, specialty_key)] = (order_id, group_name or '')

    candidate_rows = conn.execute(
        '''
        SELECT id, fio, specialty_key
        FROM enrollment_candidates
        WHERE campaign_year=?
        ''',
        (campaign_year,)
    ).fetchall()
    for candidate_id, fio, specialty_key in candidate_rows:
        match_key = make_enrollment_match_key(fio, specialty_key)
        match = order_map.get(match_key)
        if match:
            conn.execute(
                '''
                UPDATE enrollment_candidates
                SET verification_status='verified',
                    matched_order_id=?,
                    order_group_name=?,
                    verified_at=datetime('now', 'localtime')
                WHERE id=?
                ''',
                (match[0], match[1], candidate_id)
            )
        elif order_rows:
            conn.execute(
                '''
                UPDATE enrollment_candidates
                SET verification_status='missing_in_order',
                    matched_order_id=NULL,
                    order_group_name='',
                    verified_at=NULL
                WHERE id=?
                ''',
                (candidate_id,)
            )
        else:
            conn.execute(
                '''
                UPDATE enrollment_candidates
                SET verification_status='waiting_order',
                    matched_order_id=NULL,
                    order_group_name='',
                    verified_at=NULL
                WHERE id=?
                ''',
                (candidate_id,)
            )

def enrollment_candidate_status_view(status):
    if status == 'verified':
        return 'Сверен с приказом', 'status-success'
    if status == 'missing_in_order':
        return 'Нет в приказе', 'status-danger'
    return 'Ждет приказ', 'status-warning'

def get_enrollment_candidate_specialties(campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    with sqlite3.connect(DB_PATH) as conn:
        refresh_enrollment_candidate_statuses(conn, campaign_year)
        rows = conn.execute(
            '''
            SELECT specialty_key, specialty, COUNT(*)
            FROM enrollment_candidates
            WHERE campaign_year=?
            GROUP BY specialty_key, specialty
            ORDER BY specialty
            ''',
            (campaign_year,)
        ).fetchall()
    return [
        {
            'key': row[0] or '',
            'name': row[1] or get_specialty_display_name(row[0]),
            'count': row[2],
        }
        for row in rows
    ]

def get_enrollment_candidates(campaign_year, specialty_key=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    enrollment_order_required = is_enrollment_order_required()
    query = '''
        SELECT
            c.id, c.abiturient_id, c.fio, c.dogovor, c.login, c.fam, c.imotch, c.email,
            c.specialty, c.specialty_key, c.base_label, c.verification_status,
            c.matched_order_id, c.order_group_name, c.verified_at, c.created_at,
            a.id AS source_id, a.email AS source_email, a.paid AS source_paid
        FROM enrollment_candidates c
        LEFT JOIN abiturients a
            ON a.id=c.abiturient_id AND a.campaign_year=c.campaign_year
        WHERE c.campaign_year=?
    '''
    params = [campaign_year]
    if specialty_key:
        query += ' AND c.specialty_key=?'
        params.append(specialty_key)
    query += ' ORDER BY c.specialty, c.fio'
    with sqlite3.connect(DB_PATH) as conn:
        refresh_enrollment_candidate_statuses(conn, campaign_year)
        cur = conn.execute(query, params)
        columns = [desc[0] for desc in cur.description]
        rows = [dict(zip(columns, row)) for row in cur.fetchall()]

    for row in rows:
        row['email'] = (row.get('source_email') or row.get('email') or '').strip()
        readiness_issues = enrollment_candidate_readiness_issues(
            row.get('source_id'),
            row.get('email'),
            row.get('source_paid')
        )
        row['readiness_note'] = '; '.join(readiness_issues)
        row['status_label'], row['badge_class'] = enrollment_candidate_status_view(row.get('verification_status'))
        if readiness_issues:
            row['can_migrate'] = False
            row['status_label'] = readiness_issues[0]
            row['badge_class'] = 'status-danger'
        elif enrollment_order_required:
            row['can_migrate'] = row.get('verification_status') == 'verified'
        else:
            row['can_migrate'] = True
            if row.get('verification_status') != 'verified':
                row['status_label'] = 'Готов по почте и оплате'
                row['badge_class'] = 'status-info'
        row['group_hint'] = row.get('order_group_name') or 'Выберите группу по специальности'
    return rows

def sync_enrollment_candidates_from_ready_abiturients(campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    login_rules = get_login_generation_rules()
    backup_path = None

    with sqlite3.connect(DB_PATH) as conn:
        abiturient_rows = conn.execute(
            '''
            SELECT id
            FROM abiturients
            WHERE campaign_year=?
            ORDER BY fio
            ''',
            (campaign_year,)
        ).fetchall()
        backup_path = create_database_backup('before_enrollment_candidates_sync')
        summary = sync_enrollment_candidates_for_abiturients(
            conn,
            [row[0] for row in abiturient_rows],
            campaign_year,
            login_rules=login_rules
        )

        refresh_enrollment_candidate_statuses(conn, campaign_year)
        log_action(
            'enrollment_candidates_synced',
            'campaign',
            campaign_year,
            (
                f"created={summary['created']}; updated={summary['updated']}; "
                f"removed={summary['removed']}; no_email={len(summary['skipped_without_email'])}; "
                f"unpaid={len(summary['skipped_unpaid'])}; "
                f"no_specialty={len(summary['skipped_without_specialty'])}; "
                f"existing_students={len(summary['skipped_existing_students'])}; "
                f"withdrawn={len(summary['skipped_withdrawn'])}; "
                f"backup={os.path.basename(backup_path) if backup_path else ''}"
            ),
            conn
        )

    return {
        'created': summary['created'],
        'updated': summary['updated'],
        'removed': summary['removed'],
        'skipped_without_email': summary['skipped_without_email'],
        'skipped_unpaid': summary['skipped_unpaid'],
        'skipped_without_specialty': summary['skipped_without_specialty'],
        'skipped_existing_students': summary['skipped_existing_students'],
        'skipped_withdrawn': summary['skipped_withdrawn'],
    }

def order_column_aliases(field):
    return ENROLLMENT_ORDER_COLUMN_ALIASES.get(field, set())

def clean_order_date_value(value):
    if pd.isna(value):
        return ''
    if isinstance(value, (datetime, date)):
        return value.strftime('%Y-%m-%d')
    return clean_upload_text(value)

def find_specialty_in_text(text):
    compact_text = re.sub(r'[^0-9A-Za-zА-Яа-яЁё.]+', '', str(text or '')).upper().replace('Ё', 'Е')
    candidates = set(_specialty_aliases.keys()) | set(_specialty_aliases.values()) | set(SPECIALTY_DISPLAY_NAMES.values())
    for candidate in sorted(candidates, key=lambda item: len(str(item)), reverse=True):
        compact_candidate = re.sub(r'[^0-9A-Za-zА-Яа-яЁё.]+', '', str(candidate)).upper().replace('Ё', 'Е')
        if compact_candidate and compact_candidate in compact_text:
            return str(candidate)
    return ''

def find_group_in_text(text):
    match = re.search(r'\b\d{2}[A-Za-zА-Яа-яЁё]+-(?:\d{1,2}(?:[A-Za-zА-Яа-яЁё])?|[A-Za-zА-Яа-яЁё]+)-\d+\b', str(text or ''))
    return normalize_group_name(match.group(0)) if match else ''

_order_reference_re = re.compile(
    r'(?:приказ[ауе]?\s+от\s*)?'
    r'(\d{1,2}\.\d{1,2}\.\d{4})\s*(?:г\.?)?\s*'
    r'(?:№|N|No|Nо)\s*([0-9A-Za-zА-Яа-яЁё\-–—]+)',
    re.IGNORECASE
)
_order_reference_month_re = re.compile(
    r'(\d{1,2}\s+[А-Яа-яЁё]+\s+\d{4}\s*г?\.?)\s*'
    r'(?:№|N|No|Nо)\s*([0-9A-Za-zА-Яа-яЁё\-–—]+)',
    re.IGNORECASE
)
_order_specialty_heading_re = re.compile(r'(?:^|\s)СПО\s+\d{2}\.\d{2}\.\d{2}', re.IGNORECASE)
_order_specialty_noise_re = re.compile(
    r'\b(?:очная|заочная|очно-заочная)\s+форма(?:\s+обучения)?\b',
    re.IGNORECASE
)
_order_appendix_re = re.compile(r'\(\s*приложение\s*№?\s*\d+\s*\)\.?', re.IGNORECASE)
_order_appendix_start_line_re = re.compile(r'^приложение\s*№?\s*\d+', re.IGNORECASE)
_order_appendix_suffix_line_re = re.compile(r'^\(?\s*приложение\s*№?\s*\d+\s*\)?\.?$', re.IGNORECASE)
_order_context_reset_markers = (
    'приказываю',
    'директор',
    'на основании',
    'среднее общее образование',
    'основное общее образование',
)

def clean_order_specialty_heading(text):
    value = _order_specialty_noise_re.sub(' ', str(text or ''))
    value = _order_appendix_re.sub(' ', value)
    return ' '.join(value.split()).strip(' .')

def find_order_reference_in_text(text):
    line = ' '.join(str(text or '').split())
    for pattern in (_order_reference_re, _order_reference_month_re):
        match = pattern.search(line)
        if not match:
            continue
        order_date = clean_upload_text(match.group(1)).rstrip('.')
        order_number = clean_upload_text(match.group(2)).replace('–', '-').replace('—', '-')
        return {
            'order_date': order_date,
            'order_number': order_number,
        }
    return None

def find_order_specialty_heading_in_text(text):
    line = clean_order_specialty_heading(text)
    if not _order_specialty_heading_re.search(line):
        return ''
    specialty = find_specialty_in_text(line)
    return get_specialty_display_name(specialty) if specialty else ''

def order_line_resets_specialty_context(text):
    folded = str(text or '').casefold()
    return any(marker in folded for marker in _order_context_reset_markers)

def clean_order_person_line(text, group_name=''):
    source = str(text or '')
    source = re.sub(r'^\s*\d+[\).\-\s]*', '', source)
    if group_name:
        source = re.sub(re.escape(group_name), ' ', source, flags=re.IGNORECASE)
    source = re.sub(
        r'\b\d{2}[A-Za-zА-Яа-яЁё]+-(?:\d{1,2}(?:[A-Za-zА-Яа-яЁё])?|[A-Za-zА-Яа-яЁё]+)-\d+\b',
        ' ',
        source
    )
    return ' '.join(source.split()).strip(' ,;.')

def looks_like_order_person_line(text, group_name=''):
    source = clean_order_person_line(text, group_name)
    if not source or order_line_resets_specialty_context(source):
        return False
    if any(char.isdigit() for char in source) or '№' in source:
        return False
    words = source.split()
    if not 2 <= len(words) <= 4:
        return False
    for word in words:
        parts = word.split('-')
        if not parts or any(not re.fullmatch(r'[А-ЯЁа-яё]+', part) for part in parts):
            return False
        if not parts[0][0].isupper():
            return False
    return True

def find_fio_in_order_text(text, specialty='', group_name=''):
    source = clean_order_person_line(text, group_name)
    for marker in (specialty, group_name):
        marker = str(marker or '').strip()
        if marker:
            source = re.split(re.escape(marker), source, maxsplit=1, flags=re.IGNORECASE)[0]
    source = clean_order_person_line(source)
    words = re.findall(r'[А-ЯЁ][А-ЯЁа-яё]*(?:-[А-ЯЁа-яё]+)?', source)
    if len(words) >= 3:
        return ' '.join(words[:3])
    if len(words) >= 2:
        return ' '.join(words[:2])
    return ''

def dataframe_from_order_text_lines(lines, default_order_date='', default_order_number=''):
    rows = []
    current_specialty = ''
    current_order_date = clean_upload_text(default_order_date)
    current_order_number = clean_upload_text(default_order_number)
    for line in lines:
        line = ' '.join(str(line or '').split())
        if not line:
            continue

        order_reference = find_order_reference_in_text(line)
        if order_reference:
            current_order_date = order_reference['order_date']
            current_order_number = order_reference['order_number']

        if _order_appendix_start_line_re.match(line):
            current_specialty = ''
            continue
        if _order_appendix_suffix_line_re.match(line):
            continue

        specialty_heading = find_order_specialty_heading_in_text(line)
        if specialty_heading:
            current_specialty = specialty_heading
            continue

        if order_line_resets_specialty_context(line):
            current_specialty = ''
            continue

        group_name = find_group_in_text(line)
        specialty = current_specialty
        if not specialty:
            inline_specialty = find_specialty_in_text(line)
            specialty = get_specialty_display_name(inline_specialty) if inline_specialty else ''
        if not specialty or not looks_like_order_person_line(line, group_name):
            continue
        fio = find_fio_in_order_text(line, specialty, group_name)
        if not fio:
            continue
        rows.append({
            'ФИО': fio,
            'Специальность': specialty,
            'Группа': group_name,
            'Номер приказа': current_order_number,
            'Дата приказа': current_order_date,
        })
    if not rows:
        raise ValueError('Не удалось извлечь строки приказа. Для PDF/DOCX лучше использовать таблицу с колонками ФИО и Специальность.')
    return pd.DataFrame(rows)

def dataframe_has_enrollment_order_columns(df):
    available = {clean_upload_text(column).casefold() for column in df.columns}
    return all(
        available.intersection(alias.casefold() for alias in order_column_aliases(field))
        for field in ('fio', 'specialty')
    )

def read_docx_enrollment_order_dataframe(file_path):
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError('Для загрузки DOCX-приказов установите зависимость python-docx.') from exc

    document = Document(file_path)
    frames = []
    default_order_date = ''
    default_order_number = ''
    for table in document.tables:
        table_rows = [
            [clean_upload_text(cell.text) for cell in row.cells]
            for row in table.rows
        ]
        table_rows = [row for row in table_rows if any(row)]
        for row in table_rows:
            order_reference = find_order_reference_in_text(' '.join(row))
            if order_reference and not default_order_number:
                default_order_date = order_reference['order_date']
                default_order_number = order_reference['order_number']
        if len(table_rows) < 2:
            continue
        headers = table_rows[0]
        frame = pd.DataFrame(table_rows[1:], columns=headers)
        if dataframe_has_enrollment_order_columns(frame):
            frames.append(frame)
    if frames:
        return pd.concat(frames, ignore_index=True)

    lines = [paragraph.text for paragraph in document.paragraphs]
    return dataframe_from_order_text_lines(lines, default_order_date, default_order_number)

def read_pdf_enrollment_order_dataframe(file_path):
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError('Для загрузки PDF-приказов установите зависимость pypdf.') from exc

    reader = PdfReader(file_path)
    text = '\n'.join(page.extract_text() or '' for page in reader.pages)
    if not text.strip():
        raise ValueError('PDF не содержит распознаваемого текста. Нужен текстовый PDF, DOCX или таблица Excel/CSV.')
    return dataframe_from_order_text_lines(text.splitlines())

def read_enrollment_order_upload(file_path):
    extension = os.path.splitext(file_path)[1].lower()
    if extension in {'.csv', '.xls', '.xlsx'}:
        return read_tabular_upload(file_path)
    if extension == '.docx':
        return read_docx_enrollment_order_dataframe(file_path)
    if extension == '.pdf':
        return read_pdf_enrollment_order_dataframe(file_path)
    raise ValueError('Неподдерживаемый формат приказа')

def get_order_row_value(row, field):
    value = find_row_value_casefold(row, order_column_aliases(field))
    if field == 'order_date':
        return clean_order_date_value(value)
    return clean_upload_text(value)

def summarize_enrollment_order_import(df):
    action_counts = df['import_action'].value_counts().to_dict() if not df.empty else {}
    return {
        'total': int(len(df)),
        'import_count': int(action_counts.get('import', 0)),
        'matched_count': int(((df['import_action'] == 'import') & df['has_candidate']).sum()) if 'has_candidate' in df else 0,
        'unmatched_count': int(((df['import_action'] == 'import') & (~df['has_candidate'])).sum()) if 'has_candidate' in df else 0,
        'fio_suggestion_count': int(df['suggested_candidate_id'].astype(bool).sum()) if 'suggested_candidate_id' in df else 0,
        'duplicate_count': int(action_counts.get('duplicate', 0)),
        'skipped_count': int(action_counts.get('skip', 0)),
    }

def build_enrollment_order_import_plan(file_path, campaign_year=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    df = read_enrollment_order_upload(file_path)
    df.columns = [clean_upload_text(column) for column in df.columns]
    if df.empty:
        raise ValueError('Файл приказа не содержит строк для загрузки')

    missing_columns = []
    for field in ('fio', 'specialty'):
        available = {column.casefold() for column in df.columns}
        aliases = {alias.casefold() for alias in order_column_aliases(field)}
        if not available.intersection(aliases):
            missing_columns.append(ENROLLMENT_ORDER_FIELD_LABELS[field])
    if missing_columns:
        raise ValueError(f"В приказе не хватает столбцов: {', '.join(missing_columns)}")

    candidate_context = get_enrollment_candidate_match_context(campaign_year)
    candidate_keys = candidate_context['keys']
    seen_keys = set()
    rows = []

    for row_number, (_, source_row) in enumerate(df.iterrows(), start=2):
        fio = get_order_row_value(source_row, 'fio')
        specialty = get_order_row_value(source_row, 'specialty')
        group_name = normalize_group_name(get_order_row_value(source_row, 'group_name'))
        order_number = get_order_row_value(source_row, 'order_number')
        order_date = get_order_row_value(source_row, 'order_date')
        fio_key = normalize_fio_key(fio)
        specialty_key = normalize_specialty_key(specialty)
        exact_candidate = candidate_context.get('by_key', {}).get((fio_key, specialty_key))
        has_candidate = bool(exact_candidate)
        fio_suggestion = None
        action = 'import'
        status = 'Будет загружен и сверен с кандидатом' if has_candidate else 'Будет загружен, но кандидата пока нет'

        if not fio:
            action = 'skip'
            status = 'Не заполнено ФИО'
        elif not specialty_key:
            action = 'skip'
            status = 'Не заполнена специальность'
        elif (fio_key, specialty_key) in seen_keys:
            action = 'duplicate'
            status = 'ФИО и специальность повторяются в приказе'
        else:
            seen_keys.add((fio_key, specialty_key))
            source_fio = exact_candidate.get('source_fio') if exact_candidate else ''
            if has_candidate and source_fio and normalize_fio_key(source_fio) != fio_key:
                action = 'fio_review'
                has_candidate = False
                score = difflib.SequenceMatcher(
                    None,
                    normalize_fio_similarity_key(fio),
                    normalize_fio_similarity_key(source_fio)
                ).ratio()
                fio_suggestion = dict(exact_candidate)
                fio_suggestion['fio'] = source_fio
                fio_suggestion['candidate_fio'] = exact_candidate.get('fio') or ''
                fio_suggestion['similarity'] = int(round(score * 100))
                fio_suggestion['reason'] = 'source_mismatch'
                status = (
                    f"ФИО в приказе совпадает с кандидатом, но отличается от основной базы абитуриентов: "
                    f"{source_fio}. Исправьте базу по приказу или проверьте сам приказ."
                )
            elif not has_candidate:
                fio_suggestion = find_fio_candidate_suggestion(fio, specialty_key, candidate_context)
                if fio_suggestion:
                    fio_suggestion['reason'] = 'similar_fio'
                    fio_suggestion['candidate_fio'] = fio_suggestion.get('fio') or ''
                    status = (
                        f"Возможная ошибка в ФИО: в базе есть похожий кандидат "
                        f"{fio_suggestion['fio']} ({fio_suggestion['similarity']}%). "
                        f"Выберите, исправить ФИО в базе по приказу или исправить приказ и загрузить заново."
                    )

        rows.append({
            '_row_number': row_number,
            'campaign_year': campaign_year,
            'ФИО': fio,
            'Специальность': get_specialty_display_name(specialty),
            'Группа': group_name,
            'Номер приказа': order_number,
            'Дата приказа': order_date,
            'fio_key': fio_key,
            'specialty_key': specialty_key,
            'group_key': normalize_group_name(group_name).casefold() if group_name else '',
            'has_candidate': has_candidate,
            'suggested_candidate_id': fio_suggestion['id'] if fio_suggestion else '',
            'suggested_candidate_fio': fio_suggestion['fio'] if fio_suggestion else '',
            'suggested_candidate_login': fio_suggestion['login'] if fio_suggestion else '',
            'suggested_candidate_dogovor': fio_suggestion['dogovor'] if fio_suggestion else '',
            'suggested_candidate_similarity': fio_suggestion['similarity'] if fio_suggestion else '',
            'fio_review_reason': fio_suggestion.get('reason', '') if fio_suggestion else '',
            'suggested_candidate_actual_fio': fio_suggestion.get('candidate_fio', '') if fio_suggestion else '',
            'import_action': action,
            'import_status': status,
        })

    plan_df = pd.DataFrame(rows)
    return plan_df, summarize_enrollment_order_import(plan_df)

def enrollment_order_preview_rows(df):
    preview_df = df.copy()
    preview_df = preview_df.where(pd.notnull(preview_df), '')
    rows = preview_df[ENROLLMENT_ORDER_PREVIEW_COLUMNS].to_dict(orient='records')
    has_candidate_values = preview_df['has_candidate'].tolist() if 'has_candidate' in preview_df else [False] * len(rows)
    suggestion_columns = [
        'suggested_candidate_id', 'suggested_candidate_fio', 'suggested_candidate_login',
        'suggested_candidate_dogovor', 'suggested_candidate_similarity',
        'fio_review_reason', 'suggested_candidate_actual_fio'
    ]
    suggestion_rows = (
        preview_df[suggestion_columns].to_dict(orient='records')
        if all(column in preview_df for column in suggestion_columns)
        else [{} for _ in rows]
    )
    for row, has_candidate, suggestion in zip(rows, has_candidate_values, suggestion_rows):
        row.update(suggestion)
        row['has_fio_suggestion'] = bool(row.get('suggested_candidate_id'))
        action = row.get('import_action')
        if action == 'import' and has_candidate:
            row['action_label'] = 'Совпадение'
            row['badge_class'] = 'status-success'
        elif action == 'fio_review':
            row['action_label'] = 'Проверить ФИО'
            row['badge_class'] = 'status-warning'
        elif action == 'import':
            row['action_label'] = 'Нет кандидата'
            row['badge_class'] = 'status-warning'
        elif action == 'duplicate':
            row['action_label'] = 'Повтор'
            row['badge_class'] = 'status-warning'
        else:
            row['action_label'] = 'Пропуск'
            row['badge_class'] = 'status-danger'
    return rows

def build_enrollment_order_preview_report(df, summary):
    items = []
    for _, row in df.iterrows():
        action = clean_upload_text(row.get('import_action'))
        status = clean_upload_text(row.get('import_status'))
        if action == 'import' and bool(row.get('has_candidate')):
            continue
        field = 'ФИО' if 'ФИО' in status else 'Специальность' if 'специальность' in status.casefold() else 'Сверка'
        items.append(upload_report_item(int(row.get('_row_number', 0)), field, status))
    if not items:
        return None
    return build_upload_report(
        'Отчет по приказу о зачислении',
        summary['total'],
        items,
        [
            f"Загружается строк: {summary['import_count']}",
            f"Совпало с кандидатами: {summary['matched_count']}",
            f"В приказе без кандидата: {summary['unmatched_count']}",
            f"Похожих ФИО для проверки: {summary.get('fio_suggestion_count', 0)}",
            f"Повторов: {summary['duplicate_count']}",
            f"Пропущено: {summary['skipped_count']}",
        ]
    )

FILE_WORK_SECTIONS = (
    {'key': 'abiturients', 'title': 'Загрузка абитуриентов'},
    {'key': 'updates', 'title': 'Обновление почты и оплаты'},
    {'key': 'orders', 'title': 'Приказ о зачислении'},
    {'key': 'students', 'title': 'Загрузка студентов'},
)
FILE_WORK_SECTION_MAP = {section['key']: section for section in FILE_WORK_SECTIONS}

def normalize_file_work_section(section):
    section = str(section or '').strip().lower()
    return section if section in FILE_WORK_SECTION_MAP else None

def file_work_url(section=None):
    section = normalize_file_work_section(section)
    if section:
        return url_for('file_work', section=section)
    return url_for('file_work')

def file_work_redirect(section=None, code=303):
    return redirect(file_work_url(section), code=code)

def request_file_work_section(default=None):
    return normalize_file_work_section(
        request.form.get('file_section') or request.args.get('file_section') or default
    )

def render_file_work_page(campaign_year, active_section=None, **context):
    active_section = normalize_file_work_section(active_section)
    page_context = {
        'campaign_year': campaign_year,
        'file_work_sections': FILE_WORK_SECTIONS,
        'active_file_section': active_section,
        'active_file_section_title': FILE_WORK_SECTION_MAP[active_section]['title'] if active_section else None,
        'abiturients_import_result_ready': bool(session.get(ABITURIENTS_IMPORT_RESULT_SESSION_KEY)),
    }
    page_context.update(context)
    return render_template('file_work.html', **page_context)

def render_enrollment_order_preview_template(filepath, campaign_year, order_report=None):
    plan_df, summary = build_enrollment_order_import_plan(filepath, campaign_year)
    return render_file_work_page(
        campaign_year=campaign_year,
        active_section='orders',
        enrollment_order_preview=summary,
        enrollment_order_preview_rows=enrollment_order_preview_rows(plan_df),
        order_report=order_report if order_report is not None else build_enrollment_order_preview_report(plan_df, summary),
        enrollment_order_uploads=get_enrollment_order_uploads(campaign_year),
        pending_enrollment_order_import_token=os.path.basename(filepath)
    )

def fix_enrollment_candidate_fio(candidate_id, order_fio, campaign_year):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    order_fio = ' '.join(str(order_fio or '').split())
    if not candidate_id or not order_fio:
        raise ValueError('Не удалось определить кандидата или ФИО из приказа.')
    fam, imotch = split_fio_for_storage(order_fio)
    backup_path = create_database_backup('before_enrollment_candidate_fio_fix')
    with sqlite3.connect(DB_PATH) as conn:
        row = conn.execute(
            '''
            SELECT id, abiturient_id, fio, specialty_key
            FROM enrollment_candidates
            WHERE id=? AND campaign_year=?
            ''',
            (candidate_id, campaign_year)
        ).fetchone()
        if not row:
            raise ValueError('Кандидат для исправления ФИО не найден.')
        _candidate_id, abiturient_id, old_fio, specialty_key = row
        new_key = make_enrollment_match_key(order_fio, specialty_key)
        if new_key:
            duplicates = conn.execute(
                '''
                SELECT id, fio, specialty_key
                FROM enrollment_candidates
                WHERE campaign_year=? AND id<>?
                ''',
                (campaign_year, candidate_id)
            ).fetchall()
            for other_id, other_fio, other_specialty_key in duplicates:
                if make_enrollment_match_key(other_fio, other_specialty_key) == new_key:
                    raise ValueError(f'Нельзя исправить ФИО: такой кандидат уже есть в базе ({other_fio}).')
        conn.execute(
            '''
            UPDATE enrollment_candidates
            SET fio=?, fam=?, imotch=?
            WHERE id=? AND campaign_year=?
            ''',
            (order_fio, fam, imotch, candidate_id, campaign_year)
        )
        if abiturient_id is not None:
            conn.execute(
                '''
                UPDATE abiturients
                SET fio=?, fam=?, imotch=?
                WHERE id=? AND campaign_year=?
                ''',
                (order_fio, fam, imotch, abiturient_id, campaign_year)
            )
        refresh_enrollment_candidate_statuses(conn, campaign_year)
        log_action(
            'enrollment_candidate_fio_fixed_from_order',
            'campaign',
            campaign_year,
            (
                f"candidate_id={candidate_id}; old_fio={old_fio}; new_fio={order_fio}; "
                f"backup={os.path.basename(backup_path) if backup_path else ''}"
            ),
            conn
        )
    return {
        'candidate_id': candidate_id,
        'old_fio': old_fio,
        'new_fio': order_fio,
    }

def apply_enrollment_order_import(file_path, campaign_year=None, original_filename=None, uploaded_by=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    df, summary = build_enrollment_order_import_plan(file_path, campaign_year)
    if summary['import_count'] <= 0:
        raise ValueError('В приказе нет строк, которые можно загрузить. Старый приказ не изменён.')

    original_filename = os.path.basename(str(original_filename or os.path.basename(file_path) or 'order'))
    extension = os.path.splitext(file_path)[1].lower()
    stored_filename = ''
    backup_path = create_database_backup('before_enrollment_order_import')
    with sqlite3.connect(DB_PATH) as conn:
        order_rows = df[df['import_action'] == 'import'] if 'import_action' in df else df
        order_numbers = summarize_order_upload_values(order_rows, 'Номер приказа')
        order_dates = summarize_order_upload_values(order_rows, 'Дата приказа')
        cur = conn.execute(
            '''
            INSERT INTO enrollment_order_uploads
                (campaign_year, original_filename, uploaded_by, total_count, import_count,
                 matched_count, unmatched_count, duplicate_count, skipped_count,
                 order_numbers, order_dates)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''',
            (
                campaign_year, original_filename, uploaded_by or '',
                summary['total'], summary['import_count'], summary['matched_count'],
                summary['unmatched_count'], summary['duplicate_count'], summary['skipped_count'],
                order_numbers, order_dates
            )
        )
        upload_id = cur.lastrowid
        stored_filename = f'order_upload_{upload_id}_{secrets.token_hex(4)}{extension}'
        stored_path = get_stored_enrollment_order_path(stored_filename)
        shutil.copy2(file_path, stored_path)
        conn.execute(
            'UPDATE enrollment_order_uploads SET stored_filename=? WHERE id=?',
            (stored_filename, upload_id)
        )
        for _, row in df.iterrows():
            conn.execute(
                '''
                INSERT INTO enrollment_order_upload_rows
                    (upload_id, campaign_year, row_number, fio, specialty, group_name,
                     order_number, order_date, fio_key, specialty_key, group_key,
                     has_candidate, import_action, import_status)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    upload_id, row['campaign_year'], int(row['_row_number']), row['ФИО'],
                    row['Специальность'], row['Группа'], row['Номер приказа'],
                    row['Дата приказа'], row['fio_key'], row['specialty_key'],
                    row['group_key'], 1 if bool(row['has_candidate']) else 0,
                    row['import_action'], row['import_status']
                )
            )
        rebuild_enrollment_orders_from_upload_rows(conn, campaign_year)
        refresh_enrollment_candidate_statuses(conn, campaign_year)
        log_action(
            'enrollment_order_import',
            'campaign',
            campaign_year,
            (
                f"upload_id={upload_id}; filename={original_filename}; "
                f"rows={summary['total']}; imported={summary['import_count']}; "
                f"matched={summary['matched_count']}; unmatched={summary['unmatched_count']}; "
                f"duplicates={summary['duplicate_count']}; skipped={summary['skipped_count']}; "
                f"backup={os.path.basename(backup_path) if backup_path else ''}"
            ),
            conn
        )
    summary = dict(summary)
    summary['upload_id'] = upload_id
    summary['original_filename'] = original_filename
    return summary

def build_enrollment_order_template():
    output = io.BytesIO()
    template_df = pd.DataFrame(columns=['ФИО', 'Специальность', 'Группа', 'Номер приказа', 'Дата приказа'])
    help_df = pd.DataFrame([
        {
            'Поле': 'ФИО',
            'Что указать': 'ФИО человека из приказа. Используется для сверки с кандидатом.',
            'Пример': 'Иванов Иван Иванович',
        },
        {
            'Поле': 'Специальность',
            'Что указать': 'Официальное наименование из приказа или короткий код. Например: 33.02.01 «Фармация», ФМ, 34.02.01 «Сестринское дело», СД.',
            'Пример': '33.02.01 «Фармация»',
        },
        {
            'Поле': 'Группа',
            'Что указать': 'Необязательно. Если группа есть, система потребует выбрать именно ее при переносе в студенты.',
            'Пример': '26ФМ-11-1',
        },
        {
            'Поле': 'Номер приказа',
            'Что указать': 'Необязательно. Для справки и аудита.',
            'Пример': '123-у',
        },
        {
            'Поле': 'Дата приказа',
            'Что указать': 'Необязательно. Для справки и аудита.',
            'Пример': '2026-08-15',
        },
    ])
    example_df = pd.DataFrame([
        {
            'ФИО': 'Иванов Иван Иванович',
            'Специальность': '33.02.01 «Фармация»',
            'Группа': '26ФМ-11-1',
            'Номер приказа': '123-у',
            'Дата приказа': '2026-08-15',
        },
        {
            'ФИО': 'Петров Петр Петрович',
            'Специальность': '34.02.01 «Сестринское дело»',
            'Группа': '',
            'Номер приказа': '123-у',
            'Дата приказа': '2026-08-15',
        },
    ])
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        template_df.to_excel(writer, sheet_name='Шаблон', index=False)
        help_df.to_excel(writer, sheet_name='Подсказка', index=False)
        example_df.to_excel(writer, sheet_name='Пример', index=False)
    output.seek(0)
    return output

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            return redirect(url_for('login'))
        with sqlite3.connect(DB_PATH) as conn:
            cur = conn.execute('SELECT role, approved FROM users WHERE username=?', (session['user'],))
            user = cur.fetchone()
            if not user or user[0] != 'admin' or user[1] != 1:
                flash('Недостаточно прав')
                return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

def vaanedain_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session or session['user'] != 'vaanedain':
            flash('Доступ разрешён только главному администратору!')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/set_campaign', methods=['POST'])
@login_required
def set_campaign():
    campaign_year = normalize_campaign_year(request.form.get('campaign_year'), DEFAULT_CAMPAIGN_YEAR)
    session['campaign_year'] = campaign_year
    session['group_year'] = campaign_year
    next_url = request.form.get('next') or url_for('index')
    if not next_url.startswith('/') or next_url.startswith('//'):
        next_url = url_for('index')
    return redirect(next_url)

@app.route('/approve_users', methods=['GET', 'POST'])
@admin_required
def approve_users():
    with sqlite3.connect(DB_PATH) as conn:
        if request.method == 'POST':
            user_id = request.form.get('user_id')
            action = request.form.get('action')
            if action == 'approve':
                conn.execute('UPDATE users SET approved=1 WHERE id=?', (user_id,))
                log_action('user_approved', 'user', user_id, conn=conn)
            elif action == 'reject':
                conn.execute('DELETE FROM users WHERE id=?', (user_id,))
                log_action('user_rejected', 'user', user_id, conn=conn)
        cur = conn.execute('SELECT id, username, role FROM users WHERE approved=0')
        pending_users = cur.fetchall()
    return render_template('approve_users.html', pending_users=pending_users)

EXCEL_MIMETYPE = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'

def get_pending_abiturients_import_path(token):
    token = os.path.basename(str(token or ''))
    if not token.startswith(PENDING_ABITURIENTS_IMPORT_PREFIX):
        raise UploadValidationError('Временный файл импорта не найден. Загрузите файл ещё раз.')
    upload_root = os.path.abspath(app.config['UPLOAD_FOLDER'])
    import_path = os.path.abspath(os.path.join(upload_root, token))
    if os.path.commonpath([upload_root, import_path]) != upload_root or not os.path.exists(import_path):
        raise UploadValidationError('Временный файл импорта не найден. Загрузите файл ещё раз.')
    return import_path

def get_abiturients_import_result_path(token):
    token = os.path.basename(str(token or ''))
    if not token.startswith(ABITURIENTS_IMPORT_RESULT_PREFIX) or not token.endswith('.xlsx'):
        raise UploadValidationError('Файл с результатом импорта не найден.')
    upload_root = os.path.abspath(app.config['UPLOAD_FOLDER'])
    result_path = os.path.abspath(os.path.join(upload_root, token))
    if os.path.commonpath([upload_root, result_path]) != upload_root or not os.path.exists(result_path):
        raise UploadValidationError('Файл с результатом импорта не найден.')
    return result_path

def queue_abiturients_import_result(result_path):
    previous_token = session.get(ABITURIENTS_IMPORT_RESULT_SESSION_KEY)
    if previous_token:
        try:
            cleanup_temp_files(get_abiturients_import_result_path(previous_token))
        except UploadValidationError:
            pass
    session[ABITURIENTS_IMPORT_RESULT_SESSION_KEY] = os.path.basename(result_path)

def get_pending_students_import_path(token):
    token = os.path.basename(str(token or ''))
    if not token.startswith(PENDING_STUDENTS_IMPORT_PREFIX):
        raise UploadValidationError('Временный файл загрузки студентов не найден. Загрузите файл ещё раз.')
    upload_root = os.path.abspath(app.config['UPLOAD_FOLDER'])
    import_path = os.path.abspath(os.path.join(upload_root, token))
    if os.path.commonpath([upload_root, import_path]) != upload_root or not os.path.exists(import_path):
        raise UploadValidationError('Временный файл загрузки студентов не найден. Загрузите файл ещё раз.')
    return import_path

def get_pending_enrollment_order_import_path(token):
    token = os.path.basename(str(token or ''))
    if not token.startswith(PENDING_ENROLLMENT_ORDER_IMPORT_PREFIX):
        raise UploadValidationError('Временный файл приказа не найден. Загрузите файл ещё раз.')
    upload_root = os.path.abspath(app.config['UPLOAD_FOLDER'])
    import_path = os.path.abspath(os.path.join(upload_root, token))
    if os.path.commonpath([upload_root, import_path]) != upload_root or not os.path.exists(import_path):
        raise UploadValidationError('Временный файл приказа не найден. Загрузите файл ещё раз.')
    return import_path

def get_pending_email_source_import_path(token):
    token = os.path.basename(str(token or ''))
    if not token.startswith(PENDING_EMAIL_SOURCE_IMPORT_PREFIX):
        raise UploadValidationError('Временный файл источника почты не найден. Загрузите файл ещё раз.')
    upload_root = os.path.abspath(app.config['UPLOAD_FOLDER'])
    import_path = os.path.abspath(os.path.join(upload_root, token))
    if os.path.commonpath([upload_root, import_path]) != upload_root or not os.path.exists(import_path):
        raise UploadValidationError('Временный файл источника почты не найден. Загрузите файл ещё раз.')
    return import_path

def get_student_transfer_order_dir():
    order_dir = os.path.abspath(os.path.join(app.config['UPLOAD_FOLDER'], STUDENT_TRANSFER_ORDER_DIR))
    upload_root = os.path.abspath(app.config['UPLOAD_FOLDER'])
    if os.path.commonpath([upload_root, order_dir]) != upload_root:
        raise UploadValidationError('Некорректный путь хранения приказов.')
    os.makedirs(order_dir, exist_ok=True)
    return order_dir

def safe_transfer_order_original_name(file_storage):
    original_name = os.path.basename(str(file_storage.filename or 'transfer_order.pdf')).strip()
    return original_name or 'transfer_order.pdf'

def save_student_transfer_order_file(username, file_storage):
    validate_uploaded_file(file_storage, STUDENT_TRANSFER_ORDER_EXTENSIONS)
    original_name = safe_transfer_order_original_name(file_storage)
    safe_username = re.sub(r'[^a-zA-Z0-9_-]+', '_', str(username or 'student')).strip('_') or 'student'
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
    filename = f'{timestamp}_{safe_username}_{secrets.token_hex(6)}.pdf'
    order_dir = get_student_transfer_order_dir()
    order_path = os.path.join(order_dir, filename)
    file_storage.save(order_path)
    return {
        'filename': filename,
        'original_filename': original_name,
        'mime_type': file_storage.mimetype or 'application/pdf',
        'size': os.path.getsize(order_path),
        'path': order_path,
    }

def ensure_student_enrollment_movement(
    conn, username, order_map=None, rules=None, course_groups_enabled=None
):
    username = str(username or '').strip()
    if not username:
        return
    if course_groups_enabled is None:
        course_groups_enabled = are_course_groups_enabled()
    existing = conn.execute(
        '''
        SELECT 1
        FROM student_group_transfers
        WHERE username=? AND movement_type='enrollment'
        LIMIT 1
        ''',
        (username,)
    ).fetchone()
    if existing:
        return

    student = conn.execute(
        '''
        SELECT cohort1, cohort2, source_campaign_year, source_dogovor,
               source_fio, lastname, firstname
        FROM students
        WHERE username=?
        ''',
        (username,)
    ).fetchone()
    if not student:
        return
    cohort1, cohort2, source_campaign_year, source_dogovor, source_fio, lastname, firstname = student
    source_dogovor = str(source_dogovor or '').strip()
    if not source_dogovor:
        return
    fio = str(source_fio or '').strip()
    if not fio:
        fio = ' '.join(part for part in (lastname, firstname) if str(part or '').strip()).strip()
    if not fio:
        return

    campaign_year = normalize_campaign_year(source_campaign_year, infer_campaign_year(source_dogovor))
    order_match = get_enrollment_order_match_for_abiturient(
        fio,
        source_dogovor,
        campaign_year,
        order_map=order_map,
        rules=rules,
    )
    if not order_match:
        return

    first_transfer = conn.execute(
        '''
        SELECT old_cohort1, old_cohort2
        FROM student_group_transfers
        WHERE username=? AND COALESCE(movement_type, 'transfer')='transfer'
        ORDER BY created_at ASC, id ASC
        LIMIT 1
        ''',
        (username,)
    ).fetchone()
    if first_transfer and first_transfer[0]:
        target_group = normalize_group_name(first_transfer[0])
        target_cohort2 = (first_transfer[1] or derive_cohort2(target_group) or '') if course_groups_enabled else ''
    else:
        target_group = normalize_group_name(order_match.get('group_name') or cohort1)
        target_cohort2 = (derive_cohort2(target_group) or cohort2 or '') if course_groups_enabled else ''
    if not target_group:
        return

    record_student_enrollment_movement(
        conn,
        username,
        target_group,
        target_cohort2,
        order_match,
        'Система'
    )

def ensure_student_list_enrollment_movements(conn):
    rows = conn.execute(
        '''
        SELECT s.username, s.source_campaign_year, s.source_dogovor
        FROM students s
        WHERE TRIM(COALESCE(s.source_dogovor, ''))<>''
          AND NOT EXISTS (
              SELECT 1
              FROM student_group_transfers movement
              WHERE movement.username=s.username
                AND movement.movement_type='enrollment'
          )
        '''
    ).fetchall()
    if not rows:
        return

    rules = get_login_generation_rules()
    course_groups_enabled = are_course_groups_enabled()
    student_campaigns = {
        username: normalize_campaign_year(source_campaign_year, infer_campaign_year(source_dogovor))
        for username, source_campaign_year, source_dogovor in rows
    }
    order_maps = {
        campaign_year: get_enrollment_order_map(campaign_year)
        for campaign_year in set(student_campaigns.values())
    }
    for username, _source_campaign_year, _source_dogovor in rows:
        campaign_year = student_campaigns[username]
        ensure_student_enrollment_movement(
            conn,
            username,
            order_map=order_maps.get(campaign_year, {}),
            rules=rules,
            course_groups_enabled=course_groups_enabled,
        )

def get_student_transfer_orders(username):
    with sqlite3.connect(DB_PATH) as conn:
        ensure_student_enrollment_movement(conn, username)
        cur = conn.execute(
            '''
            SELECT id, username, movement_type, old_cohort1, old_cohort2, new_cohort1, new_cohort2,
                   enrollment_order_id, enrollment_order_upload_id, order_number, order_date, order_source,
                   order_filename, order_original_filename, order_mime_type, order_size,
                   created_by, created_at
            FROM student_group_transfers
            WHERE username=?
            ORDER BY created_at DESC, id DESC
            ''',
            (username,)
        )
        columns = [description[0] for description in cur.description]
        orders = [dict(zip(columns, row)) for row in cur.fetchall()]
        upload_ids = sorted({
            int(order.get('enrollment_order_upload_id'))
            for order in orders
            if str(order.get('enrollment_order_upload_id') or '').isdigit()
        })
        upload_lookup = {}
        if upload_ids:
            placeholders = ','.join('?' for _ in upload_ids)
            rows = conn.execute(
                f'''
                SELECT id, original_filename, stored_filename
                FROM enrollment_order_uploads
                WHERE id IN ({placeholders})
                ''',
                upload_ids
            ).fetchall()
            upload_lookup = {
                row[0]: {
                    'original_filename': row[1] or '',
                    'stored_filename': row[2] or '',
                }
                for row in rows
            }
    for order in orders:
        movement_type = order.get('movement_type') or 'transfer'
        upload_id = order.get('enrollment_order_upload_id')
        upload_key = int(upload_id) if str(upload_id or '').isdigit() else None
        upload_info = upload_lookup.get(upload_key, {})
        order['movement_type'] = movement_type
        order['is_enrollment'] = movement_type == 'enrollment'
        order['movement_type_label'] = 'Зачислен' if order['is_enrollment'] else 'Перевод'
        order['movement_source_label'] = 'Приказ о зачислении' if order['is_enrollment'] else 'Приказ о переводе'
        order['has_order_file'] = bool(str(order.get('order_filename') or '').strip())
        order['has_enrollment_order_file'] = order['is_enrollment'] and bool(upload_info.get('stored_filename'))
        order['order_size_text'] = format_upload_size(order.get('order_size') or 0)
        order['created_date_text'] = format_display_date(order.get('created_at'))
        order['order_date_text'] = format_display_date(order.get('order_date')) if order.get('order_date') else '-'
        order['timeline_date_text'] = (
            order['order_date_text']
            if order['is_enrollment'] and order.get('order_date')
            else order['created_date_text']
        )
        order['from_group_text'] = order.get('old_cohort1') or 'Без группы'
        order['to_group_text'] = order.get('new_cohort1') or 'Без группы'
        order['cohort2_text'] = order.get('new_cohort2') or '-'
        if order['has_order_file']:
            order['order_file_text'] = order.get('order_original_filename') or 'PDF прикреплен'
        elif order['is_enrollment']:
            order['order_file_text'] = (
                upload_info.get('original_filename')
                or order.get('order_original_filename')
                or 'Файл приказа не сохранен'
            )
        else:
            order['order_file_text'] = 'PDF не прикреплен'
        order['download_url'] = ''
        order['download_label'] = ''
        if has_request_context():
            if order['has_order_file']:
                order['download_url'] = url_for('download_student_transfer_order', transfer_id=order['id'])
                order['download_label'] = 'Скачать PDF'
            elif order['has_enrollment_order_file']:
                order['download_url'] = url_for('download_enrollment_order_upload', upload_id=upload_key)
                order['download_label'] = 'Скачать приказ'
    return orders

def record_student_enrollment_movement(conn, username, target_group, target_cohort2, order_match=None, created_by=''):
    order_match = order_match or {}
    if not order_match:
        return
    conn.execute(
        '''
        INSERT INTO student_group_transfers
            (username, movement_type, old_cohort1, old_cohort2, new_cohort1, new_cohort2,
             enrollment_order_id, enrollment_order_upload_id, order_number, order_date, order_source,
             order_filename, order_original_filename, order_mime_type, order_size, created_by)
        VALUES (?, 'enrollment', ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''',
        (
            username,
            'Абитуриенты',
            '',
            target_group,
            target_cohort2,
            order_match.get('id'),
            order_match.get('upload_id'),
            order_match.get('order_number') or '',
            order_match.get('order_date') or '',
            'enrollment_order',
            '',
            order_match.get('original_filename') or '',
            '',
            0,
            created_by,
        )
    )

def get_student_transfer_order_download(transfer_id):
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            '''
            SELECT id, username, order_filename, order_original_filename, order_mime_type
            FROM student_group_transfers
            WHERE id=?
            ''',
            (transfer_id,)
        )
        row = cur.fetchone()
    if not row:
        raise FileNotFoundError('Приказ о переводе не найден.')
    _transfer_id, username, filename, original_filename, mime_type = row
    if not str(filename or '').strip():
        raise FileNotFoundError('К переводу не прикреплен PDF.')
    order_dir = get_student_transfer_order_dir()
    order_path = os.path.abspath(os.path.join(order_dir, os.path.basename(filename)))
    if os.path.commonpath([order_dir, order_path]) != order_dir or not os.path.exists(order_path):
        raise FileNotFoundError('Файл приказа о переводе не найден.')
    return {
        'path': order_path,
        'username': username,
        'download_name': original_filename or os.path.basename(order_path),
        'mime_type': mime_type or 'application/pdf',
    }

def build_abiturients_upload_response(file_storage, campaign_year):
    upload_path = None
    result_path = None
    try:
        upload_path = save_upload_to_temp(file_storage, ABITURIENT_UPLOAD_EXTENSIONS)
        result_path, summary = apply_abiturients_import(upload_path, campaign_year)
        flash(
            (
                f"Импорт завершён: добавлено {summary['ready_count']}, "
                f"дублей {summary['duplicate_count']}, конфликтов {summary['conflict_count']}, "
                f"возможных тёзок {summary['warning_count']}."
            ),
            'success'
        )
        return send_temp_download(result_path, 'abiturients_with_logins.xlsx', EXCEL_MIMETYPE)
    except UploadValidationError as exc:
        flash(str(exc), 'error')
    except Exception as exc:
        flash(f'Ошибка обработки файла: {exc}', 'error')
    finally:
        cleanup_temp_files(upload_path, result_path)
    return None

@app.route('/', methods=['GET', 'POST'])
@login_required
def index():
    campaign_year = get_active_campaign_year()
    if request.method == 'POST':
        if not ensure_campaign_open(campaign_year):
            return redirect(url_for('index'), code=303)
        response = build_abiturients_upload_response(request.files.get('file'), campaign_year)
        if response:
            return response
        return redirect(url_for('file_work'), code=303)
    dashboard = get_dashboard_data(campaign_year)
    return render_template('index.html', dashboard=dashboard)

@app.route('/documentation')
@login_required
def documentation():
    """Render the read-only in-app user guide."""
    return render_template('documentation.html')

@app.route('/search')
@login_required
def search():
    campaign_year = get_active_campaign_year()
    query = request.args.get('q', '').strip()
    results = global_search_records(query, campaign_year) if query else []
    return render_template('search.html', query=query, results=results, campaign_year=campaign_year)

@app.route('/search_overlay')
@login_required
def search_overlay():
    campaign_year = get_active_campaign_year()
    query = request.args.get('q', '').strip()
    results = global_search_records(query, campaign_year) if query else []
    return jsonify({
        'query': query,
        'total': len(results),
        'results': results,
    })

@app.route('/data_checks')
@login_required
def data_checks():
    campaign_year = get_active_campaign_year()
    report = get_data_quality_report(campaign_year)
    return render_template('data_checks.html', report=report)

@app.route('/person/<kind>/<path:record_id>')
@login_required
def person_card(kind, record_id):
    record = get_person_record(kind, record_id)
    if not record:
        flash('Запись не найдена', 'error')
        return redirect(url_for('search'))
    if record['kind'] == 'student' and session.get('role') != 'admin':
        record['fields']['password'] = '******'
    transfer_orders = []
    if record['kind'] == 'student' and session.get('role') == 'admin':
        transfer_orders = get_student_transfer_orders(record['fields'].get('username'))
    return render_template(
        'person_card.html',
        record=record,
        card_view=build_person_card_view(record),
        transfer_orders=transfer_orders,
    )

@app.route('/file_work', defaults={'section': None}, methods=['GET', 'POST'])
@app.route('/file_work/<section>', methods=['GET', 'POST'])
@login_required
def file_work(section=None):
    if section is not None and normalize_file_work_section(section) is None:
        flash('Раздел работы с файлами не найден.', 'error')
        return file_work_redirect()
    active_section = normalize_file_work_section(section)
    campaign_year = get_active_campaign_year()
    updates_report = None
    email_source_report = None
    students_report = None
    order_report = None
    if request.method == 'GET':
        updates_report = session.pop('abiturients_updates_report', None)
        email_source_report = session.pop('abiturients_email_source_report', None)
        students_report = session.pop('students_upload_report', None)
        order_report = session.pop('enrollment_order_report', None)
    if request.method == 'POST':
        target_section = request_file_work_section(active_section) or 'abiturients'
        import_action = request.form.get('import_action', 'preview')
        if import_action == 'confirm':
            if not ensure_campaign_open(campaign_year):
                return file_work_redirect(target_section)
            pending_path = None
            result_path = None
            try:
                pending_path = get_pending_abiturients_import_path(request.form.get('pending_import'))
                result_path, summary = apply_abiturients_import(pending_path, campaign_year)
                flash(
                    (
                        f"Импорт завершён: добавлено {summary['ready_count']}, "
                        f"дублей {summary['duplicate_count']}, конфликтов {summary['conflict_count']}, "
                        f"возможных тёзок {summary['warning_count']}."
                    ),
                    'success'
                )
                queue_abiturients_import_result(result_path)
                result_path = None
                return file_work_redirect(target_section)
            except (UploadValidationError, ValueError) as exc:
                flash(str(exc), 'error')
            except Exception as exc:
                flash(f'Ошибка обработки файла: {exc}', 'error')
            finally:
                cleanup_temp_files(pending_path, result_path)
            return file_work_redirect(target_section)

        if import_action == 'cancel':
            try:
                cleanup_temp_files(get_pending_abiturients_import_path(request.form.get('pending_import')))
                flash('Предпросмотр импорта отменён.', 'info')
            except UploadValidationError:
                pass
            return file_work_redirect(target_section)

        upload_path = None
        try:
            if not ensure_campaign_open(campaign_year):
                return file_work_redirect(target_section)
            upload_path = save_upload_to_temp(
                request.files.get('file'),
                ABITURIENT_UPLOAD_EXTENSIONS,
                prefix=PENDING_ABITURIENTS_IMPORT_PREFIX
            )
            plan_df, preview_summary = build_abiturients_import_plan(upload_path, campaign_year)
            return render_file_work_page(
                campaign_year=campaign_year,
                active_section=target_section,
                abiturients_preview=preview_summary,
                abiturients_preview_rows=dataframe_preview_rows(plan_df),
                abiturients_report=build_abiturients_preview_report(plan_df, preview_summary),
                updates_report=updates_report,
                email_source_report=email_source_report,
                students_report=students_report,
                order_report=order_report,
                enrollment_order_uploads=get_enrollment_order_uploads(campaign_year),
                pending_import_token=os.path.basename(upload_path)
            )
        except UploadValidationError as exc:
            flash(str(exc), 'error')
        except Exception as exc:
            cleanup_temp_files(upload_path)
            flash(f'Ошибка обработки файла: {exc}', 'error')
        return file_work_redirect(target_section)
    return render_file_work_page(
        campaign_year=campaign_year,
        active_section=active_section,
        updates_report=updates_report,
        email_source_report=email_source_report,
        students_report=students_report,
        order_report=order_report,
        enrollment_order_uploads=get_enrollment_order_uploads(campaign_year)
    )

@app.route('/file_work/abiturients/import-result')
@login_required
def download_abiturients_import_result():
    token = session.pop(ABITURIENTS_IMPORT_RESULT_SESSION_KEY, None)
    try:
        result_path = get_abiturients_import_result_path(token)
    except UploadValidationError as exc:
        flash(str(exc), 'error')
        return file_work_redirect('abiturients')
    return send_temp_download(result_path, 'abiturients_with_logins.xlsx', EXCEL_MIMETYPE)

@app.route('/enrollment_order_upload', methods=['POST'])
@login_required
def enrollment_order_upload():
    target_section = request_file_work_section('orders')
    if session.get('role') not in {'admin', 'assistant', 'manager', 'operator'}:
        flash('Недостаточно прав', 'error')
        return file_work_redirect(target_section)
    campaign_year = get_active_campaign_year()
    order_action = request.form.get('order_import_action', 'preview')
    if order_action == 'confirm':
        if not ensure_campaign_open(campaign_year):
            return file_work_redirect(target_section)
        filepath = None
        try:
            filepath = get_pending_enrollment_order_import_path(request.form.get('pending_enrollment_order_import'))
            original_filename = session.pop('pending_enrollment_order_original_filename', os.path.basename(filepath))
            summary = apply_enrollment_order_import(
                filepath,
                campaign_year,
                original_filename=original_filename,
                uploaded_by=session.get('user')
            )
            flash(
                (
                    f"Приказ загружен: строк {summary['import_count']}, "
                    f"совпало с кандидатами {summary['matched_count']}, "
                    f"в приказе без кандидата {summary['unmatched_count']}."
                ),
                'success' if summary['unmatched_count'] == 0 else 'info'
            )
        except (UploadValidationError, ValueError) as exc:
            flash(str(exc), 'error')
        except Exception as exc:
            flash(f'Ошибка загрузки приказа: {exc}', 'error')
        finally:
            cleanup_temp_files(filepath)
        return file_work_redirect(target_section)

    if order_action == 'cancel':
        try:
            cleanup_temp_files(get_pending_enrollment_order_import_path(request.form.get('pending_enrollment_order_import')))
            session.pop('pending_enrollment_order_original_filename', None)
            flash('Предпросмотр приказа отменён.', 'info')
        except UploadValidationError:
            pass
        return file_work_redirect(target_section)

    filepath = None
    try:
        if not ensure_campaign_open(campaign_year):
            return file_work_redirect(target_section)
        order_file = request.files.get('order_file')
        original_filename = os.path.basename(order_file.filename if order_file else '')
        filepath = save_upload_to_temp(
            order_file,
            ENROLLMENT_ORDER_UPLOAD_EXTENSIONS,
            prefix=PENDING_ENROLLMENT_ORDER_IMPORT_PREFIX
        )
        session['pending_enrollment_order_original_filename'] = original_filename
        return render_enrollment_order_preview_template(filepath, campaign_year)
    except (UploadValidationError, ValueError) as exc:
        flash(str(exc), 'error')
        cleanup_temp_files(filepath)
    except Exception as exc:
        flash(f'Ошибка загрузки приказа: {exc}', 'error')
        cleanup_temp_files(filepath)
    return file_work_redirect(target_section)

@app.route('/enrollment_order_preview/fix_candidate_fio', methods=['POST'])
@login_required
def fix_enrollment_candidate_fio_from_order():
    target_section = request_file_work_section('orders')
    if session.get('role') not in {'admin', 'assistant', 'manager', 'operator'}:
        flash('Недостаточно прав', 'error')
        return file_work_redirect(target_section)
    campaign_year = get_active_campaign_year()
    if not ensure_campaign_open(campaign_year):
        return file_work_redirect(target_section)
    filepath = None
    try:
        filepath = get_pending_enrollment_order_import_path(request.form.get('pending_enrollment_order_import'))
        row_number = int(request.form.get('row_number') or 0)
        candidate_id = int(request.form.get('candidate_id') or 0)
        plan_df, _summary = build_enrollment_order_import_plan(filepath, campaign_year)
        matched_rows = plan_df[plan_df['_row_number'] == row_number]
        if matched_rows.empty:
            raise ValueError('Строка приказа в предпросмотре не найдена.')
        order_row = matched_rows.iloc[0]
        if int(order_row.get('suggested_candidate_id') or 0) != candidate_id:
            raise ValueError('Выбранный кандидат больше не совпадает с подсказкой. Обновите предпросмотр.')
        result = fix_enrollment_candidate_fio(candidate_id, order_row['ФИО'], campaign_year)
        flash(f"ФИО в базе исправлено по приказу: {result['old_fio']} → {result['new_fio']}.", 'success')
        return render_enrollment_order_preview_template(filepath, campaign_year)
    except (UploadValidationError, ValueError) as exc:
        flash(str(exc), 'error')
    except Exception as exc:
        flash(f'Ошибка исправления ФИО: {exc}', 'error')
    if filepath:
        try:
            return render_enrollment_order_preview_template(filepath, campaign_year)
        except Exception:
            pass
    return file_work_redirect(target_section)

@app.route('/enrollment_order_uploads/<int:upload_id>')
@login_required
def enrollment_order_upload_detail(upload_id):
    upload = get_enrollment_order_upload(upload_id)
    if not upload:
        flash('Загрузка приказа не найдена.', 'error')
        return file_work_redirect('orders')
    rows = get_enrollment_order_upload_rows(upload_id)
    return render_template(
        'enrollment_order_upload.html',
        upload=upload,
        rows=rows,
        campaign_year=get_active_campaign_year()
    )

@app.route('/enrollment_order_uploads/<int:upload_id>/student_roster')
@login_required
def enrollment_order_student_roster(upload_id):
    roster = build_enrollment_order_student_roster(upload_id)
    if not roster:
        flash('Загрузка приказа не найдена.', 'error')
        return redirect(url_for('data_checks'), code=303)
    return render_template(
        'enrollment_order_student_roster.html',
        roster=roster,
        upload=roster['upload'],
        groups=roster['groups'],
        summary=roster['summary'],
        max_group_students=MAX_GROUP_STUDENTS,
    )

@app.route('/enrollment_order_uploads/<int:upload_id>/student_roster/fio_review', methods=['POST'])
@login_required
def review_enrollment_order_student_roster_fio(upload_id):
    target_url = url_for('enrollment_order_student_roster', upload_id=upload_id)
    if session.get('role') not in {'admin', 'assistant', 'manager', 'operator'}:
        flash('Недостаточно прав', 'error')
        return redirect(target_url, code=303)
    upload = get_enrollment_order_upload(upload_id)
    if not upload:
        flash('Загрузка приказа не найдена.', 'error')
        return redirect(url_for('data_checks'), code=303)
    if not ensure_campaign_open(upload['campaign_year']):
        return redirect(target_url, code=303)

    try:
        action = str(request.form.get('review_action') or '').strip()
        row_id = int(request.form.get('row_id') or 0)
        abiturient_id = int(request.form.get('abiturient_id') or 0)
        roster = build_enrollment_order_student_roster(upload_id)
        roster_row = next((row for row in roster['rows'] if int(row['id']) == row_id), None)
        if not roster_row:
            raise ValueError('Строка приказа в предварительном списке не найдена.')

        if action == 'reset':
            set_enrollment_order_roster_fio_review_status(
                upload_id,
                row_id,
                '',
                updated_by=session.get('user', ''),
            )
            flash('Строка возвращена на сверку ФИО.', 'success')
        elif action in {'fix', 'skip', 'link'}:
            if roster_row.get('fio_review_kind') != 'fio_typo':
                raise ValueError('Для этой строки нет безопасного предложения по исправлению ФИО.')
            if int(roster_row.get('suggested_abiturient_id') or 0) != abiturient_id:
                raise ValueError('Предложенная запись изменилась. Обновите предварительный список.')
            if action == 'fix':
                result = fix_abiturient_fio_from_enrollment_order_roster(
                    upload_id,
                    row_id,
                    abiturient_id,
                    updated_by=session.get('user', ''),
                )
                flash(
                    f"ФИО исправлено в базе абитуриентов: {result['old_fio']} → {result['new_fio']}.",
                    'success',
                )
            elif action == 'skip':
                set_enrollment_order_roster_fio_review_status(
                    upload_id,
                    row_id,
                    'skipped',
                    abiturient_id=abiturient_id,
                    updated_by=session.get('user', ''),
                )
                flash('Исправление ФИО пропущено. Данные абитуриента не изменены.', 'info')
            else:
                set_enrollment_order_roster_fio_review_status(
                    upload_id,
                    row_id,
                    'linked',
                    abiturient_id=abiturient_id,
                    updated_by=session.get('user', ''),
                )
                flash(
                    (
                        'Опечатка отмечена как ошибка в приказе. '
                        f"Логин {roster_row.get('suggested_abiturient_login') or '-'} "
                        'присвоен строке без изменения ФИО в базе.'
                    ),
                    'success',
                )
        else:
            raise ValueError('Не выбрано действие сверки ФИО.')
    except ValueError as exc:
        flash(str(exc), 'error')
    except Exception as exc:
        flash(f'Ошибка сверки ФИО: {exc}', 'error')
    return redirect(target_url, code=303)

@app.route('/enrollment_order_uploads/<int:upload_id>/student_roster/download')
@login_required
def download_enrollment_order_student_roster(upload_id):
    roster = build_enrollment_order_student_roster(upload_id)
    if not roster:
        flash('Загрузка приказа не найдена.', 'error')
        return redirect(url_for('data_checks'), code=303)
    export_rows = enrollment_order_roster_export_rows(roster)
    output = io.BytesIO()
    pd.DataFrame(export_rows).to_excel(
        output,
        index=False,
        sheet_name='Список по приказу',
    )
    output.seek(0)
    log_action(
        'enrollment_order_student_roster_exported',
        'enrollment_order_upload',
        upload_id,
        f"rows={len(export_rows)}; campaign_year={roster['upload']['campaign_year']}",
    )
    return send_file(
        output,
        as_attachment=True,
        download_name=f'enrollment_order_students_{upload_id}.xlsx',
        mimetype=EXCEL_MIMETYPE,
    )

@app.route('/enrollment_order_uploads/<int:upload_id>/download')
@login_required
def download_enrollment_order_upload(upload_id):
    upload = get_enrollment_order_upload(upload_id)
    if not upload:
        flash('Загрузка приказа не найдена.', 'error')
        return file_work_redirect('orders')
    stored_path = get_stored_enrollment_order_path(upload.get('stored_filename'))
    if not stored_path or not os.path.exists(stored_path):
        flash('Исходный файл приказа не найден в хранилище.', 'error')
        return redirect(url_for('enrollment_order_upload_detail', upload_id=upload_id), code=303)
    return send_file(
        stored_path,
        as_attachment=True,
        download_name=upload.get('original_filename') or os.path.basename(stored_path)
    )

@app.route('/enrollment_order_uploads/<int:upload_id>/delete', methods=['POST'])
@login_required
def delete_enrollment_order_upload_route(upload_id):
    if session.get('role') not in {'admin', 'assistant', 'manager', 'operator'}:
        flash('Недостаточно прав', 'error')
        return file_work_redirect('orders')
    campaign_year = get_active_campaign_year()
    if not ensure_campaign_open(campaign_year):
        return file_work_redirect('orders')
    try:
        deleted = delete_enrollment_order_upload(upload_id, campaign_year)
        flash(f"Приказ удалён из сверки: {deleted['original_filename'] or deleted['id']}. Статусы кандидатов пересчитаны.", 'success')
    except ValueError as exc:
        flash(str(exc), 'error')
    except Exception as exc:
        flash(f'Ошибка удаления приказа: {exc}', 'error')
    return file_work_redirect('orders')

@app.route('/enrollment_order_template/download')
@login_required
def download_enrollment_order_template():
    return send_file(
        build_enrollment_order_template(),
        as_attachment=True,
        download_name='enrollment_order_template.xlsx',
        mimetype=EXCEL_MIMETYPE
    )

@app.route('/abiturients_updates_upload', methods=['POST'])
@login_required
def abiturients_updates_upload():
    target_section = request_file_work_section('updates')
    if session.get('role') not in {'admin', 'assistant', 'operator'}:
        flash('Недостаточно прав', 'error')
        return file_work_redirect(target_section)
    campaign_year = get_active_campaign_year()
    if not ensure_campaign_open(campaign_year):
        return file_work_redirect(target_section)
    filepath = None
    try:
        filepath = save_upload_to_temp(request.files.get('updates_file'), ABITURIENT_UPLOAD_EXTENSIONS)
        summary = process_abiturients_updates(filepath, campaign_year)
        report_items = (summary.get('errors') or []) + (summary.get('not_found_rows') or [])
        if report_items:
            session['abiturients_updates_report'] = build_upload_report(
                'Отчет по файлу обновлений',
                summary['total'],
                report_items,
                [
                    f"Обработано строк: {summary['total']}",
                    f"Обновлено почт: {summary['updated_email']}",
                    f"Обновлено статусов оплаты: {summary['updated_paid']}",
                ]
            )
        flash(
            (
                f"Обновления применены: почт {summary['updated_email']}, "
                f"статусов оплаты {summary['updated_paid']}, замечаний {len(report_items)}."
            ),
            'success' if not report_items else 'info'
        )
    except UploadValidationError as exc:
        flash(str(exc), 'error')
    except Exception as exc:
        flash(f'Ошибка обновления данных: {exc}', 'error')
    finally:
        cleanup_temp_files(filepath)
    return file_work_redirect(target_section)

@app.route('/abiturients_email_source_upload', methods=['POST'])
@login_required
def abiturients_email_source_upload():
    target_section = request_file_work_section('updates')
    if session.get('role') not in {'admin', 'assistant', 'operator'}:
        flash('Недостаточно прав', 'error')
        return file_work_redirect(target_section)
    campaign_year = get_active_campaign_year()
    email_source_action = request.form.get('email_source_import_action', 'preview')

    if email_source_action == 'cancel':
        try:
            cleanup_temp_files(
                get_pending_email_source_import_path(request.form.get('pending_email_source_import'))
            )
            flash('Предпросмотр источника почты отменён.', 'info')
        except UploadValidationError:
            pass
        return file_work_redirect(target_section)

    if not ensure_campaign_open(campaign_year):
        return file_work_redirect(target_section)

    if email_source_action == 'confirm':
        filepath = None
        try:
            filepath = get_pending_email_source_import_path(
                request.form.get('pending_email_source_import')
            )
            summary = apply_email_source_updates(filepath, campaign_year)
            session['abiturients_email_source_report'] = build_email_source_upload_report(
                summary,
                applied=True
            )
            flash(
                (
                    f"Источник почты применён: у абитуриентов обновлено {summary['updated_abiturients']}, "
                    f"у студентов {summary['updated_students']}, замечаний {len(summary['issues'])}."
                ),
                'success' if not summary['issues'] else 'info'
            )
        except (UploadValidationError, ValueError) as exc:
            flash(str(exc), 'error')
        except Exception as exc:
            flash(f'Ошибка применения источника почты: {exc}', 'error')
        finally:
            cleanup_temp_files(filepath)
        return file_work_redirect(target_section)

    filepath = None
    try:
        filepath = save_upload_to_temp(
            request.files.get('email_source_file'),
            ABITURIENT_UPLOAD_EXTENSIONS,
            prefix=PENDING_EMAIL_SOURCE_IMPORT_PREFIX
        )
        summary = build_email_source_update_plan(filepath, campaign_year)
        summary.pop('_operations', None)
        return render_file_work_page(
            campaign_year=campaign_year,
            active_section=target_section,
            email_source_preview=summary,
            email_source_preview_rows=summary['preview_rows'],
            email_source_report=build_email_source_upload_report(summary),
            pending_email_source_import_token=os.path.basename(filepath),
            enrollment_order_uploads=get_enrollment_order_uploads(campaign_year),
        )
    except (UploadValidationError, ValueError) as exc:
        flash(str(exc), 'error')
        cleanup_temp_files(filepath)
    except Exception as exc:
        flash(f'Ошибка обработки источника почты: {exc}', 'error')
        cleanup_temp_files(filepath)
    return file_work_redirect(target_section)

@app.route('/abiturients_updates_template/download')
@login_required
def download_abiturients_updates_template():
    return send_file(
        build_abiturients_updates_template(),
        as_attachment=True,
        download_name='abiturients_updates_template.xlsx',
        mimetype=EXCEL_MIMETYPE
    )

@app.route('/abiturients')
@login_required
def abiturients():
    campaign_year = get_active_campaign_year()
    order_by = request.args.get('order_by', 'created_at')
    order_dir = request.args.get('order_dir', 'desc')
    spec = request.args.get('spec')
    base = request.args.get('base')
    year = request.args.get('year')
    is_i = request.args.get('is_i')
    has_email = request.args.get('has_email')
    has_paid = request.args.get('has_paid')
    has_order = request.args.get('has_order')
    withdrawn = request.args.get('withdrawn')
    q = request.args.get('q', '').strip()
    abiturients = get_all_abiturients(
        order_by, order_dir, spec, base, year, is_i, campaign_year,
        has_email, has_paid, q, withdrawn, has_order
    )
    login_rules = get_login_generation_rules()
    specs = list(login_rules['spec_codes'].keys())
    bases = list(login_rules['base_codes'].keys())
    years = get_campaign_years()
    return render_template('abiturients.html', abiturients=abiturients, order_by=order_by, order_dir=order_dir, specs=specs, bases=bases, years=years, campaign_year=campaign_year)

ABITURIENT_ORDER_COLUMNS = {
    'id', 'fio', 'dogovor', 'login', 'campaign_year', 'fam', 'imotch',
    'created_at', 'email', 'paid'
}
ABITURIENT_LIST_FILTER_PARAMS = (
    'spec', 'base', 'year', 'is_i', 'has_email', 'has_paid', 'has_order',
    'withdrawn', 'q'
)

def get_abiturient_list_query_params(values):
    order_by = str(values.get('order_by') or 'created_at').strip()
    if order_by not in ABITURIENT_ORDER_COLUMNS:
        order_by = 'created_at'
    order_dir = str(values.get('order_dir') or 'desc').strip().lower()
    if order_dir not in {'asc', 'desc'}:
        order_dir = 'desc'
    params = {'order_by': order_by, 'order_dir': order_dir}
    for name in ABITURIENT_LIST_FILTER_PARAMS:
        value = str(values.get(name) or '').strip()
        if value:
            params[name] = value
    return params

def get_abiturient_edit_navigation(abiturient_id, campaign_year, list_query):
    rows = get_all_abiturients(
        list_query['order_by'],
        list_query['order_dir'],
        list_query.get('spec'),
        list_query.get('base'),
        list_query.get('year'),
        list_query.get('is_i'),
        campaign_year,
        list_query.get('has_email'),
        list_query.get('has_paid'),
        list_query.get('q'),
        list_query.get('withdrawn'),
        list_query.get('has_order'),
    )
    current_index = next(
        (index for index, row in enumerate(rows) if row['id'] == abiturient_id),
        None
    )
    if current_index is None:
        return {'previous': None, 'next': None, 'position': None, 'total': len(rows)}
    return {
        'previous': rows[current_index - 1] if current_index > 0 else None,
        'next': rows[current_index + 1] if current_index + 1 < len(rows) else None,
        'position': current_index + 1,
        'total': len(rows),
    }

def base_filter_variants(base):
    value = str(base or '').strip()
    if not value:
        return []
    variants = [value]
    canonical = canonicalize_base_label(value)
    if canonical and canonical not in variants:
        variants.append(canonical)
    if canonical in {'11И', '9И'}:
        lowercase_alias = f'{canonical[:-1]}и'
        if lowercase_alias not in variants:
            variants.append(lowercase_alias)
    return variants

def normalize_search_text(value):
    return ' '.join(str(value or '').split()).casefold()

def get_all_abiturients(
    order_by='created_at', order_dir='desc', spec=None, base=None, year=None,
    is_i=None, campaign_year=None, has_email=None, has_paid=None, q=None,
    withdrawn=None, has_order=None
):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    if order_by not in ABITURIENT_ORDER_COLUMNS:
        order_by = 'created_at'
    if order_dir.lower() not in {'asc', 'desc'}:
        order_dir = 'desc'
    query = "SELECT * FROM abiturients WHERE campaign_year=?"
    params = [campaign_year]
    if spec:
        query += " AND dogovor LIKE ?"
        params.append(f"%{spec}%")
    if base:
        variants = base_filter_variants(base)
        query += " AND (" + " OR ".join("dogovor LIKE ?" for _ in variants) + ")"
        params.extend(f"%{variant}%" for variant in variants)
    if year:
        query += " AND dogovor LIKE ?"
        params.append(f"%{year}%")
    if is_i == '1':
        query += " AND login LIKE ?"
        params.append("%i%")
    elif is_i == '0':
        query += " AND login NOT LIKE ?"
        params.append("%i%")
    if has_email == '1':
        query += " AND email IS NOT NULL AND email <> ''"
    elif has_email == '0':
        query += " AND (email IS NULL OR email = '')"
    if has_paid == '1':
        query += " AND paid = 1"
    elif has_paid == '0':
        query += " AND paid = 0"
    if withdrawn == '1':
        query += " AND LOWER(COALESCE(login, '')) LIKE 'del%'"
    elif withdrawn == '0':
        query += " AND LOWER(COALESCE(login, '')) NOT LIKE 'del%'"
    q = normalize_search_text(q)
    if q:
        query += '''
            AND (
                INSTR(NORMALIZE_SEARCH(fio), ?) > 0
                OR INSTR(NORMALIZE_SEARCH(dogovor), ?) > 0
                OR INSTR(NORMALIZE_SEARCH(login), ?) > 0
                OR INSTR(NORMALIZE_SEARCH(email), ?) > 0
            )
        '''
        params.extend([q] * 4)
    query += f" ORDER BY {order_by} {order_dir.upper()}"
    if order_by != 'id':
        query += f", id {order_dir.upper()}"
    with sqlite3.connect(DB_PATH) as conn:
        conn.create_function('NORMALIZE_SEARCH', 1, normalize_search_text)
        cur = conn.execute(query, params)
        rows = cur.fetchall()
        columns = [desc[0] for desc in cur.description]
        result = [dict(zip(columns, row)) for row in rows]

    if has_order in {'0', '1'}:
        order_map = get_enrollment_order_map(campaign_year)
        rules = get_login_generation_rules()
        expected_match = has_order == '1'
        result = [
            row for row in result
            if bool(get_enrollment_order_match_for_abiturient(
                row.get('fio'),
                row.get('dogovor'),
                campaign_year,
                order_map=order_map,
                rules=rules,
            )) == expected_match
        ]
    return result

def normalize_student_search(value):
    return normalize_search_text(value)

def student_field_matches(row, field, search_value):
    search_value = normalize_student_search(search_value)
    if not search_value:
        return True
    return search_value in normalize_student_search(row.get(field))

def make_enrollment_order_filter_value(order_number, order_date):
    return json.dumps(
        [str(order_number or '').strip(), str(order_date or '').strip()],
        ensure_ascii=False,
        separators=(',', ':')
    )

def parse_enrollment_order_filter_value(value):
    value = str(value or '').strip()
    if not value:
        return None
    try:
        decoded = json.loads(value)
    except (TypeError, ValueError, json.JSONDecodeError):
        decoded = None
    if isinstance(decoded, list) and len(decoded) == 2:
        return str(decoded[0] or '').strip(), str(decoded[1] or '').strip()
    # Поддерживаем прямой номер приказа в старых или вручную собранных ссылках.
    return value, None

def get_student_enrollment_order_options():
    with sqlite3.connect(DB_PATH) as conn:
        rows = conn.execute(
            '''
            SELECT TRIM(COALESCE(order_number, '')) AS order_number,
                   TRIM(COALESCE(order_date, '')) AS order_date,
                   COUNT(DISTINCT username) AS student_count
            FROM student_group_transfers
            WHERE movement_type='enrollment'
              AND (TRIM(COALESCE(order_number, ''))<>'' OR TRIM(COALESCE(order_date, ''))<>'')
            GROUP BY TRIM(COALESCE(order_number, '')), TRIM(COALESCE(order_date, ''))
            ORDER BY
                CASE WHEN TRIM(COALESCE(order_date, ''))='' THEN 1 ELSE 0 END,
                TRIM(COALESCE(order_date, '')) DESC,
                TRIM(COALESCE(order_number, '')) DESC
            '''
        ).fetchall()

    options = []
    for order_number, order_date, student_count in rows:
        number_text = f'№ {order_number}' if order_number else 'Без номера'
        date_text = f' от {format_display_date(order_date)}' if order_date else ''
        students_text = f'{student_count} студ.'
        options.append({
            'value': make_enrollment_order_filter_value(order_number, order_date),
            'label': f'{number_text}{date_text} — {students_text}',
        })
    return options

def enrollment_order_sort_date(value):
    text = str(value or '').strip()
    if not text:
        return (1, '')
    for date_format in ('%Y-%m-%d', '%d.%m.%Y', '%Y-%m-%d %H:%M:%S', '%d.%m.%Y %H:%M:%S'):
        try:
            return (0, datetime.strptime(text, date_format).strftime('%Y%m%d%H%M%S'))
        except ValueError:
            pass
    return (0, text.casefold())

def natural_text_sort_key(value):
    return tuple(
        int(part) if part.isdigit() else part.casefold()
        for part in re.split(r'(\d+)', str(value or '').strip())
    )

def student_enrollment_order_sort_key(student):
    order_number = str(student.get('enrollment_order_number') or '').strip()
    order_date = str(student.get('enrollment_order_date') or '').strip()
    return (
        enrollment_order_sort_date(order_date),
        natural_text_sort_key(order_number),
        natural_text_sort_key(student.get('username')),
    )

def get_all_students(
    order_by='username', order_dir='asc', cohort=None, lastname=None,
    firstname=None, username=None, enrollment_order=None
):
    valid_columns = {
        'username', 'lastname', 'firstname', 'cohort1', 'cohort2', 'email',
        'enrollment_order'
    }
    if order_by not in valid_columns:
        order_by = 'username'
    if order_dir.lower() not in {'asc', 'desc'}:
        order_dir = 'asc'
    query = '''
        SELECT s.username, s.password, s.email, s.firstname, s.lastname, s.cohort1, s.cohort2,
               TRIM(COALESCE(enrollment.order_number, '')) AS enrollment_order_number,
               TRIM(COALESCE(enrollment.order_date, '')) AS enrollment_order_date
        FROM students s
        LEFT JOIN student_group_transfers enrollment
          ON enrollment.id=(
              SELECT MIN(enrollment_row.id)
              FROM student_group_transfers enrollment_row
              WHERE enrollment_row.username=s.username
                AND enrollment_row.movement_type='enrollment'
          )
        WHERE 1=1
    '''
    params = []
    if cohort:
        query += " AND s.cohort1 = ?"
        params.append(cohort)
    parsed_order_filter = parse_enrollment_order_filter_value(enrollment_order)
    if parsed_order_filter:
        order_number, order_date = parsed_order_filter
        query += " AND TRIM(COALESCE(enrollment.order_number, '')) = ?"
        params.append(order_number)
        if order_date is not None:
            query += " AND TRIM(COALESCE(enrollment.order_date, '')) = ?"
            params.append(order_date)
    if order_by != 'enrollment_order':
        query += f" ORDER BY s.{order_by} {order_dir.upper()}, s.username ASC"
    with sqlite3.connect(DB_PATH) as conn:
        ensure_student_list_enrollment_movements(conn)
        cur = conn.execute(query, params)
        rows = cur.fetchall()
        columns = [desc[0] for desc in cur.description]
        students = [dict(zip(columns, row)) for row in rows]
    students = [
        row for row in students
        if student_field_matches(row, 'lastname', lastname)
        and student_field_matches(row, 'firstname', firstname)
        and student_field_matches(row, 'username', username)
    ]
    if order_by == 'enrollment_order':
        students.sort(
            key=student_enrollment_order_sort_key,
            reverse=order_dir.lower() == 'desc'
        )
        # Студенты без приказа всегда идут после записей с приказом.
        students.sort(
            key=lambda row: not bool(
                str(row.get('enrollment_order_number') or '').strip()
                or str(row.get('enrollment_order_date') or '').strip()
            )
        )
    return students

def get_pending_duplicates(campaign_year=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            'SELECT id, fio, dogovor, login, fam, imotch, campaign_year FROM pending_duplicates WHERE campaign_year=?',
            (campaign_year,)
        )
        return cur.fetchall()

def approve_duplicate(dup_id, campaign_year=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            'SELECT fio, dogovor, login, fam, imotch, campaign_year FROM pending_duplicates WHERE id=? AND campaign_year=?',
            (dup_id, campaign_year)
        )
        row = cur.fetchone()
        if row:
            fio, dogovor, login, fam, imotch, row_campaign_year = row
            conn.execute(
                'INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch) VALUES (?, ?, ?, ?, ?, ?)',
                (fio, dogovor, login, row_campaign_year, fam, imotch)
            )
            conn.execute('DELETE FROM pending_duplicates WHERE id=? AND campaign_year=?', (dup_id, campaign_year))

def reject_duplicate(dup_id, campaign_year=None):
    campaign_year = normalize_campaign_year(campaign_year, get_active_campaign_year())
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute('DELETE FROM pending_duplicates WHERE id=? AND campaign_year=?', (dup_id, campaign_year))

@app.route('/duplicates', methods=['GET', 'POST'])
@login_required
def duplicates():
    return render_template('duplicates.html')

def role_required(*roles):
    allowed_roles = set(roles)
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'user' not in session:
                return redirect(url_for('login'))
            with sqlite3.connect(DB_PATH) as conn:
                cur = conn.execute('SELECT role, approved FROM users WHERE username=?', (session['user'],))
                user = cur.fetchone()
                role = user[0] if user else ''
                role_allowed = role in allowed_roles or role == 'admin'
                if 'assistant' in allowed_roles and role in {'manager', 'operator'}:
                    role_allowed = True
                if not user or not role_allowed or user[1] != 1:
                    flash('Недостаточно прав')
                    return redirect(url_for('index'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator

@app.route('/abiturients/bulk', methods=['POST'])
@login_required
@role_required('admin', 'assistant', 'operator')
def bulk_abiturients():
    campaign_year = get_active_campaign_year()
    if not ensure_campaign_open(campaign_year):
        return redirect(url_for('abiturients'), code=303)
    action = request.form.get('bulk_action', '').strip()
    if action not in {'mark_paid', 'mark_unpaid', 'delete', 'export'}:
        flash('Неизвестное массовое действие.', 'error')
        return redirect(url_for('abiturients'), code=303)
    if action == 'delete' and session.get('role') != 'admin':
        flash('Удаление доступно только администратору.', 'error')
        return redirect(url_for('abiturients'), code=303)
    selected_ids = [item for item in request.form.getlist('abiturient_ids') if str(item).isdigit()]
    if not selected_ids:
        flash('Выберите хотя бы одну запись.', 'error')
        return redirect(url_for('abiturients'), code=303)

    placeholders = ','.join('?' for _ in selected_ids)
    params = selected_ids + [campaign_year]
    if action == 'export':
        with sqlite3.connect(DB_PATH) as conn:
            cur = conn.execute(
                f'SELECT * FROM abiturients WHERE id IN ({placeholders}) AND campaign_year=? ORDER BY fio',
                params
            )
            rows = cur.fetchall()
            columns = [desc[0] for desc in cur.description]
        df = pd.DataFrame([dict(zip(columns, row)) for row in rows])
        output = io.BytesIO()
        df.to_excel(output, index=False)
        output.seek(0)
        log_action('abiturients_bulk_exported', 'campaign', campaign_year, f"rows={len(rows)}")
        return send_file(output, as_attachment=True, download_name='selected_abiturients.xlsx', mimetype=EXCEL_MIMETYPE)

    backup_path = create_database_backup('before_bulk_abiturients')
    with sqlite3.connect(DB_PATH) as conn:
        if action == 'mark_paid':
            conn.execute(f'UPDATE abiturients SET paid=1 WHERE id IN ({placeholders}) AND campaign_year=?', params)
            sync_enrollment_candidates_for_abiturients(conn, selected_ids, campaign_year)
            flash(f'Отмечено оплаченных: {len(selected_ids)}', 'success')
        elif action == 'mark_unpaid':
            conn.execute(f'UPDATE abiturients SET paid=0 WHERE id IN ({placeholders}) AND campaign_year=?', params)
            sync_enrollment_candidates_for_abiturients(conn, selected_ids, campaign_year)
            flash(f'Снята отметка оплаты: {len(selected_ids)}', 'success')
        elif action == 'delete':
            conn.execute(f'DELETE FROM abiturients WHERE id IN ({placeholders}) AND campaign_year=?', params)
            conn.execute(
                f'DELETE FROM enrollment_candidates WHERE abiturient_id IN ({placeholders}) AND campaign_year=?',
                params
            )
            flash(f'Удалено записей: {len(selected_ids)}', 'success')
        log_action(
            'abiturients_bulk_action',
            'campaign',
            campaign_year,
            f"action={action}; rows={len(selected_ids)}; backup={os.path.basename(backup_path) if backup_path else ''}",
            conn
        )
    return redirect(url_for('abiturients'), code=303)

@app.route('/duplicates_abiturients', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def duplicates_abiturients():
    campaign_year = get_active_campaign_year()
    if request.method == 'POST':
        if not ensure_campaign_open(campaign_year):
            return redirect(url_for('duplicates_abiturients'), code=303)
        action = request.form.get('action')
        if action == 'reject_all':
            with sqlite3.connect(DB_PATH) as conn:
                conn.execute('DELETE FROM pending_duplicates WHERE campaign_year=?', (campaign_year,))
        else:
            dup_id = request.form.get('dup_id')
            if action == 'approve':
                approve_duplicate(dup_id, campaign_year)
            elif action == 'reject':
                reject_duplicate(dup_id, campaign_year)
    duplicates = get_pending_duplicates(campaign_year)
    return render_template('duplicates_abiturients.html', duplicates=duplicates, campaign_year=campaign_year)

@app.route('/delete_abiturient', methods=['POST'])
@login_required
@role_required('admin')
def delete_abiturient():
    campaign_year = get_active_campaign_year()
    if not ensure_campaign_open(campaign_year):
        return redirect(url_for('abiturients'), code=303)
    abiturient_id = request.form.get('id')
    login = request.form.get('login')
    if abiturient_id:
        backup_path = create_database_backup('before_delete_abiturient')
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute('DELETE FROM abiturients WHERE id=? AND campaign_year=?', (abiturient_id, campaign_year))
            conn.execute('DELETE FROM enrollment_candidates WHERE abiturient_id=? AND campaign_year=?', (abiturient_id, campaign_year))
            log_action(
                'abiturient_deleted',
                'abiturient',
                abiturient_id,
                f"campaign_year={campaign_year}; backup={os.path.basename(backup_path) if backup_path else ''}",
                conn
            )
    elif login:
        backup_path = create_database_backup('before_delete_abiturient')
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute('DELETE FROM abiturients WHERE login=? AND campaign_year=?', (login, campaign_year))
            conn.execute('DELETE FROM enrollment_candidates WHERE login=? AND campaign_year=?', (login, campaign_year))
            log_action(
                'abiturient_deleted',
                'abiturient',
                login,
                f"campaign_year={campaign_year}; backup={os.path.basename(backup_path) if backup_path else ''}",
                conn
            )
    return redirect(url_for('abiturients'))

@app.route('/abiturients/withdraw-documents', methods=['POST'])
@login_required
@role_required('admin')
def withdraw_abiturient_documents():
    campaign_year = get_active_campaign_year()
    list_query = get_abiturient_list_query_params(request.form)
    list_url = url_for('abiturients', campaign_year=campaign_year, **list_query)
    if not ensure_campaign_open(campaign_year):
        return redirect(list_url, code=303)

    abiturient_id = request.form.get('id', '').strip()
    if not abiturient_id.isdigit():
        flash('Не удалось определить абитуриента.', 'error')
        return redirect(list_url, code=303)

    with sqlite3.connect(DB_PATH) as conn:
        row = conn.execute(
            'SELECT fio, login FROM abiturients WHERE id=? AND campaign_year=?',
            (abiturient_id, campaign_year)
        ).fetchone()
    if not row:
        flash('Абитуриент не найден.', 'error')
        return redirect(list_url, code=303)

    fio, old_login = row
    if is_withdrawn_login(old_login):
        flash(f'Документы у абитуриента {fio or old_login} уже отозваны.', 'info')
        return redirect(list_url, code=303)

    login_base = str(old_login or '').strip() or f'abiturient{abiturient_id}'
    new_login = next_withdrawn_login(login_base, campaign_year)
    backup_path = create_database_backup('before_withdraw_abiturient_documents')
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            'UPDATE abiturients SET login=? WHERE id=? AND campaign_year=?',
            (new_login, abiturient_id, campaign_year)
        )
        conn.execute(
            'DELETE FROM enrollment_candidates WHERE campaign_year=? AND (abiturient_id=? OR login=?)',
            (campaign_year, abiturient_id, old_login)
        )
        refresh_enrollment_candidate_statuses(conn, campaign_year)
        log_action(
            'abiturient_documents_withdrawn',
            'abiturient',
            abiturient_id,
            (
                f'old_login={old_login}; new_login={new_login}; campaign_year={campaign_year}; '
                f'backup={os.path.basename(backup_path) if backup_path else ""}'
            ),
            conn
        )

    flash(
        f'Документы у абитуриента {fio or old_login} отозваны. Логин изменён: {old_login or "без логина"} → {new_login}.',
        'success'
    )
    return redirect(list_url, code=303)

@app.route('/toggle_abiturient_paid', methods=['POST'])
@login_required
@role_required('admin')
def toggle_abiturient_paid():
    campaign_year = get_active_campaign_year()
    if not ensure_campaign_open(campaign_year):
        return redirect(url_for('abiturients'), code=303)
    abiturient_id = request.form.get('id')
    paid = 1 if request.form.get('paid') == '1' else 0
    query_params = {
        'spec': request.form.get('spec', ''),
        'base': request.form.get('base', ''),
        'year': request.form.get('year', ''),
        'is_i': request.form.get('is_i', ''),
        'has_email': request.form.get('has_email', ''),
        'has_paid': request.form.get('has_paid', ''),
        'has_order': request.form.get('has_order', ''),
        'withdrawn': request.form.get('withdrawn', ''),
        'q': request.form.get('q', ''),
        'order_by': request.form.get('order_by', 'created_at'),
        'order_dir': request.form.get('order_dir', 'desc')
    }
    if abiturient_id:
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute('UPDATE abiturients SET paid=? WHERE id=? AND campaign_year=?', (paid, abiturient_id, campaign_year))
            sync_enrollment_candidates_for_abiturients(conn, [abiturient_id], campaign_year)
            log_action('abiturient_paid_changed', 'abiturient', abiturient_id, f"paid={paid}; campaign_year={campaign_year}", conn)
    return redirect(url_for('abiturients', **{k: v for k, v in query_params.items() if v}))

@app.route('/abiturients/download')
@login_required
def download_abiturients():
    campaign_year = get_active_campaign_year()
    order_by = request.args.get('order_by', 'created_at')
    order_dir = request.args.get('order_dir', 'desc')
    spec = request.args.get('spec')
    base = request.args.get('base')
    year = request.args.get('year')
    is_i = request.args.get('is_i')
    has_email = request.args.get('has_email')
    has_paid = request.args.get('has_paid')
    has_order = request.args.get('has_order')
    withdrawn = request.args.get('withdrawn')
    q = request.args.get('q', '').strip()
    abiturients = get_all_abiturients(
        order_by, order_dir, spec, base, year, is_i, campaign_year,
        has_email, has_paid, q, withdrawn, has_order
    )
    log_action('abiturients_exported', 'campaign', campaign_year, f"rows={len(abiturients)}")
    df = pd.DataFrame(abiturients)
    output = io.BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="abiturients_logins.xlsx", mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

@app.route('/download_template')
def download_template():
    template_path = os.path.join(app.static_folder, 'template.xlsx')
    return send_file(template_path, as_attachment=True, download_name='template.xlsx')

@app.route('/login_conflicts')
@login_required
def login_conflicts():
    campaign_year = get_active_campaign_year()
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            '''
            SELECT id, fio, dogovor, login, fam, imotch, campaign_year, conflict_time
            FROM login_conflicts
            WHERE campaign_year=?
            ORDER BY conflict_time DESC
            ''',
            (campaign_year,)
        )
        conflicts = cur.fetchall()
    return render_template('login_conflicts.html', conflicts=conflicts, campaign_year=campaign_year)

@app.route('/edit_conflict/<int:conflict_id>', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def edit_conflict(conflict_id):
    campaign_year = get_active_campaign_year()
    if request.method == 'POST':
        if not ensure_campaign_open(campaign_year):
            return redirect(url_for('login_conflicts'), code=303)
        new_login = request.form.get('login', '').strip()
        if not new_login:
            flash('Логин не может быть пустым')
            return redirect(url_for('edit_conflict', conflict_id=conflict_id))
        
        backup_path = create_database_backup('before_resolve_login_conflict')
        with sqlite3.connect(DB_PATH) as conn:
            # Проверяем уникальность логина
            if is_login_exists(new_login, campaign_year):
                flash(f'Логин {new_login} уже используется!')
                return redirect(url_for('edit_conflict', conflict_id=conflict_id))
            
            # Получаем данные конфликта
            cur = conn.execute(
                'SELECT fio, dogovor, fam, imotch, campaign_year FROM login_conflicts WHERE id=? AND campaign_year=?',
                (conflict_id, campaign_year)
            )
            conflict = cur.fetchone()
            if not conflict:
                flash('Запись не найдена')
                return redirect(url_for('login_conflicts'))
            
            fio, dogovor, fam, imotch, row_campaign_year = conflict
            
            # Сохраняем в основную таблицу абитуриентов
            try:
                conn.execute(
                    'INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch) VALUES (?, ?, ?, ?, ?, ?)',
                    (fio, dogovor, new_login, row_campaign_year, fam, imotch)
                )
                # Удаляем из конфликтов
                conn.execute('DELETE FROM login_conflicts WHERE id=? AND campaign_year=?', (conflict_id, campaign_year))
                log_action(
                    'login_conflict_resolved',
                    'login_conflict',
                    conflict_id,
                    (
                        f"new_login={new_login}; campaign_year={row_campaign_year}; "
                        f"backup={os.path.basename(backup_path) if backup_path else ''}"
                    ),
                    conn
                )
                conn.commit()
                flash(f'Абитуриент успешно добавлен с логином {new_login}')
                return redirect(url_for('login_conflicts'))
            except sqlite3.IntegrityError:
                flash(f'Логин {new_login} уже существует!')
                return redirect(url_for('edit_conflict', conflict_id=conflict_id))
    
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            'SELECT id, fio, dogovor, login, fam, imotch, campaign_year FROM login_conflicts WHERE id=? AND campaign_year=?',
            (conflict_id, campaign_year)
        )
        conflict = cur.fetchone()
    
    if not conflict:
        flash('Запись не найдена')
        return redirect(url_for('login_conflicts'))
    
    return render_template('edit_conflict.html', conflict=conflict)

@app.route('/delete_conflict/<int:conflict_id>', methods=['POST'])
@login_required
@role_required('admin')
def delete_conflict(conflict_id):
    campaign_year = get_active_campaign_year()
    if not ensure_campaign_open(campaign_year):
        return redirect(url_for('login_conflicts'), code=303)
    backup_path = create_database_backup('before_delete_login_conflict')
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute('DELETE FROM login_conflicts WHERE id=? AND campaign_year=?', (conflict_id, campaign_year))
        log_action(
            'login_conflict_deleted',
            'login_conflict',
            conflict_id,
            f"campaign_year={campaign_year}; backup={os.path.basename(backup_path) if backup_path else ''}",
            conn
        )
        conn.commit()
    flash('Запись удалена')
    return redirect(url_for('login_conflicts'))

@login_required
@role_required('admin')
def delete_database():
    if os.path.exists(DB_PATH):
        os.remove(DB_PATH)
        init_db()
    return redirect(url_for('index'))

def build_manual_login_parts(year, specialty, base, rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    year = require_campaign_year(year)
    spec_code = (rules.get('spec_codes') or {}).get(specialty)
    base_code = (rules.get('base_codes') or {}).get(base)
    if spec_code is None or base_code is None:
        return None
    return {
        'year': year,
        'year_code': year[-2:],
        'spec_label': specialty,
        'spec_code': str(spec_code),
        'base_label': base,
        'base_code': str(base_code),
    }

def get_manual_create_collisions(fio, dogovor, login, campaign_year, rules=None):
    rules = merge_login_generation_rules(rules or get_login_generation_rules())
    fio_key = normalize_fio_key(fio)
    dogovor_key = normalize_dogovor_key(dogovor)
    login_key = str(login or '').strip().casefold()
    result = {
        'fio': [],
        'dogovor': [],
        'login': [],
    }

    def add_record(source, record_fio, record_dogovor, record_login, record_campaign):
        record = {
            'source': source,
            'fio': str(record_fio or '').strip(),
            'dogovor': str(record_dogovor or '').strip(),
            'login': str(record_login or '').strip(),
            'campaign_year': str(record_campaign or '').strip(),
        }
        if fio_key and normalize_fio_key(record['fio']) == fio_key and record['campaign_year'] == campaign_year:
            result['fio'].append(record)
        if dogovor_key and normalize_dogovor_key(record['dogovor']) == dogovor_key and record['campaign_year'] == campaign_year:
            result['dogovor'].append(record)
        login_in_scope = rules['unique_scope'] == 'global' or record['campaign_year'] == campaign_year
        if login_key and record['login'].casefold() == login_key and login_in_scope:
            result['login'].append(record)

    with sqlite3.connect(DB_PATH) as conn:
        for table, label in (
            ('abiturients', 'Абитуриенты'),
            ('pending_duplicates', 'Дублирующиеся ФИО'),
            ('login_conflicts', 'Конфликты логинов'),
        ):
            rows = conn.execute(
                f'SELECT fio, dogovor, login, campaign_year FROM {table}'
            ).fetchall()
            for row in rows:
                add_record(label, *row)
        if table_exists(conn, 'students'):
            rows = conn.execute(
                '''
                SELECT source_fio, lastname, firstname, source_dogovor,
                       username, source_campaign_year
                FROM students
                '''
            ).fetchall()
            for source_fio, lastname, firstname, source_dogovor, username, source_year in rows:
                student_fio = str(source_fio or '').strip() or ' '.join(
                    part for part in (lastname, firstname) if str(part or '').strip()
                ).strip()
                add_record('Студенты', student_fio, source_dogovor, username, source_year)
    return result

@app.route('/manual_create', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def manual_create():
    message = None
    message_tone = 'info'
    collision_records = []
    fio_confirmation_required = False
    created_record = None
    campaign_year = get_active_campaign_year()
    years = get_campaign_years()
    login_rules = get_login_generation_rules()
    specs = list(login_rules['spec_codes'].keys())
    bases = list(login_rules['base_codes'].keys())
    form_data = {
        'creation_mode': 'contract',
        'fio': '',
        'dogovor': '',
        'year': campaign_year,
        'spec': specs[0] if specs else '',
        'base': bases[0] if bases else '',
        'manual_login': '',
    }
    if request.method == 'POST':
        form_data.update({
            'creation_mode': str(request.form.get('creation_mode') or 'contract').strip(),
            'fio': ' '.join(str(request.form.get('fio') or '').split()),
            'dogovor': normalize_dogovor_storage_text(request.form.get('dogovor') or ''),
            'year': str(request.form.get('year') or campaign_year).strip(),
            'spec': str(request.form.get('spec') or '').strip(),
            'base': str(request.form.get('base') or '').strip(),
            'manual_login': str(request.form.get('manual_login') or '').strip(),
        })
        mode = form_data['creation_mode'] if form_data['creation_mode'] in {'contract', 'manual'} else 'contract'
        form_data['creation_mode'] = mode
        fio = form_data['fio']
        dogovor = form_data['dogovor']
        login_parts = None
        validation_errors = []

        if not fio:
            validation_errors.append('Введите ФИО абитуриента.')
        if mode == 'contract' and not dogovor:
            validation_errors.append('Введите номер договора.')

        if mode == 'contract' and dogovor:
            login_parts = parse_dogovor_parts(dogovor, login_rules)
            if not login_parts:
                validation_errors.append(
                    'Не удалось определить год, специальность или базу образования из договора.'
                )
            else:
                campaign_year = login_parts['year']
                form_data.update({
                    'year': campaign_year,
                    'spec': login_parts['spec_label'],
                    'base': login_parts['base_label'],
                })
        elif mode == 'manual':
            try:
                campaign_year = require_campaign_year(form_data['year'])
                login_parts = build_manual_login_parts(
                    campaign_year,
                    form_data['spec'],
                    form_data['base'],
                    login_rules,
                )
                if not login_parts:
                    validation_errors.append('Выберите специальность и базу образования.')
            except ValueError as exc:
                validation_errors.append(str(exc))

        if login_parts and not validation_errors and is_campaign_archived(campaign_year):
            validation_errors.append(f'Кампания {campaign_year} находится в архиве.')

        login = ''
        existing_logins = get_used_logins(campaign_year, login_rules) if login_parts else set()
        if login_parts:
            if mode == 'manual' and form_data['manual_login']:
                login = form_data['manual_login']
                if re.search(r'\s', login):
                    validation_errors.append('Логин не должен содержать пробелы.')
            else:
                login = next_login_from_parts(login_parts, existing_logins, login_rules)

        collisions = (
            get_manual_create_collisions(fio, dogovor, login, campaign_year, login_rules)
            if fio and login_parts and login
            else {'fio': [], 'dogovor': [], 'login': []}
        )
        if collisions['dogovor']:
            validation_errors.append('Такой договор уже есть в базе.')
            collision_records.extend(collisions['dogovor'])
        if collisions['login']:
            validation_errors.append('Такой логин уже используется.')
            collision_records.extend(collisions['login'])

        exact_fio_matches = collisions['fio']
        confirmed_fio = request.form.get('confirm_fio_duplicate') == '1'
        if exact_fio_matches and not validation_errors and not confirmed_fio:
            fio_confirmation_required = True
            collision_records.extend(exact_fio_matches)
            message = 'В базе уже есть запись с точно таким ФИО. Проверьте её перед созданием.'
            message_tone = 'warning'
        elif validation_errors:
            message = ' '.join(validation_errors)
            message_tone = 'error'
        elif login_parts:
            fam, imotch = split_fio_for_storage(fio)
            try:
                save_abiturient(fio, dogovor, login, fam, imotch, campaign_year)
                session['campaign_year'] = campaign_year
                message = f'Логин успешно создан: {login}'
                message_tone = 'success'
                created_record = {
                    'fio': fio,
                    'dogovor': dogovor,
                    'login': login,
                    'campaign_year': campaign_year,
                    'specialty': login_parts['spec_label'],
                    'base': login_parts['base_label'],
                }
                log_action(
                    'manual_abiturient_create',
                    'campaign',
                    campaign_year,
                    f'mode={mode}; fio={fio}; dogovor={dogovor}; login={login}',
                )
                form_data.update({'fio': '', 'dogovor': '', 'manual_login': ''})
            except sqlite3.IntegrityError:
                message = 'Логин уже успел занять другой пользователь. Обновите форму и повторите.'
                message_tone = 'error'

    if campaign_year not in years:
        years = sorted(set(years + [campaign_year]), reverse=True)

    return render_template(
        'manual_create.html',
        message=message,
        message_tone=message_tone,
        collision_records=collision_records,
        fio_confirmation_required=fio_confirmation_required,
        created_record=created_record,
        form_data=form_data,
        years=years,
        specs=specs,
        bases=bases,
        campaign_year=campaign_year,
    )

@app.route('/login', methods=['GET', 'POST'])
def login():
    session.pop('user', None)
    session.pop('role', None)
    username = ''
    if request.method == 'POST':
        username = (request.form.get('username') or '').strip()
        password = request.form.get('password') or ''
        if not validate_login_csrf_token():
            flash('Сессия формы устарела. Попробуйте ещё раз.', 'error')
            return render_template('login.html', login_csrf_token=refresh_login_csrf_token(), username=username)
        lockout_seconds = get_login_lockout(username)
        if lockout_seconds:
            lockout_minutes = max(1, (lockout_seconds + 59) // 60)
            flash(f'Слишком много попыток входа. Попробуйте через {lockout_minutes} мин.', 'error')
            return render_template('login.html', login_csrf_token=get_login_csrf_token(), username=username)
        with sqlite3.connect(DB_PATH) as conn:
            cur = conn.execute(
                'SELECT id, password, role, approved FROM users WHERE username=?',
                (username,)
            )
            user = cur.fetchone()
            password_ok = bool(user and verify_user_password(user[1], password))
            if password_ok and not is_password_hash(user[1]):
                conn.execute('UPDATE users SET password=? WHERE id=?', (hash_user_password(password), user[0]))
            if password_ok and user[3] == 1:
                clear_login_failures(username)
                session.clear()
                session['user'] = username
                session['role'] = user[2]
                return redirect(url_for('index'))
            elif password_ok and user[3] == 0:
                flash('Ожидайте одобрения администратора.', 'error')
            else:
                record_login_failure(username)
                flash('Неверный логин или пароль.', 'error')
    return render_template('login.html', login_csrf_token=get_login_csrf_token(), username=username)

@app.route('/logout')
def logout():
    session.clear()
    flash('Вы вышли из системы.', 'success')
    return redirect(url_for('login'))

@app.route('/edit_abiturient/<login>', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def edit_abiturient(login):
    campaign_year = get_active_campaign_year()
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            'SELECT fio, dogovor, login, fam, imotch, email, comment, campaign_year, paid, id FROM abiturients WHERE login=? AND campaign_year=?',
            (login, campaign_year)
        )
        abiturient = cur.fetchone()
    if not abiturient:
        flash('Абитуриент не найден')
        return redirect(url_for('abiturients'))

    list_query = get_abiturient_list_query_params(request.values)
    navigation = get_abiturient_edit_navigation(abiturient[9], campaign_year, list_query)
    edit_context = {
        'abiturient': abiturient,
        'abiturient_navigation': navigation,
        'abiturient_list_query': list_query,
    }

    if request.method == 'POST':
        if not ensure_campaign_open(campaign_year):
            return redirect(url_for('abiturients', campaign_year=campaign_year, **list_query), code=303)
        fio, fam, imotch = split_fio(request.form.get('fio', ''))
        if not fio:
            flash('ФИО не может быть пустым')
            return render_template('edit_abiturient.html', **edit_context)
        email = request.form.get('email', '').strip()
        paid = 1 if request.form.get('paid') == '1' else 0
        new_login = request.form.get('login', '').strip()
        comment = request.form.get('comment', '').strip()
        backup_path = create_database_backup('before_edit_abiturient')
        if new_login != login:
            if is_login_exists(new_login, campaign_year):
                flash('Такой логин уже существует!')
                return render_template('edit_abiturient.html', **edit_context)
            with sqlite3.connect(DB_PATH) as conn:
                conn.execute(
                    'UPDATE abiturients SET fio=?, fam=?, imotch=?, email=?, paid=?, login=?, comment=? WHERE login=? AND campaign_year=?',
                    (fio, fam, imotch, email, paid, new_login, comment, login, campaign_year)
                )
                sync_enrollment_candidates_for_abiturients(conn, [abiturient[9]], campaign_year)
                log_action(
                    'abiturient_updated',
                    'abiturient',
                    login,
                    f"new_login={new_login}; campaign_year={campaign_year}; backup={os.path.basename(backup_path) if backup_path else ''}",
                    conn
                )
        else:
            with sqlite3.connect(DB_PATH) as conn:
                conn.execute(
                    'UPDATE abiturients SET fio=?, fam=?, imotch=?, email=?, paid=?, comment=? WHERE login=? AND campaign_year=?',
                    (fio, fam, imotch, email, paid, comment, login, campaign_year)
                )
                sync_enrollment_candidates_for_abiturients(conn, [abiturient[9]], campaign_year)
                log_action(
                    'abiturient_updated',
                    'abiturient',
                    login,
                    f"campaign_year={campaign_year}; backup={os.path.basename(backup_path) if backup_path else ''}",
                    conn
                )
        flash('Данные обновлены')
        if request.form.get('save_action') == 'stay':
            return redirect(
                url_for(
                    'edit_abiturient',
                    login=new_login,
                    campaign_year=campaign_year,
                    **list_query
                ),
                code=303
            )
        return redirect(url_for('abiturients', campaign_year=campaign_year, **list_query))

    return render_template('edit_abiturient.html', **edit_context)

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = (request.form.get('username') or '').strip()
        password = (request.form.get('password') or '').strip()
        role = 'assistant'
        if not username or len(password) < MIN_PASSWORD_LENGTH:
            flash(f'Логин обязателен, пароль должен быть не короче {MIN_PASSWORD_LENGTH} символов')
            return render_template('register.html')
        with sqlite3.connect(DB_PATH) as conn:
            cur = conn.execute('SELECT 1 FROM users WHERE username=?', (username,))
            if cur.fetchone():
                flash('Пользователь уже существует')
                return render_template('register.html')
            conn.execute('INSERT INTO users (username, password, role) VALUES (?, ?, ?)', (username, hash_user_password(password), role))
        flash('Заявка на регистрацию отправлена. Ожидайте одобрения администратора.')
        return redirect(url_for('login'))
    return render_template('register.html')

@app.route('/setup', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def login_generation_setup():
    if request.method == 'POST':
        mode = request.form.get('mode', 'custom')
        try:
            rules = build_login_rules_from_form(request.form, mode)
            save_login_generation_settings(rules, setup_completed=True, updated_by=session.get('user', ''))
            log_action('login_generation_settings_saved', 'settings', 'login_generation', f"mode={rules['mode']}")
            flash('Правила формирования логинов сохранены.', 'success')
            return redirect(url_for('index'))
        except (ValueError, re.error) as exc:
            flash(f'Не удалось сохранить правила: {exc}', 'error')
    return render_template('login_rules_setup.html', **get_login_rules_form_context())

@app.route('/admin_panel')
@admin_required
def admin_panel():
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute('SELECT id, username, role, approved FROM users')
        all_users = cur.fetchall()
    return render_template(
        'admin_panel.html',
        all_users=all_users,
        app_update_enabled=APP_UPDATE_ENABLED,
        app_update_status=get_app_update_status(),
    )


def get_app_update_status_path():
    configured_path = os.environ.get('APP_UPDATE_STATUS_FILE')
    return update_app.resolve_update_status_path(APP_DIR, configured_path)


def get_app_update_status():
    status = update_app.read_update_status(get_app_update_status_path())
    if not status:
        return {'state': 'idle', 'message': 'Обновления ещё не запускались.'}
    allowed_states = {'idle', 'queued', 'running', 'completed', 'up_to_date', 'failed'}
    if status.get('state') not in allowed_states:
        status['state'] = 'failed'
    status['message'] = str(status.get('message') or '')[:1000]
    return status


def app_update_is_active(status):
    if status.get('state') not in {'queued', 'running'}:
        return False
    try:
        updated_at = datetime.fromisoformat(str(status.get('updated_at') or '').replace('Z', '+00:00'))
        if updated_at.tzinfo is None:
            updated_at = updated_at.replace(tzinfo=timezone.utc)
        age = (datetime.now(timezone.utc) - updated_at).total_seconds()
        return age < APP_UPDATE_STALE_SECONDS
    except (TypeError, ValueError):
        return True


def build_app_update_command():
    configured_command = (os.environ.get('APP_UPDATE_COMMAND') or '').strip()
    if configured_command:
        try:
            parsed = json.loads(configured_command)
        except json.JSONDecodeError:
            parsed = shlex.split(configured_command, posix=os.name != 'nt')
        if isinstance(parsed, str):
            parsed = shlex.split(parsed, posix=os.name != 'nt')
        if not isinstance(parsed, list) or not parsed or not all(isinstance(part, str) and part for part in parsed):
            raise RuntimeError('APP_UPDATE_COMMAND должен содержать команду или JSON-массив аргументов.')
        return parsed

    if os.name == 'nt':
        raise RuntimeError('Обновление из админки предназначено для Linux-сервера с systemd.')
    sudo_path = shutil.which('sudo') or next(
        (path for path in ('/usr/bin/sudo', '/bin/sudo') if os.path.isfile(path)),
        None,
    )
    systemctl_path = shutil.which('systemctl') or next(
        (path for path in ('/usr/bin/systemctl', '/bin/systemctl') if os.path.isfile(path)),
        None,
    )
    if not sudo_path or not systemctl_path:
        raise RuntimeError('На сервере не найдены sudo и systemctl.')
    service_name = (os.environ.get('APP_UPDATE_SERVICE') or 'manticore-update.service').strip()
    if not re.fullmatch(r'[A-Za-z0-9@_.-]+', service_name):
        raise RuntimeError('В APP_UPDATE_SERVICE указано недопустимое имя службы.')
    return [sudo_path, '-n', systemctl_path, 'start', '--no-block', service_name]


def get_app_update_configuration_error():
    if not APP_UPDATE_ENABLED:
        return 'Обновление через панель управления отключено в настройках сервера.'
    if not (APP_DIR / '.git').exists():
        return 'Папка приложения не является Git-копией. Автоматическое обновление недоступно.'
    if not shutil.which('git') and not os.path.isfile('/usr/bin/git') and not os.path.isfile('/bin/git'):
        return 'На сервере не найден Git. Автоматическое обновление недоступно.'
    try:
        build_app_update_command()
    except RuntimeError as exc:
        return str(exc)
    return ''


def queue_app_update(requested_by):
    configuration_error = get_app_update_configuration_error()
    if configuration_error:
        raise RuntimeError(configuration_error)
    status_path = get_app_update_status_path()
    update_app.write_update_status(
        status_path,
        'queued',
        'Обновление поставлено в очередь.',
        current_version=APP_VERSION,
        requested_by=requested_by,
        started_at=datetime.now(timezone.utc).isoformat(),
    )
    try:
        completed = subprocess.run(
            build_app_update_command(),
            cwd=str(APP_DIR),
            text=True,
            capture_output=True,
            timeout=15,
        )
    except (OSError, subprocess.SubprocessError) as exc:
        update_app.write_update_status(
            status_path,
            'failed',
            f'Не удалось запустить службу обновления: {exc}',
            current_version=APP_VERSION,
            finished_at=datetime.now(timezone.utc).isoformat(),
        )
        raise RuntimeError(f'Не удалось запустить службу обновления: {exc}') from exc
    if completed.returncode != 0:
        error_text = (completed.stderr or completed.stdout or 'неизвестная ошибка').strip()[:800]
        update_app.write_update_status(
            status_path,
            'failed',
            f'Служба обновления не запущена: {error_text}',
            current_version=APP_VERSION,
            finished_at=datetime.now(timezone.utc).isoformat(),
        )
        raise RuntimeError(f'Служба обновления не запущена: {error_text}')


@app.route('/admin/app-update/status')
@admin_required
def app_update_status_api():
    status = get_app_update_status()
    configuration_error = get_app_update_configuration_error()
    result = {
        'enabled': not bool(configuration_error),
        'configuration_error': configuration_error,
        'current_version': update_app.normalize_version(APP_VERSION),
        'status': status,
    }
    if app_update_is_active(status):
        result.update({
            'latest_version': str(status.get('latest_version') or ''),
            'release_url': str(status.get('release_url') or ''),
            'update_available': None,
        })
        return jsonify(result)
    try:
        check_timeout = min(15.0, max(1.0, float(os.environ.get('APP_UPDATE_CHECK_TIMEOUT', '5'))))
        release = update_app.fetch_latest_release(
            APP_DIR,
            timeout=check_timeout,
        )
        latest_version = update_app.normalize_version(release['tag_name'])
        result.update({
            'latest_version': latest_version,
            'release_name': release['name'],
            'release_url': release['html_url'],
            'published_at': release['published_at'],
            'update_available': update_app.is_release_newer(latest_version, APP_VERSION),
        })
    except (ValueError, update_app.UpdateError) as exc:
        result.update({
            'latest_version': '',
            'release_url': '',
            'update_available': None,
            'check_error': str(exc),
        })
    return jsonify(result)


@app.route('/admin/app-update/start', methods=['POST'])
@admin_required
def start_app_update():
    status = get_app_update_status()
    if app_update_is_active(status):
        flash('Обновление уже запущено. Дождитесь завершения.', 'warning')
        return redirect(url_for('admin_panel'), code=303)
    try:
        queue_app_update(session.get('user', ''))
        log_action('app_update_started', 'application', APP_VERSION, 'source=admin_panel')
        flash('Обновление запущено. Сервер на короткое время станет недоступен и перезапустится автоматически.', 'success')
    except Exception as exc:
        log_action('app_update_failed_to_start', 'application', APP_VERSION, str(exc)[:500])
        flash(f'Не удалось запустить обновление: {exc}', 'error')
    return redirect(url_for('admin_panel'), code=303)

@app.route('/backups')
@admin_required
def backups():
    return render_template('backups.html', backups=list_database_backups())

@app.route('/backups/download/<backup_name>')
@admin_required
def download_backup(backup_name):
    backup_path = get_backup_path(backup_name)
    log_action('database_backup_download', 'backup', backup_name)
    return send_file(backup_path, as_attachment=True, download_name=backup_name, mimetype='application/octet-stream')

@app.route('/backups/restore', methods=['POST'])
@admin_required
def restore_backup():
    backup_name = request.form.get('backup_name')
    try:
        backup_path = get_backup_path(backup_name)
        rollback_path = create_database_backup('before_restore')
        shutil.copy2(backup_path, DB_PATH)
        init_db()
        log_action(
            'database_restore',
            'backup',
            backup_name,
            f"rollback_backup={os.path.basename(rollback_path) if rollback_path else ''}"
        )
        flash(f'База восстановлена из резервной копии {backup_name}.', 'success')
    except Exception as exc:
        flash(f'Не удалось восстановить базу: {exc}', 'error')
    return redirect(url_for('backups'))

@app.route('/audit_logs')
@admin_required
def audit_logs():
    return render_template('audit_logs.html', audit_logs=get_audit_logs())

@app.route('/campaigns', methods=['GET', 'POST'])
@admin_required
def campaigns():
    if request.method == 'POST':
        campaign_action = request.form.get('campaign_action', 'toggle_archive')
        try:
            if campaign_action == 'create':
                campaign_year = require_campaign_year(request.form.get('new_campaign_year'))
                if campaign_year in get_campaign_years():
                    raise ValueError(f'Кампания {campaign_year} уже есть в списке.')
                backup_path = create_database_backup('before_campaign_create')
                with sqlite3.connect(DB_PATH) as conn:
                    create_campaign_settings_table(conn)
                    conn.execute(
                        '''
                        INSERT INTO campaign_settings
                            (campaign_year, is_archived, archived_at, archived_by, created_at, created_by)
                        VALUES (?, 0, NULL, NULL, datetime('now', 'localtime'), ?)
                        ''',
                        (campaign_year, session.get('user', ''))
                    )
                    log_action(
                        'campaign_created',
                        'campaign',
                        campaign_year,
                        f"backup={os.path.basename(backup_path) if backup_path else ''}",
                        conn
                    )
                session['campaign_year'] = campaign_year
                session['group_year'] = campaign_year
                flash(f'Кампания {campaign_year} создана и выбрана как текущая.', 'success')
            elif campaign_action == 'set_active':
                campaign_year = require_campaign_year(request.form.get('campaign_year'))
                if campaign_year not in get_campaign_years():
                    raise ValueError(f'Кампания {campaign_year} не найдена.')
                if is_campaign_archived(campaign_year):
                    raise ValueError('Архивную кампанию нельзя закрепить как активную.')
                backup_path = create_database_backup('before_campaign_active_change')
                with sqlite3.connect(DB_PATH) as conn:
                    create_campaign_settings_table(conn)
                    conn.execute('UPDATE campaign_settings SET is_active=0 WHERE is_active=1')
                    conn.execute(
                        '''
                        INSERT INTO campaign_settings
                            (campaign_year, is_archived, created_at, created_by,
                             is_active, active_at, active_by)
                        VALUES (?, 0, datetime('now', 'localtime'), ?,
                                1, datetime('now', 'localtime'), ?)
                        ON CONFLICT(campaign_year) DO UPDATE SET
                            is_active=1,
                            active_at=excluded.active_at,
                            active_by=excluded.active_by
                        ''',
                        (campaign_year, session.get('user', ''), session.get('user', ''))
                    )
                    log_action(
                        'campaign_active_changed',
                        'campaign',
                        campaign_year,
                        f"backup={os.path.basename(backup_path) if backup_path else ''}",
                        conn
                    )
                session['campaign_year'] = campaign_year
                session['group_year'] = campaign_year
                flash(f'Кампания {campaign_year} закреплена как активная по умолчанию.', 'success')
            elif campaign_action == 'clear_active':
                pinned_campaign_year = get_pinned_campaign_year()
                if pinned_campaign_year:
                    backup_path = create_database_backup('before_campaign_active_clear')
                    with sqlite3.connect(DB_PATH) as conn:
                        create_campaign_settings_table(conn)
                        conn.execute(
                            '''
                            UPDATE campaign_settings
                            SET is_active=0, active_at=NULL, active_by=NULL
                            WHERE is_active=1
                            '''
                        )
                        log_action(
                            'campaign_active_cleared',
                            'campaign',
                            pinned_campaign_year,
                            f"backup={os.path.basename(backup_path) if backup_path else ''}",
                            conn
                        )
                automatic_year = get_latest_campaign_year()
                session['campaign_year'] = automatic_year
                session['group_year'] = automatic_year
                flash(f'Закрепление снято. Теперь автоматически выбирается последняя кампания — {automatic_year}.', 'success')
            elif campaign_action == 'toggle_archive':
                campaign_year = require_campaign_year(request.form.get('campaign_year'))
                is_archived = 1 if request.form.get('is_archived') == '1' else 0
                if is_archived and campaign_year == get_pinned_campaign_year():
                    raise ValueError(
                        'Нельзя архивировать закрепленную активную кампанию. '
                        'Сначала выберите другую активную кампанию или снимите закрепление.'
                    )
                backup_path = create_database_backup('before_campaign_archive_toggle')
                with sqlite3.connect(DB_PATH) as conn:
                    create_campaign_settings_table(conn)
                    conn.execute(
                        '''
                        INSERT INTO campaign_settings
                            (campaign_year, is_archived, archived_at, archived_by, created_at, created_by)
                        VALUES (?, ?, datetime('now', 'localtime'), ?, NULL, NULL)
                        ON CONFLICT(campaign_year) DO UPDATE SET
                            is_archived=excluded.is_archived,
                            archived_at=excluded.archived_at,
                            archived_by=excluded.archived_by
                        ''',
                        (campaign_year, is_archived, session.get('user', ''))
                    )
                    log_action(
                        'campaign_archive_changed',
                        'campaign',
                        campaign_year,
                        f"is_archived={is_archived}; backup={os.path.basename(backup_path) if backup_path else ''}",
                        conn
                    )
                flash(f"Кампания {campaign_year}: {'архивирована' if is_archived else 'открыта'}.", 'success')
            else:
                raise ValueError('Неизвестное действие с кампанией.')
        except ValueError as exc:
            flash(str(exc), 'error')
        except sqlite3.IntegrityError:
            flash('Такая кампания уже есть в списке.', 'error')
        return redirect(url_for('campaigns'))
    pinned_campaign_year = get_pinned_campaign_year()
    return render_template(
        'campaigns.html',
        campaigns=get_campaign_settings(),
        pinned_campaign_year=pinned_campaign_year,
        automatic_campaign_year=get_latest_campaign_year(),
        selected_default_campaign_year=pinned_campaign_year or get_latest_campaign_year(),
    )

@app.route('/delete_user', methods=['POST'])
@admin_required
def delete_user():
    user_id = request.form.get('user_id')
    backup_path = create_database_backup('before_delete_user')
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute('DELETE FROM users WHERE id=?', (user_id,))
        log_action(
            'user_deleted',
            'user',
            user_id,
            f"backup={os.path.basename(backup_path) if backup_path else ''}",
            conn
        )
    return redirect(url_for('admin_panel'))

@app.route('/edit_user/<int:user_id>', methods=['GET', 'POST'])
@admin_required
def edit_user(user_id):
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute('SELECT id, username, role, approved FROM users WHERE id=?', (user_id,))
        user = cur.fetchone()
        if not user:
            flash('Пользователь не найден')
            return redirect(url_for('admin_panel'))
        if request.method == 'POST':
            username = (request.form.get('username') or '').strip()
            password = (request.form.get('password') or '').strip()
            role = request.form.get('role')
            approved = int(request.form.get('approved', 0))
            if role not in ROLE_LABELS:
                role = 'viewer'
            if not username:
                flash('Логин не может быть пустым')
                return render_template('edit_user.html', user=user)
            duplicate = conn.execute(
                'SELECT 1 FROM users WHERE username=? AND id<>?',
                (username, user_id)
            ).fetchone()
            if duplicate:
                flash('Пользователь с таким логином уже существует')
                return render_template('edit_user.html', user=user)
            if password:
                if len(password) < MIN_PASSWORD_LENGTH:
                    flash(f'Новый пароль должен быть не короче {MIN_PASSWORD_LENGTH} символов')
                    return render_template('edit_user.html', user=user)
                conn.execute(
                    'UPDATE users SET username=?, password=?, role=?, approved=? WHERE id=?',
                    (username, hash_user_password(password), role, approved, user_id)
                )
            else:
                conn.execute(
                    'UPDATE users SET username=?, role=?, approved=? WHERE id=?',
                    (username, role, approved, user_id)
                )
            log_action('user_updated', 'user', user_id, f"username={username}; role={role}; approved={approved}", conn)
            return redirect(url_for('admin_panel'))
    return render_template('edit_user.html', user=user)

@app.route('/add_user', methods=['GET', 'POST'])
@admin_required
def add_user():
    if request.method == 'POST':
        username = (request.form.get('username') or '').strip()
        password = (request.form.get('password') or '').strip()
        fio = (request.form.get('fio') or '').strip()
        position = (request.form.get('position') or '').strip()
        role = request.form.get('role')
        if role not in ROLE_LABELS:
            role = 'viewer'
        approved = int(request.form.get('approved', 1))
        if not username or len(password) < MIN_PASSWORD_LENGTH:
            flash(f'Логин обязателен, пароль должен быть не короче {MIN_PASSWORD_LENGTH} символов')
            return render_template('add_user.html')
        with sqlite3.connect(DB_PATH) as conn:
            cur = conn.execute('SELECT 1 FROM users WHERE username=?', (username,))
            if cur.fetchone():
                flash('Пользователь с таким логином уже существует')
                return render_template('add_user.html')
            conn.execute(
                'INSERT INTO users (username, password, fio, position, role, approved) VALUES (?, ?, ?, ?, ?, ?)',
                (username, hash_user_password(password), fio, position, role, approved)
            )
            log_action('user_created', 'user', username, f"role={role}; approved={approved}", conn)
        flash('Пользователь успешно добавлен!')
        return redirect(url_for('admin_panel'))
    return render_template('add_user.html')

@app.route('/clear_abiturients', methods=['POST'])
@admin_required
def clear_abiturients():
    campaign_year = get_active_campaign_year()
    if not ensure_campaign_open(campaign_year):
        return redirect(url_for('admin_panel'), code=303)
    backup_path = create_database_backup('before_clear_abiturients')
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute('DELETE FROM abiturients WHERE campaign_year=?', (campaign_year,))
        conn.execute('DELETE FROM pending_duplicates WHERE campaign_year=?', (campaign_year,))
        conn.execute('DELETE FROM login_conflicts WHERE campaign_year=?', (campaign_year,))
        conn.execute('DELETE FROM enrollment_candidates WHERE campaign_year=?', (campaign_year,))
        log_action(
            'abiturients_campaign_cleared',
            'campaign',
            campaign_year,
            f"backup={os.path.basename(backup_path) if backup_path else ''}",
            conn
        )
    flash(f'Абитуриенты, кандидаты, дубли и конфликты кампании {campaign_year} успешно очищены.', 'success')
    return redirect(url_for('admin_panel'))

@app.route('/students_upload', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def students_upload():
    target_section = request_file_work_section('students')
    message = None
    if request.method == 'POST':
        students_import_action = request.form.get('students_import_action', 'preview')
        if students_import_action == 'confirm':
            filepath = None
            try:
                filepath = get_pending_students_import_path(request.form.get('pending_students_import'))
                summary = apply_students_import(filepath)
                flash(
                    (
                        f"Загрузка студентов завершена: добавлено {summary['inserted_count']}, "
                        f"дублей {summary['duplicate_count']}, пропущено {summary['skipped_count']}."
                    ),
                    'success' if not summary.get('errors') else 'info'
                )
            except (UploadValidationError, ValueError) as e:
                flash(str(e), 'error')
            except Exception as e:
                flash(f"Ошибка: {e}", 'error')
            finally:
                cleanup_temp_files(filepath)
            return file_work_redirect(target_section)

        if students_import_action == 'cancel':
            try:
                cleanup_temp_files(get_pending_students_import_path(request.form.get('pending_students_import')))
                flash('Предпросмотр загрузки студентов отменён.', 'info')
            except UploadValidationError:
                pass
            return file_work_redirect(target_section)

        filepath = None
        try:
            filepath = save_upload_to_temp(
                request.files.get('file'),
                STUDENTS_UPLOAD_EXTENSIONS,
                prefix=PENDING_STUDENTS_IMPORT_PREFIX
            )
            plan_df, summary = build_students_import_plan(filepath)
            return render_file_work_page(
                campaign_year=get_active_campaign_year(),
                active_section=target_section,
                students_preview=summary,
                students_preview_rows=student_preview_rows(plan_df),
                students_report=build_students_preview_report(summary),
                enrollment_order_uploads=get_enrollment_order_uploads(get_active_campaign_year()),
                pending_students_import_token=os.path.basename(filepath)
            )
        except (UploadValidationError, ValueError) as e:
            message = str(e)
            flash(message, 'error')
            cleanup_temp_files(filepath)
        except Exception as e:
            message = f"Ошибка: {e}"
            flash(message, 'error')
            cleanup_temp_files(filepath)
        return file_work_redirect(target_section)
    return render_template('students_upload.html', message=message)

@app.route('/students')
@login_required
def students():
    lastname = request.args.get('lastname', '').strip()
    firstname = request.args.get('firstname', '').strip()
    username = request.args.get('username', '').strip()
    cohort = request.args.get('cohort', '').strip()
    enrollment_order = request.args.get('enrollment_order', '').strip()
    order_by = request.args.get('order_by', 'username')
    order_dir = request.args.get('order_dir', 'asc')
    student_rows = get_all_students(
        order_by, order_dir, cohort, lastname, firstname, username, enrollment_order
    )

    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute('SELECT DISTINCT cohort1 FROM students ORDER BY cohort1')
        cohorts = [row[0] for row in cur.fetchall()]

    return render_template(
        'students.html',
        students=student_rows,
        cohorts=cohorts,
        enrollment_orders=get_student_enrollment_order_options(),
        order_by=order_by,
        order_dir=order_dir
    )

@app.route('/students_duplicates')
@login_required
@role_required('admin')
def students_duplicates():
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute('SELECT username, password, email, firstname, lastname, cohort1, cohort2 FROM students_duplicates')
        duplicates = cur.fetchall()
    return render_template('students_duplicates.html', duplicates=duplicates)

@app.route('/students_list')
@login_required
def students_list():
    order_by = request.args.get('order_by', 'username')
    order_dir = request.args.get('order_dir', 'asc')
    cohort = request.args.get('cohort')
    enrollment_order = request.args.get('enrollment_order')
    lastname = request.args.get('lastname')
    firstname = request.args.get('firstname')
    username = request.args.get('username')
    student_rows = get_all_students(
        order_by, order_dir, cohort, lastname, firstname, username, enrollment_order
    )
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute('SELECT DISTINCT cohort1 FROM students ORDER BY cohort1')
        cohorts = [row[0] for row in cur.fetchall()]
    return render_template(
        'students_list.html',
        students=student_rows,
        cohorts=cohorts,
        enrollment_orders=get_student_enrollment_order_options(),
        order_by=order_by,
        order_dir=order_dir
    )

@app.route('/students/download')
@login_required
def download_students():
    order_by = request.args.get('order_by', 'username')
    order_dir = request.args.get('order_dir', 'asc')
    cohort = request.args.get('cohort')
    enrollment_order = request.args.get('enrollment_order')
    lastname = request.args.get('lastname')
    firstname = request.args.get('firstname')
    username = request.args.get('username')
    student_rows = get_all_students(
        order_by, order_dir, cohort, lastname, firstname, username, enrollment_order
    )
    log_action(
        'students_exported',
        'students',
        '',
        f"rows={len(student_rows)}; enrollment_order={enrollment_order or 'all'}"
    )
    export_students = student_rows
    if session.get('role') != 'admin':
        export_students = [dict(student, password='******') for student in student_rows]
    if not are_course_groups_enabled():
        export_students = [
            {key: value for key, value in student.items() if key != 'cohort2'}
            for student in export_students
        ]
    df = pd.DataFrame(export_students)
    output = io.BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="students.xlsx", mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

@app.route('/edit_student/<username>', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def edit_student(username):
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute('SELECT username, password, email, firstname, lastname, cohort1, cohort2 FROM students WHERE username=?', (username,))
        student = cur.fetchone()
        transfer_group_year = infer_group_year(student[5], get_active_campaign_year()) if student else get_active_campaign_year()
        transfer_groups = get_groups_with_counts(conn, transfer_group_year)
    if not student:
        flash('Студент не найден')
        return redirect(url_for('students_list'))
    if request.method == 'POST':
        backup_path = create_database_backup('before_edit_student')
        password = request.form.get('password', '').strip()
        email = request.form.get('email', '').strip()
        firstname = request.form.get('firstname', '').strip()
        lastname = request.form.get('lastname', '').strip()
        cohort1 = normalize_group_name(request.form.get('cohort1', ''))
        cohort2 = get_student_course_group(cohort1)
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute('UPDATE students SET password=?, email=?, firstname=?, lastname=?, cohort1=?, cohort2=? WHERE username=?',
                         (password, email, firstname, lastname, cohort1, cohort2, username))
            log_action(
                'student_updated',
                'student',
                username,
                f"cohort1={cohort1}; cohort2={cohort2 or ''}; backup={os.path.basename(backup_path) if backup_path else ''}",
                conn
            )
        flash('Данные обновлены')
        return redirect(url_for('students_list'))
    return render_template(
        'edit_student.html',
        student=student,
        transfer_groups=transfer_groups,
        transfer_group_year=transfer_group_year,
        transfer_orders=get_student_transfer_orders(username),
    )

@app.route('/edit_student/<username>/transfer_group', methods=['POST'])
@login_required
@role_required('admin')
def transfer_student_group(username):
    selected_group = normalize_group_name(request.form.get('new_cohort1', ''))
    order_file = request.files.get('transfer_order_file')
    saved_file = None

    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute('SELECT username, cohort1, cohort2 FROM students WHERE username=?', (username,))
        student = cur.fetchone()
        if not student:
            flash('Студент не найден', 'error')
            return redirect(url_for('students_list'))

        _username, old_cohort1, old_cohort2 = student
        group_year = infer_group_year(old_cohort1, get_active_campaign_year())
        available_group = conn.execute(
            '''
            SELECT name
            FROM groups
            WHERE name=? AND group_year=? AND COALESCE(is_hidden, 0)=0
            ''',
            (selected_group, group_year)
        ).fetchone()
        if not selected_group or not available_group:
            flash('Выберите новую группу из справочника академических групп.', 'error')
            return redirect(url_for('edit_student', username=username))
        if normalize_group_name(old_cohort1).casefold() == selected_group.casefold():
            flash('Новая группа совпадает с текущей.', 'error')
            return redirect(url_for('edit_student', username=username))
        if get_group_student_count(conn, selected_group) >= MAX_GROUP_STUDENTS:
            flash(f'Выбранная группа заполнена: {MAX_GROUP_STUDENTS}/{MAX_GROUP_STUDENTS}.', 'error')
            return redirect(url_for('edit_student', username=username))

        course_groups_enabled = are_course_groups_enabled()
        new_cohort2 = (derive_cohort2(selected_group) or '') if course_groups_enabled else ''
        if course_groups_enabled and not new_cohort2:
            flash('Для выбранной группы не удалось определить глобальную группу курса.', 'error')
            return redirect(url_for('edit_student', username=username))

        if order_file and order_file.filename:
            try:
                saved_file = save_student_transfer_order_file(username, order_file)
            except UploadValidationError as exc:
                flash(str(exc), 'error')
                return redirect(url_for('edit_student', username=username))
        else:
            saved_file = {
                'filename': '',
                'original_filename': '',
                'mime_type': '',
                'size': 0,
            }

        backup_path = create_database_backup('before_student_group_transfer')
        try:
            conn.execute(
                '''
                UPDATE students
                SET cohort1=?, cohort2=?
                WHERE username=?
                ''',
                (selected_group, new_cohort2, username)
            )
            conn.execute(
                '''
                INSERT INTO student_group_transfers
                    (username, movement_type, old_cohort1, old_cohort2, new_cohort1, new_cohort2,
                     order_source,
                     order_filename, order_original_filename, order_mime_type, order_size,
                     created_by)
                VALUES (?, 'transfer', ?, ?, ?, ?, 'student_transfer', ?, ?, ?, ?, ?)
                ''',
                (
                    username, old_cohort1, old_cohort2, selected_group, new_cohort2,
                    saved_file['filename'], saved_file['original_filename'],
                    saved_file['mime_type'], saved_file['size'], session.get('user', '')
                )
            )
            log_action(
                'student_group_transferred',
                'student',
                username,
                (
                    f"old_cohort1={old_cohort1 or ''}; old_cohort2={old_cohort2 or ''}; "
                    f"new_cohort1={selected_group}; new_cohort2={new_cohort2}; "
                    f"order={saved_file['original_filename'] or 'not_attached'}; "
                    f"backup={os.path.basename(backup_path) if backup_path else ''}"
                ),
                conn
            )
        except Exception:
            if saved_file and saved_file.get('path') and os.path.exists(saved_file['path']):
                os.remove(saved_file['path'])
            raise

    if are_course_groups_enabled():
        flash(f'Студент переведен в группу {selected_group}. Глобальная группа курса: {new_cohort2}.', 'success')
    else:
        flash(f'Студент переведен в группу {selected_group}.', 'success')
    return redirect(url_for('edit_student', username=username))

@app.route('/student_transfer_orders/<int:transfer_id>/download')
@login_required
@role_required('admin')
def download_student_transfer_order(transfer_id):
    try:
        order = get_student_transfer_order_download(transfer_id)
    except FileNotFoundError as exc:
        flash(str(exc), 'error')
        return redirect(url_for('students_list'))
    return send_file(
        order['path'],
        as_attachment=True,
        download_name=order['download_name'],
        mimetype=order['mime_type']
    )

@app.route('/delete_student', methods=['POST'])
@login_required
@role_required('admin')
def delete_student():
    username = request.form.get('username')
    if username:
        with sqlite3.connect(DB_PATH) as conn:
            cur = conn.execute(
                '''
                SELECT username, email, firstname, lastname, source_campaign_year, source_dogovor, source_fio
                FROM students
                WHERE username=?
                ''',
                (username,)
            )
            student = cur.fetchone()
            if not student:
                flash('Студент не найден')
                return redirect(url_for('students_list'))

            username, email, firstname, lastname, source_campaign_year, source_dogovor, source_fio = student
            if source_campaign_year and is_campaign_archived(source_campaign_year):
                flash(ARCHIVED_CAMPAIGN_MESSAGE, 'error')
                return redirect(url_for('students_list'))

            backup_path = create_database_backup('before_delete_student')
            if source_campaign_year:
                campaign_year = normalize_campaign_year(source_campaign_year, source_campaign_year)
                fio = source_fio or ' '.join(part for part in [lastname, firstname] if part).strip()
                _, fallback_fam, fallback_imotch = split_fio(fio)
                fam = lastname or fallback_fam
                imotch = firstname or fallback_imotch

                abiturient_exists = conn.execute(
                    'SELECT 1 FROM abiturients WHERE login=? AND campaign_year=?',
                    (username, campaign_year)
                ).fetchone()
                if abiturient_exists:
                    flash(f'Абитуриент {username} уже есть в кампании {campaign_year}')
                else:
                    conn.execute(
                        '''
                        INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                        ''',
                        (fio, source_dogovor or '', username, campaign_year, fam, imotch, email)
                    )
                    flash(f'Студент {username} возвращен в абитуриенты кампании {campaign_year}')

            conn.execute('DELETE FROM students WHERE username=?', (username,))
            log_action(
                'student_deleted',
                'student',
                username,
                f"backup={os.path.basename(backup_path) if backup_path else ''}",
                conn
            )
            flash('Студент удален')
    return redirect(url_for('students_list'))

@app.route('/enrollment_candidates/sync', methods=['POST'])
@login_required
@role_required('admin', 'assistant')
def enrollment_candidates_sync():
    campaign_year = get_active_campaign_year()
    group_year = normalize_group_year(request.form.get('group_year'), campaign_year)
    if not ensure_campaign_open(campaign_year):
        return redirect(url_for('abiturients_to_students', group_year=group_year), code=303)
    summary = sync_enrollment_candidates_from_ready_abiturients(campaign_year)
    flash(
        (
            f"Кандидаты к зачислению обновлены: добавлено {summary['created']}, "
            f"обновлено {summary['updated']}, удалено неготовых {summary.get('removed', 0)}."
        ),
        'success'
    )
    if summary['skipped_without_email']:
        flash(f"Без почты не добавлены: {len(summary['skipped_without_email'])}", 'info')
    if summary['skipped_unpaid']:
        flash(f"Без оплаты не добавлены: {len(summary['skipped_unpaid'])}", 'info')
    if summary['skipped_without_specialty']:
        flash(f"Не удалось определить специальность по договору: {len(summary['skipped_without_specialty'])}", 'info')
    if summary['skipped_existing_students']:
        flash(f"Уже есть в студентах: {len(summary['skipped_existing_students'])}", 'info')
    return redirect(url_for('abiturients_to_students', group_year=group_year), code=303)

@app.route('/abiturients_to_students', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'assistant')
def abiturients_to_students():
    campaign_year = get_active_campaign_year()
    group_year = normalize_group_year(request.values.get('group_year'), campaign_year)
    specialty_filter = normalize_specialty_key(request.values.get('specialty'))
    enrollment_order_required = is_enrollment_order_required()
    group_years = get_group_years(group_year)
    with sqlite3.connect(DB_PATH) as conn:
        groups = get_groups_with_counts(conn, group_year)
        order_total = conn.execute(
            'SELECT COUNT(*) FROM enrollment_orders WHERE campaign_year=?',
            (campaign_year,)
        ).fetchone()[0]
    specialties = get_enrollment_candidate_specialties(campaign_year)
    if specialty_filter and specialty_filter not in {item['key'] for item in specialties}:
        specialty_filter = ''
    suggested_groups = get_candidate_group_options(groups, specialty_filter) if specialty_filter else groups
    candidates = get_enrollment_candidates(campaign_year, specialty_filter)
    selected_candidate_ids = [item for item in request.values.getlist('candidate_ids') if str(item).isdigit()]
    login_distribution_enabled = request.values.get('use_login_distribution') == '1'
    login_distribution_preview = None

    def render_stage(**extra_context):
        context = {
            'candidates': candidates,
            'groups': groups,
            'suggested_groups': suggested_groups,
            'specialties': specialties,
            'selected_specialty': specialty_filter,
            'order_total': order_total,
            'campaign_year': campaign_year,
            'group_year': group_year,
            'group_years': group_years,
            'selected_candidate_ids': selected_candidate_ids,
            'login_distribution_enabled': login_distribution_enabled,
            'login_distribution_preview': login_distribution_preview,
            'enrollment_order_required': enrollment_order_required,
        }
        context.update(extra_context)
        return render_template('abiturients_to_students.html', **context)

    def redirect_to_stage():
        args = {'group_year': group_year}
        if specialty_filter:
            args['specialty'] = specialty_filter
        return redirect(url_for('abiturients_to_students', **args), code=303)

    if request.method == 'POST':
        if not ensure_campaign_open(campaign_year):
            return redirect(url_for('abiturients_to_students', group_year=group_year), code=303)
        distribution_action = request.form.get('distribution_action', '').strip()
        cohort1 = normalize_group_name(request.form.get('cohort1', ''))
        ids = [item for item in request.form.getlist('candidate_ids') if str(item).isdigit()]
        selected_candidate_ids = ids
        login_distribution_enabled = request.form.get('use_login_distribution') == '1' or distribution_action == 'confirm_login_groups'
        auto_split = request.form.get('auto_split') == '1'

        if login_distribution_enabled and distribution_action != 'confirm_login_groups':
            if not ids:
                flash('Выберите хотя бы одного кандидата для распределения по логинам')
                return redirect_to_stage()
            with sqlite3.connect(DB_PATH) as conn:
                refresh_enrollment_candidate_statuses(conn, campaign_year)
                candidate_rows = get_selected_enrollment_candidate_rows(conn, campaign_year, ids)
                login_distribution_preview = build_login_group_distribution_plan(
                    conn,
                    candidate_rows,
                    campaign_year,
                    group_year
                )
            candidates = get_enrollment_candidates(campaign_year, specialty_filter)
            return render_stage(candidates=candidates, login_distribution_preview=login_distribution_preview)

        if distribution_action == 'confirm_login_groups':
            if not ids:
                flash('Выберите хотя бы одного кандидата для распределения по логинам')
                return redirect_to_stage()
            with sqlite3.connect(DB_PATH) as conn:
                refresh_enrollment_candidate_statuses(conn, campaign_year)
                candidate_rows = get_selected_enrollment_candidate_rows(conn, campaign_year, ids)
                login_distribution_preview = build_login_group_distribution_plan(
                    conn,
                    candidate_rows,
                    campaign_year,
                    group_year
                )
                if not login_distribution_preview['summary']['can_confirm']:
                    flash('Распределение не подтверждено: сначала исправьте строки со статусом "Проверить".', 'error')
                    candidates = get_enrollment_candidates(campaign_year, specialty_filter)
                    return render_stage(candidates=candidates, login_distribution_preview=login_distribution_preview)

                backup_path = create_database_backup('before_login_group_distribution_migration')
                migrated_count = 0
                touched_groups = set()
                for plan_row in login_distribution_preview['rows']:
                    target_group = normalize_group_name(plan_row['target_group'])
                    target_cohort2 = get_student_course_group(target_group)
                    conn.execute(
                        'INSERT OR IGNORE INTO groups (name, group_year) VALUES (?, ?)',
                        (target_group, group_year)
                    )
                    conn.execute(
                        '''
                        INSERT INTO students
                            (username, password, email, firstname, lastname, cohort1, cohort2, source_campaign_year, source_dogovor, source_fio)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        ''',
                        (
                            plan_row['login'], 'cron', plan_row['email'],
                            plan_row['firstname'], plan_row['lastname'], target_group, target_cohort2,
                            campaign_year, plan_row['dogovor'], plan_row['fio']
                        )
                    )
                    record_student_enrollment_movement(
                        conn,
                        plan_row['login'],
                        target_group,
                        target_cohort2,
                        plan_row.get('order_match'),
                        session.get('user', '')
                    )
                    conn.execute(
                        'DELETE FROM abiturients WHERE id=? AND campaign_year=?',
                        (plan_row['abiturient_id'], campaign_year)
                    )
                    conn.execute('DELETE FROM enrollment_candidates WHERE id=?', (plan_row['candidate_id'],))
                    migrated_count += 1
                    touched_groups.add(target_group)

                if migrated_count:
                    log_action(
                        'verified_candidates_migrated_by_login_groups',
                        'campaign',
                        campaign_year,
                        (
                            f"group_year={group_year}; count={migrated_count}; "
                            f"groups={','.join(sorted(touched_groups))}; "
                            f"backup={os.path.basename(backup_path) if backup_path else ''}"
                        ),
                        conn
                    )
            flash(f'Распределено по логинам и перенесено студентов: {migrated_count}')
            if login_distribution_preview['summary']['create_group_count']:
                flash(f"Создано групп: {login_distribution_preview['summary']['create_group_count']}")
            return redirect(url_for('students_list'))

        if not cohort1 or not ids:
            flash('Выберите группу и хотя бы одного кандидата к зачислению')
            return redirect_to_stage()
        with sqlite3.connect(DB_PATH) as conn:
            group_exists = conn.execute(
                'SELECT 1 FROM groups WHERE name=? AND group_year=? AND COALESCE(is_hidden, 0)=0',
                (cohort1, group_year)
            ).fetchone()
            if not group_exists:
                flash('Выберите видимую группу из справочника академических групп')
                return redirect_to_stage()

            placeholders = ','.join('?' for _ in ids)
            cur = conn.execute(
                f'''
                SELECT
                    c.id, c.abiturient_id, c.fio, c.dogovor, c.login, c.fam, c.imotch,
                    c.email, c.specialty, c.specialty_key, c.verification_status,
                    c.order_group_name, a.id, a.email, a.paid
                FROM enrollment_candidates c
                LEFT JOIN abiturients a
                    ON a.id=c.abiturient_id AND a.campaign_year=c.campaign_year
                WHERE c.campaign_year=? AND c.id IN ({placeholders})
                ORDER BY c.fio
                ''',
                [campaign_year] + ids
            )
            candidate_rows = cur.fetchall()

            skipped_without_order = []
            skipped_wrong_group = []
            skipped_not_ready = []
            skipped_duplicates = []
            selected_candidates = []
            for row in candidate_rows:
                (
                    candidate_id, abiturient_id, fio, dogovor, username, lastname,
                    firstname, candidate_email, specialty, specialty_key,
                    verification_status, order_group_name, source_id, source_email, source_paid
                ) = row
                display_name = fio or username or str(candidate_id)
                order_match = (
                    get_enrollment_order_match_for_abiturient(fio, dogovor, campaign_year)
                    if enrollment_order_required
                    else None
                )
                if enrollment_order_required and (verification_status != 'verified' or not order_match):
                    skipped_without_order.append(display_name)
                    continue
                if not group_matches_specialty(cohort1, specialty_key):
                    skipped_wrong_group.append(f'{display_name} ({specialty})')
                    continue
                required_group = (order_match.get('group_name') or order_group_name) if enrollment_order_required and order_match else ''
                if required_group and not group_matches_order_group(cohort1, required_group):
                    skipped_wrong_group.append(f'{display_name} (в приказе {required_group})')
                    continue
                email = (source_email or candidate_email or '').strip()
                readiness_issues = enrollment_candidate_readiness_issues(source_id, email, source_paid)
                if readiness_issues:
                    skipped_not_ready.append(f"{display_name} ({', '.join(readiness_issues)})")
                    continue
                student_exists = conn.execute('SELECT 1 FROM students WHERE username=?', (username,)).fetchone()
                if student_exists:
                    skipped_duplicates.append(username)
                    continue

                selected_candidates.append({
                    'candidate_id': candidate_id,
                    'abiturient_id': abiturient_id,
                    'username': username,
                    'email': email,
                    'firstname': firstname,
                    'lastname': lastname,
                    'fio': fio,
                    'dogovor': dogovor,
                    'required_group': required_group or '',
                    'order_match': order_match or {},
                })

            if auto_split and any(item['required_group'] for item in selected_candidates):
                flash('Автораспределение нельзя использовать для кандидатов, у которых группа указана в приказе.', 'error')
                return redirect_to_stage()

            current_count = get_group_student_count(conn, cohort1)
            free_places = MAX_GROUP_STUDENTS - current_count
            if selected_candidates and len(selected_candidates) > free_places and not auto_split:
                next_group = get_next_subgroup_name(conn, cohort1, group_year)
                flash(f'В группе {cohort1} свободно мест: {max(free_places, 0)}/{MAX_GROUP_STUDENTS}. Создайте или выберите следующую подгруппу: {next_group}')
                return redirect_to_stage()

            backup_path = create_database_backup('before_enrollment_candidates_migration') if selected_candidates else None
            migrated_count = 0
            next_group_after_full = ''
            assignments = []
            if auto_split:
                remaining_candidates = list(selected_candidates)
                target_group = cohort1
                while remaining_candidates:
                    target_count = get_group_student_count(conn, target_group)
                    free_in_group = MAX_GROUP_STUDENTS - target_count
                    if free_in_group <= 0:
                        target_group = get_next_subgroup_name(conn, target_group, group_year)
                        conn.execute(
                            'INSERT OR IGNORE INTO groups (name, group_year) VALUES (?, ?)',
                            (target_group, group_year)
                        )
                        continue
                    current_batch = remaining_candidates[:free_in_group]
                    remaining_candidates = remaining_candidates[free_in_group:]
                    assignments.extend((candidate, target_group) for candidate in current_batch)
                    if remaining_candidates:
                        next_group_name = get_next_subgroup_name(conn, target_group, group_year)
                        conn.execute(
                            'INSERT OR IGNORE INTO groups (name, group_year) VALUES (?, ?)',
                            (next_group_name, group_year)
                        )
                        target_group = next_group_name
            else:
                assignments = [(candidate, cohort1) for candidate in selected_candidates]

            touched_groups = set()
            for candidate, target_group in assignments:
                target_cohort2 = get_student_course_group(target_group)
                conn.execute(
                    '''
                    INSERT INTO students
                        (username, password, email, firstname, lastname, cohort1, cohort2, source_campaign_year, source_dogovor, source_fio)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''',
                    (
                        candidate['username'], 'cron', candidate['email'],
                        candidate['firstname'], candidate['lastname'], target_group, target_cohort2,
                        campaign_year, candidate['dogovor'], candidate['fio']
                    )
                )
                record_student_enrollment_movement(
                    conn,
                    candidate['username'],
                    target_group,
                    target_cohort2,
                    candidate.get('order_match'),
                    session.get('user', '')
                )
                conn.execute(
                    'DELETE FROM abiturients WHERE id=? AND campaign_year=?',
                    (candidate['abiturient_id'], campaign_year)
                )
                conn.execute('DELETE FROM enrollment_candidates WHERE id=?', (candidate['candidate_id'],))
                migrated_count += 1
                touched_groups.add(target_group)
            if get_group_student_count(conn, cohort1) >= MAX_GROUP_STUDENTS:
                next_group_after_full = get_next_subgroup_name(conn, cohort1, group_year)
            if migrated_count:
                log_action(
                    'verified_candidates_migrated_to_students',
                    'group',
                    cohort1,
                    (
                        f"campaign_year={campaign_year}; group_year={group_year}; "
                        f"count={migrated_count}; auto_split={int(auto_split)}; "
                        f"groups={','.join(sorted(touched_groups))}; backup={os.path.basename(backup_path) if backup_path else ''}"
                    ),
                    conn
                )

        if migrated_count:
            flash(f'Мигрировано студентов: {migrated_count}')
            if auto_split:
                flash('Автораспределение по подгруппам выполнено.')
            else:
                flash(f'Группа {cohort1}: {current_count + migrated_count}/{MAX_GROUP_STUDENTS}')
        if next_group_after_full:
            flash(f'Группа {cohort1} заполнена. Следующая подгруппа: {next_group_after_full}')
        if skipped_without_order:
            names = ', '.join(skipped_without_order[:10])
            suffix = '...' if len(skipped_without_order) > 10 else ''
            flash(f'Не перенесены без совпадения с приказом: {names}{suffix}')
        if skipped_wrong_group:
            names = ', '.join(skipped_wrong_group[:10])
            suffix = '...' if len(skipped_wrong_group) > 10 else ''
            flash(f'Не перенесены из-за неподходящей группы: {names}{suffix}')
        if skipped_not_ready:
            names = ', '.join(skipped_not_ready[:10])
            suffix = '...' if len(skipped_not_ready) > 10 else ''
            flash(f'Не перенесены, потому что исходный абитуриент уже не готов: {names}{suffix}')
        if skipped_duplicates:
            names = ', '.join(skipped_duplicates[:10])
            suffix = '...' if len(skipped_duplicates) > 10 else ''
            flash(f'Не перенесены, уже есть в студентах: {names}{suffix}')
        if not migrated_count and not skipped_without_order and not skipped_wrong_group and not skipped_not_ready and not skipped_duplicates:
            flash('Не удалось найти выбранных кандидатов для текущей кампании')

        target = 'students_list' if migrated_count and not skipped_without_order and not skipped_wrong_group and not skipped_not_ready and not skipped_duplicates else 'abiturients_to_students'
        if target == 'abiturients_to_students':
            return redirect_to_stage()
        return redirect(url_for(target))
    return render_stage()

@app.route('/migration_wizard')
@login_required
@role_required('admin', 'assistant')
def migration_wizard():
    campaign_year = get_active_campaign_year()
    group_year = normalize_group_year(request.args.get('group_year'), campaign_year)
    dashboard = get_dashboard_data(campaign_year)
    with sqlite3.connect(DB_PATH) as conn:
        groups = get_groups_with_counts(conn, group_year)
    return render_template(
        'migration_wizard.html',
        dashboard=dashboard,
        groups=groups,
        group_year=group_year,
        group_years=get_group_years(group_year),
    )

@app.route('/add_group', methods=['GET', 'POST'])
@login_required
@role_required('admin')
def add_group():
    group_year = normalize_group_year(request.values.get('group_year'), get_active_campaign_year())
    show_hidden = request.values.get('show_hidden') == '1'

    def add_group_url():
        args = {'group_year': group_year}
        if show_hidden:
            args['show_hidden'] = '1'
        return url_for('add_group', **args)

    if request.method == 'POST':
        group_action = request.form.get('group_action', '').strip()
        if group_action == 'force_subgroup':
            source_group = normalize_group_name(request.form.get('source_group', ''))
            if not source_group:
                flash('Выберите исходную группу для новой подгруппы')
                return redirect(add_group_url())

            with sqlite3.connect(DB_PATH) as conn:
                source_exists = conn.execute(
                    'SELECT 1 FROM groups WHERE name=? AND group_year=? AND COALESCE(is_hidden, 0)=0',
                    (source_group, group_year)
                ).fetchone()
                if not source_exists:
                    flash('Исходная группа не найдена или скрыта')
                    return redirect(add_group_url())

                next_group = get_next_subgroup_name(conn, source_group, group_year)
                if infer_group_year(next_group, group_year) != group_year:
                    flash(f'Для папки {group_year} новая подгруппа должна начинаться с {group_year[-2:]}')
                    return redirect(add_group_url())

                existing = {
                    row[0].casefold(): row[0]
                    for row in conn.execute('SELECT name FROM groups')
                }
                if group_exists_casefold(existing, next_group):
                    flash(f'Подгруппа {next_group} уже существует')
                    return redirect(add_group_url())

                conn.execute(
                    'INSERT INTO groups (name, group_year) VALUES (?, ?)',
                    (next_group, group_year)
                )
                flash(f'Принудительно создана подгруппа {next_group} для {source_group}')
            return redirect(add_group_url())

        if group_action in ('hide', 'show', 'delete'):
            group_name = normalize_group_name(request.form.get('group_name', ''))
            if not group_name:
                flash('Выберите группу')
                return redirect(add_group_url())

            with sqlite3.connect(DB_PATH) as conn:
                group_row = conn.execute(
                    'SELECT name FROM groups WHERE name=? AND group_year=?',
                    (group_name, group_year)
                ).fetchone()
                if not group_row:
                    flash('Группа не найдена в выбранной папке')
                    return redirect(add_group_url())

                if group_action == 'hide':
                    conn.execute(
                        'UPDATE groups SET is_hidden=1 WHERE name=? AND group_year=?',
                        (group_name, group_year)
                    )
                    flash(f'Группа {group_name} скрыта')
                elif group_action == 'show':
                    conn.execute(
                        'UPDATE groups SET is_hidden=0 WHERE name=? AND group_year=?',
                        (group_name, group_year)
                    )
                    flash(f'Группа {group_name} снова отображается')
                elif group_action == 'delete':
                    student_count = get_group_student_count(conn, group_name)
                    if student_count:
                        flash(f'Нельзя удалить группу {group_name}: в ней есть студенты ({student_count}). Можно скрыть группу.')
                    else:
                        conn.execute(
                            'DELETE FROM groups WHERE name=? AND group_year=?',
                            (group_name, group_year)
                        )
                        flash(f'Группа {group_name} удалена')
            return redirect(add_group_url())

        groups_file = request.files.get('groups_file')
        if groups_file and groups_file.filename:
            filepath = None
            if get_upload_extension(groups_file) not in GROUPS_UPLOAD_EXTENSIONS:
                flash('Загрузите файл групп в формате CSV')
                return redirect(add_group_url())

            filepath = save_upload_to_temp(groups_file, GROUPS_UPLOAD_EXTENSIONS)
            try:
                result = process_groups_csv(filepath, group_year)
            except Exception as exc:
                flash(f'Ошибка загрузки групп: {exc}')
                return redirect(add_group_url())

            finally:
                cleanup_temp_files(filepath)

            if result['created']:
                flash(f'Добавлено групп: {len(result["created"])}')
            if result['skipped']:
                flash(f'Пропущено дублей: {len(result["skipped"])}')
            if result['errors']:
                errors = '; '.join(result['errors'][:5])
                suffix = '...' if len(result['errors']) > 5 else ''
                flash(f'Ошибки в CSV: {errors}{suffix}')
            if not result['created'] and not result['errors']:
                flash('Новые группы не добавлены')
            return redirect(add_group_url())

        source_group = normalize_group_name(request.form.get('source_group', ''))
        group_name = normalize_group_name(request.form.get('group_name', ''))
        if group_name:
            if not is_valid_group_name(group_name):
                flash('Название группы должно быть в формате 26ФМ-11-1')
                return redirect(add_group_url())
            if infer_group_year(group_name, group_year) != group_year:
                flash(f'Для папки {group_year} название группы должно начинаться с {group_year[-2:]}')
                return redirect(add_group_url())
            with sqlite3.connect(DB_PATH) as conn:
                if source_group:
                    source_exists = conn.execute(
                        'SELECT 1 FROM groups WHERE name=? AND group_year=? AND COALESCE(is_hidden, 0)=0',
                        (source_group, group_year)
                    ).fetchone()
                    source_count = get_group_student_count(conn, source_group) if source_exists else 0
                    expected_group = get_next_subgroup_name(conn, source_group, group_year) if source_exists else ''
                    if not source_exists or source_count < MAX_GROUP_STUDENTS or not is_last_subgroup(conn, source_group, group_year) or group_name != expected_group:
                        flash('Дополнительную подгруппу можно создать только для последней заполненной группы')
                        return redirect(add_group_url())

                existing = {
                    row[0].casefold(): row[0]
                    for row in conn.execute('SELECT name FROM groups')
                }
                if group_exists_casefold(existing, group_name):
                    flash('Такая группа уже существует')
                    return redirect(add_group_url())
                try:
                    conn.execute('INSERT INTO groups (name, group_year) VALUES (?, ?)', (group_name, group_year))
                    if source_group:
                        flash(f'Создана дополнительная подгруппа {group_name} для {source_group}')
                    else:
                        flash('Группа добавлена')
                except sqlite3.IntegrityError:
                    flash('Такая группа уже существует')
        else:
            flash('Название группы не может быть пустым')
        return redirect(add_group_url())
    # Список всех групп для отображения
    with sqlite3.connect(DB_PATH) as conn:
        groups = get_groups_with_counts(conn, group_year, include_hidden=show_hidden)
    visible_groups = [group for group in groups if not group['is_hidden']]
    group_years = get_group_years(group_year, include_base=True)
    return render_template(
        'add_group.html',
        groups=groups,
        visible_groups=visible_groups,
        group_year=group_year,
        group_years=group_years,
        group_year_code=group_year[-2:],
        show_hidden=show_hidden,
    )

@app.route('/groups_template/download')
@login_required
@role_required('admin')
def download_groups_template():
    group_year = normalize_group_year(request.args.get('group_year'), get_active_campaign_year())
    output = io.BytesIO(build_groups_template_csv(group_year).encode('utf-8-sig'))
    output.seek(0)
    return send_file(output, as_attachment=True, download_name='groups_template.csv', mimetype='text/csv')

if __name__ == "__main__":
    app_host = os.environ.get("APP_HOST", "127.0.0.1")
    app_port = int(os.environ.get("APP_PORT", "5000"))
    app_debug = os.environ.get("APP_DEBUG", "").lower() in {"1", "true", "yes", "on"}
    app.run(host=app_host, port=app_port, debug=app_debug)
