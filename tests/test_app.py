import os
import io
import re
import shutil
import sqlite3
import sys
import tempfile
import unittest
from urllib.parse import unquote
from unittest import mock

import pandas as pd


PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TEST_UPLOAD_DIR = tempfile.mkdtemp(prefix='manticore_tests_')

os.environ['SECRET_KEY'] = 'test-secret-key'
os.environ['ADMIN_DEFAULT_PASSWORD'] = 'test-admin-password'
os.environ['UPLOAD_FOLDER'] = TEST_UPLOAD_DIR
os.environ['DB_FILENAME'] = 'test.db'
os.environ['DEFAULT_CAMPAIGN_YEAR'] = '2026'
os.environ['LEGACY_CAMPAIGN_YEAR'] = '2025'
os.environ['APP_DEBUG'] = 'false'

sys.path.insert(0, PROJECT_ROOT)
import app as manticore


manticore.app.config['TESTING'] = True


def reset_database():
    with sqlite3.connect(manticore.DB_PATH) as conn:
        for table in (
            'abiturients',
            'pending_duplicates',
            'login_conflicts',
            'enrollment_candidates',
            'enrollment_orders',
            'enrollment_order_upload_rows',
            'enrollment_order_uploads',
            'students',
            'students_duplicates',
            'audit_logs',
            'student_group_transfers',
            'login_attempts',
            'campaign_settings',
            'login_generation_settings',
        ):
            conn.execute(f'DELETE FROM {table}')


class ManticoreAppTests(unittest.TestCase):
    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(TEST_UPLOAD_DIR, ignore_errors=True)

    def setUp(self):
        reset_database()
        manticore.save_login_generation_settings(
            manticore.get_default_login_generation_rules(),
            setup_completed=True,
            updated_by='test'
        )

    def make_abiturients_file(self, rows, filename='abiturients.xlsx'):
        file_path = os.path.join(TEST_UPLOAD_DIR, filename)
        pd.DataFrame(rows).to_excel(file_path, index=False)
        return file_path

    def login_session(self, client, username='admin', role='admin'):
        with client.session_transaction() as session:
            session['user'] = username
            session['role'] = role

    def csrf_from_response(self, response):
        token_match = re.search(r'name="csrf_token" value="([^"]+)"', response.get_data(as_text=True))
        self.assertIsNotNone(token_match)
        return token_match.group(1)

    @mock.patch('app.desktop_releases.fetch_latest_release')
    def test_admin_can_approve_windows_github_release(self, fetch_latest_release):
        approval_path = os.path.join(TEST_UPLOAD_DIR, 'desktop_release_approval.json')
        if os.path.exists(approval_path):
            os.remove(approval_path)
        fetch_latest_release.return_value = {
            'repository': manticore.DESKTOP_GITHUB_REPOSITORY,
            'release_id': 987,
            'tag_name': 'v9.8.7',
            'version': '9.8.7',
            'name': 'v9.8.7',
            'notes': 'Тестовый выпуск',
            'html_url': 'https://github.com/UnicornisIT/manticore/releases/tag/v9.8.7',
            'published_at': '2026-09-01T00:00:00Z',
            'immutable': True,
            'asset_id': 654,
            'asset_name': 'Manticore-Setup-9.8.7.exe',
            'download_url': 'https://github.com/UnicornisIT/manticore/releases/download/v9.8.7/Manticore-Setup-9.8.7.exe',
            'sha256': 'a' * 64,
            'size': 12345,
        }
        client = manticore.app.test_client()
        self.login_session(client)
        token = self.csrf_from_response(client.get('/admin_panel'))

        response = client.post(
            '/admin/desktop-release/approve',
            data={'csrf_token': token},
        )

        self.assertEqual(response.status_code, 303)
        manifest_response = client.get('/api/desktop/releases/windows')
        self.assertEqual(manifest_response.status_code, 200)
        manifest = manifest_response.get_json()
        self.assertTrue(manifest['approved'])
        self.assertEqual(manifest['approval']['version'], '9.8.7')
        self.assertEqual(manifest['approval']['sha256'], 'a' * 64)
        os.remove(approval_path)

    def test_custom_login_generation_rules_are_used_for_import(self):
        rules = manticore.get_default_login_generation_rules()
        rules.update({
            'mode': 'custom',
            'template': 'u{yy}{spec}{base}{seq}',
            'number_width': 4,
            'error_prefix': 'bad',
            'duplicate_prefix': 'copy',
            'spec_codes': {'ИТ': 'it'},
            'base_codes': {'очно': 'o'},
            'base_match_mode': 'last_part',
        })
        manticore.save_login_generation_settings(rules, setup_completed=True, updated_by='test')

        file_path = self.make_abiturients_file([
            {'ФИО': 'Тестов Тест Тестович', 'Договор': '2026-ИТ-0001-очно'},
            {'ФИО': 'Ошибкин Олег Олегович', 'Договор': 'без номера'},
        ], filename='custom_rules.xlsx')

        plan_df, summary = manticore.build_abiturients_import_plan(file_path, '2026')

        self.assertEqual(summary['ready_count'], 1)
        self.assertEqual(summary['conflict_count'], 1)
        self.assertEqual(plan_df['login'].tolist(), ['u26ito0001', 'bad0001'])

    def test_lowercase_i_base_aliases_are_hidden_but_still_accepted(self):
        rules = manticore.get_default_login_generation_rules()
        rules['base_codes'] = {
            '11и': '11i',
            '11И': '11i',
            '9и': '9i',
            '9И': '9i',
            '11': '11',
            '9': '9',
        }
        manticore.save_login_generation_settings(rules, setup_completed=True, updated_by='test')

        normalized_rules = manticore.get_login_generation_rules()
        self.assertNotIn('11и', normalized_rules['base_codes'])
        self.assertNotIn('9и', normalized_rules['base_codes'])
        self.assertIn('11И', normalized_rules['base_codes'])
        self.assertIn('9И', normalized_rules['base_codes'])
        self.assertEqual(
            manticore.parse_dogovor('2026-СД-0001-11и', normalized_rules),
            manticore.parse_dogovor('2026-СД-0001-11И', normalized_rules),
        )
        self.assertEqual(
            manticore.parse_dogovor('2026-ЛД-0002-9и', normalized_rules),
            manticore.parse_dogovor('2026-ЛД-0002-9И', normalized_rules),
        )

        client = manticore.app.test_client()
        self.login_session(client)
        for path in ('/abiturients', '/manual_create'):
            response = client.get(path)
            body = response.get_data(as_text=True)
            self.assertEqual(response.status_code, 200)
            self.assertNotIn('value="11и"', body)
            self.assertNotIn('value="9и"', body)
            self.assertIn('value="11И"', body)
            self.assertIn('value="9И"', body)

        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                ('Старый Тест Запись', '2026-СД-0003-11и', 'old001', '2026', 'Старый', 'Тест Запись')
            )
        filtered_rows = manticore.get_all_abiturients(base='11И', campaign_year='2026')
        self.assertTrue(any(row['login'] == 'old001' for row in filtered_rows))

        file_path = self.make_abiturients_file([
            {'ФИО': 'Новый Один Тестович', 'Договор': '2026-СД-0001-11и'},
            {'ФИО': 'Новый Два Тестович', 'Договор': '2026-ЛД-0002-9и'},
        ], filename='lowercase_i.xlsx')

        plan_df, summary = manticore.build_abiturients_import_plan(file_path, '2026')

        self.assertEqual(summary['ready_count'], 2)
        self.assertEqual(plan_df['Договор'].tolist(), ['2026-СД-0001-11И', '2026-ЛД-0002-9И'])

        result_path, summary = manticore.apply_abiturients_import(file_path, '2026')
        self.assertEqual(summary['ready_count'], 2)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            stored_dogovors = {
                row[0]
                for row in conn.execute("SELECT dogovor FROM abiturients WHERE login IN ('26311i001', '2619i001')")
            }
        self.assertEqual(stored_dogovors, {'2026-СД-0001-11И', '2026-ЛД-0002-9И'})
        os.remove(result_path)

    def test_manual_create_builds_login_from_full_contract_and_skips_used_login(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Занятый Логин Тестович',
                    '2026-ФМ-0600-11',
                    '26611001',
                    '2026',
                    'Занятый',
                    'Логин Тестович',
                ),
            )

        client = manticore.app.test_client()
        self.login_session(client)
        page = client.get('/manual_create')
        csrf_token = self.csrf_from_response(page)
        response = client.post(
            '/manual_create',
            data={
                'csrf_token': csrf_token,
                'creation_mode': 'contract',
                'fio': 'Новый Абитуриент Тестович',
                'dogovor': '2026-ФМ-0700-11',
            },
        )

        self.assertEqual(response.status_code, 200)
        body = response.get_data(as_text=True)
        self.assertIn('Логин успешно создан: 26611002', body)
        self.assertIn('По номеру договора', body)
        self.assertIn('Настроить вручную', body)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            created = conn.execute(
                'SELECT fio, dogovor, login, campaign_year FROM abiturients WHERE dogovor=?',
                ('2026-ФМ-0700-11',),
            ).fetchone()
        self.assertEqual(
            created,
            ('Новый Абитуриент Тестович', '2026-ФМ-0700-11', '26611002', '2026'),
        )

    def test_manual_create_can_use_selected_parameters_and_custom_login(self):
        client = manticore.app.test_client()
        self.login_session(client)
        page = client.get('/manual_create')
        csrf_token = self.csrf_from_response(page)
        response = client.post(
            '/manual_create',
            data={
                'csrf_token': csrf_token,
                'creation_mode': 'manual',
                'fio': 'Ручной Режим Тестович',
                'year': '2026',
                'spec': 'ФМ',
                'base': '11',
                'manual_login': 'custom_login_77',
            },
        )

        self.assertEqual(response.status_code, 200)
        body = response.get_data(as_text=True)
        self.assertIn('Логин успешно создан: custom_login_77', body)
        self.assertIn('data-mode-panel="contract" hidden', body)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            created = conn.execute(
                'SELECT dogovor, login, campaign_year FROM abiturients WHERE fio=?',
                ('Ручной Режим Тестович',),
            ).fetchone()
        self.assertEqual(created, ('', 'custom_login_77', '2026'))

    def test_manual_create_blocks_duplicate_contract_and_login_and_confirms_exact_fio(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Иванов Иван Иванович',
                    '2026-ФМ-0701-11',
                    'taken_manual',
                    '2026',
                    'Иванов',
                    'Иван Иванович',
                ),
            )

        client = manticore.app.test_client()
        self.login_session(client)
        page = client.get('/manual_create')
        csrf_token = self.csrf_from_response(page)

        duplicate_contract = client.post(
            '/manual_create',
            data={
                'csrf_token': csrf_token,
                'creation_mode': 'contract',
                'fio': 'Другой Абитуриент Тестович',
                'dogovor': '2026-ФМ-0701-11',
            },
        )
        self.assertIn('Такой договор уже есть в базе.', duplicate_contract.get_data(as_text=True))

        duplicate_login = client.post(
            '/manual_create',
            data={
                'csrf_token': csrf_token,
                'creation_mode': 'manual',
                'fio': 'Ещё Один Тестович',
                'dogovor': 'ВНУТРЕННИЙ-78',
                'year': '2026',
                'spec': 'ФМ',
                'base': '11',
                'manual_login': 'TAKEN_MANUAL',
            },
        )
        self.assertIn('Такой логин уже используется.', duplicate_login.get_data(as_text=True))

        exact_fio = client.post(
            '/manual_create',
            data={
                'csrf_token': csrf_token,
                'creation_mode': 'contract',
                'fio': '  иванов   иван иванович ',
                'dogovor': '2026-ФМ-0702-11',
            },
        )
        exact_fio_body = exact_fio.get_data(as_text=True)
        self.assertIn('В базе уже есть запись с точно таким ФИО.', exact_fio_body)
        self.assertIn('name="confirm_fio_duplicate"', exact_fio_body)

        with sqlite3.connect(manticore.DB_PATH) as conn:
            count_before_confirmation = conn.execute(
                'SELECT COUNT(*) FROM abiturients'
            ).fetchone()[0]
        self.assertEqual(count_before_confirmation, 1)

        confirmed = client.post(
            '/manual_create',
            data={
                'csrf_token': csrf_token,
                'creation_mode': 'contract',
                'fio': 'иванов иван иванович',
                'dogovor': '2026-ФМ-0702-11',
                'confirm_fio_duplicate': '1',
            },
        )
        self.assertIn('Логин успешно создан:', confirmed.get_data(as_text=True))
        with sqlite3.connect(manticore.DB_PATH) as conn:
            rows = conn.execute(
                'SELECT dogovor, login FROM abiturients ORDER BY id'
            ).fetchall()
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[1][0], '2026-ФМ-0702-11')
        self.assertNotEqual(rows[1][1].casefold(), 'taken_manual')

    def test_admin_is_redirected_to_setup_until_login_rules_are_saved(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute('DELETE FROM login_generation_settings')

        client = manticore.app.test_client()
        self.login_session(client)

        response = client.get('/file_work')
        self.assertEqual(response.status_code, 303)
        self.assertIn('/setup', response.headers['Location'])

        setup_response = client.get('/setup')
        csrf_token = self.csrf_from_response(setup_response)
        save_response = client.post('/setup', data={'csrf_token': csrf_token, 'mode': 'standard', 'require_enrollment_order': '1'})

        self.assertEqual(save_response.status_code, 302)
        self.assertTrue(manticore.is_login_generation_setup_completed())

    def test_login_rules_setup_can_toggle_enrollment_order_requirement(self):
        client = manticore.app.test_client()
        self.login_session(client)

        setup_response = client.get('/setup')
        setup_body = setup_response.get_data(as_text=True)
        self.assertIn('Требовать сверку по приказу перед переносом', setup_body)
        csrf_token = self.csrf_from_response(setup_response)
        disabled_response = client.post('/setup', data={'csrf_token': csrf_token, 'mode': 'standard'})
        self.assertEqual(disabled_response.status_code, 302)
        self.assertFalse(manticore.is_enrollment_order_required())

        setup_response = client.get('/setup')
        csrf_token = self.csrf_from_response(setup_response)
        enabled_response = client.post('/setup', data={'csrf_token': csrf_token, 'mode': 'standard', 'require_enrollment_order': '1'})
        self.assertEqual(enabled_response.status_code, 302)
        self.assertTrue(manticore.is_enrollment_order_required())

    def test_login_rules_setup_can_toggle_course_groups(self):
        client = manticore.app.test_client()
        self.login_session(client)

        setup_response = client.get('/setup')
        setup_body = setup_response.get_data(as_text=True)
        self.assertIn('Использовать глобальные группы курса', setup_body)
        csrf_token = self.csrf_from_response(setup_response)
        disabled_response = client.post(
            '/setup',
            data={'csrf_token': csrf_token, 'mode': 'standard', 'require_enrollment_order': '1', 'use_course_groups': '0'}
        )
        self.assertEqual(disabled_response.status_code, 302)
        self.assertFalse(manticore.are_course_groups_enabled())

        setup_response = client.get('/setup')
        csrf_token = self.csrf_from_response(setup_response)
        enabled_response = client.post(
            '/setup',
            data={
                'csrf_token': csrf_token,
                'mode': 'standard',
                'require_enrollment_order': '1',
                'use_course_groups': ['0', '1'],
            }
        )
        self.assertEqual(enabled_response.status_code, 302)
        self.assertTrue(manticore.are_course_groups_enabled())

    def test_abiturients_import_plan_uses_dogovor_and_warns_about_namesakes(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                ('Иванов Иван Иванович', '2026-ФМ-0001-11', '26611001', '2026', 'Иванов', 'Иван Иванович')
            )
            conn.execute(
                '''
                INSERT INTO students
                    (username, password, email, firstname, lastname, cohort1, source_campaign_year, source_dogovor, source_fio)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'student001', 'cron', 'student@example.test', 'Семен', 'Сидоров', '26ФМ-11-1',
                    '2026', '2026-ФМ-0002-11', 'Сидоров Семен Семенович'
                )
            )
            conn.execute(
                '''
                INSERT INTO pending_duplicates (fio, dogovor, login, campaign_year, fam, imotch)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                ('Дубров Павел Павлович', '2026-ФМ-0003-11', 'dubl001', '2026', 'Дубров', 'Павел Павлович')
            )

        file_path = self.make_abiturients_file([
            {'ФИО': 'Петров Петр Петрович', 'Договор': '2026-ФМ-0004-11'},
            {'ФИО': 'Иванов Иван Иванович', 'Договор': '2026-ФМ-0005-11'},
            {'ФИО': 'Андреев Андрей Андреевич', 'Договор': '2026-ФМ-0001-11'},
            {'ФИО': 'Семенов Семен Семенович', 'Договор': '2026-ФМ-0002-11'},
            {'ФИО': 'Дубров Павел Павлович', 'Договор': '2026-ФМ-0003-11'},
            {'ФИО': 'Повторов Петр Петрович', 'Договор': '2026-ФМ-0004-11'},
        ])

        plan_df, summary = manticore.build_abiturients_import_plan(file_path, '2026')

        self.assertEqual(summary['total'], 6)
        self.assertEqual(summary['ready_count'], 2)
        self.assertEqual(summary['duplicate_count'], 1)
        self.assertEqual(summary['conflict_count'], 3)
        self.assertEqual(summary['warning_count'], 1)
        self.assertEqual(
            plan_df['import_action'].tolist(),
            ['create', 'create', 'conflict', 'conflict', 'duplicate', 'conflict']
        )
        self.assertIn('Возможный тёзка', plan_df.iloc[1]['import_status'])
        self.assertIn('Договор уже есть у абитуриента', plan_df.iloc[2]['import_status'])
        self.assertIn('Договор уже есть у студента', plan_df.iloc[3]['import_status'])
        self.assertIn('Договор уже ожидает проверки', plan_df.iloc[4]['import_status'])
        self.assertIn('Договор повторяется', plan_df.iloc[5]['import_status'])
        self.assertIn('список абитуриентов', plan_df.iloc[1]['status_hint'])
        self.assertIn('2026-ФМ-0001-11', plan_df.iloc[1]['status_hint'])
        self.assertIn('26611001', plan_df.iloc[1]['status_hint'])
        self.assertIn('список студентов', plan_df.iloc[3]['status_hint'])
        self.assertIn('Сидоров Семен Семенович', plan_df.iloc[3]['status_hint'])
        self.assertIn('раздел дублей', plan_df.iloc[4]['status_hint'])
        self.assertIn('загруженный файл, строка 2', plan_df.iloc[5]['status_hint'])

        preview_rows = manticore.dataframe_preview_rows(plan_df)
        self.assertEqual(preview_rows[1]['status_hint'], plan_df.iloc[1]['status_hint'])

    def test_apply_abiturients_import_creates_backup_and_audit_log(self):
        file_path = self.make_abiturients_file([
            {'ФИО': 'Петров Петр Петрович', 'Договор': '2026 ФМ 11'},
            {'ФИО': 'Сидоров Семен Семенович', 'Договор': 'ошибка'},
        ])

        result_path, summary = manticore.apply_abiturients_import(file_path, '2026')

        self.assertTrue(os.path.exists(result_path))
        self.assertEqual(summary['ready_count'], 1)
        self.assertEqual(summary['conflict_count'], 1)
        backups = manticore.list_database_backups()
        self.assertTrue(any('before_abiturients_import' in backup['name'] for backup in backups))

        with sqlite3.connect(manticore.DB_PATH) as conn:
            abiturients_count = conn.execute('SELECT COUNT(*) FROM abiturients').fetchone()[0]
            conflicts_count = conn.execute('SELECT COUNT(*) FROM login_conflicts').fetchone()[0]
            audit_count = conn.execute(
                "SELECT COUNT(*) FROM audit_logs WHERE action='abiturients_import'"
            ).fetchone()[0]

        self.assertEqual(abiturients_count, 1)
        self.assertEqual(conflicts_count, 1)
        self.assertEqual(audit_count, 1)
        os.remove(result_path)

    def test_students_list_hides_password_for_non_admin(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO students (username, password, email, firstname, lastname, cohort1)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                ('student001', 'visible-secret', 'student@example.test', 'Петр', 'Петров', '26ФМ-11-1')
            )

        client = manticore.app.test_client()
        with client.session_transaction() as session:
            session['user'] = 'assistant'
            session['role'] = 'assistant'

        response = client.get('/students_list')
        body = response.get_data(as_text=True)

        self.assertEqual(response.status_code, 200)
        self.assertNotIn('visible-secret', body)
        self.assertIn('••••••', body)

    def test_documentation_is_available_to_authenticated_users(self):
        client = manticore.app.test_client()

        anonymous_response = client.get('/documentation')
        self.assertEqual(anonymous_response.status_code, 302)
        self.assertIn('/login', anonymous_response.headers['Location'])

        self.login_session(client, username='viewer', role='viewer')
        response = client.get('/documentation')
        body = response.get_data(as_text=True)

        self.assertEqual(response.status_code, 200)
        self.assertIn('Как работает', body)
        self.assertIn('Путь одной приёмной кампании', body)
        self.assertIn('Снимки сделаны на изолированном тестовом стенде', body)
        self.assertIn('dashboard-test-stand.png', body)
        self.assertIn('Форматы файлов', body)
        self.assertIn('documentation.css', body)
        self.assertIn('documentation.js', body)

    def test_abiturients_sort_links_keep_current_filters(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                ('Петров Петр Петрович', '2026-ФМ-0201-11', '26611201', '2026', 'Петров', 'Петр Петрович', 'petrov@example.test', 1)
            )
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                ('Иванов Иван Иванович', '2026-ФМ-0202-11', '26611202', '2026', 'Иванов', 'Иван Иванович', '', 1)
            )

        client = manticore.app.test_client()
        self.login_session(client)
        response = client.get('/abiturients?has_email=1&has_paid=1&order_by=fam&order_dir=asc')
        body = unquote(response.get_data(as_text=True).replace('&amp;', '&'))

        self.assertEqual(response.status_code, 200)
        self.assertIn('Петров Петр Петрович', body)
        self.assertNotIn('Иванов Иван Иванович', body)
        self.assertRegex(body, r'href="/abiturients\?[^"]*has_email=1[^"]*has_paid=1[^"]*order_by=fio')

    def test_abiturients_search_is_case_insensitive_for_cyrillic(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.executemany(
                '''
                INSERT INTO abiturients
                    (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, '2026', ?, ?, ?, 1)
                ''',
                [
                    (
                        'Хачатурян Нарина Герасимовна', '2026-СД-0241-11И', 'search001',
                        'Хачатурян', 'Нарина Герасимовна', 'narina@example.test'
                    ),
                    (
                        'Иванов Иван Иванович', '2026-ФМ-0242-11', 'search002',
                        'Иванов', 'Иван Иванович', 'ivanov@example.test'
                    ),
                ]
            )

        lowercase_rows = manticore.get_all_abiturients(
            campaign_year='2026',
            q='хачатурян',
        )
        uppercase_rows = manticore.get_all_abiturients(
            campaign_year='2026',
            q='ХАЧАТУРЯН',
        )
        self.assertEqual([row['login'] for row in lowercase_rows], ['search001'])
        self.assertEqual([row['login'] for row in uppercase_rows], ['search001'])

        client = manticore.app.test_client()
        self.login_session(client)
        response = client.get('/abiturients', query_string={'q': 'хачатурян'})
        body = response.get_data(as_text=True)

        self.assertEqual(response.status_code, 200)
        self.assertIn('Хачатурян Нарина Герасимовна', body)
        self.assertNotIn('Иванов Иван Иванович', body)

        export_response = client.get(
            '/abiturients/download',
            query_string={'q': 'хачатурян'},
        )
        self.assertEqual(export_response.status_code, 200)
        exported = pd.read_excel(io.BytesIO(export_response.data))
        self.assertEqual(exported['login'].tolist(), ['search001'])

    def test_abiturients_can_filter_and_export_by_enrollment_order_presence(self):
        ordered_fio = 'Приказов Петр Петрович'
        ordered_dogovor = '2026-ФМ-0250-11'
        fio_key, specialty_key = manticore.make_abiturient_enrollment_match_key(
            ordered_fio,
            ordered_dogovor,
        )
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.executemany(
                '''
                INSERT INTO abiturients
                    (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, '2026', ?, ?, ?, 1)
                ''',
                [
                    (
                        ordered_fio, ordered_dogovor, 'order_yes', 'Приказов',
                        'Петр Петрович', 'ordered@example.test'
                    ),
                    (
                        'Безприказов Иван Иванович', '2026-ФМ-0251-11', 'order_no',
                        'Безприказов', 'Иван Иванович', 'unordered@example.test'
                    ),
                ]
            )
            conn.execute(
                '''
                INSERT INTO enrollment_orders
                    (campaign_year, fio, specialty, order_number, order_date,
                     fio_key, specialty_key)
                VALUES ('2026', ?, 'ФМ', '88-у', '2026-08-15', ?, ?)
                ''',
                (ordered_fio, fio_key, specialty_key)
            )

        with_order = manticore.get_all_abiturients(
            campaign_year='2026',
            has_order='1',
        )
        without_order = manticore.get_all_abiturients(
            campaign_year='2026',
            has_order='0',
        )
        self.assertEqual([row['login'] for row in with_order], ['order_yes'])
        self.assertEqual([row['login'] for row in without_order], ['order_no'])

        client = manticore.app.test_client()
        self.login_session(client)
        response = client.get(
            '/abiturients',
            query_string={'has_order': '1', 'order_by': 'fio', 'order_dir': 'asc'},
        )
        body = unquote(response.get_data(as_text=True).replace('&amp;', '&'))

        self.assertEqual(response.status_code, 200)
        self.assertIn('Есть приказ', body)
        self.assertIn('value="1" selected', body)
        self.assertIn(ordered_fio, body)
        self.assertNotIn('Безприказов Иван Иванович', body)
        self.assertRegex(body, r'href="/abiturients\?[^\"]*has_order=1[^\"]*order_by=dogovor')
        self.assertIn('name="has_order" value="1"', body)

        export_response = client.get(
            '/abiturients/download',
            query_string={'has_order': '1', 'order_by': 'fio', 'order_dir': 'asc'},
        )
        self.assertEqual(export_response.status_code, 200)
        exported = pd.read_excel(io.BytesIO(export_response.data))
        self.assertEqual(exported['login'].tolist(), ['order_yes'])

    def test_edit_abiturient_navigation_follows_filtered_list_order(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.executemany(
                '''
                INSERT INTO abiturients
                    (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                [
                    ('Альфа Навигация Тест', '2026-ФМ-0301-11', 'nav001', '2026', 'Альфа', 'Навигация Тест', '', 0),
                    ('Бета Навигация Тест', '2026-ФМ-0302-11', 'nav002', '2026', 'Бета', 'Навигация Тест', '', 0),
                    ('Гамма Навигация Тест', '2026-ФМ-0303-11', 'nav003', '2026', 'Гамма', 'Навигация Тест', '', 0),
                    ('Скрытый Абитуриент', '2026-ФМ-0304-11', 'nav004', '2026', 'Скрытый', 'Абитуриент', '', 0),
                ]
            )

        client = manticore.app.test_client()
        self.login_session(client)
        list_query = 'campaign_year=2026&order_by=fio&order_dir=asc&q=Навигация'

        list_response = client.get(f'/abiturients?{list_query}')
        list_body = unquote(list_response.get_data(as_text=True).replace('&amp;', '&'))
        self.assertIn('/edit_abiturient/nav002?', list_body)
        self.assertIn('order_by=fio', list_body)
        self.assertIn('order_dir=asc', list_body)
        self.assertIn('q=Навигация', list_body)

        middle_response = client.get(f'/edit_abiturient/nav002?{list_query}')
        middle_body = unquote(middle_response.get_data(as_text=True).replace('&amp;', '&'))
        self.assertEqual(middle_response.status_code, 200)
        self.assertIn('2 из 3', middle_body)
        self.assertIn('id="edit-abiturient-previous"', middle_body)
        self.assertIn('/edit_abiturient/nav001?', middle_body)
        self.assertIn('id="edit-abiturient-next"', middle_body)
        self.assertIn('/edit_abiturient/nav003?', middle_body)
        self.assertNotIn('/edit_abiturient/nav004?', middle_body)
        self.assertIn('q=Навигация', middle_body)
        self.assertIn('name="save_action" value="stay"', middle_body)
        self.assertIn('>Сохранить</button>', middle_body)
        self.assertIn('name="save_action" value="exit"', middle_body)
        self.assertIn('>Сохранить и выйти</button>', middle_body)
        self.assertLess(
            middle_body.index('id="edit-abiturient-previous"'),
            middle_body.index('id="edit-abiturient-next"')
        )
        self.assertLess(
            middle_body.index('id="edit-abiturient-next"'),
            middle_body.index('2 из 3')
        )

        first_response = client.get(f'/edit_abiturient/nav001?{list_query}')
        first_body = first_response.get_data(as_text=True)
        self.assertIn('1 из 3', first_body)
        self.assertNotIn('id="edit-abiturient-previous"', first_body)
        self.assertIn('Это первый абитуриент в списке', first_body)
        self.assertIn('id="edit-abiturient-next"', first_body)

        csrf_token = self.csrf_from_response(middle_response)
        stay_response = client.post(
            f'/edit_abiturient/nav002?{list_query}',
            data={
                'csrf_token': csrf_token,
                'campaign_year': '2026',
                'fio': 'Бета Навигация Тест',
                'email': 'beta@example.test',
                'login': 'nav002',
                'comment': 'Сохранено без выхода',
                'save_action': 'stay',
            }
        )
        stay_location = unquote(stay_response.headers['Location'])
        self.assertEqual(stay_response.status_code, 303)
        self.assertIn('/edit_abiturient/nav002?', stay_location)
        self.assertIn('order_by=fio', stay_location)
        self.assertIn('order_dir=asc', stay_location)
        self.assertIn('q=Навигация', stay_location)

        stayed_response = client.get(stay_response.headers['Location'])
        stayed_body = stayed_response.get_data(as_text=True)
        self.assertIn('value="beta@example.test"', stayed_body)
        self.assertIn('Сохранено без выхода', stayed_body)

        csrf_token = self.csrf_from_response(stayed_response)
        save_response = client.post(
            f'/edit_abiturient/nav002?{list_query}',
            data={
                'csrf_token': csrf_token,
                'campaign_year': '2026',
                'fio': 'Бета Навигация Тест',
                'email': 'beta@example.test',
                'login': 'nav002',
                'comment': '',
                'save_action': 'exit',
            }
        )
        save_location = unquote(save_response.headers['Location'])
        self.assertEqual(save_response.status_code, 302)
        self.assertIn('/abiturients?', save_location)
        self.assertIn('order_by=fio', save_location)
        self.assertIn('order_dir=asc', save_location)
        self.assertIn('q=Навигация', save_location)

    def test_withdraw_abiturient_documents_uses_next_free_del_login(self):
        original_login = '26311077'
        with sqlite3.connect(manticore.DB_PATH) as conn:
            cur = conn.execute(
                '''
                INSERT INTO abiturients
                    (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Отзывов Олег Олегович', '2026-СД-0077-11', original_login,
                    '2026', 'Отзывов', 'Олег Олегович', 'withdraw@example.test', 1
                )
            )
            abiturient_id = cur.lastrowid
            conn.execute(
                '''
                INSERT INTO pending_duplicates (fio, dogovor, login, campaign_year, fam, imotch)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                ('Первый Занятый', '2026-СД-0177-11', f'del1_{original_login}', '2026', 'Первый', 'Занятый')
            )
            conn.execute(
                '''
                INSERT INTO students (username, password, email, firstname, lastname, cohort1)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                (f'del2_{original_login}', 'secret', 'student@example.test', 'Занятый', 'Второй', '26СД-11-1')
            )
            manticore.sync_enrollment_candidates_for_abiturients(conn, [abiturient_id], '2026')
            candidate_count = conn.execute(
                'SELECT COUNT(*) FROM enrollment_candidates WHERE abiturient_id=?',
                (abiturient_id,)
            ).fetchone()[0]
        self.assertEqual(candidate_count, 1)

        client = manticore.app.test_client()
        self.login_session(client)
        list_response = client.get('/abiturients?campaign_year=2026&q=Отзывов&order_by=fio&order_dir=asc')
        list_body = list_response.get_data(as_text=True)
        self.assertIn('class="row-action-item is-warning"', list_body)
        self.assertIn('Отозвать документы', list_body)
        csrf_token = self.csrf_from_response(list_response)

        response = client.post(
            '/abiturients/withdraw-documents',
            data={
                'csrf_token': csrf_token,
                'id': str(abiturient_id),
                'campaign_year': '2026',
                'q': 'Отзывов',
                'order_by': 'fio',
                'order_dir': 'asc',
            }
        )

        self.assertEqual(response.status_code, 303)
        response_location = unquote(response.headers['Location'])
        self.assertIn('/abiturients?', response_location)
        self.assertIn('q=Отзывов', response_location)
        self.assertIn('order_by=fio', response_location)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            stored_login = conn.execute(
                'SELECT login FROM abiturients WHERE id=?',
                (abiturient_id,)
            ).fetchone()[0]
            candidate_count = conn.execute(
                'SELECT COUNT(*) FROM enrollment_candidates WHERE abiturient_id=?',
                (abiturient_id,)
            ).fetchone()[0]
            audit_count = conn.execute(
                "SELECT COUNT(*) FROM audit_logs WHERE action='abiturient_documents_withdrawn' AND entity_id=?",
                (str(abiturient_id),)
            ).fetchone()[0]
        self.assertEqual(stored_login, f'del3_{original_login}')
        self.assertEqual(candidate_count, 0)
        self.assertEqual(audit_count, 1)
        self.assertFalse(manticore.is_login_exists(original_login, '2026'))

        result_page = client.get(response.headers['Location'])
        result_body = result_page.get_data(as_text=True)
        self.assertIn(f'del3_{original_login}', result_body)
        self.assertIn('Документы уже отозваны', result_body)

    def test_withdrawn_abiturients_are_informational_not_data_issues(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients
                    (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Архивный Абитуриент', '2026-ФМ-0088-11', 'del_26611088',
                    '2026', 'Архивный', 'Абитуриент', '', 0
                )
            )

        with manticore.app.test_request_context('/data_checks?campaign_year=2026'):
            report = manticore.get_data_quality_report('2026')
            dashboard = manticore.get_dashboard_data('2026')

        withdrawn_section = next(
            section for section in report['sections']
            if section['title'] == 'Отозванные документы'
        )
        self.assertTrue(withdrawn_section['informational'])
        self.assertEqual(withdrawn_section['checks'][0]['count'], 1)
        self.assertEqual(withdrawn_section['checks'][0]['tone'], 'info')
        self.assertEqual(report['total_issues'], 0)
        self.assertEqual(dashboard['abiturients_total'], 1)
        self.assertEqual(dashboard['no_email'], 0)
        self.assertEqual(dashboard['unpaid'], 0)
        self.assertEqual(dashboard['ready'], 0)

        withdrawn_rows = manticore.get_all_abiturients(campaign_year='2026', withdrawn='1')
        active_rows = manticore.get_all_abiturients(campaign_year='2026', withdrawn='0')
        self.assertEqual([row['login'] for row in withdrawn_rows], ['del_26611088'])
        self.assertEqual(active_rows, [])

        client = manticore.app.test_client()
        self.login_session(client)
        checks_response = client.get('/data_checks?campaign_year=2026')
        checks_body = checks_response.get_data(as_text=True)
        self.assertIn('Отозванные документы', checks_body)
        self.assertIn('Справочно', checks_body)
        self.assertIn('Абитуриенты, отозвавшие документы', checks_body)
        self.assertIn('0</strong>', checks_body)

    def test_students_sort_links_keep_current_filters(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO students (username, password, email, firstname, lastname, cohort1)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                ('student_p', 'secret', 'p@example.test', 'Петр', 'Петров', '26ФМ-11-1')
            )
            conn.execute(
                '''
                INSERT INTO students (username, password, email, firstname, lastname, cohort1)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                ('student_i', 'secret', 'i@example.test', 'Иван', 'Иванов', '26СД-9-1')
            )

        client = manticore.app.test_client()
        self.login_session(client)
        response = client.get('/students_list?cohort=26ФМ-11-1&lastname=П&order_by=username&order_dir=asc')
        body = unquote(response.get_data(as_text=True).replace('&amp;', '&'))

        self.assertEqual(response.status_code, 200)
        self.assertIn('Петров', body)
        self.assertNotIn('Иванов', body)
        self.assertRegex(body, r'href="/students_list\?[^"]*cohort=26ФМ-11-1[^"]*lastname=П[^"]*order_by=lastname')

    def test_students_can_filter_sort_and_export_by_enrollment_order(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            students = [
                ('student_late', 'Поздний', 'Петр'),
                ('student_early', 'Ранний', 'Иван'),
                ('student_without_order', 'Безприказов', 'Семен'),
            ]
            for username, lastname, firstname in students:
                conn.execute(
                    '''
                    INSERT INTO students (username, password, email, firstname, lastname, cohort1)
                    VALUES (?, 'secret', ?, ?, ?, '26ФМ-11-1')
                    ''',
                    (username, f'{username}@example.test', firstname, lastname)
                )
            conn.execute(
                '''
                INSERT INTO student_group_transfers
                    (username, movement_type, old_cohort1, new_cohort1,
                     order_number, order_date, order_source)
                VALUES (?, 'enrollment', 'Абитуриенты', '26ФМ-11-1', ?, ?, 'enrollment_order')
                ''',
                ('student_late', '10-у', '2026-08-20')
            )
            conn.execute(
                '''
                INSERT INTO student_group_transfers
                    (username, movement_type, old_cohort1, new_cohort1,
                     order_number, order_date, order_source)
                VALUES (?, 'enrollment', 'Абитуриенты', '26ФМ-11-1', ?, ?, 'enrollment_order')
                ''',
                ('student_early', '2-у', '2026-08-10')
            )

        ascending = manticore.get_all_students('enrollment_order', 'asc')
        descending = manticore.get_all_students('enrollment_order', 'desc')
        self.assertEqual(
            [row['username'] for row in ascending],
            ['student_early', 'student_late', 'student_without_order']
        )
        self.assertEqual(
            [row['username'] for row in descending],
            ['student_late', 'student_early', 'student_without_order']
        )

        selected_order = manticore.make_enrollment_order_filter_value('10-у', '2026-08-20')
        client = manticore.app.test_client()
        self.login_session(client)
        response = client.get(
            '/students_list',
            query_string={
                'enrollment_order': selected_order,
                'order_by': 'enrollment_order',
                'order_dir': 'desc',
            }
        )
        body = response.get_data(as_text=True)

        self.assertEqual(response.status_code, 200)
        self.assertIn('Приказ о зачислении', body)
        self.assertIn('№ 10-у', body)
        self.assertIn('Поздний', body)
        self.assertNotIn('Ранний', body)
        self.assertNotIn('Безприказов', body)
        self.assertIn('selected', body)
        self.assertIn('class="row-action-trigger"', body)
        self.assertIn('class="row-action-item"', body)
        self.assertIn('class="row-action-item is-danger"', body)
        self.assertNotIn('action-icon-btn', body)

        export_response = client.get(
            '/students/download',
            query_string={
                'enrollment_order': selected_order,
                'order_by': 'enrollment_order',
                'order_dir': 'desc',
            }
        )
        self.assertEqual(export_response.status_code, 200)
        exported = pd.read_excel(io.BytesIO(export_response.data))
        self.assertEqual(exported['username'].tolist(), ['student_late'])
        self.assertEqual(exported.loc[0, 'enrollment_order_number'], '10-у')
        self.assertEqual(exported.loc[0, 'enrollment_order_date'], '2026-08-20')

    def test_students_list_backfills_enrollment_order_for_existing_student(self):
        fio = 'Архивов Алексей Алексеевич'
        dogovor = '2026-ФМ-0800-11'
        fio_key, specialty_key = manticore.make_abiturient_enrollment_match_key(fio, dogovor)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO students
                    (username, password, email, firstname, lastname, cohort1,
                     source_campaign_year, source_dogovor, source_fio)
                VALUES (?, 'secret', ?, ?, ?, ?, '2026', ?, ?)
                ''',
                (
                    'student_from_old_version', 'archive@example.test', 'Алексей Алексеевич',
                    'Архивов', '26ФМ-11-1', dogovor, fio
                )
            )
            conn.execute(
                '''
                INSERT INTO enrollment_orders
                    (campaign_year, fio, specialty, group_name, order_number, order_date,
                     fio_key, specialty_key, group_key)
                VALUES ('2026', ?, 'ФМ', '26ФМ-11-1', '77-у', '2026-08-01', ?, ?, ?)
                ''',
                (fio, fio_key, specialty_key, '')
            )

        rows = manticore.get_all_students(order_by='enrollment_order')

        self.assertEqual(rows[0]['enrollment_order_number'], '77-у')
        self.assertEqual(rows[0]['enrollment_order_date'], '2026-08-01')
        with sqlite3.connect(manticore.DB_PATH) as conn:
            movement = conn.execute(
                '''
                SELECT order_number, order_date
                FROM student_group_transfers
                WHERE username=? AND movement_type='enrollment'
                ''',
                ('student_from_old_version',)
            ).fetchone()
        self.assertEqual(movement, ('77-у', '2026-08-01'))

    def test_admin_backup_and_audit_pages_render(self):
        backup_path = manticore.create_database_backup('page_render_test')
        self.assertTrue(os.path.exists(backup_path))

        client = manticore.app.test_client()
        with client.session_transaction() as session:
            session['user'] = 'admin'
            session['role'] = 'admin'

        backups_response = client.get('/backups')
        audit_response = client.get('/audit_logs')

        self.assertEqual(backups_response.status_code, 200)
        self.assertEqual(audit_response.status_code, 200)
        self.assertIn('page_render_test', backups_response.get_data(as_text=True))

    def test_admin_panel_contains_server_update_controls(self):
        client = manticore.app.test_client()
        self.login_session(client)

        response = client.get('/admin_panel')
        body = response.get_data(as_text=True)

        self.assertEqual(response.status_code, 200)
        self.assertIn('Обновление серверной панели', body)
        self.assertIn('/admin/app-update/status', body)
        self.assertIn('/admin/app-update/start', body)
        self.assertIn('id="desktop-install-update-button"', body)
        self.assertIn('install_approved_update', body)

    @mock.patch.object(manticore.update_app, 'fetch_latest_release')
    def test_admin_can_check_latest_release(self, fetch_latest_release):
        fetch_latest_release.return_value = {
            'tag_name': 'v9.9.9',
            'name': 'manticore 9.9.9',
            'html_url': 'https://github.com/UnicornisIT/manticore/releases/tag/v9.9.9',
            'published_at': '2026-08-30T10:00:00Z',
        }
        client = manticore.app.test_client()
        self.login_session(client)

        response = client.get('/admin/app-update/status')
        payload = response.get_json()

        self.assertEqual(response.status_code, 200)
        self.assertEqual(payload['latest_version'], '9.9.9')
        self.assertTrue(payload['update_available'])

    def test_admin_can_queue_server_update_with_csrf_protection(self):
        client = manticore.app.test_client()
        self.login_session(client)
        page = client.get('/admin_panel')
        csrf_token = self.csrf_from_response(page)

        with mock.patch.object(manticore, 'get_app_update_status', return_value={'state': 'idle'}), \
                mock.patch.object(manticore, 'queue_app_update') as queue_update:
            response = client.post(
                '/admin/app-update/start',
                data={'csrf_token': csrf_token},
            )

        self.assertEqual(response.status_code, 303)
        queue_update.assert_called_once_with('admin')

    def test_file_work_sections_render_individual_pages(self):
        client = manticore.app.test_client()
        self.login_session(client)

        common_response = client.get('/file_work')
        common_body = common_response.get_data(as_text=True)
        orders_response = client.get('/file_work/orders')
        orders_body = orders_response.get_data(as_text=True)
        invalid_response = client.get('/file_work/unknown')

        self.assertEqual(common_response.status_code, 200)
        self.assertIn('id="file-section-abiturients"', common_body)
        self.assertIn('id="file-section-updates"', common_body)
        self.assertIn('id="file-section-orders"', common_body)
        self.assertIn('id="file-section-students"', common_body)
        self.assertIn('/file_work/orders', common_body)
        self.assertEqual(common_body.count('data-file-dropzone'), 5)
        self.assertIn('btn-download upload-template-link', common_body)
        self.assertIn('или перетащите сюда Excel/CSV/DOCX/PDF', common_body)

        self.assertEqual(orders_response.status_code, 200)
        self.assertIn('id="file-section-orders"', orders_body)
        self.assertNotIn('id="file-section-abiturients"', orders_body)
        self.assertNotIn('id="file-section-updates"', orders_body)
        self.assertNotIn('id="file-section-students"', orders_body)
        self.assertEqual(orders_body.count('data-file-dropzone'), 1)
        self.assertIn('btn-download upload-template-link', orders_body)

        self.assertEqual(invalid_response.status_code, 303)
        self.assertIn('/file_work', invalid_response.headers['Location'])

    def test_file_work_preview_renders_without_writing_to_database(self):
        client = manticore.app.test_client()
        self.login_session(client)

        get_response = client.get('/file_work')
        csrf_token = self.csrf_from_response(get_response)

        csv_bytes = 'ФИО,Договор\nПетров Петр Петрович,2026 ФМ 11\n'.encode('utf-8-sig')
        response = client.post(
            '/file_work',
            data={
                'csrf_token': csrf_token,
                'import_action': 'preview',
                'file': (io.BytesIO(csv_bytes), 'abiturients.csv'),
            },
            content_type='multipart/form-data'
        )
        body = response.get_data(as_text=True)

        self.assertEqual(response.status_code, 200)
        self.assertIn('Предпросмотр импорта абитуриентов', body)
        self.assertIn('Подтвердить импорт', body)
        self.assertIn('class="abiturients-preview-table"', body)
        self.assertIn('import-status-badge', body)
        self.assertIn('Наведите указатель на статус', body)
        self.assertIn('Расхождений с существующими записями', body)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            abiturients_count = conn.execute('SELECT COUNT(*) FROM abiturients').fetchone()[0]
        self.assertEqual(abiturients_count, 0)

    def test_file_work_preview_shows_friendly_row_report(self):
        client = manticore.app.test_client()
        self.login_session(client)

        get_response = client.get('/file_work')
        csrf_token = self.csrf_from_response(get_response)

        csv_bytes = (
            'ФИО,Договор\n'
            ',2026 ФМ 11\n'
            'Иванов Иван Иванович,не договор\n'
        ).encode('utf-8-sig')
        response = client.post(
            '/file_work',
            data={
                'csrf_token': csrf_token,
                'import_action': 'preview',
                'file': (io.BytesIO(csv_bytes), 'abiturients.csv'),
            },
            content_type='multipart/form-data'
        )
        body = response.get_data(as_text=True)

        self.assertEqual(response.status_code, 200)
        self.assertIn('Отчет по проверке абитуриентов', body)
        self.assertIn('Строка 2', body)
        self.assertIn('Не заполнено ФИО', body)
        self.assertIn('Строка 3', body)
        self.assertIn('Не удалось разобрать номер договора', body)

    def test_file_work_preview_shows_all_rows(self):
        client = manticore.app.test_client()
        self.login_session(client)

        get_response = client.get('/file_work')
        csrf_token = self.csrf_from_response(get_response)

        rows = ['ФИО,Договор']
        for index in range(1, 26):
            rows.append(f'Фамилия{index} Имя{index} Отчество{index},2026-ФМ-{index:04d}-11')
        csv_bytes = ('\n'.join(rows) + '\n').encode('utf-8-sig')
        response = client.post(
            '/file_work',
            data={
                'csrf_token': csrf_token,
                'import_action': 'preview',
                'file': (io.BytesIO(csv_bytes), 'abiturients.csv'),
            },
            content_type='multipart/form-data'
        )
        body = response.get_data(as_text=True)

        self.assertEqual(response.status_code, 200)
        self.assertIn('В таблице показаны все строки файла', body)
        self.assertNotIn('Показаны первые 20 строк', body)
        self.assertIn('Фамилия25 Имя25 Отчество25', body)

    def test_file_work_confirm_redirects_to_clean_page_before_result_download(self):
        client = manticore.app.test_client()
        self.login_session(client)

        get_response = client.get('/file_work/abiturients')
        csrf_token = self.csrf_from_response(get_response)
        csv_bytes = 'ФИО,Договор\nПетров Петр Петрович,2026-ФМ-0041-11\n'.encode('utf-8-sig')
        preview_response = client.post(
            '/file_work/abiturients',
            data={
                'csrf_token': csrf_token,
                'import_action': 'preview',
                'file_section': 'abiturients',
                'file': (io.BytesIO(csv_bytes), 'abiturients.csv'),
            },
            content_type='multipart/form-data'
        )
        preview_body = preview_response.get_data(as_text=True)
        pending_match = re.search(r'name="pending_import" value="([^"]+)"', preview_body)
        self.assertIsNotNone(pending_match)
        pending_token = pending_match.group(1)
        pending_path = os.path.join(TEST_UPLOAD_DIR, pending_token)
        self.assertTrue(os.path.exists(pending_path))

        confirm_response = client.post(
            '/file_work/abiturients',
            data={
                'csrf_token': csrf_token,
                'import_action': 'confirm',
                'file_section': 'abiturients',
                'pending_import': pending_token,
                'campaign_year': '2026',
            }
        )

        self.assertEqual(confirm_response.status_code, 303)
        self.assertTrue(confirm_response.headers['Location'].endswith('/file_work/abiturients'))
        self.assertFalse(os.path.exists(pending_path))
        with sqlite3.connect(manticore.DB_PATH) as conn:
            self.assertEqual(conn.execute('SELECT COUNT(*) FROM abiturients').fetchone()[0], 1)
        with client.session_transaction() as session:
            result_token = session.get(manticore.ABITURIENTS_IMPORT_RESULT_SESSION_KEY)
        self.assertTrue(result_token)
        result_path = os.path.join(TEST_UPLOAD_DIR, result_token)
        self.assertTrue(os.path.exists(result_path))

        clean_response = client.get(confirm_response.headers['Location'])
        clean_body = clean_response.get_data(as_text=True)
        self.assertEqual(clean_response.status_code, 200)
        self.assertNotIn('Предпросмотр импорта абитуриентов', clean_body)
        self.assertIn('Импорт завершён', clean_body)
        self.assertIn('id="abiturients-import-result-download"', clean_body)

        download_response = client.get('/file_work/abiturients/import-result')
        self.assertEqual(download_response.status_code, 200)
        self.assertIn('attachment', download_response.headers.get('Content-Disposition', ''))
        self.assertTrue(download_response.data.startswith(b'PK'))
        self.assertFalse(os.path.exists(result_path))
        with client.session_transaction() as session:
            self.assertNotIn(manticore.ABITURIENTS_IMPORT_RESULT_SESSION_KEY, session)

    def test_file_work_cancel_redirects_to_page_without_preview(self):
        client = manticore.app.test_client()
        self.login_session(client)

        get_response = client.get('/file_work/abiturients')
        csrf_token = self.csrf_from_response(get_response)
        csv_bytes = 'ФИО,Договор\nПетров Петр Петрович,2026-ФМ-0042-11\n'.encode('utf-8-sig')
        preview_response = client.post(
            '/file_work/abiturients',
            data={
                'csrf_token': csrf_token,
                'import_action': 'preview',
                'file_section': 'abiturients',
                'file': (io.BytesIO(csv_bytes), 'abiturients.csv'),
            },
            content_type='multipart/form-data'
        )
        pending_match = re.search(
            r'name="pending_import" value="([^"]+)"',
            preview_response.get_data(as_text=True)
        )
        self.assertIsNotNone(pending_match)
        pending_token = pending_match.group(1)
        pending_path = os.path.join(TEST_UPLOAD_DIR, pending_token)

        cancel_response = client.post(
            '/file_work/abiturients',
            data={
                'csrf_token': csrf_token,
                'import_action': 'cancel',
                'file_section': 'abiturients',
                'pending_import': pending_token,
            }
        )

        self.assertEqual(cancel_response.status_code, 303)
        self.assertFalse(os.path.exists(pending_path))
        clean_response = client.get(cancel_response.headers['Location'])
        clean_body = clean_response.get_data(as_text=True)
        self.assertNotIn('Предпросмотр импорта абитуриентов', clean_body)
        self.assertIn('Предпросмотр импорта отменён', clean_body)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            self.assertEqual(conn.execute('SELECT COUNT(*) FROM abiturients').fetchone()[0], 0)

    def test_dashboard_search_and_person_card_render(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                ('Петров Петр Петрович', '2026-ФМ-0100-11', '26611010', '2026', 'Петров', 'Петр Петрович', 'p@example.test', 1)
            )
            conn.execute(
                'INSERT OR IGNORE INTO groups (name, group_year) VALUES (?, ?)',
                ('99ZZZ-DASHBOARD-END-1', '2026')
            )
            abiturient_id = conn.execute('SELECT id FROM abiturients WHERE login=?', ('26611010',)).fetchone()[0]

        client = manticore.app.test_client()
        self.login_session(client)

        dashboard_response = client.get('/')
        search_response = client.get('/search?q=0100')
        overlay_response = client.get('/search_overlay?q=0100')
        card_response = client.get(f'/person/abiturient/{abiturient_id}')
        wizard_response = client.get('/migration_wizard')

        self.assertEqual(dashboard_response.status_code, 200)
        dashboard_body = dashboard_response.get_data(as_text=True)
        self.assertIn('Панель состояния', dashboard_body)
        self.assertNotIn('dashboard-search', dashboard_body)
        self.assertNotIn('quick-actions', dashboard_body)
        self.assertIn('aria-label="Поиск"', dashboard_body)
        self.assertNotIn('>Поиск</a>', dashboard_body)
        self.assertIn('nav-search', dashboard_body)
        self.assertIn('global-search-modal', dashboard_body)
        self.assertIn('data-search-overlay-url="/search_overlay"', dashboard_body)
        self.assertIn('/static/js/app.js', dashboard_body)
        self.assertIn('global-search-palette-input', dashboard_body)
        self.assertIn('Мастер миграции', dashboard_body)
        self.assertIn('Контингент', dashboard_body)
        self.assertIn('Операции', dashboard_body)
        self.assertIn('Работа с файлами', dashboard_body)
        self.assertIn('Проверка данных', dashboard_body)
        self.assertIn('Ручное создание логина', dashboard_body)
        self.assertIn('nav-dropdown', dashboard_body)
        self.assertIn('/static/css/modern.css', dashboard_body)
        self.assertNotIn('nav-menu', dashboard_body)
        self.assertNotIn('Что требует внимания', dashboard_body)
        self.assertIn('Центр задач', dashboard_body)
        self.assertIn('Готовы к кандидатам', dashboard_body)
        self.assertIn('Кандидатов к зачислению', dashboard_body)
        self.assertIn('Сверены приказом', dashboard_body)
        self.assertIn('Полная проверка данных', dashboard_body)
        self.assertNotIn('<span>Без почты</span>', dashboard_body)
        self.assertNotIn('<span>Не оплачены</span>', dashboard_body)
        self.assertNotIn('<span>Конфликтов</span>', dashboard_body)
        self.assertNotIn('<span>В дублях</span>', dashboard_body)
        self.assertIn('dashboard-groups-scroll', dashboard_body)
        self.assertIn('99ZZZ-DASHBOARD-END-1', dashboard_body)
        self.assertEqual(search_response.status_code, 200)
        self.assertIn('Петров Петр Петрович', search_response.get_data(as_text=True))
        self.assertEqual(overlay_response.status_code, 200)
        overlay_data = overlay_response.get_json()
        self.assertEqual(overlay_data['query'], '0100')
        self.assertTrue(any(item['title'] == 'Петров Петр Петрович' for item in overlay_data['results']))
        self.assertEqual(card_response.status_code, 200)
        card_body = card_response.get_data(as_text=True)
        self.assertIn('Карточка абитуриента', card_body)
        self.assertIn('Номер договора', card_body)
        self.assertIn('Логин Moodle', card_body)
        self.assertIn('Договор оплачен', card_body)
        self.assertIn('2026-ФМ-0100-11', card_body)
        self.assertNotIn('>fio<', card_body)
        self.assertNotIn('>dogovor<', card_body)
        self.assertNotIn('>paid<', card_body)
        self.assertEqual(wizard_response.status_code, 200)
        wizard_body = wizard_response.get_data(as_text=True)
        self.assertIn('Мастер миграции', wizard_body)
        self.assertIn('1. Подготовить кандидатов', wizard_body)
        self.assertIn('2. Загрузить приказ', wizard_body)
        self.assertIn('3. Перенести в студенты', wizard_body)
        self.assertIn('Академические группы', wizard_body)
        self.assertIn('Дублирующие записи студентов', wizard_body)
        self.assertIn('Дублирующие записи абитуриентов', wizard_body)
        self.assertNotIn('Что проверить перед миграцией', wizard_body)

    def test_data_checks_page_groups_actionable_issues(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                ('Иванов Иван Иванович', '2026-ФМ-0200-11', '26611020', '2026', 'Иванов', 'Иван Иванович', '', 0)
            )
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                ('Петров Петр Петрович', '', '26611021', '2026', 'Петров', 'Петр Петрович', 'bad-email', 1)
            )
            conn.execute(
                '''
                INSERT INTO students (username, password, email, firstname, lastname, cohort1, source_campaign_year, source_dogovor)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                ('student020', 'secret', 'wrong-mail', 'Анна', 'Кривец', '', '2026', '')
            )
            conn.execute(
                '''
                INSERT INTO pending_duplicates (fio, dogovor, login, campaign_year, fam, imotch)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                ('Дубль Дмитрий', '2026-СД-0201-11', '26611022', '2026', 'Дубль', 'Дмитрий')
            )
            conn.execute(
                '''
                INSERT INTO login_conflicts (fio, dogovor, login, campaign_year, fam, imotch)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                ('Конфликт Константин', '2026-СД-0202-11', 'error000', '2026', 'Конфликт', 'Константин')
            )

        client = manticore.app.test_client()
        self.login_session(client)

        dashboard_response = client.get('/')
        checks_response = client.get('/data_checks')

        self.assertEqual(dashboard_response.status_code, 200)
        dashboard_body = dashboard_response.get_data(as_text=True)
        self.assertIn('Центр задач', dashboard_body)
        self.assertIn('Абитуриенты без почты', dashboard_body)
        self.assertIn('Не оплачены', dashboard_body)
        self.assertIn('Проверка данных', dashboard_body)
        self.assertEqual(checks_response.status_code, 200)
        checks_body = checks_response.get_data(as_text=True)
        self.assertIn('Проверка данных', checks_body)
        self.assertIn('Некорректная почта', checks_body)
        self.assertIn('Без договора', checks_body)
        self.assertIn('Без академической группы', checks_body)
        self.assertIn('Без договора при поступлении', checks_body)
        self.assertIn('Дублирующие записи абитуриентов', checks_body)
        self.assertIn('Конфликты логинов', checks_body)
        self.assertIn('Иванов Иван Иванович', checks_body)
        self.assertIn('Кривец Анна', checks_body)

    def test_enrollment_order_preliminary_roster_groups_logins_and_exports(self):
        people = [
            ('Нетпочты Анна Андреевна', '2026-ФМ-0401-11', '26ФМ-11-1'),
            ('Готовов Петр Петрович', '2026-ФМ-0402-11', '26ФМ-11-1'),
            ('Студентов Сергей Сергеевич', '2026-ФМ-0403-11', '26ФМ-11-2'),
            ('Неизвестный Николай Николаевич', '2026-ФМ-0404-11', '26ФМ-11-2'),
        ]
        order_keys = {
            fio: manticore.make_abiturient_enrollment_match_key(fio, dogovor)
            for fio, dogovor, _group_name in people
        }
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO enrollment_order_uploads
                    (campaign_year, original_filename, uploaded_by, total_count,
                     import_count, matched_count, unmatched_count, order_numbers, order_dates)
                VALUES ('2026', 'приказ-список.xlsx', 'admin', 4, 4, 3, 1, '99-у', '2026-08-18')
                '''
            )
            upload_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
            for row_number, (fio, _dogovor, group_name) in enumerate(people, start=1):
                fio_key, specialty_key = order_keys[fio]
                conn.execute(
                    '''
                    INSERT INTO enrollment_order_upload_rows
                        (upload_id, campaign_year, row_number, fio, specialty, group_name,
                         order_number, order_date, fio_key, specialty_key, group_key,
                         has_candidate, import_action, import_status)
                    VALUES (?, '2026', ?, ?, 'ФМ', ?, '99-у', '2026-08-18', ?, ?, '', ?, 'import', ?)
                    ''',
                    (
                        upload_id, row_number, fio, group_name, fio_key, specialty_key,
                        0 if fio.startswith('Неизвестный') else 1,
                        'Будет загружено'
                    )
                )
            conn.execute(
                '''
                INSERT INTO abiturients
                    (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, '26811041', '2026', 'Нетпочты', 'Анна Андреевна', '', 0)
                ''',
                (people[0][0], people[0][1])
            )
            conn.execute(
                '''
                INSERT INTO abiturients
                    (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, '26811042', '2026', 'Готовов', 'Петр Петрович', 'ready@example.test', 1)
                ''',
                (people[1][0], people[1][1])
            )
            conn.execute(
                '''
                INSERT INTO students
                    (username, password, email, firstname, lastname, cohort1, cohort2,
                     source_campaign_year, source_dogovor, source_fio)
                VALUES ('student_roster', 'cron', 'student@example.test', 'Сергей Сергеевич',
                        'Студентов', '26ФМ-11-2', 'ФМ-11', '2026', ?, ?)
                ''',
                (people[2][1], people[2][0])
            )

        roster = manticore.build_enrollment_order_student_roster(upload_id)
        self.assertEqual(roster['summary']['total_count'], 4)
        self.assertEqual(roster['summary']['group_count'], 2)
        self.assertEqual(roster['summary']['student_count'], 1)
        self.assertEqual(roster['summary']['ready_count'], 1)
        self.assertEqual(roster['summary']['missing_email_count'], 1)
        self.assertEqual(roster['summary']['unpaid_count'], 1)
        self.assertEqual(roster['summary']['not_found_count'], 1)
        self.assertEqual(roster['summary']['attention_count'], 2)
        self.assertEqual(
            roster['summary']['virtual_group_count'],
            len({row['group_name'] for row in roster['rows'] if row['group_is_virtual']})
        )
        self.assertEqual(roster['summary']['unassigned_count'], 0)
        roster_by_fio = {row['fio']: row for row in roster['rows']}
        self.assertEqual(roster_by_fio[people[0][0]]['login'], '26811041')
        self.assertEqual(roster_by_fio[people[0][0]]['status'], 'Нет почты и оплаты')
        self.assertEqual(roster_by_fio[people[0][0]]['group_name'], '26ФМ-11-1')
        self.assertEqual(roster_by_fio[people[2][0]]['login'], 'student_roster')
        self.assertEqual(roster_by_fio[people[2][0]]['status'], 'Перенесён в студенты')
        self.assertEqual(roster_by_fio[people[2][0]]['group_source'], 'Фактическая группа')

        with manticore.app.test_request_context('/data_checks?campaign_year=2026'):
            report = manticore.get_data_quality_report('2026')
        roster_section = next(
            section for section in report['sections']
            if section['title'] == 'Предварительные списки по приказам'
        )
        self.assertTrue(roster_section['informational'])
        self.assertEqual(roster_section['checks'][0]['count'], 2)
        self.assertTrue(roster_section['checks'][0]['expandable'])

        client = manticore.app.test_client()
        self.login_session(client)
        checks_response = client.get('/data_checks?campaign_year=2026')
        checks_body = checks_response.get_data(as_text=True)
        self.assertEqual(checks_response.status_code, 200)
        self.assertIn('Предварительные списки по приказам', checks_body)
        self.assertIn('Предварительный список: № 99-у', checks_body)
        self.assertIn('check-item expandable warning', checks_body)
        self.assertIn(f'/enrollment_order_uploads/{upload_id}/student_roster', checks_body)
        self.assertIn(f'/enrollment_order_uploads/{upload_id}/student_roster/download', checks_body)

        roster_response = client.get(f'/enrollment_order_uploads/{upload_id}/student_roster')
        roster_body = roster_response.get_data(as_text=True)
        self.assertEqual(roster_response.status_code, 200)
        self.assertIn('Предварительный список по приказу', roster_body)
        self.assertIn('26ФМ-11-1', roster_body)
        self.assertIn('26ФМ-11-2', roster_body)
        self.assertIn('26811041', roster_body)
        self.assertIn('student_roster', roster_body)
        self.assertIn('Нет почты и оплаты', roster_body)
        self.assertIn('Только в расчёте', roster_body)
        self.assertIn('в базе они не создаются', roster_body)

        download_response = client.get(
            f'/enrollment_order_uploads/{upload_id}/student_roster/download'
        )
        self.assertEqual(download_response.status_code, 200)
        exported = pd.read_excel(io.BytesIO(download_response.data))
        self.assertEqual(len(exported), 4)
        self.assertEqual(
            exported.columns[:7].tolist(),
            ['username', 'password', 'email', 'firstname', 'lastname', 'cohort1', 'cohort2']
        )
        exported_by_fio = exported.set_index('ФИО')
        self.assertEqual(exported_by_fio.loc[people[0][0], 'username'], '26811041')
        self.assertEqual(exported_by_fio.loc[people[0][0], 'cohort1'], '26ФМ-11-1')
        self.assertEqual(exported_by_fio.loc[people[0][0], 'Состояние'], 'Нет почты и оплаты')
        self.assertEqual(exported_by_fio.loc[people[2][0], 'username'], 'student_roster')

    def test_enrollment_order_roster_simulates_next_group_without_database_writes(self):
        people = [
            ('Альфов Алексей Алексеевич', '2029-ФМ-0501-11', '29611001'),
            ('Бетов Борис Борисович', '2029-ФМ-0502-11', '29611002'),
        ]
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                "INSERT OR IGNORE INTO groups (name, group_year) VALUES ('29ФМ-11-1', '2029')"
            )
            for index in range(24):
                conn.execute(
                    '''
                    INSERT INTO students
                        (username, password, firstname, lastname, cohort1, source_campaign_year)
                    VALUES (?, 'cron', 'Занят', 'Студент', '29ФМ-11-1', '2029')
                    ''',
                    (f'occupied_{index:02d}',)
                )
            conn.execute(
                '''
                INSERT INTO enrollment_order_uploads
                    (campaign_year, original_filename, uploaded_by, total_count,
                     import_count, matched_count, unmatched_count, order_numbers, order_dates)
                VALUES ('2029', 'приказ-без-групп.xlsx', 'admin', 2, 2, 2, 0,
                        '100-у', '2029-08-19')
                '''
            )
            upload_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
            for row_number, (fio, dogovor, login) in enumerate(people, start=1):
                fio_key, specialty_key = manticore.make_abiturient_enrollment_match_key(fio, dogovor)
                conn.execute(
                    '''
                    INSERT INTO enrollment_order_upload_rows
                        (upload_id, campaign_year, row_number, fio, specialty, group_name,
                         order_number, order_date, fio_key, specialty_key, group_key,
                         has_candidate, import_action, import_status)
                    VALUES (?, '2029', ?, ?, 'ФМ', '', '100-у', '2029-08-19',
                            ?, ?, '', 1, 'import', 'Будет загружено')
                    ''',
                    (upload_id, row_number, fio, fio_key, specialty_key)
                )
                lastname, firstname = manticore.split_fio_for_storage(fio)
                conn.execute(
                    '''
                    INSERT INTO abiturients
                        (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                    VALUES (?, ?, ?, '2029', ?, ?, '', 0)
                    ''',
                    (fio, dogovor, login, lastname, firstname)
                )
            group_count_before = conn.execute('SELECT COUNT(*) FROM groups').fetchone()[0]

        roster = manticore.build_enrollment_order_student_roster(upload_id)
        roster_by_fio = {row['fio']: row for row in roster['rows']}

        self.assertEqual(roster_by_fio[people[0][0]]['group_name'], '29ФМ-11-1')
        self.assertFalse(roster_by_fio[people[0][0]]['group_is_virtual'])
        self.assertEqual(roster_by_fio[people[1][0]]['group_name'], '29ФМ-11-2')
        self.assertTrue(roster_by_fio[people[1][0]]['group_is_virtual'])
        self.assertEqual(roster['summary']['group_count'], 2)
        self.assertEqual(roster['summary']['virtual_group_count'], 1)
        self.assertEqual(roster['summary']['unassigned_count'], 0)

        export = pd.DataFrame(manticore.enrollment_order_roster_export_rows(roster))
        self.assertEqual(
            export.columns[:7].tolist(),
            ['username', 'password', 'email', 'firstname', 'lastname', 'cohort1', 'cohort2']
        )
        self.assertEqual(export['cohort1'].tolist(), ['29ФМ-11-1', '29ФМ-11-2'])

        with sqlite3.connect(manticore.DB_PATH) as conn:
            self.assertEqual(conn.execute('SELECT COUNT(*) FROM groups').fetchone()[0], group_count_before)
            self.assertIsNone(
                conn.execute("SELECT 1 FROM groups WHERE name='29ФМ-11-2'").fetchone()
            )

    def test_enrollment_order_roster_can_fix_or_skip_fio_review(self):
        people = [
            {
                'order_fio': 'Опечаткин Иван Иванович',
                'database_fio': 'Опечаткин Иван Иванвич',
                'dogovor': '2027-ФМ-0601-11',
                'login': '27611001',
                'specialty': 'ФМ',
            },
            {
                'order_fio': 'Пропускова Анна Сергеевна',
                'database_fio': 'Пропускова Ана Сергеевна',
                'dogovor': '2027-ФМ-0602-11',
                'login': '27611002',
                'specialty': 'ФМ',
            },
            {
                'order_fio': 'Конфликтов Петр Петрович',
                'database_fio': 'Конфликтов Петр Петрович',
                'dogovor': '2027-ЛД-0603-11',
                'login': '27111001',
                'specialty': 'СД',
            },
            {
                'order_fio': 'Приказов Павел Павлович',
                'database_fio': 'Приказов Павел Павловч',
                'dogovor': '2027-ФМ-0604-11',
                'login': '27611004',
                'specialty': 'ФМ',
            },
        ]
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO enrollment_order_uploads
                    (campaign_year, original_filename, uploaded_by, total_count,
                     import_count, matched_count, unmatched_count, order_numbers, order_dates)
                VALUES ('2027', 'приказ-сверка-фио.xlsx', 'admin', 4, 4, 0, 4,
                        '101-у', '2027-08-20')
                '''
            )
            upload_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
            row_ids = []
            abiturient_ids = []
            for row_number, person in enumerate(people, start=1):
                fio_key = manticore.normalize_fio_key(person['order_fio'])
                specialty_key = manticore.normalize_specialty_key(person['specialty'])
                conn.execute(
                    '''
                    INSERT INTO enrollment_order_upload_rows
                        (upload_id, campaign_year, row_number, fio, specialty, group_name,
                         order_number, order_date, fio_key, specialty_key, group_key,
                         has_candidate, import_action, import_status)
                    VALUES (?, '2027', ?, ?, ?, '', '101-у', '2027-08-20',
                            ?, ?, '', 0, 'import', 'Будет загружен, но кандидата пока нет')
                    ''',
                    (
                        upload_id, row_number, person['order_fio'], person['specialty'],
                        fio_key, specialty_key,
                    )
                )
                row_ids.append(conn.execute('SELECT last_insert_rowid()').fetchone()[0])
                fam, imotch = manticore.split_fio_for_storage(person['database_fio'])
                conn.execute(
                    '''
                    INSERT INTO abiturients
                        (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                    VALUES (?, ?, ?, '2027', ?, ?, ?, 1)
                    ''',
                    (
                        person['database_fio'], person['dogovor'], person['login'], fam, imotch,
                        f"{person['login']}@example.test",
                    )
                )
                abiturient_ids.append(conn.execute('SELECT last_insert_rowid()').fetchone()[0])

        roster = manticore.build_enrollment_order_student_roster(upload_id)
        roster_by_row = {row['row_number']: row for row in roster['rows']}
        self.assertEqual(roster_by_row[1]['fio_review_kind'], 'fio_typo')
        self.assertEqual(roster_by_row[1]['suggested_abiturient_id'], abiturient_ids[0])
        self.assertEqual(roster_by_row[2]['fio_review_kind'], 'fio_typo')
        self.assertEqual(roster_by_row[3]['fio_review_kind'], 'specialty_conflict')
        self.assertEqual(roster_by_row[4]['fio_review_kind'], 'fio_typo')
        self.assertEqual(roster['summary']['fio_review_count'], 3)
        self.assertEqual(roster['summary']['specialty_conflict_count'], 1)

        client = manticore.app.test_client()
        self.login_session(client)
        page = client.get(f'/enrollment_order_uploads/{upload_id}/student_roster')
        body = page.get_data(as_text=True)
        self.assertEqual(page.status_code, 200)
        self.assertIn('Похоже на запись в базе', body)
        self.assertIn('Исправить в базе', body)
        self.assertIn('Пропустить без исправления', body)
        self.assertIn('Пропустить и присвоить логин', body)
        self.assertIn('ФИО совпадает, специальность отличается', body)

        skip_response = client.post(
            f'/enrollment_order_uploads/{upload_id}/student_roster/fio_review',
            data={
                'csrf_token': self.csrf_from_response(page),
                'review_action': 'skip',
                'row_id': row_ids[1],
                'abiturient_id': abiturient_ids[1],
            },
            follow_redirects=True,
        )
        self.assertEqual(skip_response.status_code, 200)
        self.assertIn('Исправление ФИО пропущено', skip_response.get_data(as_text=True))
        with sqlite3.connect(manticore.DB_PATH) as conn:
            skipped_fio = conn.execute(
                'SELECT fio FROM abiturients WHERE id=?',
                (abiturient_ids[1],)
            ).fetchone()[0]
            review_status = conn.execute(
                'SELECT fio_review_status FROM enrollment_order_upload_rows WHERE id=?',
                (row_ids[1],)
            ).fetchone()[0]
        self.assertEqual(skipped_fio, people[1]['database_fio'])
        self.assertEqual(review_status, 'skipped')

        link_page = client.get(f'/enrollment_order_uploads/{upload_id}/student_roster')
        link_response = client.post(
            f'/enrollment_order_uploads/{upload_id}/student_roster/fio_review',
            data={
                'csrf_token': self.csrf_from_response(link_page),
                'review_action': 'link',
                'row_id': row_ids[3],
                'abiturient_id': abiturient_ids[3],
            },
            follow_redirects=True,
        )
        link_body = link_response.get_data(as_text=True)
        self.assertEqual(link_response.status_code, 200)
        self.assertIn('Логин 27611004 присвоен строке', link_body)
        self.assertIn('Опечатка в приказе — логин присвоен', link_body)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            linked_fio = conn.execute(
                'SELECT fio FROM abiturients WHERE id=?',
                (abiturient_ids[3],)
            ).fetchone()[0]
            linked_status = conn.execute(
                'SELECT fio_review_status FROM enrollment_order_upload_rows WHERE id=?',
                (row_ids[3],)
            ).fetchone()[0]
        self.assertEqual(linked_fio, people[3]['database_fio'])
        self.assertEqual(linked_status, 'linked')

        fix_page = client.get(f'/enrollment_order_uploads/{upload_id}/student_roster')
        fix_response = client.post(
            f'/enrollment_order_uploads/{upload_id}/student_roster/fio_review',
            data={
                'csrf_token': self.csrf_from_response(fix_page),
                'review_action': 'fix',
                'row_id': row_ids[0],
                'abiturient_id': abiturient_ids[0],
            },
            follow_redirects=True,
        )
        fix_body = fix_response.get_data(as_text=True)
        self.assertEqual(fix_response.status_code, 200)
        self.assertIn('ФИО исправлено в базе абитуриентов', fix_body)
        self.assertIn('ФИО исправлено по приказу', fix_body)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            fixed_abiturient = conn.execute(
                'SELECT fio, fam, imotch FROM abiturients WHERE id=?',
                (abiturient_ids[0],)
            ).fetchone()
        self.assertEqual(
            fixed_abiturient,
            ('Опечаткин Иван Иванович', 'Опечаткин', 'Иван Иванович'),
        )

        updated_roster = manticore.build_enrollment_order_student_roster(upload_id)
        updated_by_row = {row['row_number']: row for row in updated_roster['rows']}
        self.assertEqual(updated_by_row[1]['record_type'], 'Абитуриент')
        self.assertEqual(updated_by_row[1]['login'], '27611001')
        self.assertEqual(updated_by_row[1]['fio_review_status'], 'fixed')
        self.assertEqual(updated_by_row[2]['fio_review_status'], 'skipped')
        self.assertEqual(updated_by_row[4]['fio_review_status'], 'linked')
        self.assertEqual(updated_by_row[4]['record_type'], 'Абитуриент')
        self.assertEqual(updated_by_row[4]['login'], '27611004')
        self.assertEqual(updated_roster['summary']['fio_review_linked_count'], 1)
        export = pd.DataFrame(manticore.enrollment_order_roster_export_rows(updated_roster))
        export_by_fio = export.set_index('ФИО')
        self.assertEqual(export_by_fio.loc[people[0]['order_fio'], 'Сверка ФИО'], 'Исправлено по приказу')
        self.assertEqual(export_by_fio.loc[people[1]['order_fio'], 'Сверка ФИО'], 'Пропущено без исправления')
        self.assertEqual(
            export_by_fio.loc[people[3]['order_fio'], 'Сверка ФИО'],
            'Опечатка в приказе — логин присвоен',
        )
        self.assertEqual(export_by_fio.loc[people[3]['order_fio'], 'username'], '27611004')

    def test_data_checks_reports_missing_and_mismatched_cohort2(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            rows = [
                ('student_missing_cohort2', 'missing@example.test', 'Иван', 'Иванов', '26СД-9-1', None),
                ('student_wrong_cohort2', 'wrong@example.test', 'Петр', 'Петров', '26СД-9-1', 'ЛД-9'),
                ('student_unsupported', 'unsupported@example.test', 'Анна', 'Сидорова', '26ЛабД-11-1', None),
            ]
            for username, email, firstname, lastname, cohort1, cohort2 in rows:
                conn.execute(
                    '''
                    INSERT INTO students
                        (username, password, email, firstname, lastname, cohort1, cohort2, source_campaign_year, source_dogovor)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''',
                    (
                        username, 'cron', email, firstname, lastname,
                        cohort1, cohort2, '2026', f'2026-{username}'
                    )
                )

        with manticore.app.test_request_context('/'):
            report = manticore.get_data_quality_report('2026')

        students_section = next(section for section in report['sections'] if section['title'] == 'Студенты')
        check = next(item for item in students_section['checks'] if item['id'] == 'students-cohort2-mismatch')

        self.assertEqual(check['count'], 2)
        details = ' '.join(sample['detail'] for sample in check['samples'])
        self.assertIn('ожидалось СД-9', details)
        self.assertIn('сейчас ЛД-9', details)

    def test_abiturients_updates_import_updates_email_and_paid(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                ''',
                ('Петров Петр Петрович', '2026-ФМ-0101-11', '26611011', '2026', 'Петров', 'Петр Петрович', 0)
            )

        updates_path = os.path.join(TEST_UPLOAD_DIR, 'updates.xlsx')
        pd.DataFrame([
            {'Договор': '2026-ФМ-0101-11', 'Email': 'new@example.test', 'Оплата': 'да'}
        ]).to_excel(updates_path, index=False)

        summary = manticore.process_abiturients_updates(updates_path, '2026')

        self.assertEqual(summary['updated_email'], 1)
        self.assertEqual(summary['updated_paid'], 1)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            row = conn.execute('SELECT email, paid FROM abiturients WHERE login=?', ('26611011',)).fetchone()
        self.assertEqual(row, ('new@example.test', 1))

    def test_abiturients_updates_sync_enrollment_candidates_when_paid_changes(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Братыкина Алина Алексеевна', '2026-СтД-0046-11И', '26811i003',
                    '2026', 'Братыкина', 'Алина Алексеевна', 'alina@example.test', 1
                )
            )

        sync_summary = manticore.sync_enrollment_candidates_from_ready_abiturients('2026')
        self.assertEqual(sync_summary['created'], 1)

        updates_path = os.path.join(TEST_UPLOAD_DIR, 'updates_unpaid.xlsx')
        pd.DataFrame([
            {'Договор': '2026-СтД-0046-11И', 'Оплата': 'нет'}
        ]).to_excel(updates_path, index=False)

        summary = manticore.process_abiturients_updates(updates_path, '2026')

        self.assertEqual(summary['updated_paid'], 1)
        self.assertEqual(summary['candidate_sync']['removed'], 1)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            candidate_count = conn.execute(
                'SELECT COUNT(*) FROM enrollment_candidates WHERE login=?',
                ('26811i003',)
            ).fetchone()[0]
        self.assertEqual(candidate_count, 0)

        pd.DataFrame([
            {'Договор': '2026-СтД-0046-11И', 'Оплата': 'да'}
        ]).to_excel(updates_path, index=False)

        summary = manticore.process_abiturients_updates(updates_path, '2026')

        self.assertEqual(summary['updated_paid'], 1)
        self.assertEqual(summary['candidate_sync']['created'], 1)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            candidate = conn.execute(
                'SELECT fio, email, specialty FROM enrollment_candidates WHERE login=?',
                ('26811i003',)
            ).fetchone()
        self.assertEqual(
            candidate,
            ('Братыкина Алина Алексеевна', 'alina@example.test', '31.02.07 «Стоматологическое дело»')
        )

    def test_abiturients_updates_template_download(self):
        client = manticore.app.test_client()
        self.login_session(client)

        page_response = client.get('/file_work')
        template_response = client.get('/abiturients_updates_template/download')

        self.assertEqual(page_response.status_code, 200)
        self.assertIn('Скачать шаблон обновлений', page_response.get_data(as_text=True))
        self.assertEqual(template_response.status_code, 200)
        self.assertEqual(
            template_response.mimetype,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        workbook = pd.ExcelFile(io.BytesIO(template_response.data))
        self.assertEqual(workbook.sheet_names, ['Шаблон', 'Подсказка', 'Пример'])
        template_df = pd.read_excel(workbook, sheet_name='Шаблон')
        help_df = pd.read_excel(workbook, sheet_name='Подсказка')
        self.assertEqual(list(template_df.columns), ['Договор', 'Email', 'Оплата'])
        self.assertIn('Договор', help_df['Поле'].tolist())

    def test_abiturients_updates_upload_reports_row_errors(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                ''',
                ('Смирнова Светлана Сергеевна', '2026-ФМ-0300-11', '26611030', '2026', 'Смирнова', 'Светлана Сергеевна', 0)
            )

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/file_work')
        csrf_token = self.csrf_from_response(get_response)
        upload = io.BytesIO()
        pd.DataFrame([
            {'Договор': '2026-ФМ-0300-11', 'Email': 'wrong-mail', 'Оплата': 'оплаченно'},
            {'Договор': '2026-ФМ-9999-11', 'Email': 'ok@example.test', 'Оплата': 'да'},
            {'Договор': '', 'Email': 'empty@example.test', 'Оплата': 'да'},
        ]).to_excel(upload, index=False)
        upload.seek(0)

        response = client.post(
            '/abiturients_updates_upload',
            data={
                'csrf_token': csrf_token,
                'updates_file': (upload, 'updates.xlsx'),
            },
            content_type='multipart/form-data',
            follow_redirects=True
        )
        body = response.get_data(as_text=True)

        self.assertEqual(response.status_code, 200)
        self.assertIn('Отчет по файлу обновлений', body)
        self.assertIn('Строка 2', body)
        self.assertIn('Почта выглядит некорректно', body)
        self.assertIn('Не удалось распознать значение оплаты', body)
        self.assertIn('Строка 3', body)
        self.assertIn('Договор не найден', body)
        self.assertIn('Строка 4', body)
        self.assertIn('Не указан номер договора', body)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            row = conn.execute('SELECT email, paid FROM abiturients WHERE login=?', ('26611030',)).fetchone()
        self.assertEqual(row, (None, 0))

    def test_email_source_updates_only_email_for_exact_full_name_matches(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients
                    (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Ёлкина Анна Сергеевна', '2026-СД-0400-11', '26311040',
                    '2026', 'Ёлкина', 'Анна Сергеевна', 'old-applicant@example.test', 0
                )
            )
            conn.execute(
                '''
                INSERT INTO students
                    (username, password, email, firstname, lastname, cohort1,
                     source_campaign_year, source_fio)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'student040', 'cron', 'invalid-email', 'Анна Сергеевна', 'Ёлкина',
                    '26СД-11-1', '2026', 'Ёлкина Анна Сергеевна'
                )
            )

        source_path = os.path.join(TEST_UPLOAD_DIR, 'email_source.xlsx')
        pd.DataFrame([{
            'ID': 101,
            'Фамилия': 'Елкина',
            'Имя Отчество': 'Анна Сергеевна',
            'Почта': 'anna@example.test',
            'Телефон': '+70000000000',
        }]).to_excel(source_path, index=False)

        summary = manticore.process_email_source_updates(source_path, '2026')

        self.assertEqual(summary['matched_rows'], 1)
        self.assertEqual(summary['updated_abiturients'], 1)
        self.assertEqual(summary['updated_students'], 1)
        self.assertEqual(summary['corrected_invalid_student_emails'], 1)
        self.assertEqual(summary['issues'], [])
        with sqlite3.connect(manticore.DB_PATH) as conn:
            abiturient = conn.execute(
                'SELECT fio, email, paid FROM abiturients WHERE login=?',
                ('26311040',)
            ).fetchone()
            student = conn.execute(
                'SELECT firstname, lastname, email FROM students WHERE username=?',
                ('student040',)
            ).fetchone()
        self.assertEqual(abiturient, ('Ёлкина Анна Сергеевна', 'anna@example.test', 0))
        self.assertEqual(student, ('Анна Сергеевна', 'Ёлкина', 'anna@example.test'))

    def test_email_source_warns_about_fio_typo_without_changes(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients
                    (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Сафонова Ольга Олеговна', '2026-ЛД-0401-11', '26111041',
                    '2026', 'Сафонова', 'Ольга Олеговна', 'old@example.test', 1
                )
            )

        source_path = os.path.join(TEST_UPLOAD_DIR, 'email_source_typo.csv')
        pd.DataFrame([{
            'Фамилия': 'Сафонова',
            'Имя Отчество': 'Ольга Олговна',
            'Почта': 'new@example.test',
        }]).to_csv(source_path, index=False)

        summary = manticore.process_email_source_updates(source_path, '2026')

        self.assertEqual(summary['updated_abiturients'], 0)
        self.assertEqual(summary['fio_warning_count'], 1)
        self.assertIn('Похоже на «Сафонова Ольга Олеговна»', summary['issues'][0]['message'])
        self.assertIn('почта не изменена', summary['issues'][0]['message'])
        self.assertEqual(summary['preview_rows'][0]['action_label'], 'Проверить ФИО')
        with sqlite3.connect(manticore.DB_PATH) as conn:
            row = conn.execute(
                'SELECT fio, email FROM abiturients WHERE login=?',
                ('26111041',)
            ).fetchone()
        self.assertEqual(row, ('Сафонова Ольга Олеговна', 'old@example.test'))

    def test_email_source_requires_exact_questionnaire_column_names(self):
        source_path = os.path.join(TEST_UPLOAD_DIR, 'email_source_wrong_columns.xlsx')
        pd.DataFrame([{
            'Фамилия': 'Иванов',
            'Имя Отчество': 'Иван Иванович',
            'Email': 'ivan@example.test',
        }]).to_excel(source_path, index=False)

        with self.assertRaisesRegex(ValueError, 'Фамилия, Имя Отчество, Почта'):
            manticore.process_email_source_updates(source_path, '2026')

    def test_email_source_upload_route_previews_before_confirmation(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients
                    (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Петрова Мария Ивановна', '2026-ФМ-0402-11', '26611042',
                    '2026', 'Петрова', 'Мария Ивановна', '', 0
                )
            )

        client = manticore.app.test_client()
        self.login_session(client)
        page = client.get('/file_work/updates')
        body = page.get_data(as_text=True)
        self.assertIn('Дополнительная загрузка почты по ФИО', body)
        self.assertIn('<b>Фамилия</b>, <b>Имя Отчество</b> и <b>Почта</b>', body)
        csrf_token = self.csrf_from_response(page)
        upload = io.BytesIO()
        pd.DataFrame([{
            'Фамилия': 'Петрова',
            'Имя Отчество': 'Мария Ивановна',
            'Почта': 'maria@example.test',
        }]).to_excel(upload, index=False)
        upload.seek(0)

        preview_response = client.post(
            '/abiturients_email_source_upload',
            data={
                'csrf_token': csrf_token,
                'file_section': 'updates',
                'email_source_import_action': 'preview',
                'email_source_file': (upload, 'source.xlsx'),
            },
            content_type='multipart/form-data',
        )
        preview_body = preview_response.get_data(as_text=True)

        self.assertEqual(preview_response.status_code, 200)
        self.assertIn('Предпросмотр обновления почты', preview_body)
        self.assertIn('В базе ещё ничего не изменено', preview_body)
        self.assertIn('Петрова Мария Ивановна', preview_body)
        self.assertIn('maria@example.test', preview_body)
        self.assertIn('Подтвердить обновление почты', preview_body)
        self.assertIn('Отчет по дополнительному источнику почты', preview_body)
        self.assertIn('Точно совпало по ФИО: 1', preview_body)
        pending_match = re.search(
            r'name="pending_email_source_import" value="([^"]+)"',
            preview_body,
        )
        self.assertIsNotNone(pending_match)
        pending_token = pending_match.group(1)
        pending_path = os.path.join(TEST_UPLOAD_DIR, pending_token)
        self.assertTrue(os.path.exists(pending_path))
        with sqlite3.connect(manticore.DB_PATH) as conn:
            preview_email = conn.execute(
                'SELECT email FROM abiturients WHERE login=?',
                ('26611042',)
            ).fetchone()[0]
        self.assertEqual(preview_email, '')

        confirm_response = client.post(
            '/abiturients_email_source_upload',
            data={
                'csrf_token': csrf_token,
                'file_section': 'updates',
                'email_source_import_action': 'confirm',
                'pending_email_source_import': pending_token,
            },
            follow_redirects=True,
        )
        confirm_body = confirm_response.get_data(as_text=True)

        self.assertEqual(confirm_response.status_code, 200)
        self.assertIn('Источник почты применён', confirm_body)
        self.assertIn('Обновлено почт абитуриентов: 1', confirm_body)
        self.assertNotIn('Предпросмотр обновления почты', confirm_body)
        self.assertFalse(os.path.exists(pending_path))
        with sqlite3.connect(manticore.DB_PATH) as conn:
            email = conn.execute(
                'SELECT email FROM abiturients WHERE login=?',
                ('26611042',)
            ).fetchone()[0]
        self.assertEqual(email, 'maria@example.test')

    def test_email_source_preview_can_be_cancelled_without_changes(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients
                    (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Сидорова Анна Павловна', '2026-СД-0403-11', '26311043',
                    '2026', 'Сидорова', 'Анна Павловна', 'old@example.test', 0
                )
            )

        client = manticore.app.test_client()
        self.login_session(client)
        page = client.get('/file_work/updates')
        csrf_token = self.csrf_from_response(page)
        upload = io.BytesIO()
        pd.DataFrame([{
            'Фамилия': 'Сидорова',
            'Имя Отчество': 'Анна Павловна',
            'Почта': 'new@example.test',
        }]).to_excel(upload, index=False)
        upload.seek(0)
        preview_response = client.post(
            '/abiturients_email_source_upload',
            data={
                'csrf_token': csrf_token,
                'file_section': 'updates',
                'email_source_import_action': 'preview',
                'email_source_file': (upload, 'source.xlsx'),
            },
            content_type='multipart/form-data',
        )
        pending_match = re.search(
            r'name="pending_email_source_import" value="([^"]+)"',
            preview_response.get_data(as_text=True),
        )
        self.assertIsNotNone(pending_match)
        pending_token = pending_match.group(1)
        pending_path = os.path.join(TEST_UPLOAD_DIR, pending_token)

        cancel_response = client.post(
            '/abiturients_email_source_upload',
            data={
                'csrf_token': csrf_token,
                'file_section': 'updates',
                'email_source_import_action': 'cancel',
                'pending_email_source_import': pending_token,
            },
            follow_redirects=True,
        )

        self.assertEqual(cancel_response.status_code, 200)
        self.assertIn('Предпросмотр источника почты отменён', cancel_response.get_data(as_text=True))
        self.assertFalse(os.path.exists(pending_path))
        with sqlite3.connect(manticore.DB_PATH) as conn:
            email = conn.execute(
                'SELECT email FROM abiturients WHERE login=?',
                ('26311043',)
            ).fetchone()[0]
        self.assertEqual(email, 'old@example.test')

    def test_two_stage_enrollment_requires_order_before_student_migration(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Иванов Иван Иванович', '2026-ФМ-0500-11', '26611050',
                    '2026', 'Иванов', 'Иван Иванович', 'ivanov@example.test', 1
                )
            )
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Петров Петр Петрович', '2026-ФМ-0501-11', '26611051',
                    '2026', 'Петров', 'Петр Петрович', '', 1
                )
            )
            conn.execute(
                'INSERT OR IGNORE INTO groups (name, group_year) VALUES (?, ?)',
                ('26ФМ-11-1', '2026')
            )

        sync_summary = manticore.sync_enrollment_candidates_from_ready_abiturients('2026')
        self.assertEqual(sync_summary['created'], 1)
        self.assertEqual(len(sync_summary['skipped_without_email']), 1)

        with sqlite3.connect(manticore.DB_PATH) as conn:
            candidate = conn.execute(
                '''
                SELECT id, specialty, verification_status
                FROM enrollment_candidates
                WHERE login=?
                ''',
                ('26611050',)
            ).fetchone()
        self.assertIsNotNone(candidate)
        candidate_id, specialty, verification_status = candidate
        self.assertEqual(specialty, '33.02.01 «Фармация»')
        self.assertEqual(verification_status, 'waiting_order')

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/abiturients_to_students')
        csrf_token = self.csrf_from_response(get_response)
        blocked_response = client.post(
            '/abiturients_to_students',
            data={
                'csrf_token': csrf_token,
                'group_year': '2026',
                'cohort1': '26ФМ-11-1',
                'candidate_ids': [str(candidate_id)],
            },
            follow_redirects=True
        )
        self.assertEqual(blocked_response.status_code, 200)
        self.assertIn('Не перенесены без совпадения с приказом', blocked_response.get_data(as_text=True))
        with sqlite3.connect(manticore.DB_PATH) as conn:
            student_count = conn.execute('SELECT COUNT(*) FROM students WHERE username=?', ('26611050',)).fetchone()[0]
        self.assertEqual(student_count, 0)

        order_path = os.path.join(TEST_UPLOAD_DIR, 'enrollment_order.xlsx')
        pd.DataFrame([
            {
                'ФИО': 'Иванов Иван Иванович',
                'Специальность': '33.02.01 «Фармация»',
                'Группа': '26ФМ-11-1',
                'Номер приказа': '123-у',
                'Дата приказа': '2026-08-15',
            }
        ]).to_excel(order_path, index=False)
        order_summary = manticore.apply_enrollment_order_import(order_path, '2026')
        self.assertEqual(order_summary['matched_count'], 1)

        get_response = client.get('/abiturients_to_students?specialty=фм')
        body = get_response.get_data(as_text=True)
        self.assertIn('Сверен с приказом', body)
        self.assertIn('26ФМ-11-1', body)
        csrf_token = self.csrf_from_response(get_response)
        migrated_response = client.post(
            '/abiturients_to_students',
            data={
                'csrf_token': csrf_token,
                'group_year': '2026',
                'specialty': 'фм',
                'cohort1': '26ФМ-11-1',
                'candidate_ids': [str(candidate_id)],
            },
            follow_redirects=True
        )
        self.assertEqual(migrated_response.status_code, 200)
        self.assertIn('Иванов', migrated_response.get_data(as_text=True))
        with sqlite3.connect(manticore.DB_PATH) as conn:
            student = conn.execute(
                '''
                SELECT username, email, firstname, lastname, cohort1, cohort2, source_dogovor, source_fio
                FROM students
                WHERE username=?
                ''',
                ('26611050',)
            ).fetchone()
            abiturient_count = conn.execute('SELECT COUNT(*) FROM abiturients WHERE login=?', ('26611050',)).fetchone()[0]
            candidate_count = conn.execute('SELECT COUNT(*) FROM enrollment_candidates WHERE login=?', ('26611050',)).fetchone()[0]
            movement = conn.execute(
                '''
                SELECT movement_type, old_cohort1, old_cohort2, new_cohort1, new_cohort2,
                       order_number, order_date, order_source,
                       enrollment_order_id, enrollment_order_upload_id
                FROM student_group_transfers
                WHERE username=?
                ''',
                ('26611050',)
            ).fetchone()

        self.assertEqual(
            student,
            (
                '26611050', 'ivanov@example.test', 'Иван Иванович', 'Иванов',
                '26ФМ-11-1', 'ФМ-11', '2026-ФМ-0500-11', 'Иванов Иван Иванович'
            )
        )
        self.assertEqual(abiturient_count, 0)
        self.assertEqual(candidate_count, 0)
        self.assertIsNotNone(movement)
        self.assertEqual(
            movement[:8],
            (
                'enrollment', 'Абитуриенты', '', '26ФМ-11-1', 'ФМ-11',
                '123-у', '2026-08-15', 'enrollment_order'
            )
        )
        self.assertIsNotNone(movement[8])
        self.assertIsNotNone(movement[9])

        edit_response = client.get('/edit_student/26611050')
        self.assertEqual(edit_response.status_code, 200)
        edit_body = edit_response.get_data(as_text=True)
        self.assertIn('История перемещений', edit_body)
        self.assertIn('Зачислен', edit_body)
        self.assertIn('Приказ о зачислении', edit_body)
        self.assertIn('123-у', edit_body)
        self.assertIn('15.08.2026', edit_body)
        self.assertIn('Скачать приказ', edit_body)

    def test_student_card_backfills_initial_enrollment_movement_from_order(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Сидоров Семен Семенович', '2026-ФМ-0560-11', '26611060',
                    '2026', 'Сидоров', 'Семен Семенович', 'sidorov@example.test', 1
                )
            )
            conn.execute(
                'INSERT OR IGNORE INTO groups (name, group_year) VALUES (?, ?)',
                ('26ФМ-11-1', '2026')
            )

        sync_summary = manticore.sync_enrollment_candidates_from_ready_abiturients('2026')
        self.assertEqual(sync_summary['created'], 1)

        order_path = os.path.join(TEST_UPLOAD_DIR, 'backfill_enrollment_order.xlsx')
        pd.DataFrame([
            {
                'ФИО': 'Сидоров Семен Семенович',
                'Специальность': '33.02.01 «Фармация»',
                'Группа': '26ФМ-11-1',
                'Номер приказа': '124-у',
                'Дата приказа': '2026-08-16',
            }
        ]).to_excel(order_path, index=False)
        order_summary = manticore.apply_enrollment_order_import(order_path, '2026')
        self.assertEqual(order_summary['matched_count'], 1)

        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO students
                    (username, password, email, firstname, lastname, cohort1, cohort2,
                     source_campaign_year, source_dogovor, source_fio)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    '26611060', 'cron', 'sidorov@example.test', 'Семен Семенович', 'Сидоров',
                    '26ФМ-11-1', 'ФМ-11', '2026', '2026-ФМ-0560-11', 'Сидоров Семен Семенович'
                )
            )
            movement_count = conn.execute(
                'SELECT COUNT(*) FROM student_group_transfers WHERE username=?',
                ('26611060',)
            ).fetchone()[0]
        self.assertEqual(movement_count, 0)

        client = manticore.app.test_client()
        self.login_session(client)
        response = client.get('/edit_student/26611060')
        self.assertEqual(response.status_code, 200)
        body = response.get_data(as_text=True)
        self.assertIn('Зачислен', body)
        self.assertIn('Приказ о зачислении', body)
        self.assertIn('124-у', body)
        self.assertIn('16.08.2026', body)

        with sqlite3.connect(manticore.DB_PATH) as conn:
            movement = conn.execute(
                '''
                SELECT movement_type, old_cohort1, new_cohort1, new_cohort2, order_number, order_date
                FROM student_group_transfers
                WHERE username=?
                ''',
                ('26611060',)
            ).fetchone()
        self.assertEqual(
            movement,
            ('enrollment', 'Абитуриенты', '26ФМ-11-1', 'ФМ-11', '124-у', '2026-08-16')
        )

    def test_enrollment_order_requirement_can_be_disabled_for_manual_migration(self):
        rules = manticore.get_default_login_generation_rules()
        rules['require_enrollment_order'] = False
        manticore.save_login_generation_settings(rules, setup_completed=True, updated_by='test')

        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Смирнов Сергей Сергеевич', '2026-ФМ-0550-11', '26611055',
                    '2026', 'Смирнов', 'Сергей Сергеевич', 'smirnov@example.test', 1
                )
            )
            conn.execute(
                'INSERT OR IGNORE INTO groups (name, group_year) VALUES (?, ?)',
                ('26ФМ-11-1', '2026')
            )

        sync_summary = manticore.sync_enrollment_candidates_from_ready_abiturients('2026')
        self.assertEqual(sync_summary['created'], 1)
        with manticore.app.test_request_context('/'):
            dashboard = manticore.get_dashboard_data('2026')
        self.assertEqual(dashboard['candidate_missing_order'], 0)
        self.assertEqual(dashboard['ready_to_students'], 1)

        with sqlite3.connect(manticore.DB_PATH) as conn:
            candidate = conn.execute(
                '''
                SELECT id, verification_status
                FROM enrollment_candidates
                WHERE login=?
                ''',
                ('26611055',)
            ).fetchone()
        self.assertIsNotNone(candidate)
        candidate_id, verification_status = candidate
        self.assertEqual(verification_status, 'waiting_order')

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/abiturients_to_students')
        body = get_response.get_data(as_text=True)
        self.assertIn('Проверка по приказу отключена', body)
        self.assertIn('Готов по почте и оплате', body)
        csrf_token = self.csrf_from_response(get_response)
        migrated_response = client.post(
            '/abiturients_to_students',
            data={
                'csrf_token': csrf_token,
                'group_year': '2026',
                'cohort1': '26ФМ-11-1',
                'candidate_ids': [str(candidate_id)],
            },
            follow_redirects=True
        )
        self.assertEqual(migrated_response.status_code, 200)
        self.assertIn('Смирнов', migrated_response.get_data(as_text=True))

        with sqlite3.connect(manticore.DB_PATH) as conn:
            student_group = conn.execute(
                'SELECT cohort1, cohort2 FROM students WHERE username=?',
                ('26611055',)
            ).fetchone()
            candidate_count = conn.execute('SELECT COUNT(*) FROM enrollment_candidates WHERE login=?', ('26611055',)).fetchone()[0]
        self.assertEqual(student_group, ('26ФМ-11-1', 'ФМ-11'))
        self.assertEqual(candidate_count, 0)

    def test_enrollment_migration_respects_disabled_course_groups(self):
        rules = manticore.get_default_login_generation_rules()
        rules['require_enrollment_order'] = False
        rules['use_course_groups'] = False
        manticore.save_login_generation_settings(rules, setup_completed=True, updated_by='test')

        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Курсов Иван Иванович', '2026-ФМ-0570-11', '26611070',
                    '2026', 'Курсов', 'Иван Иванович', 'kursov@example.test', 1
                )
            )
            conn.execute(
                'INSERT OR IGNORE INTO groups (name, group_year) VALUES (?, ?)',
                ('26ФМ-11-1', '2026')
            )

        sync_summary = manticore.sync_enrollment_candidates_from_ready_abiturients('2026')
        self.assertEqual(sync_summary['created'], 1)

        with sqlite3.connect(manticore.DB_PATH) as conn:
            candidate_id = conn.execute(
                'SELECT id FROM enrollment_candidates WHERE login=?',
                ('26611070',)
            ).fetchone()[0]

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/abiturients_to_students')
        self.assertNotIn('id="cohort2_preview"', get_response.get_data(as_text=True))
        csrf_token = self.csrf_from_response(get_response)
        response = client.post(
            '/abiturients_to_students',
            data={
                'csrf_token': csrf_token,
                'group_year': '2026',
                'cohort1': '26ФМ-11-1',
                'candidate_ids': [str(candidate_id)],
            },
            follow_redirects=True
        )
        self.assertEqual(response.status_code, 200)

        with sqlite3.connect(manticore.DB_PATH) as conn:
            student_group = conn.execute(
                'SELECT cohort1, cohort2 FROM students WHERE username=?',
                ('26611070',)
            ).fetchone()
        self.assertEqual(student_group, ('26ФМ-11-1', ''))

    def test_auto_split_saves_cohort2_from_final_group(self):
        rules = manticore.get_default_login_generation_rules()
        rules['require_enrollment_order'] = False
        manticore.save_login_generation_settings(rules, setup_completed=True, updated_by='test')

        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute('INSERT OR IGNORE INTO groups (name, group_year) VALUES (?, ?)', ('26СД-9-1', '2026'))
            for index in range(manticore.MAX_GROUP_STUDENTS):
                conn.execute(
                    '''
                    INSERT INTO students (username, password, email, firstname, lastname, cohort1)
                    VALUES (?, ?, ?, ?, ?, ?)
                    ''',
                    (f'filled_sd_{index:03d}', 'cron', f'filled_sd_{index}@example.test', 'Имя', 'Фамилия', '26СД-9-1')
                )
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Сидоров Семен Семенович', '2026-СД-0901-9', '26390001',
                    '2026', 'Сидоров', 'Семен Семенович', 'sidorov@example.test', 1
                )
            )

        sync_summary = manticore.sync_enrollment_candidates_from_ready_abiturients('2026')
        self.assertEqual(sync_summary['created'], 1)

        with sqlite3.connect(manticore.DB_PATH) as conn:
            candidate_id = conn.execute(
                'SELECT id FROM enrollment_candidates WHERE login=?',
                ('26390001',)
            ).fetchone()[0]

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/abiturients_to_students')
        csrf_token = self.csrf_from_response(get_response)
        response = client.post(
            '/abiturients_to_students',
            data={
                'csrf_token': csrf_token,
                'group_year': '2026',
                'cohort1': '26СД-9-1',
                'auto_split': '1',
                'candidate_ids': [str(candidate_id)],
            },
            follow_redirects=True
        )

        self.assertEqual(response.status_code, 200)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            student_group = conn.execute(
                'SELECT cohort1, cohort2 FROM students WHERE username=?',
                ('26390001',)
            ).fetchone()

        self.assertEqual(student_group, ('26СД-9-2', 'СД-9'))

    def test_stale_enrollment_candidate_is_removed_when_source_is_unpaid(self):
        rules = manticore.get_default_login_generation_rules()
        rules['require_enrollment_order'] = False
        manticore.save_login_generation_settings(rules, setup_completed=True, updated_by='test')

        with sqlite3.connect(manticore.DB_PATH) as conn:
            cur = conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Братыкина Алина Алексеевна', '2026-СтД-0046-11И', '26811i003',
                    '2026', 'Братыкина', 'Алина Алексеевна', 'alina@example.test', 0
                )
            )
            abiturient_id = cur.lastrowid
            conn.execute(
                '''
                INSERT INTO enrollment_candidates
                    (abiturient_id, campaign_year, fio, dogovor, login, fam, imotch, email, specialty, specialty_key, base_label)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    abiturient_id, '2026', 'Братыкина Алина Алексеевна',
                    '2026-СтД-0046-11И', '26811i003', 'Братыкина',
                    'Алина Алексеевна', 'alina@example.test',
                    '31.02.07 «Стоматологическое дело»', 'стд', '11И'
                )
            )

        candidates = manticore.get_enrollment_candidates('2026')

        self.assertEqual(candidates, [])
        with sqlite3.connect(manticore.DB_PATH) as conn:
            candidate_count = conn.execute(
                'SELECT COUNT(*) FROM enrollment_candidates WHERE login=?',
                ('26811i003',)
            ).fetchone()[0]
        self.assertEqual(candidate_count, 0)

    def test_login_group_parser_builds_academic_groups(self):
        examples = {
            '26111001': '26ЛД-11-1',
            '2619001': '26ЛД-9-1',
            '26111i001': '26ЛД-11И-1',
            '2619i001': '26ЛД-9И-1',
            '261im001': '26ЛД-М-1',
        }
        for login, group_name in examples.items():
            with self.subTest(login=login):
                parsed = manticore.parse_login_group_target(login, '2026')
                self.assertTrue(parsed['ok'])
                self.assertEqual(parsed['target_group'], group_name)

    def test_derive_cohort2_from_supported_academic_groups(self):
        examples = {
            '26АД-9-1': 'АД-9',
            '26АД-11-1': 'АД-11',
            '26ЛД-9-1': 'ЛД-9',
            '26ЛД-11-2': 'ЛД-11',
            '26СД-9-1': 'СД-9',
            '26СД-9-3': 'СД-9',
            '27СД-9-1': 'СД-9',
            '26СД-11-1': 'СД-11',
            '26СтД-9-1': 'СтД-9',
            '26СтО-11-2': 'СтО-11',
            '26СтП-9-1': 'СтП-9',
            '26ФМ-11-1': 'ФМ-11',
            '26СД-9И-1': 'СД-9',
            '26СД-11И-1': 'СД-11',
            '26СД-11i-2': 'СД-11',
            '26ФМ-М-1': 'ФМ-11',
            '26ФМ-m-2': 'ФМ-11',
        }
        for cohort1, cohort2 in examples.items():
            with self.subTest(cohort1=cohort1):
                self.assertEqual(manticore.derive_cohort2(cohort1), cohort2)

    def test_derive_cohort2_returns_none_for_unsupported_groups(self):
        examples = [
            '',
            'не группа',
            '26ИТ-9-1',
            '26СД-10-1',
            '26СД-9',
            '26СД-9-А',
            '26ЛабД-11-1',
        ]
        for cohort1 in examples:
            with self.subTest(cohort1=cohort1):
                self.assertIsNone(manticore.derive_cohort2(cohort1))

        self.assertEqual(manticore.derive_cohort2(' 26СД-9-1 '), 'СД-9')

    def test_migration_page_previews_cohort2_for_selected_group(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute('DELETE FROM groups')
            for group_name in ('26СД-9И-1', '26ФМ-М-1'):
                conn.execute(
                    'INSERT INTO groups (name, group_year) VALUES (?, ?)',
                    (group_name, '2026')
                )

        client = manticore.app.test_client()
        self.login_session(client)
        response = client.get('/abiturients_to_students')
        body = response.get_data(as_text=True)

        self.assertEqual(response.status_code, 200)
        self.assertIn('id="cohort2_preview"', body)
        self.assertIn('value="26СД-9И-1" data-cohort2="СД-9"', body)
        self.assertIn('value="26ФМ-М-1" data-cohort2="ФМ-11"', body)

    def test_init_db_adds_and_backfills_cohort2_for_old_students_table(self):
        original_db_path = manticore.DB_PATH
        old_db_path = os.path.join(TEST_UPLOAD_DIR, 'old_students_schema.db')
        if os.path.exists(old_db_path):
            os.remove(old_db_path)
        try:
            manticore.DB_PATH = old_db_path
            with sqlite3.connect(manticore.DB_PATH) as conn:
                conn.execute(
                    '''
                    CREATE TABLE students (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        username TEXT UNIQUE,
                        password TEXT,
                        email TEXT,
                        firstname TEXT,
                        lastname TEXT,
                        cohort1 TEXT
                    )
                    '''
                )
                conn.execute(
                    '''
                    INSERT INTO students (username, password, email, firstname, lastname, cohort1)
                    VALUES (?, ?, ?, ?, ?, ?)
                    ''',
                    ('student_old', 'cron', 'old@example.test', 'Иван', 'Иванов', '26СД-9-1')
                )
                conn.execute(
                    '''
                    INSERT INTO students (username, password, email, firstname, lastname, cohort1)
                    VALUES (?, ?, ?, ?, ?, ?)
                    ''',
                    ('student_unknown', 'cron', 'unknown@example.test', 'Петр', 'Петров', '26ЛабД-11-1')
                )

            manticore.init_db()

            with sqlite3.connect(manticore.DB_PATH) as conn:
                columns = manticore.get_table_columns(conn, 'students')
                rows = {
                    row[0]: row[1:]
                    for row in conn.execute('SELECT username, password, cohort1, cohort2 FROM students ORDER BY username')
                }
                conn.execute('UPDATE students SET cohort2=? WHERE username=?', ('ЛД-9', 'student_old'))

            self.assertIn('cohort2', columns)
            self.assertEqual(rows['student_old'], ('cron', '26СД-9-1', 'СД-9'))
            self.assertEqual(rows['student_unknown'], ('cron', '26ЛабД-11-1', None))

            manticore.init_db()
            with sqlite3.connect(manticore.DB_PATH) as conn:
                preserved = conn.execute(
                    'SELECT cohort2 FROM students WHERE username=?',
                    ('student_old',)
                ).fetchone()[0]
                row_count = conn.execute('SELECT COUNT(*) FROM students').fetchone()[0]

            self.assertEqual(preserved, 'ЛД-9')
            self.assertEqual(row_count, 2)
        finally:
            manticore.DB_PATH = original_db_path

    def test_login_group_distribution_preview_and_confirm(self):
        people = [
            (
                'Альфов Алексей Алексеевич', '2026-ЛД-0600-11', '26111001',
                'Альфов', 'Алексей Алексеевич', 'alfov@example.test', '26ЛД-11-1'
            ),
            (
                'Бетов Борис Борисович', '2026-ЛД-0601-9', '2619001',
                'Бетов', 'Борис Борисович', 'betov@example.test', '26ЛД-9-1'
            ),
            (
                'Ветров Виктор Викторович', '2026-ЛД-0602-11И', '26111i001',
                'Ветров', 'Виктор Викторович', 'vetrov@example.test', '26ЛД-11И-1'
            ),
            (
                'Громов Глеб Глебович', '2026-ЛД-0603-9И', '2619i001',
                'Громов', 'Глеб Глебович', 'gromov@example.test', '26ЛД-9И-1'
            ),
            (
                'Дымов Денис Денисович', '2026-ЛД-0604-М', '261im001',
                'Дымов', 'Денис Денисович', 'dymov@example.test', '26ЛД-М-1'
            ),
        ]
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute('DELETE FROM groups')
            for fio, dogovor, login, fam, imotch, email, _group_name in people:
                conn.execute(
                    '''
                    INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    ''',
                    (fio, dogovor, login, '2026', fam, imotch, email, 1)
                )

        sync_summary = manticore.sync_enrollment_candidates_from_ready_abiturients('2026')
        self.assertEqual(sync_summary['created'], len(people))

        order_path = os.path.join(TEST_UPLOAD_DIR, 'login_distribution_order.xlsx')
        pd.DataFrame([
            {
                'ФИО': fio,
                'Специальность': '31.02.01 «Лечебное дело»',
                'Группа': '',
                'Номер приказа': '456-у',
                'Дата приказа': '2026-08-20',
            }
            for fio, _dogovor, _login, _fam, _imotch, _email, _group_name in people
        ]).to_excel(order_path, index=False)
        order_summary = manticore.apply_enrollment_order_import(order_path, '2026')
        self.assertEqual(order_summary['matched_count'], len(people))

        with sqlite3.connect(manticore.DB_PATH) as conn:
            candidate_ids = [
                str(row[0])
                for row in conn.execute('SELECT id FROM enrollment_candidates ORDER BY fio')
            ]

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/abiturients_to_students')
        csrf_token = self.csrf_from_response(get_response)
        preview_response = client.post(
            '/abiturients_to_students',
            data={
                'csrf_token': csrf_token,
                'group_year': '2026',
                'use_login_distribution': '1',
                'distribution_action': 'preview_login_groups',
                'candidate_ids': candidate_ids,
            }
        )
        self.assertEqual(preview_response.status_code, 200)
        preview_body = preview_response.get_data(as_text=True)
        self.assertIn('Предпросмотр распределения по логинам', preview_body)
        self.assertIn('Глобальная группа курса', preview_body)
        self.assertIn('id="login_preview_search"', preview_body)
        self.assertIn('id="login_preview_source"', preview_body)
        self.assertIn('id="login_preview_status"', preview_body)
        self.assertIn('data-sort-target="login_distribution_preview_body"', preview_body)
        for _fio, _dogovor, _login, _fam, _imotch, _email, group_name in people:
            self.assertIn(group_name, preview_body)

        csrf_token = self.csrf_from_response(preview_response)
        confirm_response = client.post(
            '/abiturients_to_students',
            data={
                'csrf_token': csrf_token,
                'group_year': '2026',
                'use_login_distribution': '1',
                'distribution_action': 'confirm_login_groups',
                'candidate_ids': candidate_ids,
            },
            follow_redirects=True
        )
        self.assertEqual(confirm_response.status_code, 200)
        self.assertIn('Альфов', confirm_response.get_data(as_text=True))

        expected_groups = {
            login: (group_name, manticore.derive_cohort2(group_name))
            for _fio, _dogovor, login, _fam, _imotch, _email, group_name in people
        }
        with sqlite3.connect(manticore.DB_PATH) as conn:
            students = {
                row[0]: (row[1], row[2])
                for row in conn.execute('SELECT username, cohort1, cohort2 FROM students')
            }
            candidate_count = conn.execute('SELECT COUNT(*) FROM enrollment_candidates').fetchone()[0]
            created_groups = {
                row[0]
                for row in conn.execute('SELECT name FROM groups WHERE group_year=?', ('2026',))
            }
            movements = {
                row[0]: row[1:]
                for row in conn.execute(
                    '''
                    SELECT username, movement_type, order_source, order_number, order_date
                    FROM student_group_transfers
                    '''
                )
            }

        self.assertEqual(students, expected_groups)
        self.assertEqual(candidate_count, 0)
        self.assertTrue({group for group, _cohort2 in expected_groups.values()}.issubset(created_groups))
        self.assertEqual(len(movements), len(people))
        for movement_type, order_source, order_number, order_date in movements.values():
            self.assertEqual(movement_type, 'enrollment')
            self.assertEqual(order_source, 'enrollment_order')
            self.assertEqual(order_number, '456-у')
            self.assertEqual(order_date, '2026-08-20')

    def test_login_group_distribution_respects_capacity_and_order_group(self):
        people = [
            (
                'Иванов Иван Иванович', '2026-ЛД-0700-11', '26111026',
                'Иванов', 'Иван Иванович', 'ivanov@example.test', ''
            ),
            (
                'Петров Петр Петрович', '2026-ЛД-0701-11', '26111027',
                'Петров', 'Петр Петрович', 'petrov@example.test', '26ЛД-9-1'
            ),
        ]
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute('DELETE FROM groups')
            conn.execute('INSERT INTO groups (name, group_year) VALUES (?, ?)', ('26ЛД-11-1', '2026'))
            for index in range(manticore.MAX_GROUP_STUDENTS):
                conn.execute(
                    '''
                    INSERT INTO students (username, password, email, firstname, lastname, cohort1)
                    VALUES (?, ?, ?, ?, ?, ?)
                    ''',
                    (f'filled{index:03d}', 'cron', f'filled{index}@example.test', 'Имя', 'Фамилия', '26ЛД-11-1')
                )
            for fio, dogovor, login, fam, imotch, email, _order_group in people:
                conn.execute(
                    '''
                    INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    ''',
                    (fio, dogovor, login, '2026', fam, imotch, email, 1)
                )

        manticore.sync_enrollment_candidates_from_ready_abiturients('2026')

        order_path = os.path.join(TEST_UPLOAD_DIR, 'login_distribution_priority_order.xlsx')
        pd.DataFrame([
            {
                'ФИО': fio,
                'Специальность': '31.02.01 «Лечебное дело»',
                'Группа': order_group,
                'Номер приказа': '789-у',
                'Дата приказа': '2026-08-21',
            }
            for fio, _dogovor, _login, _fam, _imotch, _email, order_group in people
        ]).to_excel(order_path, index=False)
        manticore.apply_enrollment_order_import(order_path, '2026')

        with sqlite3.connect(manticore.DB_PATH) as conn:
            refresh_candidates = conn.execute('SELECT id FROM enrollment_candidates ORDER BY fio').fetchall()
            candidate_rows = manticore.get_selected_enrollment_candidate_rows(
                conn,
                '2026',
                [str(row[0]) for row in refresh_candidates]
            )
            plan = manticore.build_login_group_distribution_plan(conn, candidate_rows, '2026', '2026')

        rows_by_login = {row['login']: row for row in plan['rows']}
        self.assertEqual(rows_by_login['26111026']['target_group'], '26ЛД-11-2')
        self.assertEqual(rows_by_login['26111026']['source'], 'Логин')
        self.assertEqual(rows_by_login['26111027']['target_group'], '26ЛД-9-1')
        self.assertEqual(rows_by_login['26111027']['source'], 'Приказ')
        self.assertTrue(plan['summary']['can_confirm'])

    def test_docx_enrollment_order_paragraph_blocks_use_official_specialty_names(self):
        from docx import Document

        order_path = os.path.join(TEST_UPLOAD_DIR, 'paragraph_order.docx')
        document = Document()
        table = document.add_table(rows=2, cols=3)
        table.cell(0, 0).text = '19 сентября 2025 г.'
        table.cell(0, 2).text = '№ 1909-У'
        table.cell(1, 1).text = 'Ставрополь'
        document.add_paragraph('Приложение №1')
        document.add_paragraph('к приказу от 19.09.2025 № 1909-У')
        document.add_paragraph('СПО 33.02.01 «Фармация» очная форма обучения (Приложение № 4).')
        document.add_paragraph('Иванов Иван Иванович')
        document.add_paragraph('СПО 31.02.07 «Стоматологическое дело» очная форма')
        document.add_paragraph('Петров Петр Петрович')
        document.save(order_path)

        df = manticore.read_docx_enrollment_order_dataframe(order_path)
        rows = df.to_dict(orient='records')

        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]['ФИО'], 'Иванов Иван Иванович')
        self.assertEqual(rows[0]['Специальность'], '33.02.01 «Фармация»')
        self.assertEqual(rows[0]['Номер приказа'], '1909-У')
        self.assertEqual(rows[0]['Дата приказа'], '19.09.2025')
        self.assertEqual(rows[1]['ФИО'], 'Петров Петр Петрович')
        self.assertEqual(rows[1]['Специальность'], '31.02.07 «Стоматологическое дело»')
        self.assertEqual(
            manticore.normalize_specialty_key('31.02.07 «Стоматологическое дело»'),
            manticore.normalize_specialty_key('СтД')
        )

    def test_enrollment_order_preview_can_fix_candidate_fio_from_order(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Иванов Иван Иванвич', '2026-ФМ-0550-11', '26611055',
                    '2026', 'Иванов', 'Иван Иванвич', 'ivanov@example.test', 1
                )
            )

        manticore.sync_enrollment_candidates_from_ready_abiturients('2026')
        pending_order_path = manticore.make_temp_upload_path('xlsx', prefix=manticore.PENDING_ENROLLMENT_ORDER_IMPORT_PREFIX)
        pd.DataFrame([
            {
                'ФИО': 'Иванов Иван Иванович',
                'Специальность': '33.02.01 «Фармация»',
                'Номер приказа': '055-У',
                'Дата приказа': '2026-08-25',
            }
        ]).to_excel(pending_order_path, index=False)

        plan_df, summary = manticore.build_enrollment_order_import_plan(pending_order_path, '2026')
        row = plan_df.iloc[0]
        self.assertEqual(summary['matched_count'], 0)
        self.assertEqual(summary['fio_suggestion_count'], 1)
        self.assertEqual(row['suggested_candidate_fio'], 'Иванов Иван Иванвич')

        with sqlite3.connect(manticore.DB_PATH) as conn:
            candidate_id = conn.execute(
                'SELECT id FROM enrollment_candidates WHERE login=?',
                ('26611055',)
            ).fetchone()[0]

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/file_work')
        csrf_token = self.csrf_from_response(get_response)
        fix_response = client.post(
            '/enrollment_order_preview/fix_candidate_fio',
            data={
                'csrf_token': csrf_token,
                'pending_enrollment_order_import': os.path.basename(pending_order_path),
                'row_number': str(row['_row_number']),
                'candidate_id': str(candidate_id),
            }
        )
        fix_body = fix_response.get_data(as_text=True)

        self.assertEqual(fix_response.status_code, 200)
        self.assertIn('Иванов Иван Иванович', fix_body)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            candidate = conn.execute(
                'SELECT fio, fam, imotch FROM enrollment_candidates WHERE id=?',
                (candidate_id,)
            ).fetchone()
            abiturient = conn.execute(
                'SELECT fio, fam, imotch FROM abiturients WHERE login=?',
                ('26611055',)
            ).fetchone()
        self.assertEqual(candidate, ('Иванов Иван Иванович', 'Иванов', 'Иван Иванович'))
        self.assertEqual(abiturient, ('Иванов Иван Иванович', 'Иванов', 'Иван Иванович'))

        refreshed_df, refreshed_summary = manticore.build_enrollment_order_import_plan(pending_order_path, '2026')
        self.assertEqual(refreshed_summary['matched_count'], 1)
        self.assertTrue(bool(refreshed_df.iloc[0]['has_candidate']))

    def test_enrollment_order_preview_flags_candidate_when_source_abiturient_fio_differs(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Братыпкина Алина Алексеевна', '2026-СтД-0046-11И', '26811003',
                    '2026', 'Братыпкина', 'Алина Алексеевна', 'bratypkina@example.test', 1
                )
            )

        manticore.sync_enrollment_candidates_from_ready_abiturients('2026')
        with sqlite3.connect(manticore.DB_PATH) as conn:
            candidate_id = conn.execute(
                'SELECT id FROM enrollment_candidates WHERE login=?',
                ('26811003',)
            ).fetchone()[0]
            conn.execute(
                '''
                UPDATE enrollment_candidates
                SET fio=?, fam=?, imotch=?
                WHERE id=?
                ''',
                ('Братыкина Алина Алексеевна', 'Братыкина', 'Алина Алексеевна', candidate_id)
            )

        order_path = os.path.join(TEST_UPLOAD_DIR, 'source_mismatch_order.xlsx')
        pd.DataFrame([
            {
                'ФИО': 'Братыкина Алина Алексеевна',
                'Специальность': '31.02.07 «Стоматологическое дело»',
                'Номер приказа': '3009-У',
                'Дата приказа': '30.09.2025',
            }
        ]).to_excel(order_path, index=False)

        plan_df, summary = manticore.build_enrollment_order_import_plan(order_path, '2026')
        row = plan_df.iloc[0]

        self.assertEqual(summary['matched_count'], 0)
        self.assertEqual(summary['import_count'], 0)
        self.assertEqual(summary['fio_suggestion_count'], 1)
        self.assertEqual(row['import_action'], 'fio_review')
        self.assertFalse(bool(row['has_candidate']))
        self.assertEqual(row['suggested_candidate_fio'], 'Братыпкина Алина Алексеевна')
        self.assertEqual(row['suggested_candidate_actual_fio'], 'Братыкина Алина Алексеевна')

        preview_row = manticore.enrollment_order_preview_rows(plan_df)[0]
        self.assertEqual(preview_row['action_label'], 'Проверить ФИО')
        self.assertEqual(preview_row['badge_class'], 'status-warning')

    def test_multiple_enrollment_orders_accumulate_without_clearing_previous_uploads(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Иванов Иван Иванович', '2026-ФМ-0600-11', '26611060',
                    '2026', 'Иванов', 'Иван Иванович', 'ivanov@example.test', 1
                )
            )
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Петров Петр Петрович', '2026-СД-0601-9', '26390601',
                    '2026', 'Петров', 'Петр Петрович', 'petrov@example.test', 1
                )
            )

        manticore.sync_enrollment_candidates_from_ready_abiturients('2026')
        first_order_path = os.path.join(TEST_UPLOAD_DIR, 'first_enrollment_order.xlsx')
        second_order_path = os.path.join(TEST_UPLOAD_DIR, 'second_enrollment_order.xlsx')
        pd.DataFrame([
            {
                'ФИО': 'Иванов Иван Иванович',
                'Специальность': '33.02.01 «Фармация»',
                'Номер приказа': '001-У',
                'Дата приказа': '2026-08-20',
            }
        ]).to_excel(first_order_path, index=False)
        pd.DataFrame([
            {
                'ФИО': 'Петров Петр Петрович',
                'Специальность': '34.02.01 «Сестринское дело»',
                'Номер приказа': '002-У',
                'Дата приказа': '2026-08-30',
            }
        ]).to_excel(second_order_path, index=False)

        first_summary = manticore.apply_enrollment_order_import(first_order_path, '2026')
        second_summary = manticore.apply_enrollment_order_import(second_order_path, '2026')

        self.assertEqual(first_summary['matched_count'], 1)
        self.assertEqual(second_summary['matched_count'], 1)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            order_count = conn.execute('SELECT COUNT(*) FROM enrollment_orders WHERE campaign_year=?', ('2026',)).fetchone()[0]
            statuses = conn.execute(
                '''
                SELECT fio, verification_status
                FROM enrollment_candidates
                WHERE campaign_year=?
                ORDER BY fio
                ''',
                ('2026',)
            ).fetchall()

        self.assertEqual(order_count, 2)
        self.assertEqual(statuses, [
            ('Иванов Иван Иванович', 'verified'),
            ('Петров Петр Петрович', 'verified'),
        ])
        uploads = manticore.get_enrollment_order_uploads('2026')
        self.assertEqual(len(uploads), 2)
        self.assertEqual(uploads[0]['original_filename'], 'second_enrollment_order.xlsx')
        self.assertEqual(uploads[1]['original_filename'], 'first_enrollment_order.xlsx')

        first_upload = manticore.get_enrollment_order_upload(first_summary['upload_id'])
        first_upload_path = manticore.get_stored_enrollment_order_path(first_upload['stored_filename'])
        self.assertTrue(os.path.exists(first_upload_path))

        client = manticore.app.test_client()
        self.login_session(client)
        detail_response = client.get(f"/enrollment_order_uploads/{first_summary['upload_id']}")
        detail_body = detail_response.get_data(as_text=True)
        self.assertEqual(detail_response.status_code, 200)
        self.assertIn('first_enrollment_order.xlsx', detail_body)
        self.assertIn('Иванов Иван Иванович', detail_body)

        download_response = client.get(f"/enrollment_order_uploads/{first_summary['upload_id']}/download")
        self.assertEqual(download_response.status_code, 200)
        self.assertIn('attachment', download_response.headers.get('Content-Disposition', ''))
        download_response.close()

        get_response = client.get('/file_work')
        csrf_token = self.csrf_from_response(get_response)
        delete_response = client.post(
            f"/enrollment_order_uploads/{first_summary['upload_id']}/delete",
            data={'csrf_token': csrf_token},
            follow_redirects=True
        )
        self.assertEqual(delete_response.status_code, 200)
        self.assertFalse(os.path.exists(first_upload_path))
        with sqlite3.connect(manticore.DB_PATH) as conn:
            order_count_after_delete = conn.execute(
                'SELECT COUNT(*) FROM enrollment_orders WHERE campaign_year=?',
                ('2026',)
            ).fetchone()[0]
            statuses_after_delete = conn.execute(
                '''
                SELECT fio, verification_status
                FROM enrollment_candidates
                WHERE campaign_year=?
                ORDER BY fio
                ''',
                ('2026',)
            ).fetchall()
        self.assertEqual(order_count_after_delete, 1)
        self.assertEqual(statuses_after_delete, [
            ('Иванов Иван Иванович', 'missing_in_order'),
            ('Петров Петр Петрович', 'verified'),
        ])

    def test_students_upload_reports_row_errors(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO students (username, password, email, firstname, lastname, cohort1)
                VALUES (?, ?, ?, ?, ?, ?)
                ''',
                ('student_existing', 'secret', 'old@example.test', 'Старый', 'Студент', '26ФМ-11-1')
            )

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/file_work')
        csrf_token = self.csrf_from_response(get_response)
        upload = io.BytesIO()
        pd.DataFrame([
            {
                'username': '',
                'password': 'pass',
                'email': 'missing-login@example.test',
                'firstname': 'Иван',
                'lastname': 'Иванов',
                'cohort1': '26ФМ-11-1',
            },
            {
                'username': 'student_bad_mail',
                'password': 'pass',
                'email': 'bad-mail',
                'firstname': 'Петр',
                'lastname': 'Петров',
                'cohort1': '26ФМ-11-1',
            },
            {
                'username': 'student_existing',
                'password': 'pass',
                'email': 'duplicate@example.test',
                'firstname': 'Дубль',
                'lastname': 'Студент',
                'cohort1': '26ФМ-11-1',
            },
            {
                'username': 'student_new',
                'password': 'pass',
                'email': 'new@example.test',
                'firstname': 'Новый',
                'lastname': 'Студент',
                'cohort1': '26ФМ-11-1',
            },
        ]).to_excel(upload, index=False)
        upload.seek(0)

        response = client.post(
            '/students_upload',
            data={
                'csrf_token': csrf_token,
                'file': (upload, 'students.xlsx'),
            },
            content_type='multipart/form-data'
        )
        body = response.get_data(as_text=True)

        self.assertEqual(response.status_code, 200)
        self.assertIn('Предпросмотр загрузки студентов', body)
        self.assertIn('Отчет по проверке студентов', body)
        self.assertIn('Строка 2', body)
        self.assertIn('Не заполнено: Логин', body)
        self.assertIn('Строка 3', body)
        self.assertIn('Почта выглядит некорректно', body)
        self.assertIn('Строка 4', body)
        self.assertIn('будет перенесена в дубли студентов', body)
        self.assertIn('К добавлению: 1', body)
        self.assertIn('Подтвердить загрузку студентов', body)
        self.assertIn('status-badge status-success', body)
        self.assertIn('status-badge status-warning', body)

        with sqlite3.connect(manticore.DB_PATH) as conn:
            new_count = conn.execute('SELECT COUNT(*) FROM students WHERE username=?', ('student_new',)).fetchone()[0]
            duplicate_count = conn.execute(
                'SELECT COUNT(*) FROM students_duplicates WHERE username=?',
                ('student_existing',)
            ).fetchone()[0]
        self.assertEqual(new_count, 0)
        self.assertEqual(duplicate_count, 0)

        pending_match = re.search(r'name="pending_students_import" value="([^"]+)"', body)
        self.assertIsNotNone(pending_match)
        confirm_csrf = self.csrf_from_response(response)
        confirm_response = client.post(
            '/students_upload',
            data={
                'csrf_token': confirm_csrf,
                'students_import_action': 'confirm',
                'pending_students_import': pending_match.group(1),
            },
            follow_redirects=True
        )
        confirm_body = confirm_response.get_data(as_text=True)

        self.assertEqual(confirm_response.status_code, 200)
        self.assertIn('Загрузка студентов завершена', confirm_body)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            new_student = conn.execute('SELECT cohort1, cohort2 FROM students WHERE username=?', ('student_new',)).fetchone()
            bad_count = conn.execute('SELECT COUNT(*) FROM students WHERE username=?', ('student_bad_mail',)).fetchone()[0]
            duplicate_student = conn.execute(
                'SELECT cohort1, cohort2 FROM students_duplicates WHERE username=?',
                ('student_existing',)
            ).fetchone()
        self.assertEqual(new_student, ('26ФМ-11-1', 'ФМ-11'))
        self.assertEqual(bad_count, 0)
        self.assertEqual(duplicate_student, ('26ФМ-11-1', 'ФМ-11'))

        students_response = client.get('/students_list')
        students_body = students_response.get_data(as_text=True)
        self.assertIn('status-badge status-success', students_body)
        self.assertIn('Есть почта', students_body)

    def test_students_import_rejects_conflicting_cohort2(self):
        file_path = os.path.join(TEST_UPLOAD_DIR, 'students_conflicting_cohort2.xlsx')
        pd.DataFrame([
            {
                'username': 'student_bad_cohort2',
                'password': 'pass',
                'email': 'student_bad_cohort2@example.test',
                'firstname': 'Иван',
                'lastname': 'Иванов',
                'cohort1': '26СД-9-1',
                'cohort2': 'ЛД-9',
            }
        ]).to_excel(file_path, index=False)

        plan_df, summary = manticore.build_students_import_plan(file_path)

        self.assertEqual(summary['ready_count'], 0)
        self.assertEqual(summary['skipped_count'], 1)
        self.assertEqual(plan_df.loc[0, 'import_action'], 'skip')
        self.assertIn('ожидается СД-9', summary['errors'][0]['message'])

    def test_course_groups_toggle_disables_cohort2_import_checks_and_export_column(self):
        rules = manticore.get_default_login_generation_rules()
        rules['use_course_groups'] = False
        manticore.save_login_generation_settings(rules, setup_completed=True, updated_by='test')

        file_path = os.path.join(TEST_UPLOAD_DIR, 'students_without_course_groups.xlsx')
        pd.DataFrame([
            {
                'username': 'student_no_course_group',
                'password': 'pass',
                'email': 'student_no_course_group@example.test',
                'firstname': 'Иван',
                'lastname': 'Иванов',
                'cohort1': '26СД-9-1',
                'cohort2': 'ЛД-9',
            }
        ]).to_excel(file_path, index=False)

        plan_df, summary = manticore.build_students_import_plan(file_path)

        self.assertEqual(summary['ready_count'], 1)
        self.assertEqual(summary['skipped_count'], 0)
        self.assertEqual(summary['errors'], [])
        self.assertEqual(plan_df.loc[0, 'cohort2'], '')

        import_summary = manticore.apply_students_import(file_path)
        self.assertEqual(import_summary['inserted_count'], 1)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            stored = conn.execute(
                'SELECT cohort1, cohort2 FROM students WHERE username=?',
                ('student_no_course_group',)
            ).fetchone()
        self.assertEqual(stored, ('26СД-9-1', ''))

        with manticore.app.test_request_context('/'):
            report = manticore.get_data_quality_report('2026')
        students_section = next(section for section in report['sections'] if section['title'] == 'Студенты')
        self.assertFalse(any(item['id'] == 'students-cohort2-mismatch' for item in students_section['checks']))

        client = manticore.app.test_client()
        self.login_session(client)
        card_response = client.get('/person/student/student_no_course_group')
        self.assertEqual(card_response.status_code, 200)
        card_body = card_response.get_data(as_text=True)
        self.assertNotIn('Глобальная группа курса', card_body)

        export_response = client.get('/students/download')
        self.assertEqual(export_response.status_code, 200)
        exported = pd.read_excel(io.BytesIO(export_response.data))
        self.assertNotIn('cohort2', exported.columns)

    def test_campaign_archive_page_toggles_status(self):
        client = manticore.app.test_client()
        self.login_session(client)

        get_response = client.get('/campaigns')
        csrf_token = self.csrf_from_response(get_response)
        post_response = client.post(
            '/campaigns',
            data={'csrf_token': csrf_token, 'campaign_year': '2026', 'is_archived': '1'},
            follow_redirects=True
        )

        self.assertEqual(post_response.status_code, 200)
        self.assertTrue(manticore.is_campaign_archived('2026'))

    def test_campaign_page_creates_new_campaign_for_switchers(self):
        client = manticore.app.test_client()
        self.login_session(client)

        get_response = client.get('/campaigns')
        csrf_token = self.csrf_from_response(get_response)
        post_response = client.post(
            '/campaigns',
            data={
                'csrf_token': csrf_token,
                'campaign_action': 'create',
                'new_campaign_year': '2031',
            },
            follow_redirects=True
        )
        page_body = post_response.get_data(as_text=True)

        self.assertEqual(post_response.status_code, 200)
        self.assertIn('Кампания 2031 создана', page_body)
        self.assertIn('<td>2031</td>', page_body)
        self.assertIn('2031', manticore.get_campaign_years())
        self.assertIn('2031', manticore.get_group_years(include_base=True))
        with client.session_transaction() as session:
            self.assertEqual(session.get('campaign_year'), '2031')

        abiturients_response = client.get('/abiturients')
        abiturients_body = abiturients_response.get_data(as_text=True)
        self.assertIn('<option value="2031" selected>2031</option>', abiturients_body)

    def test_campaign_page_pins_and_restores_automatic_default(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO campaign_settings
                    (campaign_year, is_archived, created_at, created_by)
                VALUES (?, 0, datetime('now', 'localtime'), ?)
                ''',
                ('2031', 'admin')
            )

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/campaigns')
        csrf_token = self.csrf_from_response(get_response)
        pin_response = client.post(
            '/campaigns',
            data={
                'csrf_token': csrf_token,
                'campaign_action': 'set_active',
                'campaign_year': '2026',
            },
            follow_redirects=True
        )
        pin_body = pin_response.get_data(as_text=True)

        self.assertEqual(pin_response.status_code, 200)
        self.assertIn('Кампания 2026 закреплена как активная', pin_body)
        self.assertIn('По умолчанию', pin_body)
        self.assertEqual(manticore.get_pinned_campaign_year(), '2026')
        with sqlite3.connect(manticore.DB_PATH) as conn:
            active_rows = conn.execute(
                'SELECT campaign_year FROM campaign_settings WHERE is_active=1'
            ).fetchall()
        self.assertEqual(active_rows, [('2026',)])

        fresh_client = manticore.app.test_client()
        self.login_session(fresh_client)
        abiturients_response = fresh_client.get('/abiturients')
        abiturients_body = abiturients_response.get_data(as_text=True)
        self.assertIn('<option value="2026" selected>2026</option>', abiturients_body)
        self.assertNotIn('<option value="2031" selected>2031</option>', abiturients_body)

        switch_csrf_token = self.csrf_from_response(abiturients_response)
        switch_response = fresh_client.post(
            '/set_campaign',
            data={
                'csrf_token': switch_csrf_token,
                'campaign_year': '2031',
                'next': '/abiturients',
            },
            follow_redirects=True
        )
        self.assertIn(
            '<option value="2031" selected>2031</option>',
            switch_response.get_data(as_text=True)
        )

        another_client = manticore.app.test_client()
        self.login_session(another_client)
        another_body = another_client.get('/abiturients').get_data(as_text=True)
        self.assertIn('<option value="2026" selected>2026</option>', another_body)

        campaigns_response = client.get('/campaigns')
        clear_csrf_token = self.csrf_from_response(campaigns_response)
        clear_response = client.post(
            '/campaigns',
            data={
                'csrf_token': clear_csrf_token,
                'campaign_action': 'clear_active',
            },
            follow_redirects=True
        )

        self.assertEqual(clear_response.status_code, 200)
        self.assertIn('Теперь автоматически выбирается последняя кампания', clear_response.get_data(as_text=True))
        self.assertIsNone(manticore.get_pinned_campaign_year())
        automatic_client = manticore.app.test_client()
        self.login_session(automatic_client)
        automatic_body = automatic_client.get('/abiturients').get_data(as_text=True)
        self.assertIn('<option value="2031" selected>2031</option>', automatic_body)

    def test_pinned_campaign_cannot_be_archived(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO campaign_settings
                    (campaign_year, is_archived, is_active, active_at, active_by)
                VALUES (?, 0, 1, datetime('now', 'localtime'), ?)
                ''',
                ('2026', 'admin')
            )

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/campaigns')
        csrf_token = self.csrf_from_response(get_response)
        response = client.post(
            '/campaigns',
            data={
                'csrf_token': csrf_token,
                'campaign_action': 'toggle_archive',
                'campaign_year': '2026',
                'is_archived': '1',
            },
            follow_redirects=True
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn('Нельзя архивировать закрепленную активную кампанию', response.get_data(as_text=True))
        self.assertFalse(manticore.is_campaign_archived('2026'))

    def test_bulk_abiturients_marks_selected_as_paid(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                ''',
                ('Петров Петр Петрович', '2026-ФМ-0102-11', '26611012', '2026', 'Петров', 'Петр Петрович', 0)
            )
            abiturient_id = conn.execute('SELECT id FROM abiturients WHERE login=?', ('26611012',)).fetchone()[0]

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/abiturients')
        csrf_token = self.csrf_from_response(get_response)
        response = client.post(
            '/abiturients/bulk',
            data={
                'csrf_token': csrf_token,
                'bulk_action': 'mark_paid',
                'abiturient_ids': [str(abiturient_id)],
            },
            follow_redirects=True
        )

        self.assertEqual(response.status_code, 200)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            paid = conn.execute('SELECT paid FROM abiturients WHERE id=?', (abiturient_id,)).fetchone()[0]
        self.assertEqual(paid, 1)

    def test_bulk_abiturients_unpaid_removes_enrollment_candidate(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO abiturients (fio, dogovor, login, campaign_year, fam, imotch, email, paid)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'Козлова Кира Кирилловна', '2026-ФМ-0103-11', '26611013',
                    '2026', 'Козлова', 'Кира Кирилловна', 'kozlova@example.test', 1
                )
            )
            abiturient_id = conn.execute('SELECT id FROM abiturients WHERE login=?', ('26611013',)).fetchone()[0]

        sync_summary = manticore.sync_enrollment_candidates_from_ready_abiturients('2026')
        self.assertEqual(sync_summary['created'], 1)

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/abiturients')
        csrf_token = self.csrf_from_response(get_response)
        response = client.post(
            '/abiturients/bulk',
            data={
                'csrf_token': csrf_token,
                'bulk_action': 'mark_unpaid',
                'abiturient_ids': [str(abiturient_id)],
            },
            follow_redirects=True
        )

        self.assertEqual(response.status_code, 200)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            paid = conn.execute('SELECT paid FROM abiturients WHERE id=?', (abiturient_id,)).fetchone()[0]
            candidate_count = conn.execute(
                'SELECT COUNT(*) FROM enrollment_candidates WHERE login=?',
                ('26611013',)
            ).fetchone()[0]
        self.assertEqual(paid, 0)
        self.assertEqual(candidate_count, 0)

    def test_student_card_and_export_hide_password_for_non_admin(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO students (username, password, email, firstname, lastname, cohort1, cohort2, source_campaign_year)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                ('student010', 'secret-pass', 'student010@example.test', 'Петр', 'Петров', '26ФМ-11-1', 'ФМ-11', '2026')
            )

        client = manticore.app.test_client()
        self.login_session(client, username='viewer', role='viewer')
        card_response = client.get('/person/student/student010')
        export_response = client.get('/students/download')

        self.assertEqual(card_response.status_code, 200)
        card_body = card_response.get_data(as_text=True)
        self.assertNotIn('secret-pass', card_body)
        self.assertIn('Скрыт для безопасности', card_body)
        self.assertIn('Глобальная группа курса', card_body)
        self.assertIn('ФМ-11', card_body)
        self.assertEqual(export_response.status_code, 200)
        exported = pd.read_excel(io.BytesIO(export_response.data))
        self.assertEqual(exported.loc[0, 'password'], '******')
        self.assertIn('cohort2', exported.columns)
        self.assertEqual(exported.loc[0, 'cohort2'], 'ФМ-11')

    def test_student_group_transfer_uploads_pdf_and_updates_groups(self):
        pdf_bytes = b'%PDF-1.4\ntransfer order\n%%EOF'
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute('DELETE FROM groups')
            conn.execute(
                'INSERT INTO groups (name, group_year) VALUES (?, ?)',
                ('26ЛД-11-1', '2026')
            )
            conn.execute(
                'INSERT INTO groups (name, group_year) VALUES (?, ?)',
                ('26ФМ-11-1', '2026')
            )
            conn.execute(
                '''
                INSERT INTO students (username, password, email, firstname, lastname, cohort1, cohort2, source_campaign_year)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                ('student_transfer', 'secret-pass', 'student_transfer@example.test', 'Анна', 'Андреева', '26ЛД-11-1', 'ЛД-11', '2026')
            )

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/edit_student/student_transfer')
        self.assertEqual(get_response.status_code, 200)
        edit_body = get_response.get_data(as_text=True)
        self.assertIn('Перевод по приказу', edit_body)
        self.assertIn('<select name="new_cohort1"', edit_body)
        self.assertIn('value="26ФМ-11-1" data-cohort2="ФМ-11"', edit_body)
        self.assertIn('class="file-picker"', edit_body)
        self.assertNotIn('26ФМ-11-1 -> ФМ-11', edit_body)
        csrf_token = self.csrf_from_response(get_response)

        response = client.post(
            '/edit_student/student_transfer/transfer_group',
            data={
                'csrf_token': csrf_token,
                'new_cohort1': '26ФМ-11-1',
                'transfer_order_file': (io.BytesIO(pdf_bytes), 'transfer_order.pdf'),
            },
            content_type='multipart/form-data',
            follow_redirects=True
        )

        self.assertEqual(response.status_code, 200)
        body = response.get_data(as_text=True)
        self.assertIn('История перемещений', body)
        self.assertIn('student-movement-timeline', body)
        self.assertIn('Дата перемещения:', body)
        self.assertIn('Из группы:', body)
        self.assertIn('26ЛД-11-1', body)
        self.assertIn('26ФМ-11-1', body)
        self.assertIn('ФМ-11', body)
        self.assertIn('Скачать PDF', body)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            student = conn.execute(
                'SELECT cohort1, cohort2 FROM students WHERE username=?',
                ('student_transfer',)
            ).fetchone()
            transfer = conn.execute(
                '''
                SELECT id, old_cohort1, old_cohort2, new_cohort1, new_cohort2,
                       order_filename, order_original_filename
                FROM student_group_transfers
                WHERE username=?
                ''',
                ('student_transfer',)
            ).fetchone()
        self.assertEqual(student, ('26ФМ-11-1', 'ФМ-11'))
        self.assertIsNotNone(transfer)
        transfer_id, old_cohort1, old_cohort2, new_cohort1, new_cohort2, filename, original_filename = transfer
        self.assertEqual((old_cohort1, old_cohort2, new_cohort1, new_cohort2), ('26ЛД-11-1', 'ЛД-11', '26ФМ-11-1', 'ФМ-11'))
        self.assertEqual(original_filename, 'transfer_order.pdf')
        saved_path = os.path.join(TEST_UPLOAD_DIR, manticore.STUDENT_TRANSFER_ORDER_DIR, filename)
        self.assertTrue(os.path.exists(saved_path))

        download_response = client.get(f'/student_transfer_orders/{transfer_id}/download')
        self.assertEqual(download_response.status_code, 200)
        self.assertEqual(download_response.data, pdf_bytes)
        download_response.close()

        card_response = client.get('/person/student/student_transfer')
        self.assertEqual(card_response.status_code, 200)
        card_body = card_response.get_data(as_text=True)
        self.assertIn('История перемещений', card_body)
        self.assertIn('student-movement-timeline', card_body)
        self.assertIn('26ФМ-11-1', card_body)

    def test_student_group_transfer_allows_missing_pdf(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute('DELETE FROM groups')
            conn.execute(
                'INSERT INTO groups (name, group_year) VALUES (?, ?)',
                ('26ЛД-11-1', '2026')
            )
            conn.execute(
                'INSERT INTO groups (name, group_year) VALUES (?, ?)',
                ('26ФМ-11-1', '2026')
            )
            conn.execute(
                '''
                INSERT INTO students (username, password, email, firstname, lastname, cohort1, cohort2, source_campaign_year)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                ('student_transfer_without_pdf', 'secret-pass', 'student_transfer_without_pdf@example.test', 'Анна', 'Андреева', '26ЛД-11-1', 'ЛД-11', '2026')
            )

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/edit_student/student_transfer_without_pdf')
        csrf_token = self.csrf_from_response(get_response)

        response = client.post(
            '/edit_student/student_transfer_without_pdf/transfer_group',
            data={
                'csrf_token': csrf_token,
                'new_cohort1': '26ФМ-11-1',
            },
            follow_redirects=True
        )

        self.assertEqual(response.status_code, 200)
        body = response.get_data(as_text=True)
        self.assertIn('История перемещений', body)
        self.assertIn('26ЛД-11-1', body)
        self.assertIn('26ФМ-11-1', body)
        self.assertIn('PDF не прикреплен', body)
        self.assertNotIn('Скачать PDF', body)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            student = conn.execute(
                'SELECT cohort1, cohort2 FROM students WHERE username=?',
                ('student_transfer_without_pdf',)
            ).fetchone()
            transfer = conn.execute(
                '''
                SELECT old_cohort1, old_cohort2, new_cohort1, new_cohort2,
                       order_filename, order_original_filename, order_size
                FROM student_group_transfers
                WHERE username=?
                ''',
                ('student_transfer_without_pdf',)
            ).fetchone()

        self.assertEqual(student, ('26ФМ-11-1', 'ФМ-11'))
        self.assertEqual(
            transfer,
            ('26ЛД-11-1', 'ЛД-11', '26ФМ-11-1', 'ФМ-11', '', '', 0)
        )

    def test_student_group_transfer_without_course_groups_stores_blank_cohort2(self):
        rules = manticore.get_default_login_generation_rules()
        rules['use_course_groups'] = False
        manticore.save_login_generation_settings(rules, setup_completed=True, updated_by='test')

        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute('DELETE FROM groups')
            conn.execute(
                'INSERT INTO groups (name, group_year) VALUES (?, ?)',
                ('26ЛД-11-1', '2026')
            )
            conn.execute(
                'INSERT INTO groups (name, group_year) VALUES (?, ?)',
                ('26ИТ-11-1', '2026')
            )
            conn.execute(
                '''
                INSERT INTO students (username, password, email, firstname, lastname, cohort1, cohort2, source_campaign_year)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                ('student_transfer_no_course_group', 'secret-pass', 'student_transfer_no_course_group@example.test', 'Анна', 'Андреева', '26ЛД-11-1', 'ЛД-11', '2026')
            )

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/edit_student/student_transfer_no_course_group')
        self.assertEqual(get_response.status_code, 200)
        edit_body = get_response.get_data(as_text=True)
        self.assertNotIn('Глобальная группа курса:', edit_body)
        csrf_token = self.csrf_from_response(get_response)

        response = client.post(
            '/edit_student/student_transfer_no_course_group/transfer_group',
            data={
                'csrf_token': csrf_token,
                'new_cohort1': '26ИТ-11-1',
            },
            follow_redirects=True
        )

        self.assertEqual(response.status_code, 200)
        body = response.get_data(as_text=True)
        self.assertIn('Студент переведен в группу 26ИТ-11-1.', body)
        self.assertNotIn('Глобальная группа курса:', body)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            student = conn.execute(
                'SELECT cohort1, cohort2 FROM students WHERE username=?',
                ('student_transfer_no_course_group',)
            ).fetchone()
            transfer = conn.execute(
                '''
                SELECT old_cohort1, old_cohort2, new_cohort1, new_cohort2
                FROM student_group_transfers
                WHERE username=?
                ''',
                ('student_transfer_no_course_group',)
            ).fetchone()

        self.assertEqual(student, ('26ИТ-11-1', ''))
        self.assertEqual(transfer, ('26ЛД-11-1', 'ЛД-11', '26ИТ-11-1', ''))

    def test_student_group_transfer_rejects_full_target_group(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute('DELETE FROM groups')
            conn.execute(
                'INSERT INTO groups (name, group_year) VALUES (?, ?)',
                ('26ЛД-11-1', '2026')
            )
            conn.execute(
                'INSERT INTO groups (name, group_year) VALUES (?, ?)',
                ('26ФМ-11-1', '2026')
            )
            conn.execute(
                '''
                INSERT INTO students (username, password, email, firstname, lastname, cohort1, cohort2, source_campaign_year)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                ('student_full_transfer', 'secret-pass', 'student_full_transfer@example.test', 'Анна', 'Андреева', '26ЛД-11-1', 'ЛД-11', '2026')
            )
            for index in range(manticore.MAX_GROUP_STUDENTS):
                conn.execute(
                    '''
                    INSERT INTO students (username, password, email, firstname, lastname, cohort1, cohort2, source_campaign_year)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    ''',
                    (
                        f'filled_fm_transfer_{index:03d}', 'secret-pass',
                        f'filled_fm_transfer_{index:03d}@example.test',
                        'Имя', 'Фамилия', '26ФМ-11-1', 'ФМ-11', '2026'
                    )
                )

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/edit_student/student_full_transfer')
        csrf_token = self.csrf_from_response(get_response)

        response = client.post(
            '/edit_student/student_full_transfer/transfer_group',
            data={
                'csrf_token': csrf_token,
                'new_cohort1': '26ФМ-11-1',
                'transfer_order_file': (io.BytesIO(b'%PDF-1.4\nfull\n%%EOF'), 'transfer_order.pdf'),
            },
            content_type='multipart/form-data',
            follow_redirects=True
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn('Выбранная группа заполнена', response.get_data(as_text=True))
        with sqlite3.connect(manticore.DB_PATH) as conn:
            student_group = conn.execute(
                'SELECT cohort1, cohort2 FROM students WHERE username=?',
                ('student_full_transfer',)
            ).fetchone()
            transfer_count = conn.execute(
                'SELECT COUNT(*) FROM student_group_transfers WHERE username=?',
                ('student_full_transfer',)
            ).fetchone()[0]

        self.assertEqual(student_group, ('26ЛД-11-1', 'ЛД-11'))
        self.assertEqual(transfer_count, 0)

    def test_delete_student_does_not_return_to_archived_campaign(self):
        with sqlite3.connect(manticore.DB_PATH) as conn:
            conn.execute(
                '''
                INSERT INTO students
                    (username, password, email, firstname, lastname, cohort1, source_campaign_year, source_dogovor, source_fio)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    'student011', 'cron', 'student011@example.test', 'Семен', 'Сидоров', '26ФМ-11-1',
                    '2026', '2026-ФМ-0111-11', 'Сидоров Семен Семенович'
                )
            )
            conn.execute(
                '''
                INSERT INTO campaign_settings (campaign_year, is_archived, archived_at, archived_by)
                VALUES (?, ?, datetime('now', 'localtime'), ?)
                ''',
                ('2026', 1, 'admin')
            )

        client = manticore.app.test_client()
        self.login_session(client)
        get_response = client.get('/students_list')
        csrf_token = self.csrf_from_response(get_response)
        response = client.post(
            '/delete_student',
            data={'csrf_token': csrf_token, 'username': 'student011'},
            follow_redirects=True
        )

        self.assertEqual(response.status_code, 200)
        with sqlite3.connect(manticore.DB_PATH) as conn:
            student_count = conn.execute('SELECT COUNT(*) FROM students WHERE username=?', ('student011',)).fetchone()[0]
            abiturient_count = conn.execute('SELECT COUNT(*) FROM abiturients WHERE login=?', ('student011',)).fetchone()[0]
        self.assertEqual(student_count, 1)
        self.assertEqual(abiturient_count, 0)


if __name__ == '__main__':
    unittest.main()
